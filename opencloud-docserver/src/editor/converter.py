"""DOCX <-> HTML conversion using python-docx.

Stoic goals: preserve text, bold/italic/underline, headings, lists,
tables, and images. We do NOT attempt pagination or print-fidelity —
the editor is a web page, not a print preview.

HTML -> DOCX is lossy by nature (web HTML is richer than we map); we map
only the subset our editor produces, plus whatever reasonable tags appear.
"""

from __future__ import annotations

import base64
import io
import logging
import re
from html import escape
from html.parser import HTMLParser
from typing import Any

from docx import Document
from docx.enum.dml import MSO_COLOR_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.parts.story import StoryPart
from docx.shared import Emu, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# --------------------------------------------------------------------------
# Image handling (package binary <-> data URI)
# --------------------------------------------------------------------------

_EMU_PER_PX = 9525  # EMU per CSS pixel at 96 dpi

_PNG_MAGIC = b"\x89PNG\r\n\x1a\n"


def _data_uri(mime: str, content: bytes) -> str:
    """Encode image bytes as a self-contained ``data:`` URI for the editor."""
    b64 = base64.b64encode(content).decode("ascii")
    return f"data:{mime};base64,{b64}"


_DATA_URI_RE = re.compile(r"^data:([^;,\s]+)(;[^,]*)?,(.*)$", re.S)


def _decode_data_uri(src: str) -> tuple[str | None, bytes | None]:
    """Decode a ``data:`` URI into (mime, bytes).

    Returns (None, None) when the value is not an embeddable data URI (e.g.
    an http(s) src that a server-side converter cannot fetch).
    """
    if not src:
        return None, None
    m = _DATA_URI_RE.match(src)
    if not m:
        return None, None
    mime = (m.group(1) or "image/png").lower()
    meta = m.group(2) or ""
    payload = m.group(3)
    try:
        if "base64" in meta:
            content = base64.b64decode(payload, validate=True)
        else:
            from urllib.parse import unquote

            content = unquote(payload).encode("utf-8")
    except Exception:
        return None, None
    if not content:
        return None, None
    return mime, content


def _sniff_mime(data: bytes) -> str:
    """Detect an image media type from its magic bytes."""
    if data[:8] == _PNG_MAGIC:
        return "image/png"
    if data[:3] == b"GIF":
        return "image/gif"
    if data[:2] == b"\xff\xd8":
        return "image/jpeg"
    if data[:2] == b"BM":
        return "image/bmp"
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    if data[:5] in (b"<?xml", b"<svg "):
        return "image/svg+xml"
    return "image/png"


_SOF_MARKERS = frozenset(
    {0xC0, 0xC1, 0xC2, 0xC3, 0xC5, 0xC6, 0xC7, 0xC9, 0xCA, 0xCB, 0xCD, 0xCE, 0xCF}
)


def _jpeg_dimensions(data: bytes) -> tuple[int, int] | None:
    """Read pixel dimensions out of a JPEG's SOF marker."""
    i = 2
    n = len(data)
    while i + 3 < n:
        if data[i] != 0xFF:
            i += 1
            continue
        marker = data[i + 1]
        if marker in _SOF_MARKERS:
            if i + 9 <= n:
                return (
                    int.from_bytes(data[i + 7:i + 9], "big"),
                    int.from_bytes(data[i + 5:i + 7], "big"),
                )
            return None
        if marker == 0xFF or 0xD0 <= marker <= 0xD7:
            i += 2
            continue
        seglen = int.from_bytes(data[i + 2:i + 4], "big")
        i += 2 + seglen
    return None


def _image_dimensions(mime: str, data: bytes) -> tuple[int, int] | None:
    """Best-effort intrinsic pixel size for common raster formats."""
    mime = (mime or "").lower()
    if mime == "image/png" and data[:8] == _PNG_MAGIC and len(data) >= 24:
        return (
            int.from_bytes(data[16:20], "big"),
            int.from_bytes(data[20:24], "big"),
        )
    if mime == "image/gif" and data[:6] in (b"GIF87a", b"GIF89a") and len(data) >= 10:
        return (
            int.from_bytes(data[6:8], "little"),
            int.from_bytes(data[8:10], "little"),
        )
    if mime == "image/jpeg" and data[:2] == b"\xff\xd8":
        return _jpeg_dimensions(data)
    if mime == "image/bmp" and data[:2] == b"BM" and len(data) >= 26:
        return (
            int.from_bytes(data[18:22], "little"),
            abs(int.from_bytes(data[22:26], "little")),
        )
    return None


def _parse_px(value) -> int | None:
    """Parse an integer pixel length from an HTML attribute (e.g. '120')."""
    if value is None:
        return None
    m = re.fullmatch(r"\s*(\d+)\s*(?:px)?\s*", value)
    return int(m.group(1)) if m else None


# --------------------------------------------------------------------------
# Footnotes and endnotes (footnotes.xml / endnotes.xml parts)
# --------------------------------------------------------------------------
#
# python-docx does not model footnotes/endnotes, so both reading and writing
# go through the raw package parts. The HTML contract (shared with the ODT
# converter): a marker <sup class="footnote-citation">[n]</sup> immediately
# followed by <span class="footnote">BODY</span> (endnotes use the
# endnote-citation / endnote classes).

_FOOTNOTES_SENTINEL_XML = (
    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:footnote w:type="separator" w:id="-1">'
    '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
    '<w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0">'
    '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
    '<w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
    '</w:footnotes>'
)

_ENDNOTES_SENTINEL_XML = (
    _FOOTNOTES_SENTINEL_XML
    .replace("<w:footnotes", "<w:endnotes")
    .replace("</w:footnotes>", "</w:endnotes>")
    .replace("<w:footnote ", "<w:endnote ")
    .replace("</w:footnote>", "</w:endnote>")
)


class _NotesPart(StoryPart):
    """A raw WordprocessingML footnotes/endnotes part.

    Carries the separator (-1) and continuationSeparator (0) sentinels Word
    and LibreOffice expect, plus the real notes appended as ``w:footnote`` /
    ``w:endnote`` children. Registered with the main document part through
    the standard relationship + content type so the package opens cleanly in
    Word/LibreOffice.
    """

    @classmethod
    def new(cls, package, kind: str):
        partname = PackURI(
            "/word/footnotes.xml" if kind == "footnote" else "/word/endnotes.xml"
        )
        content_type = CT.WML_FOOTNOTES if kind == "footnote" else CT.WML_ENDNOTES
        element = parse_xml(
            _FOOTNOTES_SENTINEL_XML if kind == "footnote" else _ENDNOTES_SENTINEL_XML
        )
        return cls(partname, content_type, element, package)


def _get_notes_part(paragraph, kind: str):
    """The footnotes/endnotes part for a document, creating + wiring it if
    missing (proper content type + relationship from the main document part)."""
    document_part = paragraph.part
    reltype = RT.FOOTNOTES if kind == "footnote" else RT.ENDNOTES
    for rel in document_part.rels.values():
        if rel.reltype == reltype and not rel.is_external:
            return rel.target_part
    part = _NotesPart.new(document_part.package, kind)
    document_part.relate_to(part, reltype)
    return part


def _next_note_id(root, kind: str) -> int:
    """Next positive w:id (sentinels use -1/0, real notes start at 1)."""
    tag = qn("w:footnote") if kind == "footnote" else qn("w:endnote")
    used = []
    for el in root.iter(tag):
        val = el.get(qn("w:id")) or ""
        if val.isdigit() and int(val) > 0:
            used.append(int(val))
    return max(used) + 1 if used else 1


def _append_note_body(part, kind: str, w_id: int, body_html: str) -> None:
    """Append a w:footnote/w:endnote element holding the body paragraphs
    (``<br/>`` separates paragraphs, same as the reader round-trips)."""
    root = part.element
    note_el = OxmlElement("w:footnote" if kind == "footnote" else "w:endnote")
    note_el.set(qn("w:id"), str(w_id))
    for frag in re.split(r"<br\s*/?>", body_html or ""):
        p_el = OxmlElement("w:p")
        _add_styled_runs(Paragraph(p_el, part), frag)
        note_el.append(p_el)
    root.append(note_el)


def _add_footnote_reference(paragraph, token: dict) -> None:
    """Emit a w:footnoteReference/w:endnoteReference run for a footnote token
    and store the note body in the notes part."""
    kind = token.get("kind") or "footnote"
    part = _get_notes_part(paragraph, kind)
    w_id = _next_note_id(part.element, kind)
    _append_note_body(part, kind, w_id, token.get("body") or "")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    rPr.append(vert)
    r.append(rPr)
    ref = OxmlElement(
        "w:footnoteReference" if kind == "footnote" else "w:endnoteReference"
    )
    ref.set(qn("w:id"), str(w_id))
    r.append(ref)
    paragraph._p.append(r)


def _note_body_html(note_el, part) -> str:
    """Inline HTML of a footnotes/endnotes part's w:p children joined with
    ``<br/>``. Nested note references are ignored (they resolve against the
    body ordering, not the note)."""
    paras: list[str] = []
    for p_el in note_el.findall(qn("w:p")):
        html = _paragraph_inline(Paragraph(p_el, part))
        if html:
            paras.append(html)
    return "<br/>".join(paras)


def _collect_notes(doc) -> dict:
    """Parse footnotes.xml/endnotes.xml into ``{id: body_html}`` maps plus the
    1-based document-order index at which each id is first referenced."""
    notes = {
        "footnote": {}, "footnote_order": {},
        "endnote": {}, "endnote_order": {},
    }
    for kind, reltype in (("footnote", RT.FOOTNOTES), ("endnote", RT.ENDNOTES)):
        part = None
        for rel in doc.part.rels.values():
            if rel.reltype == reltype and not rel.is_external:
                part = rel.target_part
                break
        if part is None:
            continue
        root = parse_xml(part.blob)
        note_tag = qn("w:footnote") if kind == "footnote" else qn("w:endnote")
        for note_el in root.iter(note_tag):
            w_id = note_el.get(qn("w:id"))
            if not w_id or not w_id.isdigit() or note_el.get(qn("w:type")):
                continue  # skip the separator/continuationSeparator sentinels
            notes[kind][w_id] = _note_body_html(note_el, part)
        ref_tag = (
            qn("w:footnoteReference") if kind == "footnote"
            else qn("w:endnoteReference")
        )
        order = 0
        for ref in doc.element.body.iter(ref_tag):
            rid = ref.get(qn("w:id"))
            if rid in notes[kind] and rid not in notes[kind + "_order"]:
                order += 1
                notes[kind + "_order"][rid] = order
    return notes


def _note_marker(kind: str, w_id, notes) -> str:
    """The ``<sup class="...-citation">[i]</sup>`` + ``<span class="...">``
    pair for a referenced note, or ``''`` when unknown/not tracked."""
    if not notes:
        return ""
    bodies = notes.get(kind) or {}
    order = notes.get(kind + "_order") or {}
    body = bodies.get(w_id)
    i = order.get(w_id)
    if body is None or i is None:
        return ""
    cls = "footnote" if kind == "footnote" else "endnote"
    return (
        f'<sup class="{cls}-citation">[{i}]</sup>'
        f'<span class="{cls}">{body}</span>'
    )


# --------------------------------------------------------------------------
# Anchored comments (word/comments.xml part)
# --------------------------------------------------------------------------
#
# python-docx does not model comments, so both reading and writing go through
# the raw package parts, mirroring the footnotes/endnotes handling. The HTML
# contract (shared with the ODT converter): a ``<span class="comment"
# data-author="AUTHOR" data-comment="BODY">ANCHORED TEXT</span>``. The
# anchored runs are wrapped in ``w:commentRangeStart``/``w:commentRangeEnd``
# and terminated by a marker run carrying ``w:commentReference``; the body +
# author live in a proper ``word/comments.xml`` part.

_COMMENTS_XML = (
    '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '</w:comments>'
)


class _CommentsPart(StoryPart):
    """A raw WordprocessingML comments part (word/comments.xml).

    Real ``w:comment`` children are appended by ``_append_comment``. The part
    is registered with the main document part through the standard comments
    relationship + content type so the package opens cleanly in Word and
    LibreOffice (content type
    ``application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml``).
    """

    @classmethod
    def new(cls, package):
        return cls(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            parse_xml(_COMMENTS_XML),
            package,
        )


def _get_comments_part(paragraph):
    """The comments part for a document, creating + wiring it if missing
    (proper content type + relationship from the main document part)."""
    document_part = paragraph.part
    for rel in document_part.rels.values():
        if rel.reltype == RT.COMMENTS and not rel.is_external:
            return rel.target_part
    part = _CommentsPart.new(document_part.package)
    document_part.relate_to(part, RT.COMMENTS)
    return part


def _next_comment_id(root) -> int:
    """Next positive w:id (comments ids are 1-based)."""
    used = []
    for el in root.iter(qn("w:comment")):
        val = el.get(qn("w:id")) or ""
        if val.isdigit() and int(val) > 0:
            used.append(int(val))
    return max(used) + 1 if used else 1


def _append_comment(part, w_id: int, author: str, body: str) -> None:
    """Append a ``w:comment`` holding one paragraph with the body text in a
    ``w:r><w:t`` (the contract: BODY is plain text)."""
    root = part.element
    comment_el = OxmlElement("w:comment")
    comment_el.set(qn("w:id"), str(w_id))
    comment_el.set(qn("w:author"), author or "")
    p_el = OxmlElement("w:p")
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = body or ""
    r_el.append(t_el)
    p_el.append(r_el)
    comment_el.append(p_el)
    root.append(comment_el)


def _add_comment_reference(paragraph, token: dict) -> None:
    """Emit a comment range for a comment token.

    ``w:commentRangeStart`` is inserted before the anchored runs, the runs
    follow (parsed from the token's inner HTML), then ``w:commentRangeEnd``
    and a marker run carrying ``w:commentReference`` attached to the end of
    the anchored range.
    """
    part = _get_comments_part(paragraph)
    w_id = _next_comment_id(part.element)
    _append_comment(part, w_id, token.get("author") or "", token.get("body") or "")
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(w_id))
    paragraph._p.append(start)
    inner = token.get("html") or ""
    if inner:
        _add_styled_runs(paragraph, inner)
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(w_id))
    paragraph._p.append(end)
    r = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(w_id))
    r.append(ref)
    paragraph._p.append(r)


def _next_track_id(body) -> int:
    """Next unique positive w:id across all w:ins/w:del in the body."""
    used: list[int] = []
    for elem in body.iter():
        if elem.tag not in (qn("w:ins"), qn("w:del")):
            continue
        val = elem.get(qn("w:id")) or ""
        if val.isdigit():
            used.append(int(val))
    return (max(used) + 1) if used else 1


def _add_track_change(paragraph, token: dict) -> None:
    """Emit a tracked change element (w:ins / w:del) for a track token.

    ``<ins class="track-insert" data-author=..>NEW</ins>`` becomes a
    ``w:ins`` holding a run with the NEW text in ``w:t``; ``<del
    class="track-delete" ...>OLD</del>`` becomes a ``w:del`` holding a run
    whose ``w:delText`` carries the removed text (so Word still renders the
    deleted run). The change elements append to the paragraph's XML in
    document order; ids are unique and increasing.
    """
    parent = paragraph._p.getparent()
    w_id = _next_track_id(parent if parent is not None else paragraph._p)
    author = token.get("author") or ""
    dt = token.get("datetime") or ""
    kind = token.get("kind") or "insert"
    text = _inline_to_text(token.get("html") or "")
    change = OxmlElement("w:ins" if kind == "insert" else "w:del")
    change.set(qn("w:id"), str(w_id))
    change.set(qn("w:author"), author)
    if dt:
        change.set(qn("w:date"), dt)
    run = OxmlElement("w:r")
    if kind == "insert":
        t_el = OxmlElement("w:t")
    else:
        t_el = OxmlElement("w:delText")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    run.append(t_el)
    change.append(run)
    paragraph._p.append(change)


def _collect_comments(doc) -> dict:
    """Parse word/comments.xml into ``{w_id: (author, body_text)}``."""
    comments: dict = {}
    part = None
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS and not rel.is_external:
            part = rel.target_part
            break
    if part is None:
        return comments
    root = parse_xml(part.blob)
    for c_el in root.iter(qn("w:comment")):
        w_id = c_el.get(qn("w:id"))
        if not w_id:
            continue
        author = c_el.get(qn("w:author")) or ""
        body = "".join(t.text or "" for t in c_el.iter(qn("w:t")))
        comments[w_id] = (author, body)
    return comments


def _comment_span(w_id, pending: list, comments: dict) -> str:
    """The anchored ``<span class="comment" data-author=.. data-comment=..>``
    for a comment range, or the raw collected HTML when the id is unknown."""
    meta = (comments or {}).get(w_id)
    if meta is None:
        return "".join(pending)
    author, body = meta
    return (
        f'<span class="comment" data-author="{escape(author)}" '
        f'data-comment="{escape(body)}">{"".join(pending)}</span>'
    )


# --------------------------------------------------------------------------
# DOCX -> HTML
# --------------------------------------------------------------------------

def parse_list_at(html: str, start: int) -> tuple[dict, int]:
    """Parse a nested <ul>/<ol> subtree starting at ``html[start]``.

    Returns (tree, index just past the matching ``</ul|ol>``). A tree is
    ``{'kind': 'ul'|'ol', 'items': [{'frag': str, 'sub': [tree, ...]}, ...]}``
    where ``frag`` is the item body with nested lists removed (they live in
    ``sub`` instead), so a downstream inline parser can consume ``frag``
    without seeing list markup.
    """
    m = re.match(r"<([uo]l)([^>]*)>", html[start:], re.I)
    if not m:
        return {"kind": "ul", "items": []}, start + 4
    kind = m.group(1).lower()
    i = start + m.end()
    items: list[dict] = []
    tok = re.compile(r"<(/)?([a-z0-9]+)\b", re.I)
    while i < len(html):
        t = tok.search(html, i)
        if not t:
            break
        closing = t.group(1)
        tag = t.group(2).lower()
        if closing:
            if tag == kind:
                # past the closing '>'
                return {"kind": kind, "items": items}, t.end() + 1
            i = t.end() + 1
            continue
        if tag in ("ul", "ol"):
            sub, ni = parse_list_at(html, t.start())
            if items:
                items[-1]["sub"].append(sub)
            i = ni
        elif tag == "li":
            li_start = t.end() + 1
            depth = 0
            j = li_start
            li_end = len(html)
            while j < len(html):
                tt = tok.search(html, j)
                if not tt:
                    break
                if tt.group(1):  # closing tag
                    if tt.group(2).lower() in ("ul", "ol", "li"):
                        depth -= 1
                        if depth < 0:
                            li_end = tt.start()
                            break
                    j = tt.end() + 1
                else:
                    if tt.group(2).lower() in ("ul", "ol", "li"):
                        depth += 1
                    j = tt.start() + 1
            frag = html[li_start:li_end]
            frag, subs = extract_sublists(frag)
            items.append({"frag": frag, "sub": subs})
            i = li_end + 5 if li_end < len(html) else len(html)
        else:
            i = t.end() + 1
    return {"kind": kind, "items": items}, i


def extract_sublists(frag: str) -> tuple[str, list[dict]]:
    """Pull nested <ul>/<ol> subtrees out of an <li> body in source order.

    Returns (frag_with_lists_removed, [trees]).
    """
    subs: list[dict] = []
    out: list[str] = []
    pos = 0
    while True:
        m = re.search(r"<[uo]l\b", frag[pos:], re.I)
        if not m:
            out.append(frag[pos:])
            break
        out.append(frag[pos:pos + m.start()])
        tree, end = parse_list_at(frag, pos + m.start())
        subs.append(tree)
        pos = end
    return "".join(out), subs


def _list_level(style: str) -> int:
    """Outline level from a list style name ("List Bullet 2" -> 2)."""
    m = re.search(r"(\d)+(?:\s*)$", style)
    return max(1, int(m.group(1))) if m else 1


def _emit_list_tree(doc, tree: dict, level: int = 1) -> None:
    """Create DOCX paragraphs for a nested list tree using the numbered
    Word list styles ("List Bullet [n]" / "List Number [n]")."""
    base = "List Bullet" if tree["kind"] == "ul" else "List Number"
    style = base if level <= 1 else f"{base} {level}"
    for item in tree["items"]:
        p = doc.add_paragraph("", style=style)
        _add_styled_runs(p, item["frag"])
        for sub in item["sub"]:
            _emit_list_tree(doc, sub, level + 1)


def _tokenize_body(body: str) -> list:
    """Split a document body into block ops in source order.

    op tuples: ('hr',) | ('pagebreak', inner) | ('p', open_attrs, inner) |
    ('h', level, open_attrs, inner) | ('blockquote', open_attrs, inner) |
    ('list', tree) | ('header', attrs, content) | ('footer', attrs, content).
    Stray text between blocks is dropped, matching the previous regex-based
    behaviour. Tables are appended separately by the callers (they are excised
    before this runs).
    """
    ops: list = []
    pos = 0
    while True:
        m = re.search(r"<", body[pos:])
        if not m:
            break
        pos += m.start()
        rest = body[pos:]
        # Header at the start of body
        mh = re.match(r'<header([^>]*)>(.*?)</header>', rest, re.S | re.I)
        if mh:
            ops.append(("header", mh.group(1), mh.group(2)))
            pos += mh.end()
            continue
        # Footer at the end of body
        mf = re.search(r'<footer([^>]*)>(.*?)</footer>', rest, re.S | re.I)
        if mf and mf.start() == 0:
            ops.append(("footer", mf.group(1), mf.group(2)))
            pos += mf.end()
            continue
        mh = re.match(r"<hr(?:\s[^>]*)?/?>", rest, re.I)
        if mh:
            ops.append(("hr",))
            pos += mh.end()
            continue
        md = re.match(r"<div([^>]*)>(.*?)</div>", rest, re.S | re.I)
        if md and "page-break" in (md.group(1) or ""):
            ops.append(("pagebreak", md.group(2)))
            pos += md.end()
            continue
        ml = re.match(r"<([uo]l)\b", rest, re.I)
        if ml:
            tree, end = parse_list_at(body, pos)
            ops.append(("list", tree))
            pos = end
            continue
        mp = re.match(r"<p([^>]*)>(.*?)</p>", rest, re.S | re.I)
        if mp:
            ops.append(("p", mp.group(1), mp.group(2)))
            pos += mp.end()
            continue
        mh6 = re.match(r"<h([1-6])([^>]*)>(.*?)</h\1>", rest, re.S | re.I)
        if mh6:
            ops.append(("h", int(mh6.group(1)), mh6.group(2), mh6.group(3)))
            pos += mh6.end()
            continue
        mb = re.match(r"<blockquote([^>]*)>(.*?)</blockquote>", rest, re.S | re.I)
        if mb:
            ops.append(("blockquote", mb.group(1), mb.group(2)))
            pos += mb.end()
            continue
        # stray markup: skip one '<' and look again (content is dropped)
        pos += 1
    return ops


def _render_list_node(node: dict) -> str:
    """Render a nested-list tree as one contiguous HTML string."""
    s = f"<{node['kind']}>"
    for item in node["items"]:
        s += "<li>" + item["frag"]
        for sub in item["sub"]:
            s += _render_list_node(sub)
        s += "</li>"
    return s + f"</{node['kind']}>"


def _list_run_tree(seq: list[tuple]) -> list[dict]:
    """Group a consecutive (kind, level, frag) list run into nested-tree
    roots. A kind switch pops the open child (fresh top-level list)."""
    roots: list[dict] = []
    open_nodes: list[dict] = []  # {'kind','level','node'}
    for kind, level, frag in seq:
        # pop only STRICTLY deeper levels: an equal level is a sibling of
        # the current node (still open), not a new child.
        while open_nodes and open_nodes[-1]["level"] > level:
            open_nodes.pop()
        if (
            open_nodes
            and open_nodes[-1]["level"] == level
            and open_nodes[-1]["kind"] != kind
        ):
            open_nodes.pop()
        if not open_nodes:
            node = {"kind": kind, "items": []}
            roots.append(node)
            open_nodes.append({"kind": kind, "level": level, "node": node})
        elif open_nodes[-1]["level"] == level:
            node = open_nodes[-1]["node"]
        else:
            # deeper level: nest under the last item of the open node
            node = {"kind": kind, "items": []}
            open_nodes[-1]["node"]["items"][-1]["sub"].append(node)
            open_nodes.append({"kind": kind, "level": level, "node": node})
        node["items"].append({"frag": frag, "sub": []})
    return roots


def docx_to_html(data: bytes) -> str:
    """Convert DOCX bytes to an HTML fragment (content only, no <html>)."""
    doc = Document(io.BytesIO(data))
    notes = _collect_notes(doc)
    comments = _collect_comments(doc)
    seq: list[tuple] = []

    def strip_li(frag: str) -> str:
        m = re.match(r"^<li[^>]*>(.*)</li>\s*$", frag, re.S)
        return m.group(1) if m else frag

    # Extract header content if present
    header_html = ""
    header_part = _find_header_part(doc)
    if header_part:
        header_html = _header_to_html(header_part)

    # Extract footer content if present
    footer_html = ""
    footer_part = _find_footer_part(doc)
    if footer_part:
        footer_html = _footer_to_html(footer_part)

    for para in doc.paragraphs:
        li, list_kind, level = _paragraph_to_html(para, notes, comments)
        if list_kind is not None:
            seq.append(("list", list_kind, level or 1, strip_li(li)))
        else:
            seq.append(("other", li or ""))

    parts: list[str] = []
    i = 0
    while i < len(seq):
        if seq[i][0] == "other":
            parts.append(seq[i][1])
            i += 1
            continue
        run: list[tuple] = []
        while i < len(seq) and seq[i][0] == "list":
            run.append(seq[i][1:])
            i += 1
        for node in _list_run_tree(run):
            parts.append(_render_list_node(node))

    for table in doc.tables:
        parts.append(_table_to_html(table, notes))

    # Build the final HTML with header/footer
    html_parts = []
    if header_html:
        html_parts.append(f'<header class="page-header">{header_html}</header>')
    html_parts.extend(parts)
    if footer_html:
        html_parts.append(f'<footer class="page-footer">{footer_html}</footer>')

    body_html = "\n".join(p for p in html_parts if p)
    # Section columns (w:sectPr/w:cols) map to a <section data-columns>
    # wrapper so the layout survives an HTML round-trip.
    try:
        sp = doc.sections[0]._sectPr
        cols_el = sp.find(qn("w:cols")) if sp is not None else None
        if cols_el is not None:
            num = cols_el.get(qn("w:num"))
            space = cols_el.get(qn("w:space"))
            if num and int(num) > 1:
                n = int(num)
                gap = int(int(space) / 15) if space else 0
                return f'<section data-columns="{n}" data-column-gap="{gap}">{body_html}</section>'
    except Exception:
        pass
    return body_html


def _find_header_part(doc) -> Any | None:
    """Find the header part from the document's relationships."""

    sectPr = doc.sections[0]._sectPr
    hdr_ref = sectPr.find(qn("w:headerReference"))
    if hdr_ref is None:
        return None

    rel_id = hdr_ref.get(qn("r:id"))
    if rel_id is None:
        return None

    for rid, rel in doc.part.rels.items():
        if rid == rel_id and "header" in rel.reltype:
            return rel.target_part
    return None


def _find_footer_part(doc) -> Any | None:
    """Find the footer part from the document's relationships."""

    sectPr = doc.sections[0]._sectPr
    ftr_ref = sectPr.find(qn("w:footerReference"))
    if ftr_ref is None:
        return None

    rel_id = ftr_ref.get(qn("r:id"))
    if rel_id is None:
        return None

    for rid, rel in doc.part.rels.items():
        if rid == rel_id and "footer" in rel.reltype:
            return rel.target_part
    return None


def _header_to_html(header_part) -> str:
    """Convert a header part to HTML."""
    # The header part is an XML element with w:hdr root
    hdr_el = header_part._element

    html_parts = []
    for p_el in hdr_el.iter(qn("w:p")):
        inner_html = _runs_to_html_from_el(p_el)
        html_parts.append(f"<p>{inner_html}</p>")

    return "\n".join(html_parts)


def _footer_to_html(footer_part) -> str:
    """Convert a footer part to HTML."""
    # The footer part is an XML element with w:ftr root
    ftr_el = footer_part._element

    html_parts = []
    for p_el in ftr_el.iter(qn("w:p")):
        inner_html = _runs_to_html_from_el(p_el)
        html_parts.append(f"<p>{inner_html}</p>")

    return "\n".join(html_parts)


def _runs_to_html_from_el(p_el) -> str:
    """Convert paragraph elements to HTML, handling PAGE fields."""
    chunks: list[str] = []

    for child in p_el.iterchildren():
        tag = child.tag
        if tag == qn("w:r"):
            # Check for PAGE field in this run
            fld_simple = child.find(qn("w:fldSimple"))
            if fld_simple is not None:
                instr = fld_simple.get(qn("w:instr"))
                if instr and " PAGE " in instr:
                    chunks.append('<span class="page-number"></span>')
                    continue

            # Regular run - process children
            for run_child in child.iterchildren():
                run_child_tag = run_child.tag
                if run_child_tag == qn("w:t"):
                    text = run_child.text or ""
                    if text:
                        chunks.append(text)
        elif tag == qn("w:br"):
            chunks.append("<br/>")

    return "".join(chunks)


def _add_hr_paragraph(doc) -> None:
    """A horizontal rule: an empty paragraph styled with a bottom border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_break(doc) -> None:
    """A page break: an empty paragraph holding ``<w:br w:type='page'/>``."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_header(doc, content_html: str) -> None:
    """Add a header with the given HTML content to the document.

    Creates word/header1.xml and adds a headerReference to the sectPr.
    """
    # Use python-docx's section header - it creates header1.xml automatically
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]

    # Parse the header content and add runs
    # Each <p> becomes a paragraph in the header
    for p_match in re.finditer(r'<p([^>]*)>(.*?)</p>', content_html, re.S | re.I):
        p_inner = p_match.group(2)

        # Add the paragraph with header style
        # Create a run with the header text and any page number fields
        _add_header_footer_runs_to_paragraph(header_para, p_inner)

        # Add additional paragraphs for multiple <p> tags
        if p_match.end() < len(content_html):
            header_para = header.add_paragraph()

    # The section.header accessor automatically adds headerReference to sectPr
    # if it's not already there


def _add_footer(doc, content_html: str) -> None:
    """Add a footer with the given HTML content to the document.

    Creates word/footer1.xml and adds a footerReference to the sectPr.
    """
    # Use python-docx's section footer - it creates footer1.xml automatically
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]

    # Parse the footer content and add runs
    for p_match in re.finditer(r'<p([^>]*)>(.*?)</p>', content_html, re.S | re.I):
        p_inner = p_match.group(2)

        # Add the paragraph with footer text and any page number fields
        _add_header_footer_runs_to_paragraph(footer_para, p_inner)

        # Add additional paragraphs for multiple <p> tags
        if p_match.end() < len(content_html):
            footer_para = footer.add_paragraph()


def _add_header_footer_runs_to_paragraph(para, content_html: str) -> None:
    """Add runs to a header/footer paragraph from HTML content.

    Handles page number fields (<span class="page-number">) specially.
    """
    # Process the content, splitting into text and page-number markers
    pos = 0
    while pos < len(content_html):
        # Check for page number span
        pn_match = re.match(r'<span\s+class="page-number"[^>]*></span>', content_html[pos:], re.I | re.S)
        if pn_match:
            # Add any preceding text
            if pn_match.start() > 0:
                text = content_html[pos:pos + pn_match.start()]
                text = re.sub(r'<[^>]+>', '', text)  # Strip any remaining tags
                if text:
                    para.add_run(text)

            # Add page number field
            run = para.add_run()
            # Use fldSimple for PAGE field
            fldSimple = OxmlElement('w:fldSimple')
            fldSimple.set(qn('w:instr'), ' PAGE ')
            run._r.append(fldSimple)

            # Add initial text
            t = OxmlElement('w:t')
            t.text = '1'
            fldSimple.append(t)

            pos += pn_match.end()
        else:
            # Regular text - find next tag or end
            tag_match = re.search(r'<', content_html[pos:])
            if tag_match:
                text = content_html[pos:pos + tag_match.start()]
                text = re.sub(r'<[^>]+>', '', text)
                if text:
                    para.add_run(text)
                pos += tag_match.start()
            else:
                # Remaining text
                text = content_html[pos:]
                text = re.sub(r'<[^>]+>', '', text)
                if text:
                    para.add_run(text)
                break


def _para_has_bottom_border(para) -> bool:
    """True when a paragraph is rendered as a bottom-border rule (<hr>)."""
    pPr = para._p.pPr
    if pPr is None:
        return False
    pBdr = pPr.find(qn("w:pBdr"))
    return pBdr is not None and pBdr.find(qn("w:bottom")) is not None


def _para_is_page_break(para) -> bool:
    """True when a paragraph holds a page-break run."""
    return any(
        (br.get(qn("w:type")) or "textWrapping") == "page"
        for br in para._p.iter(qn("w:br"))
    )


def _para_style_parts(para) -> list[str]:
    """CSS declarations for a paragraph's formatting.

    Maps the DOCX paragraph properties to the HTML contract used by both
    converters: line-height (multiple), margin-left/right, text-indent,
    margin-top/bottom (spacing before/after), direction:rtl (w:bidi) and
    page-break-before:always (w:pageBreakBefore). Alignment is NOT part
    of this list (the caller folds "text-align:*" in front).
    """
    pf = para.paragraph_format
    styles: list[str] = []
    try:
        # line_spacing is a float multiple when the paragraph uses
        # proportional spacing (python-docx may label it ONE_POINT_FIVE /
        # DOUBLE rather than MULTIPLE, so check the VALUE, not the rule).
        ls = pf.line_spacing
        if (
            ls
            and isinstance(ls, (int, float))
            and abs(float(ls) - 1.0) > 1e-6
        ):
            styles.append(f"line-height:{float(ls):g}")
    except Exception:
        pass

    def _pt(value) -> float | None:
        try:
            return float(value.pt) if value is not None else None
        except Exception:
            return None

    for css, attr in (
        ("margin-left", "left_indent"),
        ("margin-right", "right_indent"),
        ("text-indent", "first_line_indent"),
        ("margin-top", "space_before"),
        ("margin-bottom", "space_after"),
    ):
        v = _pt(getattr(pf, attr, None))
        if v:
            styles.append(f"{css}:{v:g}pt")
    try:
        ppr = para._p.get_or_add_pPr()
        if ppr.find(qn("w:bidi")) is not None:
            styles.append("direction:rtl")
        if ppr.find(qn("w:pageBreakBefore")) is not None:
            styles.append("page-break-before:always")
    except Exception:
        pass
    return styles


def _parse_len_pt(val: str) -> float | None:
    """Parse a CSS length to points (pt default; px and cm converted)."""
    m = re.match(r"^\s*(\d+(?:\.\d+)?)\s*(pt|px|cm)?\s*$", val, re.I)
    if not m:
        return None
    num = float(m.group(1))
    if num <= 0:
        return None
    unit = (m.group(2) or "pt").lower()
    if unit == "px":
        num *= 0.75
    elif unit == "cm":
        num *= 28.3465
    return num


def _parse_para_props(open_tag: str) -> dict:
    """Parse paragraph-level CSS props from a block open tag into a dict.

    Keys: ``line_height`` (float multiple), ``margin-left``/``margin-right``/
    ``margin-top``/``margin-bottom``/``text-indent`` (point floats),
    ``direction`` ("rtl"), ``page_break_before`` (True).
    """
    props: dict = {}
    m = re.search(r'style="([^"]*)"', open_tag)
    style = m.group(1) if m else ""
    for decl in style.split(";"):
        if ":" not in decl:
            continue
        prop, _, val = decl.partition(":")
        prop = prop.strip().lower()
        val = val.strip()
        if prop == "line-height":
            lm = re.match(r"^(\d+(?:\.\d+)?)$", val)
            if lm:
                v = float(lm.group(1))
                if v > 0 and abs(v - 1.0) > 1e-6:
                    props["line_height"] = v
        elif prop in ("margin-left", "margin-right", "margin-top",
                      "margin-bottom", "text-indent"):
            v = _parse_len_pt(val)
            if v is not None:
                props[prop] = v
        elif prop == "direction" and val.lower() == "rtl":
            props["direction"] = "rtl"
        elif prop == "page-break-before" and val.lower() in ("always", "page"):
            props["page_break_before"] = True
        elif prop == "text-align" and val in ("center", "right"):
            props["text-align"] = val
    return props


def _apply_para_props(p, props: dict) -> None:
    """Apply parsed paragraph props to a python-docx paragraph."""
    if not props:
        return
    pf = p.paragraph_format
    if props.get("line_height"):
        try:
            pf.line_spacing = props["line_height"]
        except Exception:
            pass
    for css, attr in (
        ("margin-left", "left_indent"),
        ("margin-right", "right_indent"),
        ("margin-top", "space_before"),
        ("margin-bottom", "space_after"),
        ("text-indent", "first_line_indent"),
    ):
        v = props.get(css)
        if v:
            try:
                setattr(pf, attr, Pt(v))
            except Exception:
                pass
    try:
        ppr = p._p.get_or_add_pPr()
        if props.get("direction") == "rtl":
            ppr.append(OxmlElement("w:bidi"))
        if props.get("page_break_before"):
            ppr.append(OxmlElement("w:pageBreakBefore"))
    except Exception:
        pass


def _paragraph_to_html(para, notes=None, comments=None) -> tuple[str | None, str | None, int | None]:
    """Return (html_fragment, list_kind, list_level).

    Non-list blocks return (html, None, None). List paragraphs return
    `[fmt]<li>..</li>` plus its kind ("ul"/"ol") and outline level so the
    caller can group/nest consecutive items. ``notes`` carries the parsed
    footnotes/endnotes maps for footnote-marker rendering; ``comments``
    carries the parsed comments map for anchored comment spans.
    """
    style = (para.style.name or "").lower()
    text = _paragraph_inline(para, notes, comments)

    # Horizontal rule: an empty paragraph with a bottom border renders as <hr/>.
    if _para_has_bottom_border(para) and not text.strip():
        return "<hr/>", None, None

    # Page break: an empty paragraph holding <w:br w:type='page'/>.
    if _para_is_page_break(para) and not text.strip():
        return '<div class="page-break"><br></div>', None, None

    # python-docx exposes list styles as "List Bullet" / "List Number".
    if style.startswith("list"):
        kind = "ul" if "bullet" in style else "ol"
        level = _list_level(style)
        return f"<li>{text}</li>", kind, level

    if style.startswith("heading") or style.startswith("titre"):
        level = _heading_level(style)
        attrs = _block_attrs(para, [])
        return f"<h{level}{attrs}>{text}</h{level}>", None, None

    styles = _para_style_parts(para)
    align = para.alignment
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        styles.insert(0, "text-align:center")
    elif align == WD_ALIGN_PARAGRAPH.RIGHT:
        styles.insert(0, "text-align:right")
    attrs = _block_attrs(para, styles)
    return f"<p{attrs}>{text}</p>", None, None


def _block_attrs(para, styles: list[str]) -> str:
    """The open-tag attribute string for a paragraph/heading block."""
    if not styles:
        return ""
    return ' style="' + ";".join(styles) + '"'


def _paragraph_inline(para, notes=None, comments=None) -> str:
    """Inline HTML for a paragraph, emitting hyperlink runs as <a href>.

    Comment ranges (``w:commentRangeStart`` … ``w:commentRangeEnd`` closed
    by a ``w:commentReference`` marker run) emit the HTML contract
    ``<span class="comment" data-author=.. data-comment=..>ANCHOR</span>``;
    the marker run itself is dropped from the output.
    """
    out: list[str] = []
    pending: list[str] = []
    in_comment: str | None = None

    def add(html: str) -> None:
        if in_comment is not None:
            pending.append(html)
        else:
            out.append(html)

    def close_comment() -> None:
        nonlocal in_comment
        if in_comment is not None:
            out.append(_comment_span(in_comment, pending, comments))
            pending.clear()
            in_comment = None

    for child in para._p.iterchildren():
        tag = child.tag
        if tag == qn("w:commentRangeStart"):
            cid = child.get(qn("w:id"))
            close_comment()  # close any malformed/unclosed previous range
            in_comment = cid
            continue
        if tag == qn("w:commentRangeEnd"):
            if in_comment is not None and child.get(qn("w:id")) == in_comment:
                close_comment()
            continue
        if tag == qn("w:ins"):
            # tracked insertion: the w:ins element wraps w:r runs carrying
            # the inserted text in w:t
            author = child.get(qn("w:author")) or ""
            date = child.get(qn("w:date")) or ""
            inner = "".join(
                _run_to_html(Run(r, para), notes)
                for r in child.findall(qn("w:r"))
            )
            outer = f'<ins class="track-insert" data-author="{escape(author)}"'
            if date:
                outer += f' data-datetime="{escape(date)}"'
            out.append(outer + ">" + inner + "</ins>")
            continue
        if tag == qn("w:del"):
            # tracked deletion: runs carry <w:delText>, never w:t — collect
            # the removed text across all runs in the change element
            author = child.get(qn("w:author")) or ""
            date = child.get(qn("w:date")) or ""
            deleted = "".join(t.text or "" for t in child.iter(qn("w:delText")))
            outer = f'<del class="track-delete" data-author="{escape(author)}"'
            if date:
                outer += f' data-datetime="{escape(date)}"'
            out.append(outer + ">" + escape(deleted) + "</del>")
            continue
        if tag == qn("w:hyperlink"):
            href = _hyperlink_href(para, child)
            inner_runs = [Run(r, para) for r in child.findall(qn("w:r"))]
            inner = _runs_to_html(inner_runs, notes)
            if href:
                inner = f'<a href="{escape(href, quote=True)}">{inner}</a>'
            add(inner)
        elif tag == qn("w:r"):
            inner = _run_to_html(Run(child, para), notes)
            if in_comment is not None and child.find(qn("w:commentReference")) is not None:
                # the marker run terminates the anchored range
                pending.append(inner)
                close_comment()
            else:
                add(inner)
        # other children (pPr etc.) are not inline content
    close_comment()
    return "".join(out)


def _hyperlink_href(para, elem) -> str:
    """Resolve a w:hyperlink element to its target URL (anchor or external)."""
    anchor = elem.get(qn("w:anchor"))
    if anchor:
        return "#" + anchor
    r_id = elem.get(qn("r:id"))
    if r_id:
        rel = para.part.rels.get(r_id)
        if rel is not None:
            return rel.target_ref
    return ""


def _heading_level(style: str) -> int:
    m = re.search(r"(\d+)", style)
    if not m:
        return 1
    return max(1, min(int(m.group(1)), 6))


def _runs_to_html(runs, notes=None) -> str:
    out: list[str] = []
    for run in runs:
        out.append(_run_to_html(run, notes))
    html = "".join(out)
    if not html:
        # paragraph with no runs (e.g. empty) still needs a newline
        return "<br/>"
    return html


def _run_to_html(run, notes=None) -> str:
    """Inline HTML for one run, keeping picture positions intact.

    A ``<w:r>`` can interleave text children (``w:t``/``w:tab``/``w:br``/``w:cr``)
    with ``w:drawing`` children; each drawing becomes a self-contained
    ``<img>`` where it sits in the run. A ``w:footnoteReference``/
    ``w:endnoteReference`` child emits the note marker + body span via the
    ``notes`` context (unknown ids are dropped).
    """
    chunks: list[str] = []
    buf: list[str] = []
    for child in run._r:
        if child.tag == qn("w:footnoteReference"):
            marker = _note_marker("footnote", child.get(qn("w:id")), notes)
            if marker:
                if buf:
                    chunks.append(_wrap_run_text(escape("".join(buf)), run))
                    buf = []
                chunks.append(marker)
        elif child.tag == qn("w:endnoteReference"):
            marker = _note_marker("endnote", child.get(qn("w:id")), notes)
            if marker:
                if buf:
                    chunks.append(_wrap_run_text(escape("".join(buf)), run))
                    buf = []
                chunks.append(marker)
        elif child.tag == qn("w:drawing"):
            img = _drawing_to_img(run, child)
            if img:
                if buf:
                    chunks.append(_wrap_run_text(escape("".join(buf)), run))
                    buf = []
                chunks.append(img)
        elif child.tag in (qn("w:t"), qn("w:tab"), qn("w:br"), qn("w:cr")):
            buf.append(_text_child_value(child))
    if buf:
        chunks.append(_wrap_run_text(escape("".join(buf)), run))
    return "".join(chunks)


def _text_child_value(child) -> str:
    """Text equivalent of one run child, mirroring ``Run.text`` semantics."""
    tag = child.tag
    if tag == qn("w:t"):
        return child.text or ""
    if tag == qn("w:tab"):
        return "\t"
    if tag == qn("w:cr"):
        return "\n"
    if tag == qn("w:br"):
        # Only line-break <w:br> counts as text; page/column breaks are
        # layout primitives the HTML fragment does not model.
        if (child.get(qn("w:type")) or "textWrapping") == "textWrapping":
            return "\n"
        return ""
    return ""


def _apply_run_style(run, token: dict) -> None:
    """Apply colour / highlight / font / vert / strike / caps tokens to a run."""
    color = token.get("color")
    if color:
        try:
            run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        except Exception:
            pass
    bg = token.get("bg")
    if bg:
        try:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), bg.lstrip("#"))
            run._r.get_or_add_rPr().append(shd)
        except Exception:
            pass
    vert = token.get("vert")
    if vert == "sup":
        run.font.superscript = True
    elif vert == "sub":
        run.font.subscript = True
    if token.get("strike"):
        run.font.strike = True
    if token.get("small_caps"):
        run.font.small_caps = True
    if token.get("all_caps"):
        run.font.all_caps = True
    if token.get("code"):
        run.font.name = "Consolas"
    fam = token.get("font_family")
    if fam:
        run.font.name = fam
    size = token.get("font_size")  # e.g. "14pt" from _parse_font_size
    if size:
        try:
            run.font.size = Pt(float(size[:-2]))
        except Exception:
            pass


def _parse_inline_style(style: str) -> dict:
    """Parse safe CSS declarations from a span style into token properties.

    Returns a dict with any of: ``color``, ``bg``, ``font_family`` (CSS
    name), ``font_size`` ("<n>pt"), ``small_caps``, ``all_caps``.
    """
    props: dict = {}
    for decl in (style or "").split(";"):
        decl = decl.strip()
        if ":" not in decl:
            continue
        prop, _, val = decl.partition(":")
        prop = prop.strip().lower()
        val = val.strip()
        if prop == "color":
            c = _normalize_color(val)
            if c:
                props["color"] = c
        elif prop in ("background-color", "background"):
            c = _normalize_color(val)
            if c:
                props["bg"] = c
        elif prop == "font-family":
            fam = val.strip().strip("'\"")
            if fam:
                props["font_family"] = fam
        elif prop == "font-size":
            size = _parse_font_size(val)
            if size:
                props["font_size"] = size
        elif prop == "font-variant" and val.lower() == "small-caps":
            props["small_caps"] = True
        elif prop == "text-transform" and val.lower() == "uppercase":
            props["all_caps"] = True
    return props


def _parse_font_size(val: str) -> str | None:
    """Normalise a CSS font-size to ``'<n>pt'`` (pt / px supported)."""
    m = re.match(r"^\s*(\d+(?:\.\d+)?)\s*(pt|px)?\s*$", val, re.I)
    if not m:
        return None
    num = float(m.group(1))
    if num <= 0:
        return None
    unit = (m.group(2) or "pt").lower()
    if unit == "px":
        num = num * 0.75
    return f"{num:g}pt"


def _normalize_color(val: str) -> str | None:
    """Normalise a CSS colour to lowercase #rrggbb, or None if unsupported."""
    val = (val or "").strip().lower()
    if val.startswith("#"):
        h = val[1:]
        if len(h) == 3:
            h = "".join(c * 2 for c in h)
        if len(h) == 6 and all(c in "0123456789abcdef" for c in h):
            return "#" + h
        return None
    if val.startswith("rgb("):
        nums = re.findall(r"\d+", val)
        if len(nums) >= 3:
            r, g, b = (int(x) & 0xFF for x in nums[:3])
            return f"#{r:02x}{g:02x}{b:02x}"
        return None
    return None


def _run_color_hex(run) -> str | None:
    """Return the run's RGB text colour as #rrggbb, or None."""
    try:
        cf = run.font.color
        if cf.type == MSO_COLOR_TYPE.RGB:
            return "#" + str(cf.rgb).lower()
    except Exception:
        return None
    return None


def _run_highlight_hex(run) -> str | None:
    """Return the run's highlight (w:shd fill) as #rrggbb, or None."""
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return None
    shd = rPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if fill:
        return "#" + fill.lower()
    return None

def _run_highlight_hex(run) -> str | None:
    """Return the run's highlight (w:shd fill) as #rrggbb, or None."""
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return None
    shd = rPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if fill:
        return "#" + fill.lower()
    return None


_MONO_FONTS = {
    "consolas", "courier new", "courier", "monospace", "liberation mono",
    "dejavu sans mono", "cascadia mono", "source code pro",
}


def _run_font_family(run) -> str | None:
    """The run's ASCII font name (w:rFonts ascii), or None."""
    try:
        return run.font.name
    except Exception:
        return None


def _wrap_run_text(text: str, run) -> str:
    """Apply a run's character formatting around already-escaped text."""
    out = text
    if run.font.superscript:
        out = f"<sup>{out}</sup>"
    elif run.font.subscript:
        out = f"<sub>{out}</sub>"
    try:
        struck = run.font.strike
    except Exception:
        struck = False
    if struck:
        out = f"<strike>{out}</strike>"
    fam = _run_font_family(run)
    if fam and fam.lower().strip() in _MONO_FONTS:
        out = f"<code>{out}</code>"
    styles: list[str] = []
    if fam and fam.lower().strip() not in _MONO_FONTS:
        styles.append(f"font-family:{fam}")
    try:
        if run.font.size and run.font.size.pt:
            styles.append(f"font-size:{run.font.size.pt:g}pt")
    except Exception:
        pass
    try:
        if run.font.small_caps:
            styles.append("font-variant:small-caps")
        if run.font.all_caps:
            styles.append("text-transform:uppercase")
    except Exception:
        pass
    color = _run_color_hex(run)
    if color:
        styles.append(f"color:{color}")
    bg = _run_highlight_hex(run)
    if bg:
        styles.append(f"background-color:{bg}")
    if styles:
        out = f'<span style="{'; '.join(styles)}">{out}</span>'
    if run.bold:
        out = f"<b>{out}</b>"
    if run.italic:
        out = f"<i>{out}</i>"
    if run.underline:
        out = f"<u>{out}</u>"
    return out


def _blip_bytes(run, embed: str) -> tuple[str | None, bytes | None]:
    """Resolve an ``r:embed`` relationship id to (mime, bytes).

    DOCX images live in ``word/media/``; the relationship maps the rId used
    by the drawing's ``a:blip`` to that part. Best-effort: a missing or
    external relationship yields (None, None) and the caller skips the
    drawing, so surrounding text still converts.
    """
    rels = getattr(run.part, "rels", None)
    if rels is None:
        return None, None
    rel = rels.get(embed)
    if rel is None:
        return None, None
    try:
        blob = rel.target_part.blob
    except Exception:
        return None, None
    if not blob:
        return None, None
    mime = getattr(rel.target_part, "content_type", None) or _sniff_mime(blob)
    return mime, blob


def _drawing_extent_px(drawing) -> tuple[int | None, int | None]:
    """Read the ``wp:extent`` of a drawing as (width_px, height_px).

    Extent is stored in EMU; we divide by the same 96-dpi EMU-per-pixel
    constant used when writing, so the round-trip is a fixed point.
    """
    for ext in drawing.iter(qn("wp:extent")):
        try:
            cx = int(ext.get("cx") or 0)
            cy = int(ext.get("cy") or 0)
        except ValueError:
            return None, None
        if cx <= 0 or cy <= 0:
            return None, None
        return cx // _EMU_PER_PX, cy // _EMU_PER_PX
    return None, None


def _drawing_alt(drawing) -> str:
    """Accessible description of a drawing (``wp:docPr`` ``descr``).

    Falls back to a non-default ``name`` for producers that store the alt
    text there, but never to python-docx/Word's auto-generated
    ``Picture N`` placeholder.
    """
    for docPr in drawing.iter(qn("wp:docPr")):
        descr = (docPr.get("descr") or "").strip()
        if descr:
            return descr
        name = (docPr.get("name") or "").strip()
        if name and not re.fullmatch(r"Picture \d+", name):
            return name
        return ""
    return ""


def _drawing_to_img(run, drawing) -> str:
    """Render a ``<w:drawing>`` as a self-contained ``<img>``.

    Returns ``''`` when the drawing holds no embeddable picture (charts,
    shapes, OLE objects, external links), so surrounding text still
    converts. Dimension attributes come from the drawing extent and alt
    text from ``wp:docPr``/``descr``.
    """
    for blip in drawing.iter(qn("a:blip")):
        embed = blip.get(qn("r:embed"))
        if not embed:
            continue
        mime, content = _blip_bytes(run, embed)
        if content is None:
            continue
        width, height = _drawing_extent_px(drawing)
        attrs = [f' src="{_data_uri(mime or "image/png", content)}"']
        if width:
            attrs.append(f' width="{width}"')
        if height:
            attrs.append(f' height="{height}"')
        alt = _drawing_alt(drawing)
        if alt:
            attrs.append(f' alt="{escape(alt)}"')
        return "<img" + "".join(attrs) + "/>"
    return ""


def _table_to_html(table, notes=None) -> str:
    """Render a python-docx table as an HTML <table> fragment.

    Cell text keeps run-level formatting (bold/italic/underline) and
    multi-paragraph cells are joined with <br/>. Horizontal merges
    (gridSpan) and vertical merges (vMerge) are preserved as colspan /
    rowspan attributes so a DOCX -> HTML -> DOCX round-trip does not
    duplicate or drop cells.
    """
    grid = _table_grid(table)
    out: list[str] = []
    for r, row in enumerate(grid):
        cells: list[str] = []
        for e in row:
            if e["vmerge"] == "cont":
                continue  # covered by the restart cell's rowspan
            tag = "th" if _row_is_header(e["tr"]) else "td"
            rowspan = 1
            if e["vmerge"] == "start":
                rowspan = _rowspan_for(grid, r, e["pos"])
            cells.append(_cell_to_html(e, rowspan, tag, table, notes))
        tr_attr = ""
        trPr = e["tr"].trPr
        if trPr is not None and trPr.find(qn("w:cantSplit")) is not None:
            tr_attr = ' data-cantsplit="1"'
        out.append(f"<tr{tr_attr}>" + "".join(cells) + "</tr>")
    head = "<table>"
    tblPr = table._tbl.tblPr
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is not None and tblW.get(qn("w:type")) == "dxa":
            w_val = tblW.get(qn("w:w"))
            if w_val:
                head = f'<table width="{round(int(w_val) / 15)}">'
    table_html = head + "".join(out) + "</table>"
    caption = None
    if tblPr is not None:
        cap = tblPr.find(qn("w:tblCaption"))
        if cap is not None:
            caption = cap.get(qn("w:val")) or cap.text or ""
    if caption:
        return (
            f"<figure>{table_html}"
            f"<figcaption>{escape(caption)}</figcaption></figure>"
        )
    return table_html


def _table_grid(table):
    """Expand every row's <w:tc> elements into grid entries.

    Each entry is ``{tr, tc, pos, width, vmerge}`` where ``pos`` is the
    0-based grid column the cell starts at (accounting for gridSpan),
    ``width`` is its gridSpan, and ``vmerge`` is None / 'start' / 'cont'
    for non-merged / vMerge-restart / vMerge-continuation cells. We walk
    the raw XML because ``row.cells`` repeats merged cells and shifts the
    grid for merged tables, which would corrupt the HTML output.
    """
    grid = []
    for tr in (row._tr for row in table.rows):
        entries = []
        pos = 0
        for tc in tr.tc_lst:
            width = 1
            vmerge = None
            if tc.tcPr is not None:
                if tc.tcPr.gridSpan is not None and tc.tcPr.gridSpan.val:
                    width = tc.tcPr.gridSpan.val
                if tc.tcPr.vMerge is not None:
                    vmerge = (
                        "start" if tc.tcPr.vMerge.val == "restart" else "cont"
                    )
            entries.append(
                {"tr": tr, "tc": tc, "pos": pos, "width": width, "vmerge": vmerge}
            )
            pos += width
        grid.append(entries)
    return grid


def _rowspan_for(grid, r, pos) -> int:
    """Count how many rows a vMerge-restart cell at (r, pos) spans."""
    span = 1
    for row in grid[r + 1:]:
        match = None
        for e in row:
            if e["pos"] > pos:
                break
            if e["pos"] == pos:
                match = e
                break
        if match is None or match["vmerge"] != "cont":
            break
        span += 1
    return span


def _row_is_header(tr) -> bool:
    """True if the row is marked as a repeating header (w:tblHeader)."""
    trPr = tr.trPr
    return trPr is not None and trPr.find(qn("w:tblHeader")) is not None


def _cell_to_html(e, rowspan: int, tag: str, table, notes=None) -> str:
    """Render one grid entry as <td|th> HTML, keeping inline formatting."""
    attrs = ""
    if e["width"] > 1:
        attrs += f' colspan="{e["width"]}"'
    if rowspan > 1:
        attrs += f' rowspan="{rowspan}"'
    sparts, cell_w = _cell_style_parts(e)
    if sparts:
        attrs += ' style="' + ";".join(sparts) + '"'
    if cell_w is not None:
        attrs += f' width="{cell_w}"'
    # Parent paragraphs with a real _Cell so Run.part resolves to the
    # document part (needed to look up drawing image relationships).
    cell = _Cell(e["tc"], table)
    paras: list[str] = []
    for p_el in e["tc"].p_lst:
        paras.append(_runs_to_html(Paragraph(p_el, cell).runs, notes))
    inner = "<br/>".join(paras)
    if inner == "<br/>":
        inner = ""  # a single empty paragraph is an empty cell
    return f"<{tag}{attrs}>{inner}</{tag}>"


def _cell_style_parts(e) -> tuple[list[str], str | None]:
    """Cell shading / explicit borders / width as HTML style + width attr."""
    parts: list[str] = []
    width = None
    tcPr = e["tc"].tcPr
    if tcPr is None:
        return parts, width
    tcW = tcPr.tcW
    if tcW is not None and tcW.type == "dxa":
        w_val = tcW.get(qn("w:w"))
        if w_val:
            width = str(round(int(w_val) / 15))
    shd = tcPr.find(qn("w:shd"))
    if shd is not None and shd.get(qn("w:val")) not in (None, "nil"):
        fill = shd.get(qn("w:fill"))
        if fill:
            parts.append("background-color:#" + fill.lstrip("#").lower())
    bord = _tc_border_style(tcPr)
    if bord:
        parts.append(bord)
    return parts, width


def _tc_border_style(tcPr) -> str | None:
    """Explicit w:tcBorders -> 'Npt solid #hex', or None."""
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        return None
    sides: list[tuple[int, str]] = []
    for side in ("top", "left", "bottom", "right"):
        el = tcBorders.find(qn(f"w:{side}"))
        if el is None:
            continue
        val = el.get(qn("w:val"))
        sz = el.get(qn("w:sz"))
        if val in (None, "nil", "none") or not sz:
            continue
        sides.append((int(sz), (el.get(qn("w:color")) or "000000")))
    if not sides:
        return None
    sz = max(s[0] for s in sides)
    color = sides[0][1].lower()
    return f"border:{sz / 8:g}pt solid #{color}"


class _TableParser(HTMLParser):
    """Parse an HTML <table> fragment into rows of cell specs.

    Each cell is ``{"tag": "td"|"th", "attrs": {...}, "html": [...]}``
    whose ``html`` keeps the raw inner markup (inline tags + <br/>) so
    run-level formatting can be re-applied when building the DOCX.
    """

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.rows: list[list[dict]] = []
        self.row_attrs: list[dict] = []  # <tr> attributes, parallel to rows
        self._row: list[dict] | None = None
        self._cell: dict | None = None
        self._nested = 0  # open inline/nested elements inside the cell

    def handle_starttag(self, tag: str, attrs) -> None:
        if self._cell is not None:
            self._cell["html"].append(self.get_starttag_text())
            if tag != "br":
                self._nested += 1
        if tag == "tr":
            if self._cell is None:
                self._row = []
                self.rows.append(self._row)
                self.row_attrs.append(dict(attrs))
        elif tag in ("td", "th") and self._cell is None:
            if self._row is None:
                self._row = []
                self.rows.append(self._row)
            cell = {"tag": tag, "attrs": dict(attrs), "html": []}
            self._row.append(cell)
            self._cell = cell
            self._nested = 0

    def handle_endtag(self, tag: str) -> None:
        if tag in ("td", "th"):
            if self._cell is not None and self._cell["tag"] == tag:
                self._cell = None
        elif tag == "tr":
            self._cell = None
            self._row = None
        elif self._cell is not None and tag != "br":
            if self._nested > 0:
                self._nested -= 1
            self._cell["html"].append(f"</{tag}>")

    def handle_data(self, data: str) -> None:
        if self._cell is not None:
            self._cell["html"].append(data)


# --------------------------------------------------------------------------
# HTML -> DOCX
# --------------------------------------------------------------------------

_TAG_TABLE = re.compile(r"<figure>.*?</figure>|<table[^>]*>.*?</table>", re.S)


def html_to_docx(html_fragment: str) -> bytes:
    """Convert an HTML fragment into DOCX bytes."""
    # A <section data-columns> wrapper carries section-column layout
    # (mapped to w:sectPr/w:cols); unwrap it before body processing.
    section_cols = None
    section_gap = None
    sec_m = re.match(r'<section([^>]*)>(.*)</section>', html_fragment, re.S | re.I)
    if sec_m:
        sattrs = sec_m.group(1)
        cm = re.search(r'data-columns\s*=\s*"?(\d+)', sattrs)
        gm = re.search(r'data-column-gap\s*=\s*"?(\d+)', sattrs)
        if cm:
            section_cols = int(cm.group(1))
        if gm:
            section_gap = int(gm.group(1))
        html_fragment = sec_m.group(2)
    # Split tables out; python-docx tables and paragraphs share the body
    # but order interleaving is complex — append tables at the end.
    tables_html = _TAG_TABLE.findall(html_fragment)
    body = _TAG_TABLE.sub("", html_fragment)

    doc = Document()
    if section_cols and section_cols > 1:
        sectPr = doc.sections[0]._sectPr
        if sectPr is None:
            sectPr = OxmlElement("w:sectPr")
            doc.sections[0]._element.append(sectPr)
        cols_el = sectPr.find(qn("w:cols"))
        if cols_el is None:
            cols_el = OxmlElement("w:cols")
            sectPr.append(cols_el)
        cols_el.set(qn("w:num"), str(section_cols))
        if section_gap:
            cols_el.set(qn("w:space"), str(int(section_gap * 15)))
        else:
            cols_el.attrib.pop(qn("w:space"), None)

    # Extract header and footer if present
    header_content = None
    footer_content = None

    # Extract header from the beginning
    header_match = re.match(r'<header([^>]*)>(.*?)</header>', body, re.S | re.I)
    if header_match:
        header_content = header_match.group(2)
        body = body[header_match.end():]

    # Extract footer from the end (search in remaining body)
    body = body.strip()
    footer_match = re.search(r'<footer([^>]*)>(.*?)</footer>', body, re.S | re.I)
    if footer_match:
        footer_content = footer_match.group(2)
        body = body[:footer_match.start()].strip()

    ops = _tokenize_body(body)

    # Process header if present
    if header_content:
        _add_header(doc, header_content)

    # Process footer if present
    if footer_content:
        _add_footer(doc, footer_content)

    for op in ops:
        kind = op[0]
        if kind == "hr":
            _add_hr_paragraph(doc)
            continue
        if kind == "pagebreak":
            _add_page_break(doc)
            # A contenteditable can leave the caret inside the break
            # marker so the div carries real content (e.g. "§"); keep
            # it as a normal paragraph after the break instead of
            # dropping it.
            inner = re.sub(r"<br\s*/?>", "", op[1]).strip()
            if inner and _inline_to_text(inner).strip():
                p = doc.add_paragraph("")
                _add_styled_runs(p, inner)
            continue
        if kind == "list":
            _emit_list_tree(doc, op[1], 1)
            continue
        if kind == "blockquote":
            # Chrome's execCommand indent wraps blocks in <blockquote>; map
            # it to a left-indented paragraph (the HTML contract for indent).
            p = doc.add_paragraph("")
            _apply_para_props(p, {"margin-left": 24.0})
            _add_styled_runs(p, op[2])
            continue
        if kind == "h":
            props = _parse_para_props(op[2])
            p = doc.add_heading("", level=op[1])
            _apply_para_props(p, props)
            _add_styled_runs(p, op[3])
            continue
        # paragraph
        props = _parse_para_props(op[1])
        p = doc.add_paragraph("")
        _apply_para_props(p, props)
        if props.get("text-align") == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif props.get("text-align") == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_styled_runs(p, op[2])

    # Tag-less input (e.g. raw text typed into an empty contenteditable):
    # keep it as a single paragraph instead of dropping it silently.
    if not ops and body.strip():
        p = doc.add_paragraph("")
        _add_styled_runs(p, body)

    for tbl_html in tables_html:
        _append_table(doc, tbl_html)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _inline_to_text(html: str) -> str:
    """Strip tags into plain text and unescape HTML entities.

    Kept as a fallback for contexts where run-level formatting is not
    supported (e.g. table cell text fallback).
    """
    from html import unescape

    text = re.sub(r"<[^>]+>", "", html)
    return unescape(text)


_VOID_TAGS = frozenset({
    "img", "br", "hr", "meta", "link", "input", "area", "base", "col",
    "embed", "frame", "param", "source", "track", "wbr",
})


class _InlineRunBuilder(HTMLParser):
    """Parse an inline HTML fragment into text and image tokens.

    Text tokens are dicts with ``type: "text"`` plus ``text``/``bold``/
    ``italic``/``underline``; image tokens have ``type: "image"`` plus
    ``src``/``alt``/``width``/``height``. The HTML contract for notes — a
    ``<sup class="footnote-citation">[n]</sup>`` immediately followed by a
    ``<span class="footnote">BODY</span>`` — becomes a single ``type:
    "footnote"`` token carrying ``kind`` / ``citation`` / ``body``.
    """

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.tokens: list[dict] = []
        self._bold = 0
        self._italic = 0
        self._underline = 0
        self._strike = 0
        self._code = 0
        self._link_href = None  # href of the <a> currently being parsed (or None)
        self._color = None      # current text colour (e.g. "#ff0000") or None
        self._bg = None         # current highlight (background) colour or None
        self._vert = None       # "sup"/"sub" or None
        self._font_family = None
        self._font_size = None
        self._small_caps = None
        self._all_caps = None
        self._span_stack: list[tuple] = []  # prior span props saved on <span> open
        self._vert_stack = []   # prev vert saved on <sup>/<sub> open
        self._note = None       # pending footnote/endnote (see _start_note)
        self._comment = None    # pending comment span (see handle_starttag)
        self._track = None      # pending track-change (see handle_starttag)
        self._buf: list[str] = []

    def _start_note(self, attrs) -> bool:
        """True when the attrs mark a footnote/endnote citation <sup>. Starts a
        pending note: the next adjacent <span class="footnote|endnote"> (or
        the end of input) decides whether it becomes a real note token or is
        emitted as a plain superscript."""
        cls = set((dict(attrs).get("class") or "").split())
        if not (cls & {"footnote-citation", "endnote-citation"}):
            return False
        self._note = {
            "kind": "endnote" if "endnote-citation" in cls else "footnote",
            "citation": "",
            "in_citation": True,   # still inside the <sup>…</sup> citation
            "in_body": False,
            "depth": 0,
            "html": [],
        }
        return True

        self._note = {
            "kind": "endnote" if "endnote-citation" in cls else "footnote",
            "citation": "",
            "in_citation": True,   # still inside the <sup>…</sup> citation
            "in_body": False,
            "depth": 0,
            "html": [],
            "token_pos": len(self.tokens),
        }
        return True

    def _orphan_note(self) -> None:
        """Emit a pending citation <sup> that was NOT followed by an adjacent
        <span class="footnote|endnote"> as a plain superscript run, keeping it
        at its original position in the token stream."""
        text = self._note["citation"]
        pos = self._note["token_pos"]
        self._note = None
        if text:
            self.tokens.insert(pos, {
                "type": "text",
                "text": text,
                "bold": self._bold > 0,
                "italic": self._italic > 0,
                "underline": self._underline > 0,
                "vert": "sup",
            })

    def handle_starttag(self, tag: str, attrs) -> None:
        if self._track is not None:
            # collect the track-change element's inner markup
            if tag not in _VOID_TAGS:
                self._track["depth"] += 1
            self._track["html"].append(self.get_starttag_text())
            return
        if tag in ("ins", "del") and "track-" in (dict(attrs).get("class") or ""):
            cls = set((dict(attrs).get("class") or "").split())
            kind = "insert" if tag == "ins" else "delete"
            marker = "track-insert" if kind == "insert" else "track-delete"
            if marker not in cls:
                return  # not actually a tracked change (e.g. unstyled del)
            self._flush()
            a = dict(attrs)
            self._track = {
                "kind": kind,
                "author": a.get("data-author") or "",
                "datetime": a.get("data-datetime") or "",
                "html": [],
                "depth": 1,
            }
            return
        if self._comment is not None:
            # collect the comment span's inner markup until its closing tag
            if tag not in _VOID_TAGS:
                self._comment["depth"] += 1
            self._comment["html"].append(self.get_starttag_text())
            return
        if tag == "span" and "comment" in (dict(attrs).get("class") or "").split():
            # <span class="comment" data-author=.. data-comment=..>TEXT</span>
            self._flush()
            a = dict(attrs)
            self._comment = {
                "author": a.get("data-author") or "",
                "body": a.get("data-comment") or "",
                "html": [],
                "depth": 1,
            }
            return
        if self._note is not None and not self._note["in_body"]:
            # An open citation <sup>: the confirming <span class=...> directly
            # adjacent starts the body; anything else orphans the citation
            # back into a plain superscript run.
            cls = (dict(attrs).get("class") or "").split()
            body_cls = "footnote" if self._note["kind"] == "footnote" else "endnote"
            if tag == "span" and body_cls in cls and not self._buf:
                self._note["in_body"] = True
                self._note["depth"] = 1
                self._note["html"].append(self.get_starttag_text())
                return
            self._orphan_note()
            # fall through to normal handling for this tag
        elif self._note is not None and self._note["in_body"]:
            if tag not in _VOID_TAGS:
                self._note["depth"] += 1
            self._note["html"].append(self.get_starttag_text())
            return
        if tag == "img":
            self._flush()
            a = dict(attrs)
            self.tokens.append({
                "type": "image",
                "src": a.get("src", ""),
                "alt": a.get("alt", ""),
                "width": _parse_px(a.get("width")),
                "height": _parse_px(a.get("height")),
            })
        elif tag in ("b", "strong"):
            self._flush()
            self._bold += 1
        elif tag in ("i", "em"):
            self._flush()
            self._italic += 1
        elif tag == "u":
            self._flush()
            self._underline += 1
        elif tag == "a":
            self._flush()  # leading text must not merge into the link token
            a = dict(attrs)
            href = a.get("href", "")
            # Only carry safe schemes; javascript:/vbscript: etc. are dropped
            # (the editor sanitizer would strip them anyway).
            if href.startswith("#") or not re.search(r"^[a-z]+:", href, re.I):
                self._link_href = href or None
            elif href.startswith(("https://", "http://", "mailto:", "tel:", "/", "./", "../")):
                self._link_href = href
            else:
                self._link_href = None
        elif tag == "span":
            self._flush()
            a = dict(attrs)
            props = _parse_inline_style(a.get("style", ""))
            self._span_stack.append(
                (self._color, self._bg, self._font_family, self._font_size,
                 self._small_caps, self._all_caps)
            )
            if "color" in props:
                self._color = props["color"]
            if "bg" in props:
                self._bg = props["bg"]
            if "font_family" in props:
                self._font_family = props["font_family"]
            if "font_size" in props:
                self._font_size = props["font_size"]
            if props.get("small_caps"):
                self._small_caps = True
            if props.get("all_caps"):
                self._all_caps = True
        elif tag == "sup":
            self._flush()
            if self._start_note(attrs):
                return  # citation sups do not touch the vertical-align state
            self._vert_stack.append(self._vert)
            self._vert = "sup"
        elif tag == "sub":
            self._flush()
            self._vert_stack.append(self._vert)
            self._vert = "sub"
        elif tag in ("strike", "s", "del"):
            self._flush()
            self._strike += 1
        elif tag == "code":
            self._flush()
            self._code += 1
        else:
            self._flush()

    def handle_startendtag(self, tag: str, attrs) -> None:
        """Record self-closing void tags inside a note/comment body without
        skewing the nesting depth (they have no matching end tag; the body's
        closing span is tracked by depth, so an unclosed <br/> must not
        swallow it)."""
        if self._comment is not None:
            if tag in _VOID_TAGS:
                self._comment["html"].append(self.get_starttag_text())
            return
        if self._note is not None and self._note["in_body"]:
            if tag in _VOID_TAGS:
                self._note["html"].append(self.get_starttag_text())
                return
        super().handle_startendtag(tag, attrs)

    def handle_endtag(self, tag: str) -> None:
        if self._track is not None:
            if tag not in _VOID_TAGS:
                self._track["depth"] = max(0, self._track["depth"] - 1)
            self._track["html"].append(f"</{tag}>")
            if self._track["depth"] == 0:
                track = self._track
                self._track = None
                self.tokens.append({
                    "type": "track",
                    "kind": track["kind"],
                    "author": track["author"],
                    "datetime": track["datetime"],
                    "html": "".join(track["html"]),
                })
            return
        if self._comment is not None:
            if tag not in _VOID_TAGS:
                self._comment["depth"] = max(0, self._comment["depth"] - 1)
            self._comment["html"].append(f"</{tag}>")
            if self._comment["depth"] == 0:
                comment = self._comment
                self._comment = None
                self.tokens.append({
                    "type": "comment",
                    "author": comment["author"],
                    "body": comment["body"],
                    "html": "".join(comment["html"]),
                })
            return
        if self._note is not None and self._note["in_body"]:
            if tag not in _VOID_TAGS:
                self._note["depth"] = max(0, self._note["depth"] - 1)
            self._note["html"].append(f"</{tag}>")
            if self._note["depth"] == 0 and tag == "span":
                note = self._note
                self._note = None
                self.tokens.append({
                    "type": "footnote",
                    "kind": note["kind"],
                    "citation": (note["citation"] or "").strip(),
                    "body": "".join(note["html"]),
                })
            return
        if tag in ("b", "strong"):
            self._flush()
            self._bold = max(0, self._bold - 1)
        elif tag in ("i", "em"):
            self._flush()
            self._italic = max(0, self._italic - 1)
        elif tag == "u":
            self._flush()
            self._underline = max(0, self._underline - 1)
        elif tag == "a":
            self._flush()  # emits a link token (if href set) then clears context
            self._link_href = None
        elif tag == "span":
            self._flush()
            if self._span_stack:
                (self._color, self._bg, self._font_family, self._font_size,
                 self._small_caps, self._all_caps) = self._span_stack.pop()
        elif tag == "sup":
            self._flush()
            if self._note is not None and self._note.get("in_citation"):
                self._note["in_citation"] = False
            elif self._vert_stack:
                self._vert = self._vert_stack.pop()
        elif tag == "sub":
            self._flush()
            if self._vert_stack:
                self._vert = self._vert_stack.pop()
        elif tag in ("strike", "s", "del"):
            self._flush()
            self._strike = max(0, self._strike - 1)
        elif tag == "code":
            self._flush()
            self._code = max(0, self._code - 1)
        else:
            self._flush()

    def handle_data(self, data: str) -> None:
        if self._track is not None:
            # keep the inner text escaped in the capture (the writer re-uses
            # _inline_to_text on the inner HTML)
            self._track["html"].append(escape(data))
            return
        if self._comment is not None:
            self._comment["html"].append(escape(data))
            return
        if (
            self._note is not None
            and not self._note["in_body"]
            and self._note["in_citation"]
        ):
            self._note["citation"] += data
        elif self._note is not None and self._note["in_body"]:
            self._note["html"].append(escape(data))
        else:
            self._buf.append(data)

    def _finish(self) -> None:
        """Close out any unterminated footnote/endnote/comment at end of input."""
        if self._comment is not None:
            comment = self._comment
            self._comment = None
            self.tokens.append({
                "type": "comment",
                "author": comment["author"],
                "body": comment["body"],
                "html": "".join(comment["html"]),
            })
        if self._note is not None:
            if not self._note["in_body"]:
                self._orphan_note()
            else:
                note = self._note
                self._note = None
                self.tokens.append({
                    "type": "footnote",
                    "kind": note["kind"],
                    "citation": (note["citation"] or "").strip(),
                    "body": "".join(note["html"]),
                })
        self._flush()

    def _flush(self) -> None:
        text = "".join(self._buf)
        if text:
            token = {
                "type": "text",
                "text": text,
                "bold": self._bold > 0,
                "italic": self._italic > 0,
                "underline": self._underline > 0,
            }
            if self._link_href:
                token["type"] = "link"
                token["href"] = self._link_href
            if self._color:
                token["color"] = self._color
            if self._bg:
                token["bg"] = self._bg
            if self._vert:
                token["vert"] = self._vert
            if self._strike:
                token["strike"] = True
            if self._code:
                token["code"] = True
            if self._font_family:
                token["font_family"] = self._font_family
            if self._font_size:
                token["font_size"] = self._font_size
            if self._small_caps:
                token["small_caps"] = True
            if self._all_caps:
                token["all_caps"] = True
            self.tokens.append(token)
        self._buf = []


def _inline_tokens(html: str) -> list[dict]:
    builder = _InlineRunBuilder()
    builder.feed(html)
    builder._finish()
    return builder.tokens


def _add_styled_runs(paragraph, html: str) -> None:
    """Add runs parsed from an inline HTML fragment to a paragraph.

    Image tokens are embedded as inline pictures (``data:`` URIs only);
    http(s)/relative src values are skipped server-side. Footnote/endnote
    tokens emit a reference run and record the note body in the package.
    """
    for token in _inline_tokens(html):
        if token["type"] == "image":
            _add_image_run(paragraph, token)
            continue
        if token["type"] == "footnote":
            _add_footnote_reference(paragraph, token)
            continue
        if token["type"] == "comment":
            _add_comment_reference(paragraph, token)
            continue
        if token["type"] == "track":
            _add_track_change(paragraph, token)
            continue
        if token["type"] == "link":
            _add_hyperlink(paragraph, token)
            continue
        run = paragraph.add_run(token["text"])
        if token["bold"]:
            run.bold = True
        if token["italic"]:
            run.italic = True
        if token["underline"]:
            run.underline = True
        _apply_run_style(run, token)


def _add_hyperlink(paragraph, token: dict) -> None:
    """Append a hyperlink run (w:hyperlink + external relationship)."""
    url = token.get("href", "")
    part = paragraph.part
    try:
        # Generate a free rId for the external hyperlink relationship.
        used = set(part.rels.keys())
        i = 1
        while f"rId{i}" in used:
            i += 1
        r_id = f"rId{i}"
        part.rels.add_relationship(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            url,
            r_id,
            True,
        )
    except Exception as exc:  # malformed URL must not crash conversion
        logging.getLogger(__name__).warning("skipping un-embeddable hyperlink: %s", exc)
        run = paragraph.add_run(token["text"])
        if token.get("bold"):
            run.bold = True
        if token.get("italic"):
            run.italic = True
        if token.get("underline"):
            run.underline = True
        return
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if token.get("bold"):
        rPr.append(OxmlElement("w:b"))
    if token.get("italic"):
        rPr.append(OxmlElement("w:i"))
    if token.get("underline"):
        rPr.append(OxmlElement("w:u"))
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = token["text"]
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)



def _add_image_run(paragraph, token: dict) -> None:
    """Embed a data-URI ``<img>`` into the paragraph as an inline picture.

    Only ``data:`` URIs are embeddable server-side; http(s)/relative src
    values are skipped (no network fetching in a converter). Explicit
    width/height attributes become the drawing extent (px -> EMU); missing
    ones fall back to the image's intrinsic pixel size. The alt text is
    stored on the drawing's ``wp:docPr`` ``descr`` attribute so it
    round-trips back out.
    """
    mime, content = _decode_data_uri(token.get("src") or "")
    if content is None:
        return
    width = token.get("width")
    height = token.get("height")
    if not width or not height:
        dims = _image_dimensions(mime, content)
        if dims:
            width = width or dims[0]
            height = height or dims[1]
    kwargs: dict = {}
    if width and height:
        kwargs["width"] = Emu(width * _EMU_PER_PX)
        kwargs["height"] = Emu(height * _EMU_PER_PX)
    try:
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(content), **kwargs)
    except Exception as exc:  # corrupt/unsupported image bytes must not crash conversion
        logging.getLogger(__name__).warning("skipping un-embeddable image: %s", exc)
        return
    alt = (token.get("alt") or "").strip()
    if alt:
        _set_drawing_alt(run._r, alt)


def _set_drawing_alt(r, alt: str) -> None:
    """Record alt text on the ``wp:docPr`` of the run's picture."""
    drawing = r.find(qn("w:drawing"))
    if drawing is None:
        return
    for docPr in drawing.iter(qn("wp:docPr")):
        docPr.set("descr", alt)
        break


def _parse_border(style: str) -> str | None:
    """Extract a ``border`` declaration as ``'Npt solid #hex'`` or None.

    Supports pt/px (px converts at 3/4); named colours fall back to
    ``_normalize_color`` (basic names only). This is the HTML side of the
    cell-border contract shared with the ODT converter.
    """
    m = re.search(
        r"border\s*:\s*(\d+(?:\.\d+)?)\s*(pt|px)?\s*solid\s*"
        r"(#[0-9a-fA-F]{3,8}|[a-zA-Z]+)",
        style or "",
    )
    if not m:
        return None
    num = float(m.group(1))
    if m.group(2) == "px":
        num *= 0.75
    if num <= 0:
        return None
    color = m.group(3)
    if not color.startswith("#"):
        norm = _normalize_color(color)
        if not norm:
            return None
        color = norm
    return f"{num:g}pt solid {color}"


def _apply_cell_props(cell, attrs) -> None:
    """Apply background-color / border / width parsed from a cell tag."""
    style = attrs.get("style", "")
    tcPr = cell._tc.get_or_add_tcPr()
    width = (attrs.get("width") or "").strip()
    tcWel = tcPr.find(qn("w:tcW"))
    if re.fullmatch(r"\d+(\.\d+)?", width):
        try:
            cell.width = Emu(float(width) * 9525)
        except Exception:
            pass
    elif tcWel is not None:
        # python-docx fills every cell with a default w:tcW; drop it so the
        # reader does not invent a width for a cell the HTML left unsized.
        tcPr.remove(tcWel)
    bg = _parse_inline_style(style).get("bg")
    if bg:
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg.lstrip("#"))
    bord = _parse_border(style)
    if bord:
        _set_tc_borders(tcPr, bord)


def _set_tc_borders(tcPr, border_style: str) -> None:
    """Set all four explicit w:tcBorders from 'Npt solid #hex'."""
    m = re.match(r"([\d.]+)\s*pt\s*solid\s*(#[0-9a-fA-F]{6})", border_style)
    if not m:
        return
    sz = max(1, round(float(m.group(1)) * 8))
    color = m.group(2).lstrip("#")
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)


def _set_table_width(table, px: float) -> None:
    """Set an absolute (centered) table width w:tblW + w:jc."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(round(px * 15))))
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")


def _append_table(doc: Document, tbl_html: str) -> None:
    """Append an HTML <table> to the document as a python-docx table.

    Honors <th> (row becomes a repeating header), colspan (gridSpan),
    rowspan (vMerge) and <br/> (extra paragraphs inside a cell).
    """
    # A <figure> may wrap the <table> with a <figcaption>; map the
    # figcaption to w:tblCaption (the OOXML caption property).
    caption = None
    fm = re.search(r"<figure[^>]*>(.*?)</figure>", tbl_html, re.S)
    if fm:
        inner = fm.group(1)
        fc = re.search(r"<figcaption[^>]*>(.*?)</figcaption>", inner, re.S)
        if fc:
            caption = _inline_to_text(fc.group(1)).strip()
        tm = re.search(r"<table[^>]*>.*?</table>", inner, re.S)
        tbl_html = tm.group(0) if tm else tbl_html
    parser = _TableParser()
    parser.feed(tbl_html)
    rows = parser.rows
    if not rows:
        return
    ncols = 0
    for cells in rows:
        width = 0
        for c in cells:
            width += int(c["attrs"].get("colspan", 1) or 1)
        ncols = max(ncols, width)
    table = doc.add_table(rows=len(rows), cols=ncols or 1)
    table.style = "Table Grid"
    if caption:
        tblPr = table._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)
        cap = OxmlElement("w:tblCaption")
        cap.set(qn("w:val"), caption)
        tblPr.append(cap)
    m = re.search(r"<table[^>]*\bwidth\s*=\s*[\"']?(\d+(\.\d+)?)", tbl_html)
    if m:
        _set_table_width(table, float(m.group(1)))
    pending = [0] * ncols  # remaining rows covered by a rowspan, per grid column
    for r, cells in enumerate(rows):
        row_attrs = parser.row_attrs[r] if r < len(parser.row_attrs) else {}
        if row_attrs.get("data-cantsplit") == "1":
            trPr = table.rows[r]._tr.get_or_add_trPr()
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(OxmlElement("w:cantSplit"))
        if any(c["tag"] == "th" for c in cells):
            _mark_header_row(table.rows[r]._tr)
        pos = 0
        for c in cells:
            # Skip grid columns still covered by a rowspan from above.
            while pos < ncols and pending[pos] > 0:
                pending[pos] -= 1
                pos += 1
            if pos >= ncols:
                break
            colspan = int(c["attrs"].get("colspan", 1) or 1)
            rowspan = int(c["attrs"].get("rowspan", 1) or 1)
            cell = table.cell(r, pos)
            _fill_cell(cell, "".join(c["html"]))
            _apply_cell_props(cell, c["attrs"])
            if colspan > 1 or rowspan > 1:
                bottom = min(r + rowspan - 1, len(rows) - 1)
                right = min(pos + colspan - 1, ncols - 1)
                if (bottom, right) != (r, pos):
                    # Best effort: a malformed colspan/rowspan grid must
                    # not crash the conversion.
                    try:
                        cell.merge(table.cell(bottom, right))
                    except Exception:
                        pass
            for cc in range(pos, pos + colspan):
                pending[cc] = max(pending[cc], rowspan - 1)
            pos += colspan
        # A trailing rowspan may extend past this row in malformed input;
        # drain it so later rows are laid out from a clean state.
        for cc in range(pos, ncols):
            if pending[cc] > 0:
                pending[cc] -= 1


def _fill_cell(cell, html: str) -> None:
    """Fill a table cell; <br/>-separated fragments become paragraphs."""
    parts = re.split(r"<br\s*/?>", html)
    first = True
    for part in parts:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        _add_styled_runs(paragraph, part)


def _mark_header_row(tr) -> None:
    """Mark a table row as a repeating header (w:tblHeader)."""
    trPr = tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


