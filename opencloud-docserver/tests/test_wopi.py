"""Tests for WOPI host endpoints and editor API (integration style)."""

from __future__ import annotations

import io
import urllib.parse as urlparse
from contextlib import asynccontextmanager
from html.parser import HTMLParser

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from odf.opendocument import load

from src.config import Config
from src.editor.odt_converter import html_to_odt
from src.editor.router import router as editor_router
from src.editor.sanitize import sanitize_html
from src.editor.session import SessionRegistry
from src.lib.store import DocumentStore, wipe_db, wipe_dir
from src.wopi.router import router as wopi_router


def _make_app(tmp_path) -> tuple[FastAPI, DocumentStore]:
    db = str(tmp_path / "t.db")
    content = str(tmp_path / "content")
    store = DocumentStore(db, content)
    cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

    @asynccontextmanager
    async def lifespan(app: FastAPI):
        app.state.store = store
        app.state.sessions = SessionRegistry()
        app.state.config = cfg
        yield

    app = FastAPI(lifespan=lifespan)
    app.include_router(wopi_router)
    app.include_router(editor_router)
    return app, store


def _docx_bytes(text: str = "Test body") -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def client(tmp_path):
    app, store = _make_app(tmp_path)
    with TestClient(app) as c:
        c.test_store = store  # type: ignore[attr-defined]
        yield c
    wipe_db(tmp_path / "t.db")
    wipe_dir(tmp_path / "content")


def _seed_doc(client, doc_id="doc1", name="hello.docx", data=None):
    store = client.test_store  # type: ignore[attr-defined]
    store.init(doc_id, name)
    store.put_content(doc_id, data or _docx_bytes())
    return doc_id


# ----------------------------------------------------------------------
# WOPI host endpoints
# ----------------------------------------------------------------------

def test_check_file_info(client):
    _seed_doc(client)
    res = client.get("/wopi/files/doc1")
    assert res.status_code == 200
    body = res.json()
    assert body["BaseFileName"] == "hello.docx"
    assert body["SupportsLocks"] is True


def test_check_file_info_missing(client):
    res = client.get("/wopi/files/ghost")
    assert res.status_code == 404


def test_get_file(client):
    data = _docx_bytes()
    _seed_doc(client, data=data)
    res = client.get("/wopi/files/doc1/contents")
    assert res.status_code == 200
    assert res.content == data
    assert "X-WOPI-ItemVersion" in res.headers


def test_put_file(client):
    _seed_doc(client)
    new_data = _docx_bytes("Updated content")
    res = client.post("/wopi/files/doc1/contents", content=new_data)
    assert res.status_code == 200
    assert client.test_store.get_content("doc1") == new_data  # type: ignore[attr-defined]


def test_put_file_respects_lock(client):
    store = client.test_store  # type: ignore[attr-defined]
    _seed_doc(client)
    store.set_lock("doc1", "LOCK-123", "alice")

    # wrong lock -> 409
    res = client.post(
        "/wopi/files/doc1/contents",
        content=b"x",
        headers={"X-WOPI-Lock": "WRONG"},
    )
    assert res.status_code == 409

    # correct lock -> 200
    res = client.post(
        "/wopi/files/doc1/contents",
        content=b"y",
        headers={"X-WOPI-Lock": "LOCK-123"},
    )
    assert res.status_code == 200


def test_lock_unlock_cycle(client):
    _seed_doc(client)
    res = client.post("/wopi/files/doc1/lock", headers={"X-WOPI-Lock": "L1"})
    assert res.status_code == 200
    assert client.test_store.get_lock("doc1") == "L1"  # type: ignore[attr-defined]

    # unlock with wrong token -> 409
    res = client.post("/wopi/files/doc1/unlock", headers={"X-WOPI-Lock": "BAD"})
    assert res.status_code == 409

    # unlock with right token -> 200
    res = client.post("/wopi/files/doc1/unlock", headers={"X-WOPI-Lock": "L1"})
    assert res.status_code == 200
    assert client.test_store.get_lock("doc1") == ""  # type: ignore[attr-defined]

    # getlock returns current token
    res = client.post("/wopi/files/doc1/getlock")
    assert res.status_code == 200


def test_refresh_lock(client):
    store = client.test_store  # type: ignore[attr-defined]
    _seed_doc(client)
    store.set_lock("doc1", "L1", "bob")
    res = client.post("/wopi/files/doc1/refreshlock", headers={"X-WOPI-Lock": "L1"})
    assert res.status_code == 200


# ----------------------------------------------------------------------
# Extended WOPI API (contents endpoint + metadata) — gate: -k "extension"
# ----------------------------------------------------------------------
# The editor router exposes raw-bytes contents + extended metadata mirroring
# the WOPI GetFile/PutFile/CheckFileInfo protocol (host mode uses the local
# store; client mode forwards to the OCIS host). Every test lives inside
# TestWopiApiExtension so the whole battery runs under `pytest -k extension`.


class TestWopiApiExtension:
    def test_extension_contents_get_roundtrip(self, client):
        data = _docx_bytes("Extension body")
        _seed_doc(client, data=data)
        res = client.get("/api/documents/doc1/contents")
        assert res.status_code == 200
        assert res.content == data
        assert "X-WOPI-ItemVersion" in res.headers
        assert res.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_extension_contents_get_missing(self, client):
        res = client.get("/api/documents/ghost/contents")
        assert res.status_code == 404

    def test_extension_contents_put_roundtrip(self, client):
        _seed_doc(client)
        new_data = _docx_bytes("Replaced via contents PUT")
        res = client.put("/api/documents/doc1/contents", content=new_data)
        assert res.status_code == 200
        assert res.json() == {"ok": True, "size": len(new_data)}
        assert client.test_store.get_content("doc1") == new_data  # type: ignore[attr-defined]

    def test_extension_contents_put_via_wopi_override_post(self, client):
        """POST with X-WOPI-Override: PUT is the convention the OCIS
        wopiserver itself requires — the extension must accept it."""
        _seed_doc(client)
        new_data = _docx_bytes("Via override POST")
        res = client.post(
            "/api/documents/doc1/contents",
            content=new_data,
            headers={"X-WOPI-Override": "PUT"},
        )
        assert res.status_code == 200
        assert client.test_store.get_content("doc1") == new_data  # type: ignore[attr-defined]

    def test_extension_contents_post_without_override_rejected(self, client):
        _seed_doc(client)
        res = client.post("/api/documents/doc1/contents", content=b"x")
        assert res.status_code == 400

    def test_extension_contents_put_respects_lock(self, client):
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)
        store.set_lock("doc1", "LOCK-456", "alice")

        # wrong lock -> 409, echoing the current lock token
        res = client.put(
            "/api/documents/doc1/contents",
            content=b"x",
            headers={"X-WOPI-Lock": "WRONG"},
        )
        assert res.status_code == 409
        assert res.headers.get("X-WOPI-Lock") == "LOCK-456"

        # correct lock -> 200 and content replaced
        new_data = _docx_bytes("Unlocked write")
        res = client.put(
            "/api/documents/doc1/contents",
            content=new_data,
            headers={"X-WOPI-Lock": "LOCK-456"},
        )
        assert res.status_code == 200
        assert client.test_store.get_content("doc1") == new_data  # type: ignore[attr-defined]

    def test_extension_contents_put_missing_file(self, client):
        res = client.put("/api/documents/ghost/contents", content=b"x")
        assert res.status_code == 404

    def test_extension_metadata_extra_fields(self, client):
        _seed_doc(client, name="report.docx")
        res = client.get("/api/documents/doc1")
        assert res.status_code == 200
        body = res.json()
        assert body["base_file_name"] == "report.docx"
        assert body["format"] == "docx"
        assert body["mime_type"].startswith(
            "application/vnd.openxmlformats-officedocument"
        )
        assert body["version"] == str(body["updated_at"])
        assert body["editable"] is True
        assert body["writable"] is True
        assert body["contents_url"] == "/api/documents/doc1/contents"
        # existing fields are preserved (backwards compatible)
        assert body["id"] == "doc1"
        assert body["name"] == "report.docx"
        assert body["locked"] is False

    def test_extension_metadata_odt_format(self, client):
        store = client.test_store  # type: ignore[attr-defined]
        store.init("odt-2", "notes.odt")
        store.put_content("odt-2", b"odt-bytes")
        res = client.get("/api/documents/odt-2")
        assert res.status_code == 200
        body = res.json()
        assert body["format"] == "odt"
        assert body["mime_type"] == "application/vnd.oasis.opendocument.text"

    def test_extension_metadata_missing(self, client):
        res = client.get("/api/documents/ghost")
        assert res.status_code == 404


# ----------------------------------------------------------------------
# Editor API
# ----------------------------------------------------------------------

def test_upload_then_html(client):
    res = client.post(
        "/api/upload",
        files={"file": ("myfile.docx", _docx_bytes("Uploaded body"), "application/octet-stream")},
    )
    assert res.status_code == 200
    doc_id = res.json()["id"]

    res = client.get(f"/api/documents/{doc_id}/html")
    assert res.status_code == 200
    assert "Uploaded body" in res.json()["html"]


def test_save_document(client):
    _seed_doc(client)
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": "<p>Typed in the editor</p>"},
    )
    assert res.status_code == 200
    assert res.json()["ok"] is True

    # content is now a valid docx containing the new text
    docx_bytes = client.test_store.get_content("doc1")  # type: ignore[attr-defined]
    doc = Document(io.BytesIO(docx_bytes))
    assert "Typed in the editor" in "\n".join(p.text for p in doc.paragraphs)


def test_save_invalid_json(client):
    _seed_doc(client)
    res = client.post("/api/documents/doc1/save", content=b"not json", headers={"Content-Type": "application/json"})
    assert res.status_code == 400


def test_document_list(client):
    _seed_doc(client, doc_id="a", name="a.docx")
    _seed_doc(client, doc_id="b", name="b.docx")
    res = client.get("/api/documents")
    assert res.status_code == 200
    ids = {d["id"] for d in res.json()}
    assert ids == {"a", "b"}


def test_editor_page_served(client):
    _seed_doc(client)
    res = client.get("/editor/doc1")
    assert res.status_code == 200
    assert "contenteditable" in res.text


def test_hosting_discovery_clean_urlsrc_no_access_token(client):
    """Validated against real OpenCloud 7.3.0: OpenCloud appends WOPISrc to
    the urlsrc itself and POSTs a form with the real access_token in the
    body. urlsrc must therefore contain NO access_token placeholder/param."""
    r = client.get("/hosting/discovery")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/xml")
    xml = r.text
    assert "access_token" not in xml, "urlsrc must not carry an access_token param"
    assert 'urlsrc="http://localhost:8000/editor"' in xml
    assert 'ext="docx"' in xml


def test_hosting_discovery_includes_odt_actions(client):
    """WOPI discovery must advertise ODT view/edit actions."""
    r = client.get("/hosting/discovery")
    assert r.status_code == 200
    xml = r.text
    assert 'ext="odt"' in xml
    assert 'action name="view" ext="odt"' in xml
    assert 'action name="edit" ext="odt"' in xml


def test_odt_file_routes_to_odt_converter(client):
    """Files with .odt extension must use the ODT converter."""
    store = client.test_store  # type: ignore[attr-defined]
    store.init("odt-1", "document.odt")
    odt_bytes = html_to_odt("<p>Hello ODT</p>")
    store.put_content("odt-1", odt_bytes)

    # GET /html should convert ODT -> HTML
    r = client.get("/api/documents/odt-1/html")
    assert r.status_code == 200
    assert "Hello ODT" in r.json()["html"]

    # SAVE should convert HTML -> ODT
    res = client.post(
        "/api/documents/odt-1/save",
        json={"html": "<p>Updated ODT</p>"},
    )
    assert res.status_code == 200
    assert res.json()["ok"] is True

    # Verify the stored bytes are a valid ODT
    stored = store.get_content("odt-1")  # type: ignore[attr-defined]
    doc = load(io.BytesIO(stored))
    from odf import teletype
    text = teletype.extractText(doc.text)
    assert "Updated ODT" in text


def test_editor_launch_accepts_ocis_form_post(client):
    """OpenCloud launches the app by POSTing an urlencoded form with the real
    access_token in the body and WOPISrc in the query string (WOPI handshake)."""
    resp = client.post(
        "/editor",
        data={"access_token": "tok-123", "file_id": "doc-client", "embedded": "true"},
        params={"WOPISrc": "http://collaboration:9300/wopi/files/abc123"},
    )
    assert resp.status_code == 200
    # form `file_id` takes precedence as the session id
    session = client.app.state.sessions.get("doc-client")
    assert session is not None, "client-mode session must be registered from POST body"
    assert session.remote_host == "http://collaboration:9300"
    assert session.access_token == "tok-123"
    # Editor page must be wired to the resolved doc id (root path has no id):
    assert '"doc-client"' in resp.text


def test_editor_launch_get_uses_wopisrc_doc_id(client):
    """GET launches (dev/local) carry everything in the query string; the
    doc id is derived from the last segment of WOPISrc."""
    resp = client.get("/editor", params={
        "access_token": "tok-9",
        "WOPISrc": "http://collaboration:9300/wopi/files/fid-77",
    })
    assert resp.status_code == 200
    session = client.app.state.sessions.get("fid-77")
    assert session is not None
    assert session.access_token == "tok-9"
    assert session.remote_host == "http://collaboration:9300"
    assert '"fid-77"' in resp.text


def test_document_html_empty_file_returns_blank(client):
    """A 0-byte file must open as a blank document (start-typing UX), not
    fail with 'File is not a zip file' (US-3)."""
    store = client.test_store  # type: ignore[attr-defined]
    store.init("e1", "empty.docx")
    store.put_content("e1", b"")
    r = client.get("/api/documents/e1/html")
    assert r.status_code == 200
    body = r.json()
    assert body["html"] == ""
    assert body["blank"] is True


def test_document_html_corrupt_content_degrades_gracefully(client):
    """Fault-injection contract: corrupt (non-zip) content must NOT produce a
    server error — the read path degrades to an empty document instead of a
    500, so a damaged file can never take the docserver down (formerly 500)."""
    store = client.test_store  # type: ignore[attr-defined]
    store.init("e2", "corrupt.docx")
    store.put_content("e2", b"this is not a zip file, just text bytes")
    r = client.get("/api/documents/e2/html")
    assert r.status_code == 200
    body = r.json()
    assert body["html"] == ""
    assert body["name"] == "corrupt.docx"


# ----------------------------------------------------------------------
# XSS sanitizer
# ----------------------------------------------------------------------

def test_save_document_sanitizes_script_tag(client):
    """Script tags must be stripped before storage."""
    _seed_doc(client)
    malicious = '<p>Hello</p><script>alert("xss")</script><p>World</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200
    assert res.json()["ok"] is True

    docx_bytes = client.test_store.get_content("doc1")  # type: ignore[attr-defined]
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    # The script tag and its content should be removed
    assert "<script>" not in text
    assert "alert" not in text
    # Safe content should remain
    assert "Hello" in text
    assert "World" in text


def test_save_document_sanitizes_event_handler_attributes(client):
    """Event handler attributes (onclick, onerror, etc.) must be stripped."""
    _seed_doc(client)
    malicious = '<p onclick="alert(1)">Click me</p><img src="x" onerror="alert(1)">'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")  # type: ignore[attr-defined]
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "onclick" not in text
    assert "onerror" not in text
    assert "Click me" in text


def test_save_document_sanitizes_iframe(client):
    """iframe elements must be stripped (potential XSS vector)."""
    _seed_doc(client)
    malicious = '<p>Safe content</p><iframe src="https://evil.com"></iframe><p>More safe</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")  # type: ignore[attr-defined]
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "iframe" not in text
    assert "Safe content" in text
    assert "More safe" in text


def test_save_document_sanitizes_style_with_url(client):
    """Style attributes with url() or data: URIs must be stripped."""
    _seed_doc(client)
    malicious = '<p style="background-image:url(\'javascript:alert(1)\')">Test</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")  # type: ignore[attr-defined]
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    # The style with dangerous content should be removed
    assert "alert" not in text
    assert "Test" in text


def test_save_document_sanitizes_empty_string(client):
    """Empty string input should return empty HTML."""
    _seed_doc(client)
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": ""},
    )
    assert res.status_code == 200


def test_save_document_preserves_safe_formatting(client):
    """Safe formatting (bold, italic, underline, headings) should be preserved."""
    _seed_doc(client)
    formatted = "<h1>Heading</h1><p><b>Bold</b> and <i>italic</i> and <u>underline</u></p>"
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": formatted},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")  # type: ignore[attr-defined]
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Heading" in text
    assert "Bold" in text
    assert "italic" in text
    assert "underline" in text


def test_save_document_sanitizes_nested_script(client):
    """Nested script tags with mixed case must be stripped."""
    _seed_doc(client)
    malicious = '<p>Safe</p><SCRIPT>alert(1)</SCRIPT><p>Also safe</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")  # type: ignore[attr-defined]
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "<script>" not in text.lower() or "SCRIPT" not in text
    assert "alert" not in text
    assert "Safe" in text
    assert "Also safe" in text


# ----------------------------------------------------------------------
# XSS Evasion Tests (US-44)
# ----------------------------------------------------------------------

def test_save_document_sanitizes_html_encoded_script(client):
    """HTML-encoded script tags must be stripped (&#60;script&#62;)."""
    _seed_doc(client)
    # HTML entity encoded script tag
    malicious = '<p>Safe</p>&#60;script&#62;alert(1)&#60;/script&#62;<p>End</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "alert" not in text
    assert "Safe" in text
    assert "End" in text


def test_save_document_sanitizes_hex_encoded_script(client):
    """Hex-encoded script tags must be stripped (&#x3c;script&#x3e;)."""
    _seed_doc(client)
    # Hex entity encoded script tag
    malicious = '<p>Safe</p>&#x3c;script&#x3e;alert(1)&#x3c;/script&#x3e;<p>End</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "alert" not in text
    assert "Safe" in text


def test_save_document_sanitizes_mixed_case_script(client):
    """Mixed case script tags (ScRiPt) must be stripped."""
    _seed_doc(client)
    malicious = '<p>Safe</p><ScRiPt>alert(1)</sCrIpT><p>End</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "alert" not in text.lower()
    assert "Safe" in text


def test_save_document_sanitizes_javascript_href(client):
    """javascript: URLs in href must be stripped."""
    _seed_doc(client)
    malicious = '<p><a href="javascript:alert(1)">Click</a></p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "javascript" not in text.lower()
    assert "Click" in text


def test_save_document_sanitizes_javascript_src(client):
    """javascript: URLs in src must be stripped."""
    _seed_doc(client)
    malicious = '<img src="javascript:alert(1)">'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "javascript" not in text.lower()


def test_save_document_sanitizes_vbscript(client):
    """vbscript: URLs must be stripped."""
    _seed_doc(client)
    malicious = '<img src="vbscript:msgbox(1)">'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "vbscript" not in text.lower()


def test_save_document_sanitizes_css_expression(client):
    """CSS expression() must be stripped (IE legacy)."""
    _seed_doc(client)
    malicious = '<p style="width:expression(alert(1))">Test</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "expression" not in text.lower()
    assert "Test" in text


def test_sanitize_keeps_safe_color_and_highlight():
    """A #rrggbb color/highlight span must survive the sanitizer."""
    out = sanitize_html(
        '<span style="color:#ff0000;background-color:#ffff00">x</span>'
    )
    assert "color: #ff0000" in out
    assert "background-color: #ffff00" in out


def test_sanitize_drops_expression_color():
    """A colour declared as expression(alert(1)) must lose its style."""
    out = sanitize_html('<span style="color:expression(alert(1))">x</span>')
    assert "expression" not in out
    # the text survives as a span, but no style is emitted
    assert "x" in out
    assert "style=" not in out


def test_sanitize_keeps_superscript_and_subscript():
    """<sup>/<sub> (and strikethrough) must survive the sanitizer."""
    out = sanitize_html("<p>x<sup>2</sup> H<sub>2</sub>O <s>old</s></p>")
    assert "<sup>2</sup>" in out
    assert "<sub>2</sub>" in out
    assert "<s>old</s>" in out


def test_sanitize_keeps_strike_del_and_code():
    """Strike/del/code survive the sanitizer (inline-text parity)."""
    out = sanitize_html(
        "<p><strike>gone</strike> <del>old</del> <code>plain</code></p>"
    )
    assert "<strike>gone</strike>" in out
    assert "<del>old</del>" in out
    assert "<code>plain</code>" in out
    # arbitrary inactive tags remain dropped (figure/figcaption are now
    # supported for table captions, so use a genuinely inactive tag)
    assert "marquee" not in sanitize_html("<p><marquee>x</marquee></p>")


def test_save_document_sanitizes_css_behavior(client):
    """CSS behavior: must be stripped."""
    _seed_doc(client)
    malicious = '<p style="behavior:url(evil.htc)">Test</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "behavior" not in text.lower()
    assert "Test" in text


def test_save_document_sanitizes_moz_binding(client):
    """-moz-binding: must be stripped."""
    _seed_doc(client)
    malicious = '<p style="-moz-binding:url(evil.xml#xss)">Test</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "moz-binding" not in text.lower()
    assert "Test" in text


def test_save_document_sanitizes_data_uri_style(client):
    """data: URIs in style must be stripped."""
    _seed_doc(client)
    malicious = '<p style="background:url(data:text/html,<script>alert(1)</script>)">Test</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "data:" not in text.lower()
    assert "alert" not in text
    assert "Test" in text


def test_save_document_sanitizes_svg_script(client):
    """SVG with script elements must be stripped."""
    _seed_doc(client)
    malicious = '<p>Safe</p><svg><script>alert(1)</script></svg><p>End</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "alert" not in text
    assert "svg" not in text.lower()
    assert "Safe" in text


def test_save_document_sanitizes_object_embed(client):
    """object and embed tags must be stripped."""
    _seed_doc(client)
    malicious = '<p>Safe</p><object data="evil.swf"></object><embed src="evil.swf"><p>End</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "object" not in text.lower()
    assert "embed" not in text.lower()
    assert "Safe" in text
    assert "End" in text


def test_save_document_sanitizes_base_tag(client):
    """base tag must be stripped (can redirect relative URLs)."""
    _seed_doc(client)
    malicious = '<base href="javascript:alert(1)"><p>Safe</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "base" not in text.lower()
    assert "javascript" not in text.lower()
    assert "Safe" in text


def test_save_document_sanitizes_form_input(client):
    """form and input tags must be stripped (phishing vector)."""
    _seed_doc(client)
    malicious = '<form action="evil.com"><input name="credit"></form>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "form" not in text.lower()
    assert "input" not in text.lower()


def test_save_document_sanitizes_meta_refresh(client):
    """meta refresh tags must be stripped."""
    _seed_doc(client)
    malicious = '<meta http-equiv="refresh" content="0;url=evil.com"><p>Safe</p>'
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": malicious},
    )
    assert res.status_code == 200

    docx_bytes = client.test_store.get_content("doc1")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "meta" not in text.lower()
    assert "refresh" not in text.lower()
    assert "Safe" in text


# ----------------------------------------------------------------------
# Direct sanitizer tests (US-44: functional preservation + evasion)
# Test the sanitizer directly — not through the DOCX pipeline, which
# masks attribute-level bugs (only extracts text).
# ----------------------------------------------------------------------

def test_sanitize_preserves_img_with_data_url():
    """img with data:image src must be preserved (editor image feature)."""
    out = sanitize_html('<img src="data:image/png;base64,AAAA">')
    assert "<img" in out
    assert 'src="data:image/png;base64,AAAA"' in out


def test_sanitize_preserves_a_href_https():
    """a with https href must be preserved."""
    out = sanitize_html('<a href="https://example.com">Link</a>')
    assert "<a" in out
    assert 'href="https://example.com"' in out
    assert "Link" in out


def test_sanitize_preserves_a_href_relative():
    """a with relative href must be preserved."""
    out = sanitize_html('<a href="/doc/123">Link</a>')
    assert 'href="/doc/123"' in out


def test_sanitize_preserves_a_href_mailto():
    """a with mailto: href must be preserved."""
    out = sanitize_html('<a href="mailto:a@b.de">Mail</a>')
    assert "mailto:" in out


def test_sanitize_strips_javascript_href_direct():
    """javascript: href must be removed directly by the sanitizer."""
    out = sanitize_html('<a href="javascript:alert(1)">Click</a>')
    assert "javascript" not in out.lower()


def test_sanitize_strips_javascript_src_direct():
    """javascript: src on img must be removed directly."""
    out = sanitize_html('<img src="javascript:alert(1)">')
    assert "javascript" not in out.lower()


def test_sanitize_strips_mixed_case_javascript_direct():
    """JaVaScRiPt: scheme must be removed (case-insensitive)."""
    out = sanitize_html('<img src="JaVaScRiPt:alert(1)">')
    assert "javascript" not in out.lower()


def test_sanitize_strips_data_text_html_direct():
    """data:text/html src must be removed (only data:image allowed)."""
    out = sanitize_html('<img src="data:text/html,<script>alert(1)</script>">')
    assert "data:text/html" not in out.lower()


def test_sanitize_strips_onerror_direct():
    """onerror attribute must be removed but img kept."""
    out = sanitize_html('<img src="data:image/png;base64,x" onerror="alert(1)">')
    assert "onerror" not in out
    assert "<img" in out


def test_sanitize_strips_vbscript_direct():
    """vbscript: scheme must be removed."""
    out = sanitize_html('<img src="vbscript:msgbox(1)">')
    assert "vbscript" not in out.lower()


def test_sanitize_strips_css_expression_direct():
    """CSS expression() in style must be removed directly."""
    out = sanitize_html('<p style="width:expression(alert(1))">Test</p>')
    assert "expression" not in out.lower()
    assert "Test" in out


def test_sanitize_strips_html_encoded_script_direct():
    """HTML entity encoded script must not survive as a real HTML element."""
    out = sanitize_html("&#60;script&#62;alert(1)&#60;/script&#62;")
    # No real <script> tag may survive as a parseable element
    assert "<script" not in out.lower()
    # The angle brackets must be escaped in the data (round-trip safe)
    assert "&lt;" in out


def test_sanitize_preserves_safe_styles():
    """Safe inline styles (color, font-size) must be preserved."""
    out = sanitize_html('<p style="color:red; font-size:14pt">Text</p>')
    assert "color" in out
    assert "font-size" in out
    assert "Text" in out


def test_sanitize_strips_style_url_direct():
    """url() in style must be removed directly."""
    out = sanitize_html('<p style="background:url(https://evil.com/x.png)">Test</p>')
    assert "url(" not in out.lower()


def test_sanitize_strips_svg_direct():
    """SVG tags must be removed (not in safe list)."""
    out = sanitize_html('<svg><script>alert(1)</script></svg>')
    assert "svg" not in out.lower()
    assert "script" not in out.lower()


def test_sanitize_preserves_formatting():
    """Core formatting must survive: bold, italic, headings, lists."""
    out = sanitize_html('<p><b>Bold</b> <i>Italic</i> <u>Under</u></p><h2>Head</h2><ul><li>Item</li></ul>')
    assert "<b>Bold</b>" in out
    assert "<i>Italic</i>" in out
    assert "<u>Under</u>" in out
    assert "<h2>Head</h2>" in out
    assert "<ul><li>Item</li></ul>" in out


def test_sanitize_strips_base_direct():
    """base tag must be removed directly."""
    out = sanitize_html('<base href="https://evil.com">')
    assert "base" not in out.lower()


def test_sanitize_strips_form_direct():
    """form/input must be removed directly (phishing vector)."""
    out = sanitize_html('<form action="https://evil.com"><input name="x"></form>')
    assert "form" not in out.lower()
    assert "input" not in out.lower()


def test_sanitize_strips_meta_refresh_direct():
    """meta refresh must be removed directly."""
    out = sanitize_html('<meta http-equiv="refresh" content="0;url=https://evil.com">')
    assert "meta" not in out.lower()


# ----------------------------------------------------------------------
# XSS Sanitizer Evasion (US-44) — acceptance gate: -k "sanitize_evasion"
# ----------------------------------------------------------------------
# Every test below lives inside class TestSanitizeEvasion so the whole
# battery runs under `pytest -k sanitize_evasion`. Each case asserts the
# dangerous payload is neutralized AND that surrounding safe content
# survives, so the sanitizer proves both secure and functional.


class _SanitizedStructure(HTMLParser):
    """Re-parse sanitized output to check structural safety.

    Raw-substring assertions are misleading here: an escaped attribute value
    like `title="&quot; onmouseover=&quot;alert(1)"` legitimately still
    contains the letters `onmouseover` — but as inert text inside a quoted
    value that re-encodes the quote, never as a real attribute. The security
    property is structural: the re-parsed tree must contain no handler
    attributes and no script-bearing tags.
    """

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.tags: list[str] = []
        self.attr_names: list[str] = []
        self.text: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:
        self.tags.append(tag.lower())
        self.attr_names.extend((name or "").lower() for name, _ in attrs)

    def handle_startendtag(self, tag: str, attrs) -> None:
        self.handle_starttag(tag, attrs)

    def handle_data(self, data: str) -> None:
        self.text.append(data)


def _structure(html: str) -> _SanitizedStructure:
    probe = _SanitizedStructure()
    probe.feed(html)
    return probe


class TestSanitizeEvasion:
    # ---- event handlers (most common XSS vector) ----

    def test_sanitize_evasion_event_handlers_stripped(self):
        out = sanitize_html(
            '<img src="data:image/png;base64,x" onerror="alert(1)" '
            'onload="alert(2)" onmouseover="alert(3)" onclick="alert(4)">'
        )
        for handler in ("onerror", "onload", "onmouseover", "onclick"):
            assert handler not in out
        assert "alert" not in out
        assert "<img" in out  # element itself preserved

    def test_sanitize_evasion_mixed_case_event_handlers_stripped(self):
        out = sanitize_html('<p OnError="alert(1)" ONCLICK="alert(2)">Safe</p>')
        assert "onerror" not in out.lower()
        assert "onclick" not in out.lower()
        assert "Safe" in out

    def test_sanitize_evasion_space_handler_on_safe_tag_stripped(self):
        # Event handlers are not restricted to tags that normally use them.
        out = sanitize_html('<b onfocus="alert(1)" onblur="alert(2)">Bold</b>')
        assert "onfocus" not in out
        assert "onblur" not in out
        assert "<b>Bold</b>" in out

    # ---- URL schemes: javascript: / vbscript: / data:text/html ----

    def test_sanitize_evasion_javascript_href_stripped(self):
        out = sanitize_html('<p><a href="javascript:alert(1)">Click</a></p>')
        assert "javascript" not in out.lower()
        assert "alert" not in out
        assert "Click" in out

    def test_sanitize_evasion_entity_encoded_javascript_href_stripped(self):
        # &#106; = 'j' — browsers decode entities in attribute values, so an
        # attacker can hide the scheme; the sanitizer must still block it.
        out = sanitize_html('<a href="&#106;avascript:alert(1)">Click</a>')
        assert "javascript" not in out.lower()
        assert "alert" not in out
        assert "Click" in out

    def test_sanitize_evasion_hex_entity_javascript_src_stripped(self):
        out = sanitize_html('<img src="jav&#x61;script:alert(1)">')
        assert "javascript" not in out.lower()
        assert "alert" not in out

    def test_sanitize_evasion_mixed_case_javascript_scheme_stripped(self):
        out = sanitize_html('<a href="JaVaScRiPt:alert(1)">Click</a>')
        assert "javascript" not in out.lower()
        assert "Click" in out

    def test_sanitize_evasion_whitespace_obfuscated_scheme_stripped(self):
        # Tabs/newlines inside the scheme are ignored by lenient browsers.
        out = sanitize_html('<a href="java\tscript:alert(1)">Click</a>')
        assert "javascript" not in out.lower()
        assert "alert" not in out

    def test_sanitize_evasion_null_byte_scheme_stripped(self):
        out = sanitize_html('<a href="java\x00script:alert(1)">Click</a>')
        assert "script:" not in out
        assert "alert" not in out

    def test_sanitize_evasion_control_character_scheme_stripped(self):
        out = sanitize_html('<a href="\x01javascript:alert(1)">Click</a>')
        assert "javascript" not in out.lower()
        assert "Click" in out

    def test_sanitize_evasion_fullwidth_unicode_scheme_stripped(self):
        # Full-width confusables must not survive either (whitelist rejects).
        out = sanitize_html('<img src="\uff4a\uff41\uff56\uff41\uff53\uff43\uff52\uff49\uff50\uff54:alert(1)">')
        assert "alert" not in out
        assert "http" not in out

    def test_sanitize_evasion_vbscript_scheme_stripped(self):
        out = sanitize_html('<img src="vbscript:msgbox(1)"><a href="vBsCrIpT:msgbox(1)">x</a>')
        assert "vbscript" not in out.lower()
        assert "msgbox" not in out

    def test_sanitize_evasion_data_text_html_src_stripped(self):
        # data:text/html is an inline-execution vector; only data:image/ allowed.
        out = sanitize_html('<img src="data:text/html;base64,PHNjcmlwdD5hbGVydCgxKTwvc2NyaXB0Pg==">')
        assert "data:text/html" not in out.lower()
        assert "<script" not in out.lower()

    def test_sanitize_evasion_data_text_html_href_stripped(self):
        out = sanitize_html('<a href="data:text/html,<script>alert(1)</script>">Click</a>')
        assert "data:text/html" not in out.lower()
        assert "<script" not in out.lower()
        assert "Click" in out

    # ---- lesser-known URL-bearing attributes ----

    def test_sanitize_evasion_srcset_javascript_stripped(self):
        out = sanitize_html('<img src="x.png" srcset="javascript:alert(1) 1x, x2.png 2x">')
        assert "srcset" not in out
        assert "javascript" not in out.lower()
        assert "alert" not in out
        assert "<img" in out

    def test_sanitize_evasion_srcset_encoded_candidate_stripped(self):
        out = sanitize_html('<img src="x.png" srcset="java&#x73;cript:alert(1) 1x">')
        assert "srcset" not in out
        assert "javascript" not in out.lower()
        assert "alert" not in out

    def test_sanitize_evasion_srcset_data_text_html_stripped(self):
        out = sanitize_html('<img src="x.png" srcset="data:text/html;base64,PHNjcmlwdD4= 1x">')
        assert "data:text/html" not in out.lower()
        assert "srcset" not in out

    def test_sanitize_evasion_dynsrc_lowsrc_stripped(self):
        # Legacy IE attributes on <img> that load and execute URLs.
        out = sanitize_html('<img src="x.png" dynsrc="javascript:alert(1)" lowsrc="vbscript:msgbox(1)">')
        assert "dynsrc" not in out
        assert "lowsrc" not in out
        assert "javascript" not in out.lower()
        assert "vbscript" not in out.lower()

    def test_sanitize_evasion_background_attribute_stripped(self):
        # Legacy background attribute on table/td can carry javascript:.
        out = sanitize_html(
            '<table background="javascript:alert(1)"><tr>'
            '<td background="java&#x73;cript:alert(2)">Cell</td></tr></table>'
        )
        assert "background" not in out
        assert "javascript" not in out.lower()
        assert "alert" not in out
        assert "Cell" in out

    def test_sanitize_evasion_poster_data_html_stripped(self):
        out = sanitize_html('<img poster="data:text/html,<script>alert(1)</script>">')
        assert "data:text/html" not in out.lower()
        assert "<script" not in out.lower()

    # ---- attribute / tag breakout via escaped quotes & brackets ----

    def test_sanitize_evasion_attribute_breakout_neutralized(self):
        # &quot; decodes to a real quote: an attacker tries to terminate the
        # title attribute and forge onmouseover. The re-emitted value must
        # stay quoted so no new attribute can be forged on re-parse.
        out = sanitize_html('<p title="&quot; onmouseover=&quot;alert(1)">X</p>')
        probe = _structure(out)
        assert not any(a.startswith("on") for a in probe.attr_names)
        assert probe.tags == ["p"]
        # the quote is re-encoded inside the value, never emitted raw
        assert '<p title="&quot;' in out

    def test_sanitize_evasion_attribute_breakout_onerror_neutralized(self):
        out = sanitize_html('<img alt="&quot; onerror=&quot;alert(1)" src="data:image/png;base64,x">')
        probe = _structure(out)
        assert not any(a.startswith("on") for a in probe.attr_names)
        assert probe.tags == ["img"]
        assert "<img" in out

    def test_sanitize_evasion_tag_breakout_neutralized(self):
        # &quot;&gt;&lt;script&gt; decodes to "><script> — must not become
        # a real script element in the sanitized output.
        out = sanitize_html('<p title="&quot;&gt;&lt;script&gt;alert(1)&lt;/script&gt;">X</p>')
        probe = _structure(out)
        assert "script" not in probe.tags
        assert probe.tags == ["p"]
        assert "<script" not in out.lower()

    def test_sanitize_evasion_anchor_breakout_neutralized(self):
        out = sanitize_html('<a href="https://ok.com" title="&quot; onclick=&quot;alert(1)">L</a>')
        probe = _structure(out)
        assert not any(a.startswith("on") for a in probe.attr_names)
        assert probe.tags == ["a"]
        assert 'href="https://ok.com"' in out

    # ---- CSS / inline style injection ----

    def test_sanitize_evasion_style_url_stripped(self):
        out = sanitize_html('<p style="background:url(javascript:alert(1))">T</p>')
        assert "url(" not in out.lower()
        assert "javascript" not in out.lower()
        assert "alert" not in out
        assert "T" in out

    def test_sanitize_evasion_style_expression_stripped(self):
        out = sanitize_html('<p style="width:expression(alert(1))">T</p>')
        assert "expression" not in out.lower()
        assert "alert" not in out

    def test_sanitize_evasion_style_behavior_stripped(self):
        out = sanitize_html('<p style="behavior:url(evil.htc)">T</p>')
        assert "behavior" not in out.lower()

    def test_sanitize_evasion_style_moz_binding_stripped(self):
        out = sanitize_html('<p style="-moz-binding:url(evil.xml#xss)">T</p>')
        assert "moz-binding" not in out.lower()

    def test_sanitize_evasion_style_import_stripped(self):
        out = sanitize_html('<p style="@import url(https://evil.com/x.css)">T</p>')
        assert "@import" not in out.lower()
        assert "evil.com" not in out

    def test_sanitize_evasion_style_css_escape_stripped(self):
        # \65 is CSS-escaped 'e': \65 xpression(..) evades an expression check
        # in naive sanitizers and executes in legacy CSS-expression browsers.
        out = sanitize_html('<p style="color:\\65 xpression(alert(1))">T</p>')
        assert "expression" not in out.lower()
        assert "alert" not in out

    def test_sanitize_evasion_style_entity_obfuscation_stripped(self):
        out = sanitize_html('<p style="color:&#101;xpression(alert(1))">T</p>')
        assert "expression" not in out.lower()
        assert "alert" not in out

    def test_sanitize_evasion_style_entity_url_stripped(self):
        # u&#114;l(...) decodes to url(...) — must still be blocked.
        out = sanitize_html('<p style="background-color:u&#114;l(javascript:alert(1))">T</p>')
        assert "url(" not in out.lower()
        assert "javascript" not in out.lower()

    def test_sanitize_evasion_style_tag_removed(self):
        out = sanitize_html('<p>Safe</p><style>@import url(https://evil.com/x.css);</style><p>End</p>')
        assert "<style" not in out.lower()
        # content of the dropped <style> must not leak as visible text either
        assert "evil.com" not in out
        assert "@import" not in out
        assert _structure(out).tags == ["p", "p"]
        assert "Safe" in out
        assert "End" in out

    # ---- encoded / tampered element names ----

    def test_sanitize_evasion_script_entities_in_text_escaped(self):
        # &#60;script&#62; must come back as text, not a real script element.
        out = sanitize_html("&#60;script&#62;alert(1)&#60;/script&#62;")
        assert "<script" not in out.lower()
        assert "&lt;" in out

    def test_sanitize_evasion_hex_entities_in_text_escaped(self):
        out = sanitize_html("&#x3c;img src=x onerror=alert(1)&#x3e;")
        assert "<img" not in out.lower()
        assert "&lt;" in out
        # re-parsed, it is pure text — no img element, no handler attribute
        probe = _structure(out)
        assert probe.tags == []
        assert not any(a.startswith("on") for a in probe.attr_names)

    def test_sanitize_evasion_mixed_case_script_removed(self):
        out = sanitize_html('<ScRiPt src="https://evil.com/x.js"></ScRiPt><p>Safe</p>')
        assert "script" not in out.lower()
        assert "evil.com" not in out
        assert "Safe" in out

    def test_sanitize_evasion_script_src_external_removed(self):
        out = sanitize_html('<script src="https://evil.com/x.js"></script><p>Safe</p>')
        assert "script" not in out.lower()
        assert "evil.com" not in out
        assert "Safe" in out

    # ---- dangerous containers / elements ----

    def test_sanitize_evasion_svg_math_removed(self):
        out = sanitize_html('<p>S</p><svg onload="alert(1)"><script>alert(2)</script></svg><math><mi>x</mi></math><p>E</p>')
        assert "svg" not in out.lower()
        assert "math" not in out.lower()
        assert "script" not in out.lower()
        assert "onload" not in out
        assert "alert" not in out
        assert "S" in out and "E" in out

    def test_sanitize_evasion_iframe_object_embed_removed(self):
        out = sanitize_html(
            '<p>Safe</p><iframe src="javascript:alert(1)"></iframe>'
            '<object data="evil.swf"></object><embed src="evil.swf">'
        )
        assert "iframe" not in out.lower()
        assert "object" not in out.lower()
        assert "embed" not in out.lower()
        assert "alert" not in out
        assert "Safe" in out

    def test_sanitize_evasion_form_input_removed(self):
        out = sanitize_html('<form action="javascript:alert(1)"><input name="x" onfocus="alert(1)"></form>')
        assert "form" not in out
        assert "input" not in out
        assert "alert" not in out

    def test_sanitize_evasion_meta_base_link_removed(self):
        out = sanitize_html(
            '<meta http-equiv="refresh" content="0;url=javascript:alert(1)">'
            '<base href="https://evil.com"><link rel="stylesheet" href="https://evil.com/x.css">'
            '<p>Safe</p>'
        )
        assert "meta" not in out
        assert "base" not in out
        assert "link" not in out
        assert "evil.com" not in out
        assert "Safe" in out

    def test_sanitize_evasion_entity_encoded_tags_never_reconstruct(self):
        # A charref-encoded script/iframe embedded in a safe element's text.
        out = sanitize_html('<p>&#60;iframe src=&#34;javascript:alert(1)&#34;&#62;</p>')
        probe = _structure(out)
        assert probe.tags == ["p"]
        assert "iframe" not in probe.tags
        assert not any(a.startswith("on") for a in probe.attr_names)
        assert "&lt;" in out

    # ---- functional preservation (no over-stripping) ----

    def test_sanitize_evasion_safe_markup_preserved(self):
        out = sanitize_html(
            '<h2>Title</h2><p><b>Bold</b> and <i>italic</i></p>'
            '<a href="https://example.com/doc">Link</a>'
            '<img src="data:image/png;base64,AAAA">'
        )
        assert "<h2>Title</h2>" in out
        assert "<b>Bold</b>" in out
        assert "<i>italic</i>" in out
        assert 'href="https://example.com/doc"' in out
        assert 'src="data:image/png;base64,AAAA"' in out

    def test_sanitize_evasion_safe_styles_preserved(self):
        out = sanitize_html('<p style="color:#ff0000; font-weight:bold; margin:4px">Safe</p>')
        assert "color:#ff0000" in out or "color: #ff0000" in out
        assert "font-weight" in out
        assert "Safe" in out

    def test_sanitize_evasion_real_script_typed_as_text_stripped(self):
        # A literal (non-encoded) script tag is dropped — element AND its
        # inner payload — while the ordinary text around it stays.
        out = sanitize_html('<p>Hello</p><script>alert("xss")</script><p>World</p>')
        assert "script" not in out.lower()
        assert "alert" not in out
        assert _structure(out).tags == ["p", "p"]
        assert "Hello" in out
        assert "World" in out

    def test_sanitize_evasion_simple_plain_text_passthrough(self):
        out = sanitize_html("<p>1 < 2 and 3 > 2 and A & B</p>")
        assert "1 &lt; 2" in out
        assert "3 &gt; 2" in out
        assert "A &amp; B" in out


# ----------------------------------------------------------------------

# WOPI Lock Contention — acceptance gate: -k "lock_contention"
# ----------------------------------------------------------------------
# The WOPI lock is what lets several editors fight over one document while
# keeping a single writer.  These tests exercise the contention surface of
# the host endpoints: two editors ("alice"/"bob") race for the same file.
#
# Invariants covered (WOPI spec):
#   * Exactly one Lock wins; losers get 409 with the winner's token echoed
#     in X-WOPI-Lock (so the loser can adopt or back off).
#   * A Lock with the same token as the current lock is a refresh (200) and
#     keeps the lock; Lock responses must echo the lock token.
#   * Lock tokens MUST be non-empty; an empty token is rejected (400).
#   * PutFile honours the lock: a wrong/missing token is rejected with 409
#     and the saved content is never clobbered; a matching token succeeds
#     and keeps the lock.
#   * RefreshLock / Unlock with the wrong token are rejected (409) and never
#     steal or drop the winner's lock.
#   * GetLock reports the current holder; once the winner unlocks, the next
#     editor can take the lock with a fresh token.


class TestWopiLockContention:
    @staticmethod
    def _headers(token: str) -> dict:
        return {"X-WOPI-Lock": token}

    def test_lock_contention_single_winner(self, client):
        """Only one of two contenders may acquire the lock."""
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)

        res_a = client.post("/wopi/files/doc1/lock", headers=self._headers("wo:alice:A"))
        assert res_a.status_code == 200
        assert res_a.headers.get("X-WOPI-Lock") == "wo:alice:A"

        # bob's simultaneous attempt loses — 409 echoing alice's token
        res_b = client.post("/wopi/files/doc1/lock", headers=self._headers("wo:bob:B"))
        assert res_b.status_code == 409
        assert res_b.headers.get("X-WOPI-Lock") == "wo:alice:A"

        # the lock still belongs to alice
        assert store.get_lock("doc1") == "wo:alice:A"

    def test_lock_contention_same_token_is_a_refresh(self, client):
        """A Lock with the holder's token refreshes (200) and keeps the lock."""
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)
        store.set_lock("doc1", "wo:alice:A", "alice")

        res = client.post("/wopi/files/doc1/lock", headers=self._headers("wo:alice:A"))
        assert res.status_code == 200
        # spec: Lock responses must echo the lock token
        assert res.headers.get("X-WOPI-Lock") == "wo:alice:A"
        assert store.get_lock("doc1") == "wo:alice:A"

    def test_lock_contention_loser_cannot_put(self, client):
        """Bob's PutFile (missing or wrong lock) is rejected, content kept."""
        store = client.test_store  # type: ignore[attr-defined]
        original = _docx_bytes("Alice original")
        _seed_doc(client, data=original)
        store.set_lock("doc1", "wo:alice:A", "alice")

        # no lock header at all
        res = client.post("/wopi/files/doc1/contents", content=_docx_bytes("Bob clobber"))
        assert res.status_code == 409
        assert res.headers.get("X-WOPI-Lock") == "wo:alice:A"

        # wrong lock token
        bob = _docx_bytes("Bob clobber 2")
        res = client.post(
            "/wopi/files/doc1/contents", content=bob, headers=self._headers("wo:bob:B")
        )
        assert res.status_code == 409
        assert res.headers.get("X-WOPI-Lock") == "wo:alice:A"

        # neither attempt reached the document
        assert store.get_content("doc1") == original

        # alice with the matching token still saves
        alice_final = _docx_bytes("Alice final")
        res = client.post(
            "/wopi/files/doc1/contents",
            content=alice_final,
            headers=self._headers("wo:alice:A"),
        )
        assert res.status_code == 200
        assert store.get_content("doc1") == alice_final

    def test_lock_contention_refresh_rejected_for_loser(self, client):
        """Only the holder can extend the lock lease."""
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)
        store.set_lock("doc1", "wo:alice:A", "alice")

        res = client.post(
            "/wopi/files/doc1/refreshlock", headers=self._headers("wo:bob:B")
        )
        assert res.status_code == 409
        assert res.headers.get("X-WOPI-Lock") == "wo:alice:A"
        assert store.get_lock("doc1") == "wo:alice:A"

        res = client.post(
            "/wopi/files/doc1/refreshlock", headers=self._headers("wo:alice:A")
        )
        assert res.status_code == 200
        assert res.headers.get("X-WOPI-Lock") == "wo:alice:A"
        assert store.get_lock("doc1") == "wo:alice:A"

    def test_lock_contention_unlock_rejected_for_loser(self, client):
        """A loser cannot unlock — the winner's lock survives the attack."""
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)
        store.set_lock("doc1", "wo:alice:A", "alice")

        res = client.post("/wopi/files/doc1/unlock", headers=self._headers("wo:bob:B"))
        assert res.status_code == 409
        assert res.headers.get("X-WOPI-Lock") == "wo:alice:A"
        assert store.get_lock("doc1") == "wo:alice:A"

        # the holder unlocks cleanly
        res = client.post("/wopi/files/doc1/unlock", headers=self._headers("wo:alice:A"))
        assert res.status_code == 200
        assert store.get_lock("doc1") == ""

    def test_lock_contention_getlock_reports_holder(self, client):
        """GetLock must surface the current lock token to contenders."""
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)
        store.set_lock("doc1", "wo:alice:A", "alice")

        res = client.post("/wopi/files/doc1/getlock")
        assert res.status_code == 200
        assert res.headers.get("X-WOPI-Lock") == "wo:alice:A"

    def test_lock_contention_empty_token_rejected(self, client):
        """WOPI lock tokens must be non-empty — a blank Lock is refused."""
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)

        res = client.post("/wopi/files/doc1/lock", headers=self._headers(""))
        assert res.status_code == 400
        assert store.get_lock("doc1") == ""

    def test_lock_contention_relock_after_unlock(self, client):
        """Once the winner unlocks, the next editor takes the lock."""
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)

        res = client.post("/wopi/files/doc1/lock", headers=self._headers("wo:alice:A"))
        assert res.status_code == 200
        store.release_lock("doc1")

        res = client.post("/wopi/files/doc1/lock", headers=self._headers("wo:bob:B"))
        assert res.status_code == 200
        assert store.get_lock("doc1") == "wo:bob:B"

    def test_lock_contention_put_keeps_lock(self, client):
        """A successful save does not drop the lock mid-edit-session."""
        store = client.test_store  # type: ignore[attr-defined]
        _seed_doc(client)
        store.set_lock("doc1", "wo:alice:A", "alice")

        res = client.post(
            "/wopi/files/doc1/contents",
            content=_docx_bytes("More work"),
            headers=self._headers("wo:alice:A"),
        )
        assert res.status_code == 200
        assert store.get_lock("doc1") == "wo:alice:A"

    def test_lock_contention_handoff_write_cycle(self, client):
        """End-to-end: alice locks/edits/unlocks, then bob takes over."""
        store = client.test_store  # type: ignore[attr-defined]
        v0 = _docx_bytes("v0")
        _seed_doc(client, data=v0)
        assert store.get_content("doc1") == v0

        # alice: lock -> write -> unlock
        res = client.post("/wopi/files/doc1/lock", headers=self._headers("wo:alice:A"))
        assert res.status_code == 200
        v1 = _docx_bytes("v1 alice")
        res = client.post(
            "/wopi/files/doc1/contents", content=v1, headers=self._headers("wo:alice:A")
        )
        assert res.status_code == 200
        assert store.get_content("doc1") == v1
        res = client.post("/wopi/files/doc1/unlock", headers=self._headers("wo:alice:A"))
        assert res.status_code == 200
        assert store.get_lock("doc1") == ""

        # bob: lock (fresh token) -> write -> unlock
        res = client.post("/wopi/files/doc1/lock", headers=self._headers("wo:bob:B"))
        assert res.status_code == 200
        v2 = _docx_bytes("v2 bob")
        res = client.post(
            "/wopi/files/doc1/contents", content=v2, headers=self._headers("wo:bob:B")
        )
        assert res.status_code == 200
        assert store.get_content("doc1") == v2
        res = client.post("/wopi/files/doc1/unlock", headers=self._headers("wo:bob:B"))
        assert res.status_code == 200
        assert store.get_lock("doc1") == ""

# ----------------------------------------------------------------------
# File-Path-Traversal Defense (gate: -k "path_traversal")
# ----------------------------------------------------------------------
# Content bytes live at {content_dir}/{doc_id}.bin, so a WOPI file id is an
# opaque host id that must NEVER contain path separators or traversal
# segments. An attacker can smuggle separators into the URL path param
# URI-encoded (%2F, %5C, %2E) — FastAPI/Starlette decodes the segment before
# the handler runs — producing ids like "../secret" that the store's
# content_path() would join straight onto the content directory, escaping it
# and letting a crafted id read or write arbitrary files. The WOPI host must
# reject such ids outright. These tests prove both the rejection and that no
# file outside the content directory can be read or written through WOPI.


# Raw (pre-encoding) ids that must be rejected as traversal attempts. These
# are the DECODED values a handler receives (dot-dot, separators, NUL).
_TRAVERSAL_IDS = [
    "../secret",            # POSIX dot-dot escape
    "..\\..\\secret",       # Windows separator form
    "/etc/passwd",          # absolute path
    "a/../b",               # embedded traversal
    "sub/../../secret",     # deep escape
    "..",                   # bare dot-dot
    ".",                    # current directory
    "../\x00",              # NUL byte
    "C:\\Windows\\x",       # drive-qualified absolute (Windows)
]

# Opaque host ids that must keep working (no over-blocking).
_SAFE_IDS = [
    "doc1",
    "ghost",
    "doc-123",
    "doc_1",
    "550e8400-e29b-41d4-a716-446655440000",
    "hello.world",
    "subdoc.odt",
]

# Handlers only ever see SINGLE-segment ids (a URL cannot carry a raw '/'
# inside a path param — FastAPI/Starlette percent-decodes the path before
# routing, so an encoded %2F turns the id into extra segments and the route
# 404s before any handler runs). The id forms that DO reach a handler are
# backslash-encoded (%5C), dot-encoded (%2E) and NUL-encoded (%00):
#   ..%5C..%5Csecret / %2E%2E%5Csecret  ->  handler receives ..\..\secret
#   %2E%2E                                ->  handler receives ..
#   %2E                                   ->  handler receives .
#   ..%00x                                ->  handler receives ..\x00x
# Every one of these must be rejected with 400 by the WOPI host guard.
_HANDLER_REACHING = [
    "..%5C..%5Csecret",
    "..%5Csecret",
    "%2E%2E%5Csecret",
    "%2E%2E",
    "%2E",
    "..%00x",
]

# Encoded '/' forms are rejected by ROUTING itself (404, store untouched).
_ROUTING_REJECTED = [
    "..%2Fsecret",
    "%2E%2E%2Fsecret",
    "..%2F..%2Fsecret",
    "%2Fetc%2Fpasswd",
    "a%2F..%2Fb",
    "sub%2F..%2F..%2Fx",
]


def _url_enc(raw: str) -> str:
    """URI-encode an id exactly the way a malicious client URL would."""
    return urlparse.quote(raw, safe="")


def _content_dir(store) -> object:
    """The store's content directory (parent of any content file)."""
    return store.content_path("x").parent


class TestPathTraversal:
    def test_path_traversal_helper_rejects_dangerous_ids(self):
        """Unit level: the id guard rejects every traversal shape and accepts
        every ordinary opaque id."""
        from src.wopi.router import _invalid_doc_id

        for bad in _TRAVERSAL_IDS + [""]:
            assert _invalid_doc_id(bad), f"{bad!r} must be rejected"
        for good in _SAFE_IDS:
            assert not _invalid_doc_id(good), f"{good!r} must be accepted"

    def test_path_traversal_wopi_endpoints_all_reject_traversal_id(self, client):
        """Every WOPI host endpoint rejects a traversal id that reaches the
        handler — even when a matching store row was planted directly
        (simulating a malicious upload that registered the id before
        routing)."""
        store = client.test_store  # type: ignore[attr-defined]
        store.init("..\\..\\secret", "secret.docx")

        base = f"/wopi/files/{_url_enc('..\\..\\secret')}"
        cases = [
            ("GET", base, None),
            ("GET", f"{base}/contents", None),
            ("POST", f"{base}/contents", b"x"),
            ("POST", f"{base}/lock", None),
            ("POST", f"{base}/unlock", None),
            ("POST", f"{base}/refreshlock", None),
            ("POST", f"{base}/getlock", None),
        ]
        for method, url, body in cases:
            res = client.request(method, url, content=body)
            assert res.status_code == 400, (
                f"expected WOPI host guard 400, got {res.status_code}: {method} {url}"
            )

    def test_path_traversal_get_file_never_leaks_outside_content(self, client):
        """A traversal id must never return bytes of a file outside content_dir.

        A decoy 'secret.bin' is placed one directory ABOVE the content dir and
        a store row with id '../secret' is planted — content_path('../secret')
        resolves exactly onto that decoy, so the unguarded host WOULD leak it.
        The WOPI host must refuse at every layer and reveal nothing."""
        store = client.test_store  # type: ignore[attr-defined]
        content_dir = _content_dir(store)
        decoy = content_dir.parent / "secret.bin"
        decoy.write_bytes(b"TOP-SECRET-CONTENTS")

        store.init("../secret", "secret.docx")
        # Prove this is a real escape, not a vacuous setup:
        assert store.content_path("../secret").resolve() == decoy.resolve()
        assert decoy.resolve() != content_dir.resolve()

        # %2F forms are bounced by routing (404); %5C / %2E forms reach the
        # handler and must be bounced by the id guard (400). Neither may
        # return the decoy's bytes.
        for enc in _ROUTING_REJECTED + _HANDLER_REACHING:
            res = client.get(f"/wopi/files/{enc}/contents")
            assert res.status_code in (400, 404), enc
            assert b"TOP-SECRET" not in res.content, enc
        # CheckFileInfo/GetFile metadata must reject it too, silently.
        res = client.get("/wopi/files/..%2Fsecret")
        assert res.status_code in (400, 404)
        res = client.get("/wopi/files/..%5Csecret")
        assert res.status_code == 400

    def test_path_traversal_put_file_never_writes_outside_content(self, client):
        """A traversal id must not be able to write a file outside content_dir."""
        store = client.test_store  # type: ignore[attr-defined]
        content_dir = _content_dir(store)
        victim = content_dir.parent / "pwned.bin"
        store.init("../pwned", "pwned.docx")

        # %2F: blocked at routing. %5C: blocked by the id guard.
        for enc in ["..%2Fpwned", "..%5Cpwned"]:
            res = client.post(f"/wopi/files/{enc}/contents", content=b"owned")
            assert res.status_code in (400, 404), enc
            assert not victim.exists(), "no file may be written outside content_dir"
            # and nothing was stored under the traversal id in the store either
            assert store.get_content("../pwned") is None

    def test_path_traversal_rejects_absolute_and_windows_ids(self, client):
        """Absolute paths, Windows separators and NUL bytes are rejected by
        the WOPI host (either at routing or by the id guard)."""
        for enc in ["%2Fetc%2Fpasswd", "a%2F..%2Fb", "sub%2F..%2F..%2Fx"]:
            res = client.get(f"/wopi/files/{enc}")
            assert res.status_code in (400, 404), f"{enc} -> {res.status_code}"
        for enc in ["..%5C..%5Cetc%5Cpasswd", "..%00x", "%2E%2E"]:
            res = client.get(f"/wopi/files/{enc}")
            assert res.status_code == 400, f"{enc} -> {res.status_code}"

    def test_path_traversal_legitimate_opaque_ids_still_work(self, client):
        """The defense must not over-block ordinary opaque host ids."""
        for doc_id in ["doc-123", "550e8400-e29b-41d4-a716-446655440000", "notes.v2"]:
            _seed_doc(client, doc_id=doc_id, name=f"{doc_id}.docx")
            res = client.get(f"/wopi/files/{doc_id}")
            assert res.status_code == 200, doc_id
            assert res.json()["BaseFileName"] == f"{doc_id}.docx"


def test_sanitize_lifts_list_out_of_paragraph():
    """A contenteditable can nest <ul> inside <p>; the block must be lifted."""
    out = sanitize_html("<p><ul><li>x</li></ul></p>")
    assert "<ul><li>x</li></ul>" in out
    assert "<p><ul>" not in out
    # surrounding text is preserved as a sibling paragraph
    out2 = sanitize_html("<p>before<ul><li>x</li></ul>after</p>")
    assert "before" in out2 and "after" in out2


def test_sanitize_nests_stray_list_into_preceding_li():
    """A <ul> directly inside another <ul> (Chromium Tab quirk) re-nests."""
    out = sanitize_html(
        "<ul><li>first item</li><ul><li>second item</li></ul></ul>"
    )
    norm = out.replace("\n", "")
    assert ("<ul><li>first item<ul><li>second item</li></ul></li></ul>"
            in norm), norm
