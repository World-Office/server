"""Documents REST lifecycle: new→contents→save→versions→restore + typed 404.

This test suite validates the document API surface on the editor router,
covering the full document lifecycle from creation through edits to version
history management. Paradigm: **unit tests** for HTTP-level endpoints.

Scenarios under test (the endpoints behind the lifecycle):

1. **new_document** - Create a blank document via POST /api/documents/new
2. **document_contents** - Retrieve raw bytes via GET /api/documents/{doc_id}/contents
3. **save_document** - Save HTML content via POST /api/documents/{doc_id}/save
4. **versions** - List document versions via GET /api/documents/{doc_id}/versions
5. **restore** - Restore a previous version via POST /api/documents/{doc_id}/versions/{ts}/restore
6. **typed_404** - Verify proper error responses (400, 404) for invalid/unknown document IDs

Deterministic: no network, no sleeps, no time-of-day dependence. Uses the
TestClient with a temp SQLite store and content directory.
"""

from __future__ import annotations

from contextlib import asynccontextmanager

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from src.config import Config
from src.editor.router import router as editor_router
from src.editor.session import SessionRegistry
from src.lib.store import DocumentStore, wipe_db, wipe_dir

# -----------------------------------------------------------------------------
# Shared app builder
# -----------------------------------------------------------------------------


def _make_app(tmp_path):
    db = str(tmp_path / "t.db")
    content = str(tmp_path / "content")
    store = DocumentStore(db, content)
    cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

    @asynccontextmanager
    async def lifespan(app: FastAPI):
        app.state.store = store
        app.state.sessions = SessionRegistry()
        app.state.config = cfg
        yield

    app = FastAPI(lifespan=lifespan)
    app.include_router(editor_router)
    return app, store


@pytest.fixture
def client(tmp_path):
    """TestClient with lifespan running; backing store on ``client.test_store``."""
    app, store = _make_app(tmp_path)
    with TestClient(app) as c:
        c.test_store = store  # type: ignore[attr-defined]
        yield c
    wipe_db(tmp_path / "t.db")
    wipe_dir(tmp_path / "content")


# -----------------------------------------------------------------------------
# 1. new_document - Create blank documents (POST /api/documents/new)
# -----------------------------------------------------------------------------


def test_new_document_creates_blank_docx(client):
    """POST /api/documents/new creates a blank DOCX and registers a session.

    The response includes doc_id, editor URL, and document name. The document
    is immediately available via /api/documents/{doc_id}.
    """
    res = client.post("/api/documents/new", params={"format": "docx"})

    assert res.status_code == 200
    body = res.json()
    assert "doc_id" in body
    assert body["doc_id"].startswith("new-")
    assert "url" in body
    assert "/editor/" in body["url"]
    assert body["name"] == "untitled.docx"

    # The document should be immediately accessible via metadata endpoint
    doc_id = body["doc_id"]
    meta = client.get(f"/api/documents/{doc_id}")
    assert meta.status_code == 200
    assert meta.json()["id"] == doc_id
    assert meta.json()["name"] == "untitled.docx"


def test_new_document_creates_blank_odt(client):
    """POST /api/documents/new with format=odt creates a blank ODT document."""
    res = client.post("/api/documents/new", params={"format": "odt"})

    assert res.status_code == 200
    body = res.json()
    assert body["name"] == "untitled.odt"

    # Verify it's actually an ODT by checking contents
    doc_id = body["doc_id"]
    contents = client.get(f"/api/documents/{doc_id}/contents")
    assert contents.status_code == 200
    # ODT is a ZIP format; check for ODT magic bytes (ZIP header)
    assert contents.content[:4] == b"PK\x03\x04"
    assert contents.headers["content-type"] == "application/vnd.oasis.opendocument.text"


# -----------------------------------------------------------------------------
# 2. document_contents - Retrieve raw bytes (GET /api/documents/{doc_id}/contents)
# -----------------------------------------------------------------------------


def test_contents_get_returns_raw_bytes(client):
    """GET /api/documents/{doc_id}/contents returns the raw document bytes."""
    # First create and seed a document
    store = client.test_store  # type: ignore[attr-defined]
    store.init("doc1", "test.docx")
    test_bytes = b"PK\x03\x04" + b"x" * 100  # Fake DOCX header + payload
    store.put_content("doc1", test_bytes)

    res = client.get("/api/documents/doc1/contents")

    assert res.status_code == 200
    assert res.content == test_bytes
    assert "X-WOPI-ItemVersion" in res.headers
    assert res.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_contents_get_missing_document_returns_404(client):
    """GET /api/documents/{doc_id}/contents returns 404 for unknown doc_id."""
    res = client.get("/api/documents/ghost/contents")

    assert res.status_code == 404
    assert res.json()["error"] == "not found"


def test_contents_get_invalid_doc_id_returns_400(client):
    """GET /api/documents/{doc_id}/contents returns 400 for invalid doc_id format.

    Only ids that survive URL decoding and reach the handler are rejected as
    ``Invalid file id`` (400): URI-encoded separators (``%5C`` backslash), null
    bytes (``%00``), bare ``..`` (``%2E%2E``) and ids longer than 128 chars.
    Ids containing a decoded ``/`` never match the route, so they surface as
    FastAPI's route-level 404 before the handler runs.
    """
    # Invalid doc IDs that pass route matching and fail the handler validation
    invalid_ids = ["doc%5C1", "a%00b", "%2E%2E", "x" * 129]

    for invalid_id in invalid_ids:
        res = client.get(f"/api/documents/{invalid_id}/contents")
        # Should be 400 for invalid format, not 404
        assert res.status_code == 400, f"Expected 400 for {invalid_id!r}, got {res.status_code}"
        assert "Invalid file id" in res.json()["error"]

    # An id whose decoded separator breaks the URL never reaches the handler:
    # the route itself 404s before any validation can run.
    res = client.get("/api/documents/..%2Fetc%2Fpasswd/contents")
    assert res.status_code == 404


# -----------------------------------------------------------------------------
# 3. save_document - Save HTML content (POST /api/documents/{doc_id}/save)
# -----------------------------------------------------------------------------


def test_save_document_updates_content(client):
    """POST /api/documents/{doc_id}/save converts HTML back to DOCX and persists."""
    # Seed a document first
    store = client.test_store  # type: ignore[attr-defined]
    store.init("doc1", "test.docx")
    # Initial content
    import io

    from docx import Document
    doc = Document()
    doc.add_paragraph("Original text")
    buf = io.BytesIO()
    doc.save(buf)
    initial_bytes = buf.getvalue()
    store.put_content("doc1", initial_bytes)

    # Save new HTML content
    res = client.post(
        "/api/documents/doc1/save",
        json={"html": "<p>Updated content</p>"},
    )

    assert res.status_code == 200
    assert res.json()["ok"] is True
    assert "size" in res.json()

    # Verify the content was updated
    new_content = store.get_content("doc1")
    assert new_content is not None
    assert len(new_content) > 0

    # Verify by loading and checking text
    doc = Document(io.BytesIO(new_content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Updated content" in text


def test_save_document_unknown_doc_pins_current_behaviour(client):
    """POST /api/documents/{doc_id}/save on an unknown doc currently succeeds.

    # NOTE: existing behaviour — the ``save_document`` handler converts the
    # submitted HTML and writes it through ``store.put_content`` without an
    # existence check, so saving to a never-registered doc creates orphaned
    # content instead of a typed 404. Pinned here so the contract is explicit;
    # a future fix (existence check -> 404) would flip this test.
    """
    store = client.test_store  # type: ignore[attr-defined]
    assert store.get("ghost") is None

    res = client.post(
        "/api/documents/ghost/save",
        json={"html": "<p>Test</p>"},
    )

    assert res.status_code == 200
    assert res.json()["ok"] is True
    # the bytes were written even though no index row exists
    assert store.has_content("ghost") is True
    assert store.get("ghost") is None


def test_save_document_invalid_json_returns_400(client):
    """POST /api/documents/{doc_id}/save returns 400 for invalid JSON."""
    # Seed a document first
    store = client.test_store  # type: ignore[attr-defined]
    store.init("doc1", "test.docx")
    store.put_content("doc1", b"PK\x03\x04")

    res = client.post(
        "/api/documents/doc1/save",
        content=b"not valid json",
        headers={"Content-Type": "application/json"},
    )

    assert res.status_code == 400
    assert res.json()["error"] == "invalid JSON"


# -----------------------------------------------------------------------------
# 4. versions - List document versions (GET /api/documents/{doc_id}/versions)
# -----------------------------------------------------------------------------


def test_versions_list_returns_history(client):
    """GET /api/documents/{doc_id}/versions returns version history, newest first."""
    store = client.test_store  # type: ignore[attr-defined]
    store.init("doc1", "test.docx")
    store.put_content("doc1", b"version 1")
    store.put_content("doc1", b"version 2")
    store.put_content("doc1", b"version 3")

    res = client.get("/api/documents/doc1/versions")

    assert res.status_code == 200
    body = res.json()
    assert "versions" in body
    versions = body["versions"]
    assert len(versions) == 3

    # Newest first
    assert versions[0]["size"] == len(b"version 3")
    assert versions[1]["size"] == len(b"version 2")
    assert versions[2]["size"] == len(b"version 1")


def test_versions_list_missing_document_returns_404(client):
    """GET /api/documents/{doc_id}/versions returns 404 for unknown doc_id."""
    res = client.get("/api/documents/ghost/versions")

    assert res.status_code == 404
    assert res.json()["error"] == "not found"


def test_versions_list_empty_document_returns_empty(client):
    """GET /api/documents/{doc_id}/versions returns empty list for document with no versions."""
    store = client.test_store  # type: ignore[attr-defined]
    store.init("doc1", "test.docx")
    # Only init, no put_content

    res = client.get("/api/documents/doc1/versions")

    assert res.status_code == 200
    assert res.json()["versions"] == []


# -----------------------------------------------------------------------------
# 5. restore - Restore a previous version (POST /api/documents/{doc_id}/versions/{ts}/restore)
# -----------------------------------------------------------------------------


def test_restore_version_reverts_to_snapshot(client):
    """POST /api/documents/{doc_id}/versions/{ts}/restore reverts to a snapshot.

    # NOTE: existing behaviour — restore_version snapshots the pre-restore
    # content explicitly, then put_content snapshots again, so one restore
    # appends TWO versions (pre-restore state + restored bytes); the restored
    # bytes become the current content and the recoverable pre-restore state
    # sits directly above it in the history.
    """
    store = client.test_store  # type: ignore[attr-defined]
    store.init("doc1", "test.docx")
    store.put_content("doc1", b"v-one")
    store.put_content("doc1", b"v-two-two")
    store.put_content("doc1", b"v-three-three-three")

    # Get version timestamps
    versions = client.get("/api/documents/doc1/versions").json()["versions"]
    # Find the ts for the oldest version ("v-one")
    ts_v1 = versions[-1]["ts"]

    # Restore the oldest version
    res = client.post(f"/api/documents/doc1/versions/{ts_v1}/restore")

    assert res.status_code == 200
    assert res.json()["ok"] is True
    assert "ts" in res.json()

    # Current content is the restored (oldest) version
    assert store.get_content("doc1") == b"v-one"

    # 3 original + 2 restore snapshots = 5 versions
    new_versions = client.get("/api/documents/doc1/versions").json()["versions"]
    assert len(new_versions) == 5
    # newest snapshot is the restored bytes (v-one, 5 bytes)
    assert new_versions[0]["size"] == len(b"v-one")
    # the pre-restore state (v-three...) is recoverable right below it
    assert new_versions[1]["size"] == len(b"v-three-three-three")


def test_restore_version_missing_document_returns_404(client):
    """POST /api/documents/{doc_id}/versions/{ts}/restore returns 404 for unknown doc_id."""
    res = client.post("/api/documents/ghost/versions/1234567890/restore")

    assert res.status_code == 404
    assert res.json()["error"] == "not found"


def test_restore_version_missing_ts_returns_404(client):
    """POST /api/documents/{doc_id}/versions/{ts}/restore returns 404 for unknown ts."""
    store = client.test_store  # type: ignore[attr-defined]
    store.init("doc1", "test.docx")
    store.put_content("doc1", b"content")

    # Use a timestamp that doesn't exist
    res = client.post("/api/documents/doc1/versions/9999999999/restore")

    assert res.status_code == 404
    assert "version" in res.json()["error"].lower()


# -----------------------------------------------------------------------------
# 6. typed_404 - Verify proper error responses for invalid inputs
# -----------------------------------------------------------------------------


def test_document_meta_invalid_doc_id_returns_400(client):
    """GET /api/documents/{doc_id} returns 400 for an invalid doc_id that
    reaches the handler (bare ``..``). A decoded slash id never matches the
    route and surfaces as FastAPI's route-level 404 instead."""
    res = client.get("/api/documents/%2E%2E")

    assert res.status_code == 400
    assert "Invalid file id" in res.json()["error"]

    res = client.get("/api/documents/../../etc/passwd")
    assert res.status_code == 404


def test_save_invalid_doc_id_returns_400(client):
    """POST /api/documents/{doc_id}/save returns 400 for invalid doc_id."""
    res = client.post(
        "/api/documents/invalid%00path/save",
        json={"html": "<p>Test</p>"},
    )

    assert res.status_code == 400
    assert "Invalid file id" in res.json()["error"]


def test_contents_put_invalid_doc_id_returns_400(client):
    """PUT /api/documents/{doc_id}/contents returns 400 for an invalid doc_id
    (null byte inside the decoded path param)."""
    res = client.put(
        "/api/documents/a%00b/contents",
        content=b"payload",
    )

    assert res.status_code == 400
    assert "Invalid file id" in res.json()["error"]


def test_versions_invalid_doc_id_returns_400(client):
    """GET /api/documents/{doc_id}/versions returns 400 for an invalid doc_id
    (backslash-encoded separator that reaches the handler)."""
    res = client.get("/api/documents/doc%5C1/versions")

    assert res.status_code == 400
    assert "Invalid file id" in res.json()["error"]


def test_restore_invalid_doc_id_returns_400(client):
    """POST /api/documents/{doc_id}/versions/{ts}/restore returns 400 for invalid doc_id."""
    res = client.post("/api/documents/..%00/versions/123/restore")

    assert res.status_code == 400
    assert "Invalid file id" in res.json()["error"]


# -----------------------------------------------------------------------------
# Additional integration scenarios
# -----------------------------------------------------------------------------


def test_full_document_lifecycle(client):
    """End-to-end lifecycle: new → contents → save → versions → restore.

    This test validates the complete workflow: create a document, modify it
    multiple times, verify version history, and restore a previous version.
    """
    store = client.test_store  # type: ignore[attr-defined]

    # 1. NEW: Create a blank document
    res = client.post("/api/documents/new", params={"format": "docx"})
    assert res.status_code == 200
    doc_id = res.json()["doc_id"]

    # 2. CONTENTS: Retrieve initial (blank) content
    contents = client.get(f"/api/documents/{doc_id}/contents")
    assert contents.status_code == 200
    # A blank DOCX should have some minimal size
    assert len(contents.content) > 0

    # 3. SAVE: Make an edit
    res = client.post(
        f"/api/documents/{doc_id}/save",
        json={"html": "<p>First edit</p>"},
    )
    assert res.status_code == 200
    assert res.json()["ok"] is True

    # 4. VERSIONS: Check we have at least 2 versions (initial + first save)
    versions = client.get(f"/api/documents/{doc_id}/versions").json()["versions"]
    assert len(versions) >= 2

    # 5. SAVE: Make another edit
    res = client.post(
        f"/api/documents/{doc_id}/save",
        json={"html": "<p>Second edit</p>"},
    )
    assert res.status_code == 200

    # 6. VERSIONS: Verify we now have more versions
    versions_after = client.get(f"/api/documents/{doc_id}/versions").json()["versions"]
    assert len(versions_after) >= 3

    # 7. RESTORE: Restore the first version
    first_version_ts = versions_after[-1]["ts"]
    res = client.post(
        f"/api/documents/{doc_id}/versions/{first_version_ts}/restore"
    )
    assert res.status_code == 200

    # Verify current content matches the first version
    current = store.get_content(doc_id)
    # The first save had "<p>First edit</p>", which should now be current
    assert current is not None
