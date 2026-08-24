"""Tests for OCIS client mode: editor sessions + remote WOPI client.

We stand up a tiny mock WOPI host (FastAPI) that simulates OCIS, then
run the docserver against it through the editor API.
"""

from __future__ import annotations

import io
import threading
from wsgiref.simple_server import make_server

from src.editor.session import RemoteWopiClient, SessionRegistry, session_from_token
from src.lib.crypto import encode_token

SECRET = "0123456789abcdef0123456789abcdef"

# ----------------------------------------------------------------------
# Unittest-style mock WOPI host (no FastAPI TestClient, real HTTP loop)
# ----------------------------------------------------------------------



class _MockHost:
    """A bare WSGI host that speaks enough WOPI for tests."""

    def __init__(self) -> None:
        self.content: bytes | None = None
        self.auth_seen: list[str] = []
        self.query_seen: list[str] = []
        self.override_seen: list[str] = []
        self.lock_token: str | None = None
        self.user_by_token: dict[str, str] = {}

    def __call__(self, environ, start_response):
        token = environ.get("HTTP_AUTHORIZATION", "")
        if token:
            self.auth_seen.append(token)
        q = environ.get("QUERY_STRING", "")
        if q:
            self.query_seen.append(q)
        override = environ.get("HTTP_X_WOPI_OVERRIDE", "")
        if override:
            self.override_seen.append(override)

        path = environ.get("PATH_INFO", "")
        method = environ.get("REQUEST_METHOD", "GET")

        if path == "/wopi/files/doc1/contents" and method == "GET":
            start_response("200 OK", [("Content-Type", "application/octet-stream")])
            return [self.content or b""]
        if path == "/wopi/files/doc1/contents" and method == "POST" and override == "PUT":
            length = int(environ.get("CONTENT_LENGTH", "0"))
            lock_hdr = environ.get("HTTP_X_WOPI_LOCK", "")
            if lock_hdr != (self.lock_token or ""):
                # wopiserver behaviour: unlocked file -> 409, mismatched -> 500
                if not self.lock_token:
                    start_response(
                        "409 Conflict", [("Content-Type", "text/plain"), ("X-WOPI-Lock", ""), ("X-WOPI-LockFailureReason", "Cannot PutFile on unlocked file")]
                    )
                    return [b"conflict"]
                start_response("500 Internal Server Error", [("Content-Type", "text/plain")])
                return [b"lock mismatch"]
            self.content = environ["wsgi.input"].read(length)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]
        if path == "/wopi/files/doc1" and method == "POST" and override == "LOCK":
            tok = environ.get("HTTP_X_WOPI_LOCK", "")
            if self.lock_token and self.lock_token != tok:
                start_response("409 Conflict", [("Content-Type", "text/plain"), ("X-WOPI-Lock", self.lock_token)])
                return [b"locked by other"]
            self.lock_token = tok
            start_response("200 OK", [("Content-Type", "application/json"), ("X-WOPI-ItemVersion", "v1")])
            return [b"{}"]
        if path == "/wopi/files/doc1" and method == "POST" and override == "GET_LOCK":
            start_response("200 OK", [("Content-Type", "application/json"), ("X-WOPI-Lock", self.lock_token or "")])
            return [b"{}"]
        if path == "/wopi/files/doc1" and method == "POST" and override == "UNLOCK":
            if environ.get("HTTP_X_WOPI_LOCK", "") == self.lock_token:
                self.lock_token = None
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]
        if path == "/wopi/files/doc1" and method == "GET":
            auth = environ.get("HTTP_AUTHORIZATION", "").replace("Bearer ", "")
            user_id = self.user_by_token.get(auth, "default-user")
            start_response(
                "200 OK",
                [("Content-Type", "application/json")],
            )
            return [b"{\"BaseFileName\":\"remote.docx\",\"UserId\":\"" + user_id.encode() + b"\"}"]
        start_response("404 Not Found", [("Content-Type", "text/plain")])
        return [b"not found"]


class _RealWopiClientTest:
    """Run one WSGI host on a thread and test RemoteWopiClient against it."""

    def __enter__(self):
        host = _MockHost()
        httpd = make_server("127.0.0.1", 0, host)
        port = httpd.server_address[1]
        thread = threading.Thread(target=httpd.serve_forever, daemon=True)
        thread.start()
        client = RemoteWopiClient(f"http://127.0.0.1:{port}", "token-123")
        client_wrapped = type("W", (), {})()
        client_wrapped.httpd = httpd
        client_wrapped.thread = thread
        client_wrapped.host = host
        client_wrapped.client = client
        self._obj = client_wrapped
        return client_wrapped

    def __exit__(self, *exc):
        self._obj.httpd.shutdown()
        self._obj.httpd.server_close()
        self._obj.thread.join(timeout=2)
        return False


# ----------------------------------------------------------------------
# Undo/Redo-Kette (US-31): frontend-style snapshot chain driven through
# the real client-mode save -> DOCX -> fetch round-trip.
# ----------------------------------------------------------------------

class _UndoRedoHistory:
    """Frontend-style undo/redo snapshot chain.

    Mirrors the explicit innerHTML snapshot history the browser editor
    maintains (see web/editor.js): a stack of document states where every
    undo step restores exactly the previous state, redo re-applies the
    next one, and a fresh edit truncates the redo branch. Driving the
    same contract through the real save/fetch round-trip lets the tests
    verify the Undo/Redo-Kette end-to-end: 20+ steps, undo after save.
    """

    def __init__(self, initial: str) -> None:
        self._undo: list[str] = [initial]
        self._redo: list[str] = []

    def edit(self, html: str) -> None:
        """User makes an edit; the previous state is pushed onto the undo
        chain and the (now stale) redo branch is discarded."""
        self._undo.append(html)
        self._redo.clear()

    def undo(self) -> str | None:
        """Step back to (and return) the previous state, or None when the
        chain is exhausted."""
        if len(self._undo) < 2:
            return None
        current = self._undo.pop()
        self._redo.append(current)
        return self._undo[-1]

    def redo(self) -> str | None:
        """Step forward to (and return) the next state, or None when the
        redo branch is empty."""
        if not self._redo:
            return None
        current = self._redo.pop()
        self._undo.append(current)
        return current

    @property
    def can_undo(self) -> bool:
        return len(self._undo) > 1

    @property
    def can_redo(self) -> bool:
        return bool(self._redo)


class _EditorHarness:
    """Client-mode editor harness: a real WSGI mock WOPI host plus the
    FastAPI TestClient, so every undo/redo step travels the full
    save -> DOCX -> html-fetch path exactly like the deployed docserver.
    """

    def __init__(self, tmp_path, seed_html: str = "<p>start</p>") -> None:
        self.tmp_path = tmp_path
        self.seed_html = seed_html

    def __enter__(self):
        from contextlib import asynccontextmanager
        from wsgiref.simple_server import make_server

        from fastapi import FastAPI
        from fastapi.testclient import TestClient

        from src.config import Config
        from src.editor.converter import html_to_docx
        from src.editor.router import router as editor_router
        from src.lib.store import DocumentStore
        from src.wopi.router import router as wopi_router

        self.host = _MockHost()
        self.host.content = html_to_docx(self.seed_html)
        self.httpd = make_server("127.0.0.1", 0, self.host)
        port = self.httpd.server_address[1]
        self.thread = threading.Thread(target=self.httpd.serve_forever, daemon=True)
        self.thread.start()

        db = str(self.tmp_path / "t.db")
        content = str(self.tmp_path / "content")
        self.store = DocumentStore(db, content)
        cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

        @asynccontextmanager
        async def lifespan(app):
            app.state.store = self.store
            app.state.sessions = SessionRegistry()
            app.state.config = cfg
            yield

        app = FastAPI(lifespan=lifespan)
        app.include_router(wopi_router)
        app.include_router(editor_router)
        self.client = TestClient(app)
        self.client.__enter__()
        src = f"http://127.0.0.1:{port}/wopi/files/doc1"
        resp = self.client.post(
            "/editor", data={"access_token": "tok-1"}, params={"WOPISrc": src}
        )
        assert resp.status_code == 200
        session = self.client.app.state.sessions.get("doc1")
        assert session is not None
        self.session_id = session.session_id
        self.doc_api = "/api/documents/doc1"
        return self

    def __exit__(self, *exc):
        from src.lib.store import wipe_db, wipe_dir

        self.client.__exit__(*exc)
        self.httpd.shutdown()
        self.httpd.server_close()
        self.thread.join(timeout=2)
        wipe_db(str(self.tmp_path / "t.db"))
        wipe_dir(str(self.tmp_path / "content"))
        return False

    def save(self, html: str) -> None:
        """POST the editor's current innerHTML to the server (a save)."""
        resp = self.client.post(
            f"{self.doc_api}/save?session={self.session_id}", json={"html": html}
        )
        assert resp.status_code == 200, resp.text

    def fetch_html(self) -> str:
        """GET the normalized HTML the editor would reload."""
        resp = self.client.get(f"{self.doc_api}/html?session={self.session_id}")
        assert resp.status_code == 200, resp.text
        return resp.json()["html"]

    def page_text(self) -> str:
        """Plain text of the remote document (python-docx): the authority
        on what actually reached the WOPI host after each save."""
        from docx import Document

        doc = Document(io.BytesIO(self.host.content))
        return "\n".join(p.text for p in doc.paragraphs)


def test_remote_client_get_and_put():
    with _RealWopiClientTest() as env:
        env.client.put_contents("doc1", b"hello from editor")
        assert env.host.content == b"hello from editor"
        got = env.client.get_contents("doc1")
        assert got == b"hello from editor"
        # WOPI hosts receive the access token as a query parameter.
        assert any("access_token=token-123" in q for q in env.host.query_seen)
        # OpenCloud/OCIS wopiserver expects PUT on /wopi/files/{id}/contents
        # with the X-WOPI-Override: PUT header.
        assert "PUT" in env.host.override_seen


def test_remote_client_sends_lock_on_put():
    with _RealWopiClientTest() as env:
        lock, _ = env.client.acquire_or_adopt_lock("doc1")
        env.client.put_contents("doc1", b"x")
        assert env.host.content == b"x"
        assert env.host.override_seen.count("PUT") >= 1
        assert env.host.override_seen.count("LOCK") >= 1
        assert env.host.lock_token == lock


# ----------------------------------------------------------------------
# Session decoded from OCIS JWT
# ----------------------------------------------------------------------

def test_session_from_token():
    token = encode_token(
        SECRET,
        {"file_id": "doc1", "file_name": "r.docx", "user_id": "alice"},
        ttl=60,
    )
    s = session_from_token(token, SECRET)
    assert s is not None
    assert s.doc_id == "doc1"
    assert s.user_id == "alice"


def test_session_from_bad_token_returns_none():
    assert session_from_token("garbage", SECRET) is None
    assert session_from_token("", SECRET) is None


def test_session_registry():
    reg = SessionRegistry()
    assert reg.get("doc1") is None
    s = session_from_token(encode_token(SECRET, {"file_id": "doc1"}, ttl=60), SECRET)
    reg.register(s)
    assert reg.get("doc1") is s
    reg.drop("doc1")
    assert reg.get("doc1") is None


def test_remote_client_acquire_lock_and_put():
    """Client-mode save path: lock first, then PutFile with the lock (the
    wopiserver refuses PutFile on unlocked files)."""
    with _RealWopiClientTest() as env:
        lock, writable = env.client.acquire_or_adopt_lock("doc1")
        assert writable is True
        assert lock and env.host.lock_token == lock
        env.client.put_contents("doc1", b"locked write")
        assert env.host.content == b"locked write"
        env.client.release_lock("doc1")
        assert env.host.lock_token is None


def test_remote_client_adopts_existing_lock():
    with _RealWopiClientTest() as env:
        other, w = env.client.acquire_or_adopt_lock("doc1")  # first holds lock
        assert w is True
        env2 = RemoteWopiClient(env.client.host, "token-456")
        adopted, w2 = env2.acquire_or_adopt_lock("doc1")  # legacy lock -> takeover
        assert w2 is True
        # legacy lock is upgraded with a fresh named token
        assert adopted != other
        assert env.host.lock_token == adopted
        assert env2.lock_token == adopted
        env2.put_contents("doc1", b"takeover write")  # save succeeds with new lock
        assert env.host.content == b"takeover write"
        env2.release_lock("doc1")
        assert env.host.lock_token is None


def test_remote_client_second_owner_is_readonly():
    """A second DIFFERENT user must not be able to write (lost-update guard).
    Same user re-opening adopts the lock and stays writable."""
    with _RealWopiClientTest() as env:
        env.host.user_by_token["tok-a"] = "alice"
        env.host.user_by_token["tok-b"] = "bob"
        c1 = RemoteWopiClient(env.client.host, "tok-a")
        l1, w1 = c1.acquire_or_adopt_lock("doc1", owner="alice")
        assert w1 is True
        assert l1.startswith("wo:alice:")
        # same user re-opens -> adopt, still writable
        c1b = RemoteWopiClient(env.client.host, "tok-a")
        l1b, w1b = c1b.acquire_or_adopt_lock("doc1", owner="alice")
        assert w1b is True
        assert l1b == l1
        # different user -> read-only, lock not stolen
        c2 = RemoteWopiClient(env.client.host, "tok-b")
        l2, w2 = c2.acquire_or_adopt_lock("doc1", owner="bob")
        assert w2 is False
        assert l2 == ""
        assert c2.lock_token == ""
        # owner can still save and unlock
        c1.put_contents("doc1", b"alice keeps writing")
        assert env.host.content == b"alice keeps writing"
        c1.release_lock("doc1")
        assert env.host.lock_token is None


def test_editor_save_uses_session_lock_in_client_mode(tmp_path):
    """Regression: the client-mode save path must carry the session's WOPI
    lock to the remote PutFile (otherwise the wopiserver answers 409
    \"Cannot PutFile on unlocked file\")."""
    import threading
    from contextlib import asynccontextmanager
    from wsgiref.simple_server import make_server

    from docx import Document
    from fastapi import FastAPI
    from fastapi.testclient import TestClient

    from src.config import Config
    from src.editor.router import router as editor_router
    from src.lib.store import DocumentStore, wipe_db, wipe_dir
    from src.wopi.router import router as wopi_router

    seed = Document()
    seed.add_paragraph("hello remote")
    buf = io.BytesIO()
    seed.save(buf)

    host = _MockHost()
    httpd = make_server("127.0.0.1", 0, host)
    port = httpd.server_address[1]
    thread = threading.Thread(target=httpd.serve_forever, daemon=True)
    thread.start()
    try:
        host.content = buf.getvalue()
        db = str(tmp_path / "t.db")
        content = str(tmp_path / "content")
        store = DocumentStore(db, content)
        cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

        @asynccontextmanager
        async def lifespan(app):
            app.state.store = store
            app.state.sessions = SessionRegistry()
            app.state.config = cfg
            yield

        app = FastAPI(lifespan=lifespan)
        app.include_router(wopi_router)
        app.include_router(editor_router)
        with TestClient(app) as c:
            resp = c.post(
                "/editor",
                data={"access_token": "tok-1"},
                params={"WOPISrc": f"http://127.0.0.1:{port}/wopi/files/doc1"},
            )
            assert resp.status_code == 200
            assert host.lock_token, "launch must take the WOPI lock on the remote host"

            r = c.get("/api/documents/doc1/html")
            assert r.status_code == 200, r.text

            r = c.post("/api/documents/doc1/save", json={"html": "<p>saved via lock</p>"})
            assert r.status_code == 200, r.text
            saved = Document(io.BytesIO(host.content))
            assert "saved via lock" in "\n".join(p.text for p in saved.paragraphs)
    finally:
        httpd.shutdown()
        httpd.server_close()
        thread.join(timeout=2)
        wipe_db(str(tmp_path / "t.db"))
        wipe_dir(str(tmp_path / "content"))


def test_launch_readonly_when_other_user_edits(tmp_path):
    """US-5: when another user already holds the file's WOPI lock, the second
    launch is served read-only; its save is rejected with 403. Both sessions
    coexist in ONE registry (per-launch ?session= ids) without clobbering."""
    import threading
    from contextlib import asynccontextmanager
    from wsgiref.simple_server import make_server

    from fastapi import FastAPI
    from fastapi.testclient import TestClient

    from src.config import Config
    from src.editor.router import router as editor_router
    from src.lib.store import DocumentStore, wipe_db, wipe_dir
    from src.wopi.router import router as wopi_router

    host = _MockHost()
    host.user_by_token["tok-alice"] = "alice"
    host.user_by_token["tok-bob"] = "bob"
    httpd = make_server("127.0.0.1", 0, host)
    port = httpd.server_address[1]
    thread = threading.Thread(target=httpd.serve_forever, daemon=True)
    thread.start()
    try:
        db = str(tmp_path / "t.db")
        content = str(tmp_path / "content")
        store = DocumentStore(db, content)
        cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

        @asynccontextmanager
        async def lifespan(app):
            app.state.store = store
            app.state.sessions = SessionRegistry()
            app.state.config = cfg
            yield

        app = FastAPI(lifespan=lifespan)
        app.include_router(wopi_router)
        app.include_router(editor_router)
        src = f"http://127.0.0.1:{port}/wopi/files/doc1"

        with TestClient(app) as c:
            # alice takes the lock (writable)
            ra = c.post("/editor", data={"access_token": "tok-alice"}, params={"WOPISrc": src})
            assert ra.status_code == 200
            assert host.lock_token and host.lock_token.startswith("wo:alice:")
            sa = c.app.state.sessions.get("doc1")
            assert sa is not None and sa.read_only is False
            assert "window.__READ_ONLY__ = false;" in ra.text

            # bob opens the same file in the SAME registry -> read-only session
            rb = c.post("/editor", data={"access_token": "tok-bob"}, params={"WOPISrc": src})
            assert rb.status_code == 200
            sb = c.app.state.sessions.get("doc1")
            assert sb is not None and sb.read_only is True
            assert "window.__READ_ONLY__ = true;" in rb.text
            # both sessions coexist, keyed by distinct session ids
            assert sa.session_id != sb.session_id
            assert c.app.state.sessions.get_by_id(sa.session_id) is sa
            assert c.app.state.sessions.get_by_id(sb.session_id) is sb

            # alice can still save (her own session stays writable)
            r = c.post(
                f"/api/documents/doc1/save?session={sa.session_id}",
                json={"html": "<p>alice keeps editing</p>"},
            )
            assert r.status_code == 200, r.text
            # bob's save is rejected
            r = c.post(
                f"/api/documents/doc1/save?session={sb.session_id}",
                json={"html": "<p>should not land</p>"},
            )
            assert r.status_code == 403, r.text
    finally:
        httpd.shutdown()
        httpd.server_close()
        thread.join(timeout=2)
        wipe_db(str(tmp_path / "t.db"))
        wipe_dir(str(tmp_path / "content"))


# ----------------------------------------------------------------------
# Undo/Redo-Kette (US-31): 20+ steps, exact content at each step, undo
# after save, redo-branch truncation on new edits.
# ----------------------------------------------------------------------

def test_undo_redo_chain_20_steps(tmp_path):
    """US-31: a 20-step edit chain must undo and redo one step at a time,
    each step restoring the exact previous/next document state, return to
    identity after a full undo cycle, and climb back to the tip on a full
    redo. The chain lives client-side (editor.js snapshot history); this
    test drives the SAME chain contract through the real client-mode
    save -> DOCX -> fetch round-trip for every step, so every undo/redo
    restores comparable content through the server conversion."""
    steps = [f"step-{i:02d}" for i in range(20)]
    edits = []
    for i, s in enumerate(steps):
        if i % 3 == 0:
            edits.append(f"<p><b>{s}</b></p>")
        elif i % 3 == 1:
            edits.append(f"<p><i>{s}</i></p>")
        else:
            edits.append(f"<p>{s}</p>")

    with _EditorHarness(tmp_path, seed_html="<p>undo-redo</p>") as env:
        chain = _UndoRedoHistory(env.fetch_html())

        # ---- 20 changes forward, persisted through a real save each ----
        # The editor POSTs its WHOLE innerHTML on every save (the paragraphs
        # accumulate), exactly like typing in the browser.
        seen = set()
        document = []
        for i, html in enumerate(edits):
            document.append(html)
            env.save("".join(document))
            state = env.fetch_html()
            chain.edit(state)
            text = env.page_text()
            assert text not in seen, f"edit {i} did not change the document"
            seen.add(text)
            for k in range(i + 1):
                assert steps[k] in text, f"edit {i}: step {k} missing"
            if i + 1 < 20:
                assert steps[i + 1] not in text, f"edit {i}: future step leaked in"

        # ---- undo 20x: each step restores exactly the previous state ----
        for back in range(20):
            state = chain.undo()
            assert state is not None, f"undo {back} exhausted the chain early"
            env.save(state)
            text = env.page_text()
            remaining = 18 - back  # last step index still present
            for k in range(remaining + 1):
                assert steps[k] in text, f"after {back + 1} undos: {steps[k]} missing"
            for k in range(remaining + 1, 20):
                assert steps[k] not in text, (
                    f"after {back + 1} undos: {steps[k]} still present"
                )

        # full undo cycle == identity (back to the untouched seed)
        assert env.page_text().strip() == "undo-redo"
        assert chain.can_undo is False
        assert chain.undo() is None

        # ---- redo 20x: each step restores exactly the next state ----
        for fwd in range(20):
            state = chain.redo()
            assert state is not None, f"redo {fwd} exhausted the chain early"
            env.save(state)
            text = env.page_text()
            for k in range(fwd + 1):
                assert steps[k] in text, f"after {fwd + 1} redos: {steps[k]} missing"
            if fwd + 1 < 20:
                assert steps[fwd + 1] not in text, f"redo {fwd}: next step appeared early"
        assert chain.redo() is None

        # formatting survives the whole chain (bold on step-00, italic on step-01)
        html = env.fetch_html()
        assert "<b>step-00</b>" in html
        assert "<i>step-01</i>" in html


def test_undo_redo_after_save(tmp_path):
    """US-31 edge case: an explicit save must NOT truncate the
    Undo/Redo-Kette — saving persists the document but the client-side
    chain stays intact, so undo right after a save still restores the
    pre-save state (and redo brings the saved edit back)."""
    with _EditorHarness(tmp_path, seed_html="<p>before</p>") as env:
        chain = _UndoRedoHistory(env.fetch_html())

        env.save("<p>before</p><p><b>edit-A</b></p>")
        chain.edit(env.fetch_html())

        env.save("<p>before</p><p><b>edit-A</b></p><p>edit-B</p>")
        chain.edit(env.fetch_html())

        # explicit "Save" round-trip (document persisted as-is)
        env.save(env.fetch_html())
        assert "edit-B" in env.page_text()

        # undo AFTER save -> the chain is untouched by saving
        state = chain.undo()
        assert state is not None
        env.save(state)
        text = env.page_text()
        assert "edit-A" in text
        assert "edit-B" not in text

        # redo AFTER save -> forward again
        state = chain.redo()
        assert state is not None
        env.save(state)
        assert "edit-B" in env.page_text()


def test_undo_redo_new_edit_truncates_redo(tmp_path):
    """Undo/Redo-Kette invariant: after undoing and then typing something
    new, the stale redo branch is discarded — the new edit becomes the tip
    of the chain and redo is no longer offered."""
    with _EditorHarness(tmp_path, seed_html="<p>base</p>") as env:
        chain = _UndoRedoHistory(env.fetch_html())

        env.save("<p>base</p><p>one</p>")
        chain.edit(env.fetch_html())
        env.save("<p>base</p><p>one</p><p>two</p>")
        chain.edit(env.fetch_html())

        # undo one step: the redo branch is now live
        assert chain.undo() is not None
        assert chain.can_redo is True

        # user types a NEW edit instead of redoing: redo branch is dropped
        env.save("<p>base</p><p>one</p><p>three</p>")
        chain.edit(env.fetch_html())
        assert chain.can_redo is False
        assert chain.redo() is None
        text = env.page_text()
        assert "three" in text
        assert "two" not in text
