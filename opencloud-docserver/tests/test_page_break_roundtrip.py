"""Page break roundtrip: editor marker <-> DOCX/ODT serialization (F-076).

The editor represents a page break as ``<div class="page-break"><br></div>``
(inserted by the Insert > Page break menu item via ``runCommand("insertPageBreak")``).
Saving must serialize it to a real ``<w:br w:type="page"/>`` run, and opening
a DOCX with a page break must restore the marker — no silent drops in either
direction.
"""

from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from docx.enum.text import WD_BREAK

from src.editor.converter import docx_to_html, html_to_docx
from src.editor.odt_converter import html_to_odt


def _docx_with_page_break() -> bytes:
    doc = Document()
    doc.add_paragraph("before the break")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("after the break")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


PAGE_BREAK_HTML = '<p>before the break</p><div class="page-break"><br></div><p>after the break</p>'


def test_html_to_docx_serializes_page_break_marker():
    out = html_to_docx(PAGE_BREAK_HTML)
    zf = zipfile.ZipFile(io.BytesIO(out))
    xml = zf.read("word/document.xml")
    assert b'<w:br w:type="page"/>' in xml, "page-break marker lost on save"


def test_docx_to_html_restores_page_break_marker():
    html = docx_to_html(_docx_with_page_break())
    assert 'class="page-break"' in html or "page-break" in html, (
        "DOCX page break not projected into the editor HTML"
    )


def test_page_break_roundtrip_docx():
    """DOCX -> HTML -> DOCX keeps exactly one page-break run."""
    html = docx_to_html(_docx_with_page_break())
    out = html_to_docx(html)
    zf = zipfile.ZipFile(io.BytesIO(out))
    xml = zf.read("word/document.xml")
    assert xml.count(b'<w:br w:type="page"/>') == 1
    assert b"before the break" in xml and b"after the break" in xml


def test_html_to_odt_serializes_page_break_marker():
    out = html_to_odt(PAGE_BREAK_HTML)
    zf = zipfile.ZipFile(io.BytesIO(out))
    content = zf.read("content.xml")
    assert b'fo:break-before="page"' in content or b"page" in content and b"break" in content, (
        "page-break marker lost on ODT save"
    )


def test_page_break_survives_sanitization():
    """The structural marker must not be stripped by the HTML sanitizer."""
    from src.editor.sanitize import sanitize_html

    cleaned = sanitize_html(PAGE_BREAK_HTML)
    assert "page-break" in cleaned
