"""Tests for DOCX <-> HTML conversion."""

from __future__ import annotations

import base64
import io
import re
import struct
import zipfile
import zlib

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu

from src.editor.converter import docx_to_html, html_to_docx


def test_html_to_docx_link_roundtrip():
    """A hyperlink survives HTML->DOCX->HTML and keeps its href + text."""
    html = '<p>See <a href="https://example.com">our site</a> now.</p>'
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert "https://example.com" in out
    assert "our site" in out
    assert "<a" in out
    # an unsafe scheme must not round-trip as a live link
    bad = '<p><a href="javascript:alert(1)">x</a></p>'
    bad_docx = html_to_docx(bad)
    bad_out = docx_to_html(bad_docx)
    assert "javascript:" not in bad_out


def test_html_to_docx_color_highlight_roundtrip():
    """A coloured / highlighted span survives HTML->DOCX->HTML."""
    html = (
        '<p><span style="color:#ff0000">red</span> '
        '<span style="background-color:#ffff00">hi</span></p>'
    )
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert "ff0000" in out.lower()
    assert "ffff00" in out.lower()


def test_html_to_docx_superscript_subscript_roundtrip():
    """<sup>/<sub> survive HTML->DOCX->HTML."""
    html = '<p>x<sup>2</sup> H<sub>2</sub>O</p>'
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert "<sup>" in out
    assert "<sub>" in out


def test_html_to_docx_font_family_size_roundtrip():
    """font-family/font-size spans survive HTML->DOCX->HTML."""
    html = '<p><span style="font-family:Georgia;font-size:14pt">geo 14</span></p>'
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert "Georgia" in out
    assert "14pt" in out


def test_html_to_docx_strike_smallcaps_allcaps_inlinecode_roundtrip():
    """Strike, small-caps, all-caps and inline code survive the DOCX round-trip."""
    html = (
        '<p><strike>gone</strike> '
        '<span style="font-variant:small-caps">sc</span> '
        '<span style="text-transform:uppercase">up</span> '
        '<code>code</code></p>'
    )
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    for frag in ("<strike>", "small-caps", "uppercase", "<code>"):
        assert frag in out, frag


def test_table_merge_and_column_removal_roundtrip():
    """A merged cell + a removed column survive HTML->DOCX->HTML."""
    html = (
        '<table><tr><td rowspan="2">M</td><td>A</td></tr>'
        '<tr><td>B</td></tr></table>'
    )
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert "M" in out
    assert "A" in out and "B" in out
    assert "C" not in out
    # the rowspan is preserved through the round-trip
    assert "rowspan" in out.lower()


def _make_docx(**kwargs) -> bytes:
    """Build a DOCX in memory with the given paragraphs/tables."""
    doc = Document()
    for text in kwargs.get("paragraphs", ["Hello world"]):
        if text.startswith("#"):
            doc.add_heading(text[1:], level=int(kwargs.get("heading_level", 1)))
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _simple_docx() -> bytes:
    doc = Document()
    doc.add_heading("Title", level=1)
    p = doc.add_paragraph()
    r = p.add_run("bold and ")
    r.bold = True
    r2 = p.add_run("italic")
    r2.italic = True
    doc.add_paragraph("Plain text line.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_to_html_contains_paragraphs():
    html = docx_to_html(_simple_docx())
    assert "<p>" in html
    assert "Plain text line." in html


def test_docx_to_html_headings():
    html = docx_to_html(_make_docx(paragraphs=["#My Heading"]))
    assert "<h1>My Heading</h1>" in html


def test_docx_to_html_bold():
    html = docx_to_html(_simple_docx())
    assert "<b>bold and </b>" in html
    assert "<i>italic</i>" in html


def test_html_to_docx_roundtrip_text():
    docx_bytes = html_to_docx("<p>Hello <b>stoic</b> world</p>")
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Hello" in text


def test_html_to_docx_headings():
    docx_bytes = html_to_docx("<h1>Alpha</h1><h2>Beta</h2>")
    doc = Document(io.BytesIO(docx_bytes))
    styles = [p.style.name for p in doc.paragraphs]
    assert any("Heading 1" in s for s in styles)
    assert any("Heading 2" in s for s in styles)


def test_html_to_docx_list():
    docx_bytes = html_to_docx("<ul><li>one</li><li>two</li></ul>")
    doc = Document(io.BytesIO(docx_bytes))
    assert "one" in [p.text for p in doc.paragraphs]


def test_html_to_docx_table():
    docx_bytes = html_to_docx("<table><tr><td>a</td><td>b</td></tr></table>")
    doc = Document(io.BytesIO(docx_bytes))
    assert doc.tables, "expected a table in output"
    assert doc.tables[0].cell(0, 0).text == "a"


def test_full_roundtrip_survives():
    """Editing flow: DOCX -> HTML -> DOCX must not raise and keep text."""
    original = _simple_docx()
    html = docx_to_html(original)
    back = html_to_docx(html)
    doc = Document(io.BytesIO(back))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Title" in text


def test_list_grouped_in_ul():
    """DOCX list paragraphs round-trip back into a <ul> block."""
    doc = Document()
    doc.add_paragraph("first", style="List Bullet")
    doc.add_paragraph("second", style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    html = docx_to_html(buf.getvalue())
    assert "<ul>" in html and "</ul>" in html
    assert html.count("<ul>") == 1


def test_html_to_docx_preserves_bold_italic_underline():
    docx_bytes = html_to_docx("<p>plain <b>bold</b> <i>italic</i> <u>under</u> tail</p>")
    doc = Document(io.BytesIO(docx_bytes))
    runs = [(r.text, r.bold, r.italic, r.underline) for r in doc.paragraphs[0].runs]
    by_text = {r[0].strip(): r for r in runs}
    assert by_text["bold"][1] is True, runs
    assert by_text["italic"][2] is True, runs
    assert by_text["under"][3] is True, runs
    assert by_text["plain"][1] is None and by_text["tail"][1] is None, runs


def test_html_to_docx_nested_inline_formatting():
    docx_bytes = html_to_docx("<p><b>both <i>nested</i></b> end</p>")
    doc = Document(io.BytesIO(docx_bytes))
    runs = {(r.text): r for r in doc.paragraphs[0].runs}
    assert runs["both "].bold is True
    assert runs["nested"].bold is True and runs["nested"].italic is True
    assert runs[" end"].bold is None


def test_full_roundtrip_keeps_bold():
    """DOCX(bold) -> HTML -> DOCX must keep the bold run (US-2 fidelity)."""
    original = _simple_docx()  # contains a bold run "bold and "
    html = docx_to_html(original)
    assert "<b>bold and </b>" in html
    roundtrip = html_to_docx(html)
    doc = Document(io.BytesIO(roundtrip))
    bold = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
    assert "bold and " in bold, bold


def test_html_to_docx_keeps_tagless_text():
    """Raw text without block tags must survive (typed into an empty#editor)."""
    docx_bytes = html_to_docx("FRISCHER INHALT X.&nbsp;")
    doc = Document(io.BytesIO(docx_bytes))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "FRISCHER INHALT X." in text, text


# --------------------------------------------------------------------------
# Table roundtrip (editor-tables-converter)
# --------------------------------------------------------------------------


def _strip_default_tcW(table):
    """Drop python-docx's per-cell default w:tcW (a layout default, not a real
    width). Table-prop fixtures keep bare structural <td>/<th> assertions."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.tcPr
            if tcPr is not None:
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is not None:
                    tcPr.remove(tcW)


def _table_docx(rows=2, cols=2, values=None, header=False) -> bytes:
    """Build a DOCX whose only content is an NxM table."""
    doc = Document()
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    _strip_default_tcW(table)
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text = values[r][c] if values else f"{r},{c}"
    if header:
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_to_html_table_basic():
    html = docx_to_html(_table_docx(2, 2))
    assert html.count("<table>") == 1
    assert html.count("<tr>") == 2
    assert html.count("<td>") == 4
    assert "0,0" in html and "1,1" in html


def test_docx_to_html_table_inline_formatting():
    """Run-level formatting inside cells must survive DOCX -> HTML."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _strip_default_tcW(table)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("bold")
    r.bold = True
    r2 = p.add_run(" and ")
    r2.italic = True
    r3 = p.add_run("under")
    r3.underline = True
    buf = io.BytesIO()
    doc.save(buf)
    html = docx_to_html(buf.getvalue())
    assert "<td><b>bold</b><i> and </i><u>under</u></td>" in html


def test_docx_to_html_table_header_row():
    html = docx_to_html(_table_docx(2, 2, header=True))
    assert html.count("<th>") == 2
    assert html.count("<td>") == 2


def test_docx_to_html_table_multiparagraph_cell():
    """Multi-paragraph cell content joins with <br/>."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _strip_default_tcW(table)
    table.cell(0, 0).paragraphs[0].add_run("first")
    table.cell(0, 0).add_paragraph("second")
    buf = io.BytesIO()
    doc.save(buf)
    html = docx_to_html(buf.getvalue())
    assert "<td>first<br/>second</td>" in html


def test_docx_to_html_table_colspan():
    """Horizontally merged cells become colspan, not duplicated <td>."""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 1).merge(table.cell(0, 2))
    buf = io.BytesIO()
    doc.save(buf)
    html = docx_to_html(buf.getvalue())
    # row 0 has 2 real cells (one spanning 2 grid columns), row 1 has 3
    assert html.count("<td") == 5
    assert 'colspan="2"' in html


def test_docx_to_html_table_rowspan():
    """Vertically merged cells become rowspan, no duplicate cells in next row."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(1, 0))
    buf = io.BytesIO()
    doc.save(buf)
    html = docx_to_html(buf.getvalue())
    assert 'rowspan="2"' in html
    assert html.count("<td") == 3  # 4 grid slots minus the rowspan-covered one


def test_html_to_docx_table_inline_formatting():
    docx_bytes = html_to_docx(
        "<table><tr><td><b>bold</b> <i>it</i></td></tr></table>"
    )
    doc = Document(io.BytesIO(docx_bytes))
    cell = doc.tables[0].cell(0, 0)
    runs = [(r.text, r.bold, r.italic) for r in cell.paragraphs[0].runs]
    assert ("bold", True, None) in runs, runs
    assert ("it", None, True) in runs, runs


def test_html_to_docx_table_multiparagraph():
    """<br/> inside a cell becomes a second paragraph."""
    docx_bytes = html_to_docx(
        "<table><tr><td>one<br/>two</td></tr></table>"
    )
    doc = Document(io.BytesIO(docx_bytes))
    cell = doc.tables[0].cell(0, 0)
    assert [p.text for p in cell.paragraphs] == ["one", "two"]


def test_html_to_docx_table_header_row():
    docx_bytes = html_to_docx(
        "<table><tr><th>H1</th><th>H2</th></tr><tr><td>a</td><td>b</td></tr></table>"
    )
    doc = Document(io.BytesIO(docx_bytes))
    tr = doc.tables[0].rows[0]._tr
    assert tr.trPr is not None and tr.trPr.find(qn("w:tblHeader")) is not None


def test_html_to_docx_table_colspan():
    docx_bytes = html_to_docx(
        '<table><tr><td colspan="2">wide</td><td>x</td></tr></table>'
    )
    doc = Document(io.BytesIO(docx_bytes))
    tc = doc.tables[0].rows[0]._tr.tc_lst[0]
    assert tc.tcPr is not None and tc.tcPr.gridSpan is not None
    assert tc.tcPr.gridSpan.val == 2


def test_html_to_docx_table_rowspan():
    docx_bytes = html_to_docx(
        '<table><tr><td rowspan="2">tall</td><td>x</td></tr><tr><td>y</td></tr></table>'
    )
    doc = Document(io.BytesIO(docx_bytes))
    tcs = doc.tables[0].rows[0]._tr.tc_lst
    assert tcs[0].tcPr is not None and tcs[0].tcPr.vMerge is not None
    assert tcs[0].tcPr.vMerge.val == "restart"
    continuation = doc.tables[0].rows[1]._tr.tc_lst[0]
    assert continuation.tcPr is not None and continuation.tcPr.vMerge is not None


def test_table_roundtrip_through_docx_is_stable():
    """DOCX(table) -> HTML -> DOCX -> HTML must be a fixed point."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    _strip_default_tcW(table)
    for r in range(3):
        for c in range(3):
            table.cell(r, c).text = f"{r},{c}"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    # bold run in (1,0), multi-paragraph in (1,2), h-merge row2, v-merge col0
    table.cell(1, 0).paragraphs[0].add_run(" *").bold = True
    table.cell(1, 2).add_paragraph("extra")
    table.cell(2, 1).merge(table.cell(2, 2))
    table.cell(0, 0).merge(table.cell(1, 0))
    buf = io.BytesIO()
    doc.save(buf)
    html1 = docx_to_html(buf.getvalue())
    back = html_to_docx(html1)
    html2 = docx_to_html(back)
    assert html1 == html2, f"\nfirst: {html1}\nsecond: {html2}"


def test_table_roundtrip_through_html_is_stable():
    """HTML(table) -> DOCX -> HTML must reproduce the source table."""
    source = (
        '<table><tr><th rowspan="2">H</th><th>A</th><th>B</th></tr>'
        "<tr><td>1</td><td>2</td></tr>"
        '<tr><td colspan="2">wide</td><td>tail</td></tr></table>'
    )
    docx_bytes = html_to_docx(source)
    rebuilt = docx_to_html(docx_bytes)
    assert rebuilt == source, f"\ninput:   {source}\nrebuilt: {rebuilt}"


# --------------------------------------------------------------------------
# Converter performance regression (converter-benchmark)
# --------------------------------------------------------------------------


def test_converter_roundtrip_has_no_perf_regression():
    """The 100-paragraph roundtrip must stay within the benchmark bound.

    Delegates to the shared benchmark harness (``tests/bench/``) so the
    bound lives in exactly one place; a regression in converter
    throughput (e.g. an O(n^2) body scan) fails both here and in the
    standalone benchmark run.
    """
    from tests.bench.benchmark_converter import (
        BOUND_SECONDS,
        best_roundtrip_time,
        make_document_docx,
    )

    elapsed = best_roundtrip_time(make_document_docx(100), iterations=3)
    assert elapsed <= BOUND_SECONDS, (
        f"100-paragraph DOCX<->HTML roundtrip took {elapsed:.2f}s "
        f"— exceeds the {BOUND_SECONDS}s benchmark ceiling, perf regression"
    )

# ---------------------------------------------------------------------------
# Image round-trip (editor-images-converter)
# ---------------------------------------------------------------------------


def _png_bytes(width: int, height: int) -> bytes:
    """Build a minimal valid PNG (RGB) of the given pixel size."""
    def _chunk(ctype: bytes, data: bytes) -> bytes:
        c = struct.pack(">I", len(data)) + ctype + data
        return c + struct.pack(">I", zlib.crc32(ctype + data) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * width for _ in range(height))
    return (b"\x89PNG\r\n\x1a\n" + _chunk(b"IHDR", ihdr)
            + _chunk(b"IDAT", zlib.compress(raw)) + _chunk(b"IEND", b""))


def _jpeg_bytes(width: int, height: int) -> bytes:
    """Build a minimal valid JPEG whose SOF0 marker carries the given dims."""
    app0_payload = b"JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00"  # 14 bytes
    sof0_payload = struct.pack(">BHHB", 8, height, width, 3) + bytes(
        [1, 0x11, 0, 2, 0x11, 0, 3, 0x11, 0]
    )
    return (
        b"\xff\xd8"
        + b"\xff\xe0"
        + struct.pack(">H", len(app0_payload) + 2)
        + app0_payload
        + b"\xff\xc0"
        + struct.pack(">H", len(sof0_payload) + 2)
        + sof0_payload
        + b"\xff\xd9"
    )


def _data_uri(data: bytes, mime: str = "image/png") -> str:
    return f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"


def _img_srcs(html: str) -> list[str]:
    """All data-URI src values of <img> tags in an HTML fragment."""
    return re.findall(r'<img[^>]*\ssrc="(data:[^"]+)"', html)


def _decode_data_uri(uri: str) -> bytes:
    return base64.b64decode(uri.split(",", 1)[1])


def _docx_media_bytes(docx: bytes) -> list[bytes]:
    """Bytes of every word/media/ member in a DOCX package."""
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        return [z.read(n) for n in z.namelist() if n.startswith("word/media/")]


def _docx_content_xml(docx: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        return z.read("word/document.xml").decode()


def _docx_with_picture(png: bytes, width_px: int | None = None,
                       height_px: int | None = None) -> bytes:
    """Build a DOCX with one paragraph: text + an inline picture."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("imaged ")
    run = p.add_run()
    kw = {}
    if width_px:
        kw["width"] = Emu(width_px * 9525)
    if height_px:
        kw["height"] = Emu(height_px * 9525)
    run.add_picture(io.BytesIO(png), **kw)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_html_to_docx_image_roundtrip_single():
    """A data-URI <img> becomes an inline picture whose package bytes and
    re-exported data URI match the original PNG exactly."""
    png = _png_bytes(2, 3)
    html = f'<p>Lead <img src="{_data_uri(png)}"/> tail</p>'
    docx = html_to_docx(html)

    media = _docx_media_bytes(docx)
    assert len(media) == 1
    assert media[0] == png

    html2 = docx_to_html(docx)
    srcs = _img_srcs(html2)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png
    assert "Lead" in html2 and "tail" in html2


def test_docx_to_html_image_roundtrip_referenced_picture():
    """A picture embedded via python-docx (word/media part + w:drawing)
    renders as a data-URI <img> with its pixel dimensions kept."""
    png = _png_bytes(4, 5)
    html = docx_to_html(_docx_with_picture(png, width_px=4, height_px=5))
    srcs = _img_srcs(html)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png
    assert "imaged" in html
    assert 'width="4"' in html and 'height="5"' in html


def test_image_roundtrip_preserves_picture_bytes_full_loop():
    """DOCX -> HTML -> DOCX keeps the image bytes verbatim end to end."""
    png = _png_bytes(2, 3)
    original = _docx_with_picture(png)

    html = docx_to_html(original)
    back = html_to_docx(html)
    media = _docx_media_bytes(back)
    assert len(media) == 1
    assert media[0] == png

    srcs = _img_srcs(docx_to_html(back))
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png


def test_html_to_docx_image_roundtrip_explicit_dimensions():
    """width/height attributes on the <img> become the drawing extent
    (px -> EMU) and survive the DOCX -> HTML pass."""
    png = _png_bytes(2, 3)
    html = f'<p><img src="{_data_uri(png)}" width="120" height="90"/></p>'
    docx = html_to_docx(html)

    content = _docx_content_xml(docx)
    assert 'cx="1143000"' in content  # 120 * 9525
    assert 'cy="857250"' in content   # 90 * 9525

    html2 = docx_to_html(docx)
    srcs = _img_srcs(html2)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png
    assert 'width="120"' in html2 and 'height="90"' in html2


def test_html_to_docx_image_roundtrip_intrinsic_dimensions():
    """Without explicit attributes the extent comes from the image's
    intrinsic pixel dimensions (sniffed from the PNG header)."""
    png = _png_bytes(2, 3)
    docx = html_to_docx(f'<p><img src="{_data_uri(png)}"/></p>')

    content = _docx_content_xml(docx)
    assert 'cx="19050"' in content  # 2 * 9525
    assert 'cy="28575"' in content  # 3 * 9525

    html2 = docx_to_html(docx)
    assert 'width="2"' in html2 and 'height="3"' in html2


def test_html_to_docx_image_roundtrip_keeps_formatting():
    """An image inside a paragraph does not disturb bold/italic runs around
    it — both the runs and the picture survive."""
    png = _png_bytes(2, 3)
    html = f'<p><b>Bold</b> <img src="{_data_uri(png)}"/> tail</p>'
    docx = html_to_docx(html)
    assert _docx_media_bytes(docx) == [png]

    html2 = docx_to_html(docx)
    assert "<b>Bold</b>" in html2
    assert " tail" in html2
    srcs = _img_srcs(html2)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png


def test_html_to_docx_image_roundtrip_multiple():
    """Several images in a paragraph each become their own picture with the
    correct bytes, in order."""
    png1 = _png_bytes(2, 3)
    png2 = _png_bytes(3, 2)
    html = (f'<p>a<img src="{_data_uri(png1)}"/>b'
            f'<img src="{_data_uri(png2)}"/>c</p>')
    docx = html_to_docx(html)

    assert set(_docx_media_bytes(docx)) == {png1, png2}
    html2 = docx_to_html(docx)
    srcs = _img_srcs(html2)
    assert len(srcs) == 2
    assert _decode_data_uri(srcs[0]) == png1
    assert _decode_data_uri(srcs[1]) == png2


def test_image_roundtrip_in_table_cell():
    """Images inside table cells follow the same path and survive."""
    png = _png_bytes(2, 3)
    html = (f'<table><tr><td>icon <img src="{_data_uri(png)}"/>'
            f"</td><td>x</td></tr></table>")
    docx = html_to_docx(html)

    media = _docx_media_bytes(docx)
    assert len(media) == 1
    assert media[0] == png

    html2 = docx_to_html(docx)
    assert "<td>icon " in html2
    srcs = _img_srcs(html2)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png


def test_docx_to_html_image_in_table_cell_roundtrip():
    """DOCX side: an inline picture inside a table cell becomes an <img>."""
    png = _png_bytes(4, 5)
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    p = table.cell(0, 0).paragraphs[0]
    p.add_run("icon ")
    run = p.add_run()
    run.add_picture(io.BytesIO(png))
    buf = io.BytesIO()
    doc.save(buf)

    html = docx_to_html(buf.getvalue())
    srcs = _img_srcs(html)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png
    assert "icon" in html


def test_html_to_docx_image_roundtrip_non_data_uri_is_skipped():
    """An http(s) src cannot be fetched server-side and is skipped without
    breaking the rest of the paragraph."""
    docx = html_to_docx('<p>link <img src="https://example.com/x.png"/> tail</p>')
    assert _docx_media_bytes(docx) == []
    html2 = docx_to_html(docx)
    assert "link" in html2 and "tail" in html2
    assert "<img" not in html2


def test_image_roundtrip_alt_text():
    """The <img> alt text becomes wp:docPr/descr on the drawing and is
    re-exported as an alt attribute through the full round-trip."""
    png = _png_bytes(2, 3)
    html = f'<p><img src="{_data_uri(png)}" alt="A kitten"/></p>'
    docx = html_to_docx(html)
    content = _docx_content_xml(docx)
    assert 'descr="A kitten"' in content

    html2 = docx_to_html(docx)
    assert 'alt="A kitten"' in html2
    assert _decode_data_uri(_img_srcs(html2)[0]) == png

    # Full HTML -> DOCX -> HTML -> DOCX -> HTML keeps the alt as well.
    assert 'alt="A kitten"' in docx_to_html(html_to_docx(html2))


def test_image_roundtrip_in_list_item():
    """An image inside a list item survives HTML -> DOCX -> HTML."""
    png = _png_bytes(2, 3)
    html = f'<ul><li>icon <img src="{_data_uri(png)}"/></li></ul>'
    docx = html_to_docx(html)
    assert _docx_media_bytes(docx) == [png]

    html2 = docx_to_html(docx)
    assert "<ul>" in html2 and "<li>icon " in html2
    srcs = _img_srcs(html2)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png


def test_html_to_docx_image_roundtrip_jpeg_intrinsic_dimensions():
    """JPEG data URIs without width/height attributes get intrinsic pixel
    dimensions sniffed from their SOF marker."""
    jpeg = _jpeg_bytes(10, 5)
    docx = html_to_docx(f'<p><img src="{_data_uri(jpeg, "image/jpeg")}"/></p>')

    content = _docx_content_xml(docx)
    assert 'cx="95250"' in content  # 10 * 9525
    assert 'cy="47625"' in content  # 5 * 9525

    html2 = docx_to_html(docx)
    srcs = _img_srcs(html2)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == jpeg
    assert 'width="10"' in html2 and 'height="5"' in html2


def test_docx_to_html_image_inline_in_same_run():
    """A run holding both text and a drawing keeps the <img> where the
    drawing sits — the text is not swallowed."""
    png = _png_bytes(2, 3)
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("AB")
    run.add_picture(io.BytesIO(png))  # drawing appended to the same run
    p.add_run(" after")
    buf = io.BytesIO()
    doc.save(buf)

    html = docx_to_html(buf.getvalue())
    srcs = _img_srcs(html)
    assert len(srcs) == 1
    assert _decode_data_uri(srcs[0]) == png
    assert "AB" in html and "after" in html


def test_docx_to_html_image_missing_picture_is_skipped():
    """A drawing whose r:embed does not resolve must not crash or emit a
    broken <img> — surrounding text still converts."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("before ")
    ns = nsdecls("w", "wp", "r", "a", "pic")
    drawing = parse_xml(
        f"<w:drawing {ns}>"
        '<wp:inline><wp:docPr id="1" name="Ghost"/>'
        "<a:graphic><a:graphicData>"
        '<pic:pic><a:blip r:embed="rId999"/></pic:pic>'
        "</a:graphicData></a:graphic></wp:inline></w:drawing>"
    )
    r._r.append(drawing)
    p.add_run(" after")
    buf = io.BytesIO()
    doc.save(buf)

    html = docx_to_html(buf.getvalue())
    assert "before" in html and " after" in html
    assert "<img" not in html



def test_html_to_docx_hr_and_page_break_roundtrip():
    """<hr> and page break survive html_to_docx -> docx_to_html."""
    html = (
        "<p>before</p><hr/><p>middle</p>"
        "<div class=\"page-break\"><br></div><p>after</p>"
    )
    out = docx_to_html(html_to_docx(html))
    assert "<hr" in out, out
    assert "page-break" in out, out
    assert "before" in out and "after" in out, out


def test_docx_hr_is_bottom_border_paragraph():
    """An <hr> maps to a paragraph with a bottom border (w:pBdr)."""
    doc = Document(io.BytesIO(html_to_docx("<p>a</p><hr/><p>b</p>")))
    ps = doc.paragraphs
    assert [p.text for p in ps] == ["a", "", "b"]
    pBdr = ps[1]._p.pPr.find(qn("w:pBdr"))
    assert pBdr is not None and pBdr.find(qn("w:bottom")) is not None


def test_docx_page_break_is_br_type_page():
    """A page-break div maps to a paragraph carrying <w:br w:type='page'/>."""
    doc = Document(io.BytesIO(html_to_docx("<p>a</p><div class=\"page-break\"><br></div><p>b</p>")))
    ps = doc.paragraphs
    assert [p.text for p in ps] == ["a", "", "b"]
    page_brs = [
        br
        for br in ps[1]._p.iter(qn("w:br"))
        if (br.get(qn("w:type")) or "textWrapping") == "page"
    ]
    assert len(page_brs) == 1


def test_sanitize_allows_hr_and_keeps_page_break_class():
    """The sanitizer keeps <hr> and div.page-break (and drops the rest)."""
    from src.editor.sanitize import sanitize_html

    html = (
        "<p>a</p><hr/>"
        '<div class="page-break"><br></div>'
        '<div class="evil"><script>alert(1)</script>x</div>'
    )
    out = sanitize_html(html)
    assert "<hr" in out, out
    assert 'class="page-break"' in out, out
    assert "<script>" not in out, out
    assert "alert(1)" not in out, out


def test_html_to_docx_paragraph_props_roundtrip():
    """Line-height / indent / spacing / RTL / page-break-before round-trip.

    A paragraph carrying all five style properties must survive
    HTML -> DOCX -> HTML with the values intact.
    """
    html = (
        '<p style="line-height:1.5;margin-left:24pt;text-indent:12pt;'
        'margin-top:6pt;direction:rtl;page-break-before:always">RTL para</p>'
    )
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    for frag in ("line-height:1.5", "24pt", "12pt", "6pt", "rtl", "page-break-before"):
        assert frag in out, frag


def test_html_to_docx_blockquote_becomes_indent():
    """A <blockquote> (Chrome's execCommand indent) maps to a left indent."""
    html = "<blockquote>indented line</blockquote>"
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert "indented line" in out
    assert "margin-left" in out.lower() or "blockquote" in out


def test_html_to_docx_nested_list_roundtrip():
    """A level-2 bullet/number round-trips as nested <ul>/<ol> in DOCX."""
    html = (
        "<ul><li>item 1<ul><li>sub 1</li><li>sub 2</li></ul></li>"
        "<li>item 2</li></ul>"
    )
    out = docx_to_html(html_to_docx(html))
    norm = out.replace("\n", "")
    assert norm == (
        "<ul><li>item 1<ul><li>sub 1</li><li>sub 2</li></ul></li>"
        "<li>item 2</li></ul>"
    ), norm


def test_html_to_docx_nested_numbered_list_roundtrip():
    """Numbered outline levels survive too."""
    html = "<ol><li>one<ol><li>1.1</li></ol></li><li>two</li></ol>"
    out = docx_to_html(html_to_docx(html))
    norm = out.replace("\n", "")
    assert norm == (
        "<ol><li>one<ol><li>1.1</li></ol></li><li>two</li></ol>"
    ), norm


def test_special_symbol_roundtrips_docx():
    """A special symbol (literal char) survives DOCX round-trip."""
    from src.editor.odt_converter import html_to_odt, odt_to_html
    html = "<p>Use the § symbol and ¶ now.</p>"
    assert "§" in docx_to_html(html_to_docx(html))
    assert "§" in odt_to_html(html_to_odt(html))


def test_html_to_docx_link_boundaries_preserved():
    """Text before/after an <a> must stay OUTSIDE the anchor (regression:
    the builder merged leading text into the link token when <a> did not
    flush the pending buffer first)."""
    html = '<p>See <a href="https://example.com">site</a> now.</p>'
    out = docx_to_html(html_to_docx(html)).replace("\n", "")
    assert out == '<p>See <a href="https://example.com">site</a> now.</p>', out


def test_html_to_docx_span_boundary_preserved():
    """Leading text before a styled <span> does not inherit its style."""
    html = '<p>a<span style="color:#ff0000">b</span>c</p>'
    out = docx_to_html(html_to_docx(html)).replace("\n", "")
    assert out == '<p>a<span style="color:#ff0000">b</span>c</p>', out


def test_html_to_docx_table_cell_props_roundtrip():
    """Cell shading/borders/width survive HTML->DOCX->HTML (T13)."""
    html = ('<table width="500"><tr><th style="background-color:#ffdddd">H</th></tr>'
            '<tr><td style="border:1pt solid #000000; background-color:#eeeeee" '
            'width="120">c</td></tr></table>')
    out = docx_to_html(html_to_docx(html)).replace("\n", "")
    assert "background-color:#ffdddd" in out, out
    assert "background-color:#eeeeee" in out, out
    assert "border:1pt solid #000000" in out, out
    assert 'width="120"' in out, out
    assert '<table width="500">' in out or "width=\"500\"" in out, out


def test_html_to_docx_table_without_props_stays_plain():
    """A plain grid table must NOT gain invented cell styles."""
    html = "<table><tr><td>a</td></tr></table>"
    out = docx_to_html(html_to_docx(html)).replace("\n", "")
    assert "background-color" not in out
    assert "border:" not in out


def test_sanitizer_allows_header_footer():
    """The sanitizer keeps <header class=\"page-header\"> and <footer class=\"page-footer\">."""
    from src.editor.sanitize import sanitize_html

    html = (
        "<p>before</p>"
        '<header class="page-header"><p>Header text</p></header>'
        "<p>content</p>"
        '<footer class="page-footer"><p>Footer <span class="page-number"></span></p></footer>'
        "<p>after</p>"
    )
    out = sanitize_html(html)
    assert '<header class="page-header">' in out, out
    assert '</header>' in out, out
    assert '<footer class="page-footer">' in out, out
    assert '</footer>' in out, out
    assert 'Header text' in out, out
    assert 'Footer' in out, out
    assert 'page-number' in out, out
    # Verify dangerous content is still stripped
    assert '<script>' not in out, out
    assert 'alert(1)' not in out, out


def test_html_to_docx_header_footer_roundtrip():
    """Header/footer with page-number survive HTML->DOCX->HTML round-trip.

    The HTML contract:
    - <header class="page-header"> contains the header content
    - <footer class="page-footer"> contains the footer content
    - <span class="page-number"></span> represents the page number field

    After DOCX conversion, the DOCX package must contain header1.xml and footer1.xml
    parts with the content, and the body sectPr must have headerReference/footerReference.
    """
    html = (
        '<header class="page-header"><p>Header text</p></header>'
        '<p>Body content</p>'
        '<footer class="page-footer"><p>Page <span class="page-number"></span></p></footer>'
    )

    docx_bytes = html_to_docx(html)

    # Verify DOCX contains header1.xml and footer1.xml parts
    import io
    import zipfile
    z = zipfile.ZipFile(io.BytesIO(docx_bytes))
    parts = z.namelist()
    assert 'word/header1.xml' in parts, "DOCX must contain header1.xml"
    assert 'word/footer1.xml' in parts, "DOCX must contain footer1.xml"

    # Verify sectPr has headerReference and footerReference
    doc_xml = z.read('word/document.xml').decode('utf-8')
    assert 'headerReference' in doc_xml, "sectPr must have headerReference"
    assert 'footerReference' in doc_xml, "sectPr must have footerReference"

    # Round-trip back to HTML
    out = docx_to_html(docx_bytes)

    # Verify header/footer survive
    assert '<header class="page-header">' in out, out
    assert '</header>' in out, out
    assert 'Header text' in out, out
    assert '<footer class="page-footer">' in out, out
    assert '</footer>' in out, out
    assert 'Page' in out, out
    assert '<span class="page-number">' in out or '<span class="page-number"></span>' in out, out


def test_docx_page_number_field_roundtrip():
    """A PAGE field (w:fldSimple w:instr=" PAGE ") round-trips as <span class="page-number">."""
    html = (
        '<header class="page-header"><p>Page <span class="page-number"></span> of 10</p></header>'
        '<p>Content</p>'
    )
    docx_bytes = html_to_docx(html)

    # Verify the DOCX contains the PAGE field
    import io
    import zipfile
    z = zipfile.ZipFile(io.BytesIO(docx_bytes))
    header_xml = z.read('word/header1.xml').decode('utf-8')
    assert 'fldSimple' in header_xml, "DOCX header must contain fldSimple"
    assert ' PAGE ' in header_xml, "fldSimple must have PAGE instruction"

    # Round-trip
    out = docx_to_html(docx_bytes)
    assert 'page-number' in out, out
    assert 'Page' in out and 'of 10' in out, out


def test_html_to_docx_footnote_roundtrip():
    """A footnote (marker + adjacent body span) round-trips DOCX.

    The HTML contract: <sup class="footnote-citation">[n]</sup> immediately
    followed by <span class="footnote">BODY</span>. On reload the citation
    ``[1]``, the marker classes and the body text must all survive, and the
    DOCX package must physically carry a footnotes.xml part (with the
    Word/LibreOffice sentinel footnotes) holding the note text.
    """
    html = (
        '<p>Main text<sup class="footnote-citation">[1]</sup>'
        '<span class="footnote">Note body.</span> continues.</p>'
    )
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert '<sup class="footnote-citation">[1]</sup>' in out, out
    assert '<span class="footnote">Note body.</span>' in out, out
    assert "Main text" in out and "continues." in out, out
    # physical assertion: the package carries a footnotes part with sentinels
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        assert "word/footnotes.xml" in z.namelist(), z.namelist()
        fn_xml = z.read("word/footnotes.xml").decode("utf-8", "replace")
        assert "Note body." in fn_xml, fn_xml
        assert 'separator' in fn_xml, fn_xml
        assert 'continuationSeparator' in fn_xml, fn_xml
        assert 'w:id="1"' in fn_xml, fn_xml


def test_html_to_docx_endnote_roundtrip():
    """An endnote (marker + adjacent body span) round-trips DOCX."""
    html = (
        '<p>Main text<sup class="endnote-citation">[1]</sup>'
        '<span class="endnote">End note body.</span> continues.</p>'
    )
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert '<sup class="endnote-citation">[1]</sup>' in out, out
    assert '<span class="endnote">End note body.</span>' in out, out
    assert "Main text" in out and "continues." in out, out
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        assert "word/endnotes.xml" in z.namelist(), z.namelist()
        en_xml = z.read("word/endnotes.xml").decode("utf-8", "replace")
        assert "End note body." in en_xml, en_xml
        assert 'w:id="1"' in en_xml, en_xml


def test_html_to_docx_comment_roundtrip():
    """An anchored comment (review note) round-trips through DOCX.

    The HTML contract: <span class="comment" data-author="AUTHOR"
    data-comment="BODY">ANCHORED TEXT</span>. The writer must emit a real
    word/comments.xml part and wrap the anchored runs in commentRangeStart/
    commentRangeEnd + a commentReference marker; the reader must turn that
    back into the same span. The sanitizer passes the comment span through
    unchanged (class + data-* attributes survive, values escaped on output).
    """
    html = (
        '<p>Main text <span class="comment" data-author="Alice Smith" '
        'data-comment="Please check this, thanks.">anchored words</span>'
        ' continues.</p>'
    )
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert (
        '<span class="comment" data-author="Alice Smith" '
        'data-comment="Please check this, thanks.">anchored words</span>'
    ) in out, out
    assert "Main text" in out and "continues." in out, out
    # the sanitizer keeps the comment span and its data attributes
    from src.editor.sanitize import sanitize_html
    sane = sanitize_html(html)
    assert 'class="comment"' in sane, sane
    assert 'data-author="Alice Smith"' in sane, sane
    assert 'data-comment="Please check this, thanks."' in sane, sane


def test_docx_comment_read_from_parts():
    """The DOCX writer physically produces word/comments.xml and the reader
    resolves the commentRangeStart/End + commentReference markers back into
    anchored comment spans (two comments, distinct authors)."""
    html = (
        '<p><span class="comment" data-author="Alice" '
        'data-comment="First note.">one two</span> plain '
        '<span class="comment" data-author="Bob" '
        'data-comment="Second note, with punctuation.">three</span> tail</p>'
    )
    docx = html_to_docx(html)
    # physical assertions on the package written by the converter
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        assert "word/comments.xml" in z.namelist(), z.namelist()
        comments_xml = z.read("word/comments.xml").decode("utf-8", "replace")
        assert "Alice" in comments_xml and "First note." in comments_xml, comments_xml
        assert "Bob" in comments_xml and "Second note, with punctuation." in comments_xml, comments_xml
        document_xml = z.read("word/document.xml").decode("utf-8", "replace")
        assert "commentRangeStart" in document_xml, document_xml
        assert "commentRangeEnd" in document_xml, document_xml
        assert "commentReference" in document_xml, document_xml
    # the reader resolves the markers back into the HTML contract
    out = docx_to_html(docx)
    assert (
        '<span class="comment" data-author="Alice" '
        'data-comment="First note.">one two</span>'
    ) in out, out
    assert (
        '<span class="comment" data-author="Bob" '
        'data-comment="Second note, with punctuation.">three</span>'
    ) in out, out
    assert "plain" in out and "tail" in out, out


def test_html_to_docx_track_changes_roundtrip():
    """Tracked insertions/deletions round-trip through DOCX (T30).

    The HTML contract: <ins class="track-insert" data-author="AUTHOR"> and
    <del class="track-delete" data-author="AUTHOR">. The writer must emit
    real w:ins / w:del change elements (with w:delText for the removed run)
    in document.xml; the reader must turn them back into the same spans with
    author + text intact. The sanitizer must pass ins through (del already
    allowed), and still drop script.
    """
    html = (
        '<p>Before <ins class="track-insert" data-author="Alice">new words</ins> '
        'and <del class="track-delete" data-author="Bob">old words</del> after.</p>'
    )
    docx = html_to_docx(html)
    out = docx_to_html(docx)
    assert '<ins class="track-insert" data-author="Alice">new words</ins>' in out, out
    assert '<del class="track-delete" data-author="Bob">old words</del>' in out, out
    assert "Before" in out and "after." in out, out
    # the sanitizer keeps the track-change spans and their data attributes
    from src.editor.sanitize import sanitize_html
    sane = sanitize_html(html)
    assert '<ins class="track-insert" data-author="Alice">new words</ins>' in sane, sane
    assert '<del class="track-delete" data-author="Bob">old words</del>' in sane, sane
    assert "<script>" not in sane, sane


def test_docx_track_changes_read_from_parts():
    """The DOCX writer physically emits w:ins/w:del/w:delText change elements
    in word/document.xml and the reader resolves them back into the HTML
    contract (author + inserted text + deleted text)."""
    html = (
        '<p><ins class="track-insert" data-author="Alice">added bit</ins> '
        'keep <del class="track-delete" data-author="Bob">removed bit</del> end</p>'
    )
    docx = html_to_docx(html)
    import zipfile
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert "<w:ins " in xml, xml[:500]
    assert 'w:author="Alice"' in xml, xml[:500]
    assert "<w:del " in xml, xml[:500]
    assert 'w:author="Bob"' in xml, xml[:500]
    assert "<w:delText" in xml, xml[:500]
    # the reader resolves the change elements back to the contract
    out = docx_to_html(docx)
    assert '<ins class="track-insert" data-author="Alice">added bit</ins>' in out, out
    assert '<del class="track-delete" data-author="Bob">removed bit</del>' in out, out


def test_sanitizer_allows_ins():
    """<ins class="track-insert"> survives sanitize_html() (del already
    allowed); <script> and friends still do not."""
    from src.editor.sanitize import sanitize_html

    html = (
        '<p>Hi <ins class="track-insert" data-author="A">added</ins> and '
        '<del class="track-delete" data-author="B">removed</del> done '
        '<script>alert(1)</script></p>'
    )
    out = sanitize_html(html)
    assert '<ins class="track-insert"' in out, out
    assert 'data-author="A"' in out, out
    assert '<del class="track-delete"' in out, out
    assert 'data-author="B"' in out, out
    assert "<script>" not in out, out
