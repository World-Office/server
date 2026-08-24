"""Performance-regression tests for the DOCX <-> HTML converter.

These reuse the benchmark fixtures and bounds from
``tests.bench.benchmark_converter`` so the regular ``uv run pytest`` run
gets a cheap conversion-throughput gate: an absolute ceiling on a
100-paragraph roundtrip and a sub-quadratic scaling check. The added
correctness assertions guarantee the timed work is a real, content-
preserving conversion (a benchmark that drops everything would be a
cheat).
"""

from __future__ import annotations

import io

from docx import Document

from src.editor.converter import docx_to_html, html_to_docx
from tests.bench.benchmark_converter import (
    BOUND_SECONDS,
    SCALE_FACTOR,
    SCALE_TIME_RATIO_MAX,
    best_roundtrip_time,
    make_document_docx,
)


def test_roundtrip_100_paragraphs_within_bound():
    """A 100-paragraph document must round-trip well below the ceiling.

    Measured ~0.09 s on a dev box; the 10 s ceiling only trips on
    catastrophic regressions (infinite/quadratic XML walks), not on slow
    CI hardware.
    """
    data = make_document_docx(100)
    elapsed = best_roundtrip_time(data, iterations=1)
    assert elapsed <= BOUND_SECONDS, (
        f"100-paragraph roundtrip took {elapsed:.2f}s (ceiling {BOUND_SECONDS}s) "
        "— converter performance regression, benchmark would fail"
    )
    assert b"PK\x03\x04" in data  # sanity: the fixture is a real zip/docx


def test_roundtrip_scales_better_than_quadratic():
    """4x the paragraphs must cost under 8x the time.

    A linear converter lands at a ratio of ~SCALE_FACTOR (4); a quadratic
    body scan lands at ~SCALE_FACTOR**2 (16). The 8.0 threshold sits in
    between with generous noise headroom.
    """
    small = make_document_docx(100)
    large = make_document_docx(100 * SCALE_FACTOR)

    t_small = best_roundtrip_time(small, iterations=3)
    t_large = best_roundtrip_time(large, iterations=3)

    if t_small <= 0:
        raise AssertionError("timing returned non-positive duration")
    ratio = t_large / t_small
    assert ratio <= SCALE_TIME_RATIO_MAX, (
        f"{SCALE_FACTOR}x input took {ratio:.2f}x the time "
        f"(ceiling {SCALE_TIME_RATIO_MAX}x) — scaling regression, "
        "likely an accidental O(n^2) path in the converter"
    )


def test_benchmark_fixture_roundtrip_preserves_content():
    """The timed conversion must preserve every paragraph's text."""
    data = make_document_docx(100)
    html = docx_to_html(data)
    assert html.count("<p>") >= 95, "missing body paragraphs in HTML"

    back = html_to_docx(html)
    doc = Document(io.BytesIO(back))
    texts = [p.text for p in doc.paragraphs]
    assert any("Paragraph 49:" in t for t in texts), "fixture text lost"
    assert any("bullet point 1" in t for t in texts), "list item lost"
    assert doc.tables, "fixture table must survive the roundtrip"
    assert doc.tables[0].cell(1, 2).text == "1,2"


def test_benchmark_harness_reports_ok():
    """The standalone harness must pass its own bounds (exit 0)."""
    from tests.bench.benchmark_converter import run_benchmark

    ok = run_benchmark(n_paragraphs=100, iterations=2)
    assert ok, "run_benchmark() reported bounds violated"
