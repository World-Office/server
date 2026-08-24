"""Benchmark harness for the DOCX <-> HTML converter.

Measures the full conversion roundtrip (``docx_to_html`` + ``html_to_docx``)
on a synthetic N-paragraph fixture that mirrors what the editor produces
(headings, bold/italic/underline runs, bullet lists, one table). It
reports milliseconds per op and conversions/sec, then asserts two
bounds:

1. an absolute ceiling for the default 100-paragraph fixture, and
2. a sub-quadratic scaling check (4x the paragraphs must not cost more
   than ~8x the time — a linear converter lands at ~4x, a quadratic one
   at ~16x).

The bounds are deliberately generous: they are not micro-benchmarks, they
are a regression gate that fails loudly only when the converter stops
scaling (e.g. an accidental O(n^2) body scan or a runaway XML walk).

Run directly::

    uv run python tests/bench/benchmark_converter.py [--paragraphs 100]
        [--iterations 5]
"""

from __future__ import annotations

import argparse
import io
import sys
import time
from pathlib import Path

# Allow running straight from a checkout: ``uv run python
# tests/bench/benchmark_converter.py`` resolves ``src`` via the project
# root, like pytest's ``pythonpath = ["src"]`` does for tests.
_PROJECT_ROOT = Path(__file__).resolve().parents[2]
for _path in (_PROJECT_ROOT, _PROJECT_ROOT / "src"):
    if str(_path) not in sys.path:
        sys.path.insert(0, str(_path))

from docx import Document  # noqa: E402

from src.editor.converter import docx_to_html, html_to_docx  # noqa: E402

#: Paragraph count of the default benchmark fixture.
DEFAULT_PARAGRAPHS = 100
#: Timing iterations per measurement (best-of is reported).
DEFAULT_ITERATIONS = 5
#: Absolute bound (seconds) for a single default-fixture roundtrip.
#: Measured ~0.09 s on a dev box; the ceiling is ~100x headroom so it
#: only trips on genuinely broken/regressed converters, not on slow CI.
BOUND_SECONDS = 10.0
#: Input growth factor for the scaling check (small -> large).
SCALE_FACTOR = 4
#: Max allowed time ratio for the scaling check. Linear conversion
#: yields ~SCALE_FACTOR, quadratic ~SCALE_FACTOR**2.
SCALE_TIME_RATIO_MAX = 8.0

_WORDS = "the quick brown fox jumps over the lazy dog "


def make_document_docx(n_paragraphs: int = DEFAULT_PARAGRAPHS) -> bytes:
    """Build a realistic-ish DOCX fixture with *n_paragraphs* blocks.

    Content is drawn from the editor's production mix: an H1 heading, a
    few bullet-list items, a small table, and body paragraphs with
    bold/italic/underline runs.
    """
    doc = Document()
    doc.add_heading("Benchmark fixture", level=1)

    for i in range(3):
        doc.add_paragraph(f"bullet point {i}", style="List Bullet")

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for r in range(2):
        for c in range(3):
            table.cell(r, c).text = f"{r},{c}"

    for i in range(max(0, n_paragraphs - 5)):
        p = doc.add_paragraph()
        lead = p.add_run(f"Paragraph {i}: ")
        lead.bold = True
        body = p.add_run(_WORDS * 2)
        if i % 3 == 0:
            body.italic = True
        if i % 7 == 0:
            p.add_run("emphasis.").underline = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def roundtrip(data: bytes) -> tuple[str, bytes]:
    """Convert DOCX -> HTML -> DOCX; returns (html, docx_back)."""
    html = docx_to_html(data)
    back = html_to_docx(html)
    return html, back


def best_roundtrip_time(data: bytes, iterations: int = DEFAULT_ITERATIONS) -> float:
    """Return the best (minimum) full-roundtrip wall time in seconds."""
    best = float("inf")
    for _ in range(max(1, iterations)):
        t0 = time.perf_counter()
        roundtrip(data)
        elapsed = time.perf_counter() - t0
        if elapsed < best:
            best = elapsed
    return best


def run_benchmark(
    n_paragraphs: int = DEFAULT_PARAGRAPHS, iterations: int = DEFAULT_ITERATIONS
) -> bool:
    """Run the benchmark, print a report, and return True if bounds hold."""
    small = make_document_docx(n_paragraphs)
    large = make_document_docx(n_paragraphs * SCALE_FACTOR)

    t_small = best_roundtrip_time(small, iterations)
    t_large = best_roundtrip_time(large, iterations)
    ratio = t_large / t_small if t_small > 0 else float("inf")

    ok_abs = t_small <= BOUND_SECONDS
    ok_scale = ratio <= SCALE_TIME_RATIO_MAX

    print(f"fixture paragraphs      : {n_paragraphs} (+ {n_paragraphs * SCALE_FACTOR} for scaling)")
    print(f"iterations (best-of)    : {iterations}")
    print(f"small roundtrip         : {t_small * 1000:9.1f} ms   ({1.0 / t_small:8.0f} conv/s)" if t_small > 0 else "small roundtrip         : n/a")
    print(f"large roundtrip         : {t_large * 1000:9.1f} ms   ({1.0 / t_large:8.0f} conv/s)" if t_large > 0 else "large roundtrip         : n/a")
    print(f"scaling ratio ({SCALE_FACTOR}x input) : {ratio:7.2f}   (max {SCALE_TIME_RATIO_MAX})")
    print(f"absolute bound          : {t_small:.2f}s <= {BOUND_SECONDS}s -> {'OK' if ok_abs else 'FAIL'}")
    print(f"scaling bound           : {'OK' if ok_scale else 'FAIL'}")
    return ok_abs and ok_scale


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--paragraphs", type=int, default=DEFAULT_PARAGRAPHS,
        help="paragraph count of the small fixture (default: %(default)s)",
    )
    parser.add_argument(
        "--iterations", type=int, default=DEFAULT_ITERATIONS,
        help="best-of timing iterations (default: %(default)s)",
    )
    args = parser.parse_args(argv)
    ok = run_benchmark(args.paragraphs, args.iterations)
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
