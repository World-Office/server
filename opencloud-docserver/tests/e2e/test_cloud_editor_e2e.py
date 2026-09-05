"""End-to-end cloud editor test (T5).

Drives the REAL editor in a real browser (Playwright/chromium) through a mock
OpenCloud/Nextcloud WOPI host:

  * the editor loads a DOCX from the mock host (client / WOPI mode);
  * TWO browser sessions edit the SAME document and the characters converge
    in real time (server-side character CRDT + poll push);
  * a save forwards the converted bytes back to the mock host (open -> edit
    -> save -> host loop proven in a browser, not just at the API level);
  * the editor notifies its embedding host via postMessage (woopi bridge);
  * Insert > page break via the toolbar button lands in the saved document.

feature register: F-076 (page break surface + serialization)

Run: pytest tests/e2e/test_cloud_editor_e2e.py
"""

from __future__ import annotations

import base64
import io
import json
import socket
import threading
import time
import urllib.parse
import urllib.request
from contextlib import asynccontextmanager
from pathlib import Path

import pytest
import uvicorn
from docx import Document
from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles

from src.config import Config
from src.editor.converter import docx_to_html
from src.editor.router import router as editor_router
from src.editor.session import SessionRegistry
from src.lib.store import DocumentStore
from src.wopi.testhost import app as mock_host_app
from src.wopi.testhost import reset_store

WEB_DIR = Path(__file__).resolve().parent.parent.parent / "web"


def _docx_bytes(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_app(tmp_path: Path) -> FastAPI:
    db = str(tmp_path / "t.db")
    content = str(tmp_path / "content")
    store = DocumentStore(db, content)
    cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

    @asynccontextmanager
    async def lifespan(app: FastAPI):
        app.state.store = store
        app.state.sessions = SessionRegistry()
        app.state.config = cfg
        yield

    app = FastAPI(lifespan=lifespan)
    app.include_router(editor_router)
    app.mount("/static", StaticFiles(directory=str(WEB_DIR)), name="static")
    return app


def _free_port() -> int:
    s = socket.socket()
    s.bind(("127.0.0.1", 0))
    port = s.getsockname()[1]
    s.close()
    return port


# A trivial host page that embeds the editor in an <iframe> and records every
# postMessage the editor sends upward — this is how OpenCloud/Nextcloud would
# receive the "woopi" bridge messages.
PARENT_HTML = """<!doctype html><html><head><meta charset="utf-8"></head><body>
<iframe id="ed" src="__EDITOR_URL__" style="width:100%;height:600px;border:0"></iframe>
<script>
  window.__msgs = [];
  window.addEventListener('message', function (e) {
    if (e.data && e.data.type === 'woopi') window.__msgs.push(e.data);
  });
</script>
</body></html>"""

parent_app = FastAPI()


@parent_app.get("/")
async def parent_index(request: Request) -> HTMLResponse:
    editor_url = request.query_params.get("editor", "")
    return HTMLResponse(PARENT_HTML.replace("__EDITOR_URL__", editor_url))


@pytest.fixture(scope="module")
def servers(tmp_path_factory):
    tmp = tmp_path_factory.mktemp("e2e")
    doc_port = _free_port()
    host_port = _free_port()
    parent_port = _free_port()

    doc_srv = uvicorn.Server(uvicorn.Config(_make_app(tmp), host="127.0.0.1", port=doc_port, log_level="error"))
    host_srv = uvicorn.Server(uvicorn.Config(mock_host_app, host="127.0.0.1", port=host_port, log_level="error"))
    parent_srv = uvicorn.Server(uvicorn.Config(parent_app, host="127.0.0.1", port=parent_port, log_level="error"))

    for s in (doc_srv, host_srv, parent_srv):
        threading.Thread(target=s.run, daemon=True).start()

    for port in (doc_port, host_port, parent_port):
        deadline = time.time() + 10
        while time.time() < deadline:
            try:
                with socket.create_connection(("127.0.0.1", port), timeout=0.5):
                    break
            except OSError:
                time.sleep(0.05)

    reset_store()
    seed = json.loads(
        urllib.request.urlopen(
            urllib.request.Request(
                f"http://127.0.0.1:{host_port}/_host/files",
                data=json.dumps(
                    {"name": "e2e.docx", "data": base64.b64encode(_docx_bytes("E2E base text")).decode()}
                ).encode(),
                headers={"Content-Type": "application/json"},
                method="POST",
            ),
            timeout=10,
        ).read()
    )
    yield {
        "doc_port": doc_port,
        "host_port": host_port,
        "parent_port": parent_port,
        "doc_id": seed["id"],
        "token": seed["access_token"],
    }
    doc_srv.should_exit = True
    host_srv.should_exit = True
    parent_srv.should_exit = True


def _editor_url(servers: dict, seed: dict | None = None) -> str:
    seed = seed or servers
    wopi_src = f"http://127.0.0.1:{servers['host_port']}/wopi/files/{seed['doc_id']}"
    return (
        f"http://127.0.0.1:{servers['doc_port']}/editor/{seed['doc_id']}"
        f"?access_token={seed['token']}&WOPISrc={urllib.parse.quote(wopi_src, safe='')}"
    )


def _parent_url(servers: dict, seed: dict | None = None) -> str:
    return f"http://127.0.0.1:{servers['parent_port']}/?editor={urllib.parse.quote(_editor_url(servers, seed), safe='')}"


def _seed_doc(servers: dict, name: str = "t.docx", text: str = "E2E base text") -> dict:
    """Create a fresh document in the mock host and return its seed."""
    resp = json.loads(
        urllib.request.urlopen(
            urllib.request.Request(
                f"http://127.0.0.1:{servers['host_port']}/_host/files",
                data=json.dumps(
                    {"name": name, "data": base64.b64encode(_docx_bytes(text)).decode()}
                ).encode(),
                headers={"Content-Type": "application/json"},
                method="POST",
            ),
            timeout=10,
        ).read()
    )
    return {"doc_id": resp["id"], "token": resp["access_token"]}


def _frame_text(frame) -> str:
    return frame.locator("#editor").inner_text()


def _post_sync(servers: dict, seed: dict, text: str) -> None:
    urllib.request.urlopen(
        urllib.request.Request(
            f"http://127.0.0.1:{servers['doc_port']}/api/documents/{seed['doc_id']}/collab/sync",
            data=json.dumps({"client_id": "probe", "text": text}).encode(),
            headers={"Content-Type": "application/json"},
            method="POST",
        ),
        timeout=10,
    ).read()


def _wait(predicate, timeout: float = 20.0) -> None:
    deadline = time.time() + timeout
    while time.time() < deadline:
        if predicate():
            return
        time.sleep(0.3)
    raise AssertionError(f"condition not met within {timeout}s")


def _host_text(servers: dict, seed: dict | None = None) -> str:
    seed = seed or servers
    url = (
        f"http://127.0.0.1:{servers['host_port']}/wopi/files/{seed['doc_id']}/contents"
        f"?access_token={seed['token']}"
    )
    data = urllib.request.urlopen(url, timeout=10).read()
    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)


def _host_html(servers: dict, seed: dict | None = None) -> str:
    """Convert the host DOCX back to HTML so markup (color, <sup>) is visible."""
    seed = seed or servers
    url = (
        f"http://127.0.0.1:{servers['host_port']}/wopi/files/{seed['doc_id']}/contents"
        f"?access_token={seed['token']}"
    )
    data = urllib.request.urlopen(url, timeout=10).read()
    return docx_to_html(data)


def test_two_users_collaborate_save_and_notify_host(servers):
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "collab.docx")
            ctx_a = browser.new_context()
            ctx_b = browser.new_context()
            parent_a = ctx_a.new_page()
            parent_b = ctx_b.new_page()

            errors = []
            parent_a.on("pageerror", lambda e: errors.append(str(e)))

            parent_a.goto(_parent_url(servers, seed))
            parent_b.goto(_parent_url(servers, seed))

            frame_a = parent_a.frame("ed")
            frame_b = parent_b.frame("ed")
            frame_a.locator("#editor").wait_for(state="visible", timeout=15000)
            frame_b.locator("#editor").wait_for(state="visible", timeout=15000)
            assert "E2E base text" in _frame_text(frame_a)

            # Live push: a hub change converges in BOTH browsers (real-time).
            _post_sync(servers, seed, "LIVEPROBE_X")
            _wait(
                lambda: "LIVEPROBE_X" in _frame_text(frame_a) and "LIVEPROBE_X" in _frame_text(frame_b)
            )

            # Presence: both editors show each other as collaborators (chips),
            # and a remote caret is rendered for the peer.
            assert frame_a.locator("#collab-peers .peer-chip").count() >= 2
            assert frame_b.locator("#collab-peers .peer-chip").count() >= 2
            assert frame_a.locator(".remote-caret").count() >= 1

            # User A types -> User B converges.
            frame_a.locator("#editor").click()
            frame_a.locator("#editor").press("End")
            frame_a.locator("#editor").press_sequentially(" from A")
            _wait(lambda: "from A" in _frame_text(frame_a))
            _wait(lambda: "from A" in _frame_text(frame_b))

            # User B types -> User A converges (bidirectional).
            frame_b.locator("#editor").click()
            frame_b.locator("#editor").press("End")
            frame_b.locator("#editor").press_sequentially(" +B")
            _wait(lambda: "+B" in _frame_text(frame_a))

            # Save -> converted bytes reach the mock WOPI host.
            frame_a.locator("#btn-save").click()
            _wait(lambda: "from A" in _host_text(servers, seed) and "+B" in _host_text(servers, seed))
            host_text = _host_text(servers, seed)
            assert "from A" in host_text and "+B" in host_text

            # Editor notified its embedding host via postMessage (woopi bridge).
            msgs = parent_a.evaluate("window.__msgs")
            assert any(
                m.get("type") == "woopi" and m.get("action") in ("editing", "saved") for m in msgs
            )
        finally:
            ctx_a.close()
            ctx_b.close()
            browser.close()


def _word_count(frame):
    import re
    txt = frame.locator("#word-count").inner_text()
    m = re.search(r"(\d+)", txt)
    return int(m.group(1)) if m else 0


def test_status_bar_word_count_and_save_indicator(servers):
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "status.docx")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)

            # The status bar shows a live word count for the loaded document.
            wc0 = _word_count(frame)
            assert wc0 > 0, "status bar must show a word count"

            # Typing updates the count live.
            frame.locator("#editor").click()
            frame.locator("#editor").press("End")
            frame.locator("#editor").press_sequentially(" extra words here")
            _wait(lambda: _word_count(frame) > wc0 + 2)

            # Typing flips the save indicator away from a saved state.
            status0 = frame.locator("#status").inner_text().strip().lower()
            assert status0 != "ready", f"status should show unsaved, got {status0!r}"

            # Saving returns the indicator to saved/ready.
            frame.locator("#btn-save").click()
            _wait(
                lambda: frame.locator("#status").inner_text().strip().lower()
                in ("saved", "ready")
            )
        finally:
            ctx.close()
            browser.close()


def test_view_controls_zoom_theme_fullscreen(servers):
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "view.docx")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)

            # Zoom in scales only the editing surface (inline zoom grows).
            z0 = float(frame.locator("#editor").evaluate("el => parseFloat(el.style.zoom || '1')"))
            frame.locator("#btn-zoom-in").click()
            z1 = float(frame.locator("#editor").evaluate("el => parseFloat(el.style.zoom || '1')"))
            assert z1 > z0, f"zoom should increase: {z0} -> {z1}"
            # Content is unaffected by zoom.
            assert _frame_text(frame).strip()

            # Theme toggle flips the page background colour.
            bg0 = frame.evaluate("getComputedStyle(document.body).backgroundColor")
            frame.locator("#btn-theme").click()
            bg1 = frame.evaluate("getComputedStyle(document.body).backgroundColor")
            assert bg0 != bg1, f"theme should change bg: {bg0} -> {bg1}"

            # Fullscreen toggles the body class (browser Fullscreen API is
            # best-effort; the class drives the layout expansion).
            assert "fullscreen" not in frame.evaluate("document.body.className")
            frame.locator("#btn-fullscreen").click()
            assert "fullscreen" in frame.evaluate("document.body.className")
        finally:
            ctx.close()
            browser.close()


def test_insert_link_roundtrip(servers):
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "link.docx")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)

            frame.locator("#editor").click()
            frame.locator("#editor").press("End")
            frame.locator("#editor").press_sequentially("Visit our site")
            frame.locator("#editor").select_text()  # select all document text
            frame.locator("#btn-link").click()
            frame.locator("#link-url").fill("https://example.com")
            frame.locator("#btn-link-ok").click()
            frame.locator("#editor a[href='https://example.com']").wait_for(state="attached", timeout=5000)
            assert frame.locator("#editor a").first.get_attribute("href") == "https://example.com"

            # Save -> the link survives the round-trip back to the host.
            frame.locator("#btn-save").click()
            _wait(lambda: "example.com" in _host_text(servers, seed))
        finally:
            ctx.close()
            browser.close()


def test_format_color_highlight_superscript(servers):
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "color.docx")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)

            frame.locator("#editor").click()
            frame.locator("#editor").press("End")
            frame.locator("#editor").press_sequentially("Color me x2")

            # Select the word "Color" and apply red via the colour picker.
            frame.locator("#editor").select_text()
            frame.locator("#text-color").fill("#ff0000")
            frame.locator("#text-color").dispatch_event("change")
            # Select the whole contenteditable so superscript applies visibly,
            # then check the styling spans survive.
            frame.locator("#editor").select_text()
            frame.locator("button[data-cmd='superscript']").click()

            # The styling is visible in the DOM.
            colored = frame.locator("#editor span[style*='color']").count()
            assert colored >= 1, "expected a colored span in the editor"
            assert frame.locator("#editor sup").count() >= 1, "expected sup in the editor"

            # Save -> the styling persists back to the host bytes.
            frame.locator("#btn-save").click()
            _wait(lambda: "ff0000" in _host_html(servers, seed) and "<sup>" in _host_html(servers, seed))
        finally:
            ctx.close()
            browser.close()


def test_table_merge_and_column_ops(servers):
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "table.docx")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)

            frame.locator("#editor").click()
            frame.locator("#editor").press("End")
            frame.locator("#btn-table").click()
            frame.locator("#table-rows").fill("2")
            frame.locator("#table-cols").fill("2")
            frame.locator("#btn-table-ok").click()
            frame.locator("#editor table").wait_for(state="attached", timeout=5000)
            assert frame.locator("#editor table tr").count() == 2

            # Select from the 1st cell of row 1 to the 2nd cell, then merge.
            frame.evaluate("""() => {
              const rows = document.querySelectorAll('#editor table tr');
              const c0 = rows[0].cells[0];
              const c1 = rows[0].cells[1];
              const endNode = (c1.firstChild && c1.firstChild.nodeType === 3)
                ? c1.firstChild : c1;
              const r = document.createRange();
              r.setStart(c0.firstChild || c0, 0);
              r.setEnd(endNode, endNode.nodeType === 3 ? endNode.length : 1);
              const s = window.getSelection();
              s.removeAllRanges();
              s.addRange(r);
            }""")
            frame.locator("#btn-table-ops").click()
            frame.locator("#op-merge").click()
            merged = frame.locator("#editor tr:first-child td").first
            merged.wait_for(state="attached", timeout=5000)
            colspan = merged.get_attribute("colspan")
            assert colspan == "2", f"expected merged colspan=2, got {colspan}"

            # Delete the 2nd column: click its cell in the second row, then act.
            frame.locator("#editor tr").nth(1).locator("td").nth(1).click()
            frame.locator("#btn-table-ops").click()
            frame.locator("#op-del-col").click()
            assert frame.locator("#editor tr").nth(1).locator("td").count() == 1

            # Save -> the merged colspan survives round-trip to the host.
            frame.locator("#btn-save").click()
            _wait(lambda: "colspan" in _host_html(servers, seed).lower())
        finally:
            ctx.close()
            browser.close()


def test_insert_hr_pagebreak_symbol(servers):
    """Insert HR, page break and a symbol; all survive save via the host."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "insert-misc.docx")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)

            frame.locator("#editor").click()
            frame.locator("#editor").press("End")

            # Horizontal rule.
            frame.locator("#btn-hr").click()
            frame.locator("#editor hr").wait_for(state="attached", timeout=5000)
            assert frame.locator("#editor hr").count() == 1

            # Page break marker.
            frame.locator("#btn-page-break").click()
            frame.locator("#editor div.page-break").wait_for(state="attached", timeout=5000)
            assert frame.locator("#editor div.page-break").count() == 1

            # Symbol picker -> first symbol (§) inserted as text.
            frame.locator("#btn-symbol").click()
            frame.locator("#symbol-dialog .symbol-btn").first.click()
            _wait(lambda: "§" in _frame_text(frame))

            # Date/time insert -> ISO date appears as text.
            frame.locator("#btn-datetime").click()
            _wait(lambda: "2026-" in _frame_text(frame))

            # Save -> markers + symbol + date reach the host DOCX.
            frame.locator("#btn-save").click()
            _wait(lambda: (
                "<hr" in _host_html(servers, seed)
                and "page-break" in _host_html(servers, seed)
                and "§" in _host_html(servers, seed)
                and "2026-" in _host_html(servers, seed)
            ))

            # Reload -> they come back from the host into the editor.
            parent.reload()
            _wait(lambda: parent.frame("ed") is not None)
            frame2 = parent.frame("ed")
            frame2.locator("#editor").wait_for(state="visible", timeout=15000)
            _wait(lambda: frame2.locator("#editor hr").count() == 1)
            assert frame2.locator("#editor div.page-break").count() == 1
            assert "§" in _frame_text(frame2)
        finally:
            ctx.close()
            browser.close()


def test_file_menu_export_odt_and_new_document(servers):
    """Export ODT from the File menu; New clears + persists the editor."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "file-ops.docx", "Exportable body")
            ctx = browser.new_context(accept_downloads=True)
            page = ctx.new_page()
            # Accept the New-document confirm().
            page.on("dialog", lambda d: d.accept())
            page.goto(_parent_url(servers, seed))
            frame = page.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)

            # --- Export ODT via File > Export > ODT -> downloadable archive.
            frame.locator("#btn-file").click()
            frame.locator("#btn-export").hover()
            with page.expect_download(timeout=10000) as dl_info:
                frame.locator("button[data-export='odt']").click()
            dl = dl_info.value
            assert dl.suggested_filename.endswith(".odt"), dl.suggested_filename
            path = dl.path()
            data = path.read_bytes()
            assert data[:2] == b"PK", "ODT export must be a zip"
            import zipfile as _zipfile
            with _zipfile.ZipFile(io.BytesIO(data)) as zf:
                content = zf.read("content.xml").decode("utf-8", "replace")
            assert "Exportable body" in content

            # --- New: confirm -> editor cleared.
            frame.locator("#btn-file").click()
            frame.locator("#btn-new").click()
            _wait(lambda: frame.locator("#editor").inner_text().strip() == "")
            assert frame.locator("#editor").inner_text().strip() == ""

            # Save -> the blank document reaches the host too.
            frame.locator("#btn-save").click()
            _wait(lambda: _host_text(servers, seed).strip() == "")
        finally:
            ctx.close()
            browser.close()


def test_offline_queue_and_resync(servers):
    """An offline save queues a local snapshot; it flushes on reconnect."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "offline.docx", "Offline seed")
            ctx = browser.new_context()
            page = ctx.new_page()
            page.goto(_parent_url(servers, seed))
            frame = page.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)
            _wait(lambda: "Offline seed" in _frame_text(frame))

            frame.locator("#editor").click()
            frame.locator("#editor").press("End")
            frame.locator("#editor").press_sequentially(" queued edits")

            # Go offline -> Save fails -> snapshot queued locally + indicator.
            ctx.set_offline(True)
            frame.locator("#btn-save").click()
            _wait(lambda: not frame.locator("#offline-indicator").evaluate(
                "el => el.hidden"))
            queued = frame.evaluate(
                "JSON.parse(localStorage.getItem('wo-offline-queue') || 'null')")
            assert queued and queued["docId"] == seed["doc_id"], queued
            assert "queued edits" in queued["html"]

            # Back online -> the queued snapshot flushes to the host.
            ctx.set_offline(False)
            frame.evaluate("window.dispatchEvent(new Event('online'))")
            _wait(lambda: "queued edits" in _host_text(servers, seed))
            _wait(lambda: frame.evaluate(
                "localStorage.getItem('wo-offline-queue') === null"))
            assert frame.locator("#offline-indicator").evaluate("el => el.hidden")
        finally:
            ctx.close()
            browser.close()


def test_inline_format_commands_code_caps_strike(servers):
    """Code / small-caps / all-caps / strike round-trip in the browser."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "inlinefmt.docx", "plain base")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)
            _wait(lambda: "plain base" in _frame_text(frame))

            # Type a line, then wrap one word in inline code (monospace).
            frame.locator("#editor").click()
            frame.locator("#editor").press("End")
            frame.locator("#editor").press_sequentially(" code SC UP strike ")
            frame.locator("#editor").evaluate("""() => {
              const ed = document.getElementById('editor');
              const t = ed.querySelector('p:last-of-type');
              const range = document.createRange();
              range.selectNodeContents(t);
              const sel = window.getSelection();
              sel.removeAllRanges(); sel.addRange(range);
            }""")
            frame.locator("button[data-cmd='code']").click()
            frame.locator("#editor").evaluate("""() => {
              const ed = document.getElementById('editor');
              const t = ed.querySelector('p:last-of-type');
              const range = document.createRange();
              range.selectNodeContents(t);
              const sel = window.getSelection();
              sel.removeAllRanges(); sel.addRange(range);
            }""")
            frame.locator("button[data-cmd='allCaps']").click()
            frame.locator("#editor").evaluate("""() => {
              const ed = document.getElementById('editor');
              const t = ed.querySelector('p:last-of-type');
              const range = document.createRange();
              range.selectNodeContents(t);
              const sel = window.getSelection();
              sel.removeAllRanges(); sel.addRange(range);
            }""")
            frame.locator("button[data-cmd='strikeThrough']").click()

            html = frame.evaluate("document.getElementById('editor').innerHTML")
            assert "Consolas" in html or "monospace" in html.lower(), html
            assert "uppercase" in html, html
            assert "<s>" in html or "<strike>" in html, html

            # Save -> all three survive to the host DOCX.
            frame.locator("#btn-save").click()
            _wait(lambda: (
                ("Consolas" in _host_html(servers, seed) or
                 "<code>" in _host_html(servers, seed))
                and "uppercase" in _host_html(servers, seed)
                and ("<s>" in _host_html(servers, seed) or "<strike>" in _host_html(servers, seed))
            ))

        finally:
            ctx.close()
            browser.close()


def test_paragraph_rtl_and_line_spacing_roundtrip(servers):
    """RTL + line-spacing paragraph props persist through save + reload."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "parafmt.docx", "para base")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)
            _wait(lambda: "para base" in _frame_text(frame))

            frame.locator("#editor").click()
            frame.locator("#editor").press("End")
            frame.locator("#editor").press_sequentially(" RTL line")

            # Apply 1.5 line spacing via the dropdown then RTL on the same block.
            frame.select_option("#line-spacing", "1.5")
            frame.locator("button[data-cmd='directionRtl']").click()

            html = frame.evaluate("document.getElementById('editor').innerHTML")
            assert 'line-height: 1.5' in html, html
            assert 'direction: rtl' in html, html

            # Save -> both reach the host DOCX.
            frame.locator("#btn-save").click()
            _wait(lambda: (
                "line-height:1.5" in _host_html(servers, seed)
                and "direction:rtl" in _host_html(servers, seed)
            ))
            # Let any in-flight collab poll settle so it can't clobber the
            # just-saved DOCX before the reload reads it back (save vs poll
            # race), then confirm the host still holds the properties.
            _t_settle = time.time()
            while time.time() - _t_settle < 1.5:
                time.sleep(0.2)
            assert "line-height:1.5" in _host_html(servers, seed)
            assert "direction:rtl" in _host_html(servers, seed)

            # Reload from the host -> props survive.
            parent.reload()
            frame2 = parent.frame("ed")
            frame2.locator("#editor").wait_for(state="visible", timeout=15000)
            html2 = frame2.evaluate("document.getElementById('editor').innerHTML")
            assert 'line-height' in html2 and 'rtl' in html2.lower(), html2
        finally:
            ctx.close()
            browser.close()


def test_nested_list_tab_indent_roundtrip(servers):
    """Tab indents a list item into a nested list that survives save."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
        try:
            seed = _seed_doc(servers, "lists.docx", "list base")
            ctx = browser.new_context()
            parent = ctx.new_page()
            parent.goto(_parent_url(servers, seed))
            frame = parent.frame("ed")
            frame.locator("#editor").wait_for(state="visible", timeout=15000)
            _wait(lambda: "list base" in _frame_text(frame))

            frame.locator("#editor").click()
            frame.locator("#editor").press("End")
            # Start a bullet list on a fresh line: Enter, then the bullet
            # toolbar button (toggleList -> native execCommand) turns the new
            # paragraph into an <li>; type two items; Tab indents the second.
            frame.locator("#editor").press("Enter")
            frame.locator("#editor").press_sequentially("first item")
            frame.locator("button[data-cmd='insertUnorderedList']").click()
            frame.locator("#editor").press("End")
            frame.locator("#editor").press("Enter")
            frame.locator("#editor").press_sequentially("second item")
            frame.locator("#editor").press("Tab")
            # The live editor DOM may be transiently invalid (Chromium can
            # wrap the list in a <p> and put the nested <ul> beside the <li>);
            # the sanitizer normalises it on save, so assert on the host.
            frame.locator("#btn-save").click()
            _wait(lambda: (
                "<ul><li>second item</li></ul>" in _host_html(servers, seed).replace("\n", "")
                or "List Bullet 2" in _host_html(servers, seed)
            ))
        finally:
            ctx.close()
            browser.close()
