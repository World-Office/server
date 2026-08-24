"""E2E — Vollständiger Workflow (complete production workflow).

This suite drives every hop of the production document path through real
HTTP, exactly as the docserver is used in production (deployed against
real OpenCloud / OCIS, e.g. cloud.graphwiz.ai):

    WOPI discovery announces the editor
      -> OpenCloud launches /editor with access_token + WOPISrc
      -> the editor loads the document as HTML      (docx/odt -> HTML)
      -> the user edits the HTML (pastes the persistence marker)
      -> the editor saves the HTML                  (HTML -> docx/odt)
      -> the docserver PUTs the bytes to the remote host   (PutFile)
      -> reload: the edit is still there
      -> close (unlock) and reopen: the edit survives
      -> a second editor of the same file is served read-only
      -> XSS payloads in the submitted HTML are neutralized before
         anything is persisted

A real HTTP WOPI host stands in for OCIS (same wire protocol, same
lock-before-PutFile semantics the wopiserver enforces). The real
docserver app, editor router, RemoteWopiClient and both converters
(DOCX + ODT) are exercised end to end — nothing is mocked on the
docserver side.

The persistence contract: text typed in the editor survives every hop of
the production path, and the bytes stored back on the remote host are a
real office document (DOCX paragraphs / ODT content.xml) that contains it.

Marker emitted when the complete workflow passes:
FULL-WORKFLOW: OK
"""

from __future__ import annotations

import io
import json
import re
import threading
import zipfile
from wsgiref.simple_server import make_server

import pytest
from fastapi.testclient import TestClient

from src.config import Config
from src.lib.store import wipe_db, wipe_dir
from src.main import create_app

# Markers the "user" types into the document; a save is never treated as
# successful until the marker is found in the stored remote bytes.
MARKER_DOCX = "FULL-WORKFLOW-DOCX"
MARKER_ODT = "FULL-WORKFLOW-ODT"


# ----------------------------------------------------------------------
# Production OpenCloud (OCIS) wopiserver stand-in
# ----------------------------------------------------------------------

class _ProdOcisHost:
    """Minimal OCIS wopiserver over WSGI — what cloud.graphwiz.ai does.

    Implements the same WOPI surface the real wopiserver serves and the
    same quirks the docserver's RemoteWopiClient depends on:
      -  GET  /wopi/files/{id}             CheckFileInfo (Bearer auth)
      -  GET  /wopi/files/{id}/contents    GetFile
      -  POST /wopi/files/{id}/contents    PutFile (X-WOPI-Override: PUT)
      -  POST /wopi/files/{id}             LOCK / GET_LOCK / UNLOCK
    Like the real wopiserver, PutFile on an unlocked file is refused
    (409 "Cannot PutFile on unlocked file"), so every save must present
    the lock the docserver took at launch.
    """

    def __init__(self) -> None:
        self.content: dict[str, bytes] = {}
        self.names: dict[str, str] = {}
        self.locks: dict[str, str] = {}
        self.lock_events: list[str] = []
        self.put_count = 0
        self.put_overrides: list[str] = []
        self.put_lock_headers: list[str] = []
        self.getfile_count = 0

    def seed(self, doc_id: str, name: str, data: bytes) -> None:
        self.content[doc_id] = data
        self.names[doc_id] = name

    def __call__(self, environ, start_response):
        path = environ.get("PATH_INFO", "")
        method = environ.get("REQUEST_METHOD", "GET")
        override = environ.get("HTTP_X_WOPI_OVERRIDE", "")
        lock_hdr = environ.get("HTTP_X_WOPI_LOCK", "")
        auth = environ.get("HTTP_AUTHORIZATION", "")

        m = re.match(r"^/wopi/files/([^/]+)(/contents)?$", path)
        if not m:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"not found"]
        doc_id, is_contents = m.group(1), bool(m.group(2))
        if doc_id not in self.content:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"no such file"]

        # GetFile — the docserver reads the raw bytes with the token.
        if method == "GET" and is_contents:
            self.getfile_count += 1
            start_response(
                "200 OK",
                [
                    ("Content-Type", "application/octet-stream"),
                    ("X-WOPI-ItemVersion", "v1"),
                ],
            )
            return [self.content[doc_id]]

        # CheckFileInfo — the docserver reads BaseFileName (routes the
        # converter by extension) and UserId (names the WOPI lock).
        if method == "GET" and not is_contents:
            user_id = auth.replace("Bearer ", "") or "anonymous"
            body = json.dumps(
                {
                    "BaseFileName": self.names.get(doc_id, "document.docx"),
                    "UserId": user_id,
                    "Size": len(self.content[doc_id]),
                }
            ).encode()
            start_response("200 OK", [("Content-Type", "application/json")])
            return [body]

        # PutFile — only accepted with the lock taken at launch.
        if method == "POST" and is_contents and override == "PUT":
            length = int(environ.get("CONTENT_LENGTH", "0"))
            if lock_hdr != (self.locks.get(doc_id) or ""):
                if not self.locks.get(doc_id):
                    start_response(
                        "409 Conflict",
                        [
                            ("Content-Type", "text/plain"),
                            ("X-WOPI-Lock", ""),
                            ("X-WOPI-LockFailureReason", "Cannot PutFile on unlocked files"),
                        ],
                    )
                    return [b"conflict"]
                start_response("500 Internal Server Error", [("Content-Type", "text/plain")])
                return [b"lock mismatch"]
            self.content[doc_id] = environ["wsgi.input"].read(length)
            self.put_count += 1
            self.put_overrides.append(override)
            self.put_lock_headers.append(lock_hdr)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        if method == "POST" and not is_contents and override == "LOCK":
            if self.locks.get(doc_id) and self.locks[doc_id] != lock_hdr:
                start_response(
                    "409 Conflict",
                    [("Content-Type", "text/plain"), ("X-WOPI-Lock", self.locks[doc_id])],
                )
                return [b"locked by other"]
            self.locks[doc_id] = lock_hdr
            self.lock_events.append(f"LOCK:{doc_id}:{lock_hdr}")
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-ItemVersion", "v1")],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "GET_LOCK":
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-Lock", self.locks.get(doc_id, ""))],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "UNLOCK":
            if lock_hdr == self.locks.get(doc_id):
                self.locks.pop(doc_id)
            self.lock_events.append(f"UNLOCK:{doc_id}")
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        start_response("404 Not Found", [("Content-Type", "text/plain")])
        return [b"not found"]


# ----------------------------------------------------------------------
# The docserver wired against a live remote OpenCloud host
# ----------------------------------------------------------------------

class _E2EStack:
    """The docserver plus a remote OpenCloud host, wired like production.

    The docserver runs as a FastAPI TestClient; the remote host is a real
    WSGI server on 127.0.0.1 so the docserver's RemoteWopiClient makes real
    HTTP calls (urllib) exactly like it would against cloud.graphwiz.ai.
    """

    def __init__(self, tmp_path) -> None:
        self.host = _ProdOcisHost()
        self._httpd = make_server("127.0.0.1", 0, self.host)
        self.port = self._httpd.server_address[1]
        self.wopi_host = f"http://127.0.0.1:{self.port}"
        self._thread = threading.Thread(target=self._httpd.serve_forever, daemon=True)
        self._thread.start()

        db = str(tmp_path / "t.db")
        content = str(tmp_path / "content")
        cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")
        # The REAL production app (src.main.create_app) — WOPI + editor
        # routers, /health, /, static files — exactly what a deployed
        # docserver runs.
        app = create_app(cfg)
        self.client = TestClient(app)
        self.client.__enter__()  # run lifespan (app.state.*)
        self._db, self._content = db, content

    def close(self) -> None:
        try:
            self.client.__exit__(None, None, None)
        except Exception:
            pass
        self._httpd.shutdown()
        self._httpd.server_close()
        self._thread.join(timeout=2)
        wipe_db(self._db)
        wipe_dir(self._content)

    def launch_editor(self, doc_id: str, user: str = "alice") -> str:
        """OpenCloud launches /editor (POST form) for a user's file.

        Returns the rendered editor page so the test can read the session
        id (`__SESSION__`) and readonly flag (`__READ_ONLY__`) the page car
        ries. The WOPI lock must be taken on the remote host.
        """
        wopi_src = f"{self.wopi_host}/wopi/files/{doc_id}"
        resp = self.client.post(
            "/editor",
            params={"WOPISrc": wopi_src},
            data={"file_id": doc_id, "access_token": user},
        )
        assert resp.status_code == 200
        assert self.host.locks.get(doc_id), "launch must have taken the WOPI lock"
        return resp.text

    def load_html(self, doc_id: str) -> str:
        resp = self.client.get(f"/api/documents/{doc_id}/html")
        assert resp.status_code == 200, resp.text
        return resp.json()["html"]

    def save_html(self, doc_id: str, html: str, session: str = "") -> object:
        url = f"/api/documents/{doc_id}/save"
        params = {"session": session} if session else {}
        resp = self.client.post(url, params=params, json={"html": html})
        return resp

    def unlock(self, doc_id: str, session: str = "") -> None:
        url = f"/api/documents/{doc_id}/unlock"
        params = {"session": session} if session else {}
        resp = self.client.post(url, params=params)
        assert resp.status_code == 200, resp.text

    def remote_doc(self, doc_id: str) -> bytes:
        """Raw bytes the OpenCloud host stored after the docserver's PutFile."""
        return self.host.content[doc_id]


@pytest.fixture
def stack(tmp_path):
    s = _E2EStack(tmp_path)
    yield s
    s.close()


def _page_session(page: str) -> str:
    """Extract window.__SESSION__ = "..." from the rendered editor page."""
    m = re.search(r"window\.__SESSION__\s*=\s*(\"[^\"]*\"|null)", page)
    assert m, "editor page must expose __SESSION__"
    return json.loads(m.group(1))


# ----------------------------------------------------------------------
# DOCX / ODT builders + stored-bytes inspection
# ----------------------------------------------------------------------

def _build_docx(title: str = "Stoic Title", body: str = "Original Stoic body") -> bytes:
    """Build a DOCX with a heading, a bold run, a list and a plain paragraph."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    r = p.add_run("bold and ")
    r.bold = True
    r2 = p.add_run("italic")
    r2.italic = True
    doc.add_paragraph(body)
    doc.add_paragraph("first item", style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _build_odt(title: str = "Stoic Title", body: str = "Original Stoic body") -> bytes:
    """Build an ODT with a heading, a bold run and a plain paragraph."""
    from odf.opendocument import OpenDocumentText
    from odf.style import Style, TextProperties
    from odf.text import H, P, Span

    doc = OpenDocumentText()
    doc.text.addElement(H(outlinelevel=1, text=title))
    bold = Style(name="WO_B", family="text")
    bold.addElement(TextProperties(fontweight="bold"))
    doc.automaticstyles.addElement(bold)
    p = P()
    p.addElement(Span(text=body, stylename="WO_B"))
    doc.text.addElement(p)
    doc.text.addElement(P(text="Plain text line."))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _odt_text(data: bytes) -> str:
    from odf import teletype
    from odf.opendocument import load

    doc = load(io.BytesIO(data))
    return teletype.extractText(doc.text)


def _content_xml_has(data: bytes, needle: str) -> bool:
    """True when an ODT's content.xml (the persisted body) contains needle."""
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        content_xml = zf.read("content.xml").decode("utf-8", "replace")
    return needle in content_xml


# ----------------------------------------------------------------------
# The E2E tests — the complete production workflow
# ----------------------------------------------------------------------

def test_complete_docx_lifecycle_through_production_opencloud(stack):
    """The vollständiger Workflow for DOCX: discovery -> OpenCloud launches
    the editor -> load as HTML -> edit -> save -> remote PutFile -> reload
    -> unlock -> reopen. Emits the FULL-WORKFLOW marker only when the DOCX
    stored back on the OpenCloud host contains the marker text."""
    # 1. OpenCloud asks the docserver what it can edit (WOPI discovery).
    disc = stack.client.get("/hosting/discovery")
    assert disc.status_code == 200
    assert "docx" in disc.text and "odt" in disc.text
    assert 'name="edit" ext="docx"' in disc.text
    assert "{public_url}" not in disc.text  # placeholder must be substituted

    # 2. OpenCloud seeds the file and launches the editor for alice.
    stack.host.seed("wf-docx", "report.docx", _build_docx("Stoic Report", "Draft paragraph"))
    page = stack.launch_editor("wf-docx", user="alice")
    session = _page_session(page)
    assert "alice" in stack.host.locks["wf-docx"], "lock must be owner-named"
    assert 'window.__READ_ONLY__ = false' in page
    assert 'window.__DOC_ID__ = "wf-docx"' in page

    # 3. The editor loads the document as HTML.
    html = stack.load_html("wf-docx")
    assert "Stoic Report" in html
    assert "Draft paragraph" in html
    assert "<b>bold and </b>" in html and "<i>italic</i>" in html
    assert "<ul>" in html and "<li>first item</li>" in html

    # 4. The user types the marker and saves.
    edited = html + f"\n<p>{MARKER_DOCX}</p>"
    resp = stack.save_html("wf-docx", edited, session=session)
    assert resp.status_code == 200, resp.text
    assert resp.json().get("ok") is True

    # 5. The docserver must have PUT the DOCX back to the OpenCloud host,
    #    presenting the WOPI lock it took at launch (409 otherwise).
    assert stack.host.put_count == 1, "edited doc must be PUT back to OpenCloud"
    assert stack.host.put_overrides == ["PUT"]
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0], (
        "save must present the WOPI lock taken at launch"
    )

    # 6. Reload from the remote host: the edited text is still there.
    reloaded = stack.load_html("wf-docx")
    assert MARKER_DOCX in reloaded

    # 7. The stored bytes are a real DOCX whose paragraphs contain the marker.
    stored = stack.remote_doc("wf-docx")
    text = _docx_text(stored)
    assert MARKER_DOCX in text, "stored DOCX must contain the marker"
    assert "Stoic Report" in text and "Draft paragraph" in text
    assert "first item" in text

    # 8. Close (unlock on the remote host) and reopen as another user: the
    #    edit survives the full close/reopen cycle.
    stack.unlock("wf-docx", session=session)
    assert not stack.host.locks.get("wf-docx"), "unlock must release the remote lock"
    stack.launch_editor("wf-docx", user="bob")
    again = stack.load_html("wf-docx")
    assert MARKER_DOCX in again
    assert "Stoic Report" in again

    print("FULL-WORKFLOW: OK")


def test_complete_odt_lifecycle_through_production_opencloud(stack):
    """The vollständiger Workflow for ODT — same production path, same
    persistence contract, verified in the stored ODT's content.xml."""
    stack.host.seed("wf-odt", "notes.odt", _build_odt("ODT Notes", "First draft line"))
    page = stack.launch_editor("wf-odt", user="alice")
    session = _page_session(page)
    assert 'window.__DOC_NAME__ = "notes.odt"' in page

    html = stack.load_html("wf-odt")
    assert "ODT Notes" in html and "First draft line" in html

    edited = html + f"\n<p>{MARKER_ODT}</p>"
    resp = stack.save_html("wf-odt", edited, session=session)
    assert resp.status_code == 200, resp.text

    assert stack.host.put_count == 1
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0]

    reloaded = stack.load_html("wf-odt")
    assert MARKER_ODT in reloaded

    stored = stack.remote_doc("wf-odt")
    assert _content_xml_has(stored, MARKER_ODT), "content.xml must contain the marker"
    assert MARKER_ODT in _odt_text(stored)
    assert "ODT Notes" in _odt_text(stored)
    assert "First draft line" in _odt_text(stored)

    # Close and reopen: the edit survives.
    stack.unlock("wf-odt", session=session)
    stack.launch_editor("wf-odt", user="bob")
    again = stack.load_html("wf-odt")
    assert MARKER_ODT in again


def test_second_editor_is_read_only_and_cannot_clobber(stack):
    """US-5 in the full workflow: while alice edits a remote file, bob's
    editor is served read-only; his save is rejected with 403 and the
    remote bytes stay alice's. Both sessions coexist in ONE registry."""
    stack.host.seed("wf-share", "shared.docx", _build_docx("Shared", "alice owns this"))
    page_a = stack.launch_editor("wf-share", user="alice")
    session_a = _page_session(page_a)
    assert 'window.__READ_ONLY__ = false' in page_a
    assert stack.host.locks["wf-share"].startswith("wo:alice:")

    # bob opens the same file -> read-only, no lock stolen.
    page_b = stack.launch_editor("wf-share", user="bob")
    session_b = _page_session(page_b)
    assert 'window.__READ_ONLY__ = true' in page_b
    assert stack.host.locks["wf-share"].startswith("wo:alice:"), (
        "bob must not steal alice's lock"
    )

    # bob saves -> 403, remote bytes unchanged.
    resp = stack.save_html("wf-share", "<p>bob clobbers</p>", session=session_b)
    assert resp.status_code == 403, resp.text
    assert "bob clobbers" not in _docx_text(stack.remote_doc("wf-share"))

    # alice still saves fine through her own session.
    resp = stack.save_html(
        "wf-share", "<p>alice keeps editing</p>", session=session_a
    )
    assert resp.status_code == 200, resp.text
    assert "alice keeps editing" in _docx_text(stack.remote_doc("wf-share"))

    # both sessions coexist, keyed by distinct session ids.
    reg = stack.client.app.state.sessions
    assert reg.get_by_id(session_a) is not None
    assert reg.get_by_id(session_b) is not None
    assert session_a != session_b


def test_xss_payload_in_save_is_sanitized_before_persistence(stack):
    """The save path sanitizes before converting/persisting: a <script>,
    an onerror image and a javascript: link never reach the stored bytes
    nor the reloaded HTML."""
    stack.host.seed("wf-xss", "evil.docx", _build_docx("Clean", "benign text"))
    session = _page_session(stack.launch_editor("wf-xss", user="alice"))

    evil = (
        "<p>benign text</p>"
        "<script>alert('pwned')</script>"
        '<p><img src="x" onerror="alert(1)"/></p>'
        '<a href="javascript:alert(1)">clickme</a>'
    )
    resp = stack.save_html("wf-xss", evil, session=session)
    assert resp.status_code == 200, resp.text

    stored_text = _docx_text(stack.remote_doc("wf-xss"))
    assert "benign text" in stored_text
    assert "pwned" not in stored_text, "script content must not persist"
    assert "alert" not in stored_text
    assert "javascript:" not in stored_text

    reloaded = stack.load_html("wf-xss")
    assert "<script" not in reloaded
    assert "onerror" not in reloaded
    assert "javascript:" not in reloaded


# ----------------------------------------------------------------------
# Standalone local-host mode — the complete workflow without OCIS
# ----------------------------------------------------------------------

def test_local_host_mode_complete_workflow(stack, tmp_path):
    """The vollständiger Workflow with the docserver as its own WOPI host:
    upload -> list -> metadata -> edit -> save -> GetFile -> PutFile with
    lock -> unlock -> health. Exercises the store-backed surface on top of
    the same editor + WOPI routers."""
    c = stack.client

    # 1. Upload a document (what the admin panel / curl does).
    data = _build_docx("Local Doc", "local body")
    resp = c.post(
        "/api/upload",
        files={"file": ("local.docx", io.BytesIO(data), "application/octet-stream")},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["id"] == "local.docx"

    # 2. List + metadata.
    listing = c.get("/api/documents").json()
    assert any(d["id"] == "local.docx" for d in listing)
    meta = c.get("/api/documents/local.docx").json()
    assert meta["name"] == "local.docx" and meta["size"] == len(data)

    # 3. Edit as HTML, save locally.
    html = c.get("/api/documents/local.docx/html").json()["html"]
    assert "Local Doc" in html and "local body" in html
    resp = c.post("/api/documents/local.docx/save", json={"html": html + "<p>LOCAL-WF</p>"})
    assert resp.status_code == 200, resp.text

    # 4. Environment-mirroring WOPI host surface: GetFile returns the saved
    #    bytes; PutFile honours the lock (409 on mismatch).
    gf = c.get("/wopi/files/local.docx/contents")
    assert gf.status_code == 200
    assert "LOCAL-WF" in _docx_text(gf.content)
    assert "X-WOPI-ItemVersion" in gf.headers

    c.post("/wopi/files/local.docx/lock", headers={"X-WOPI-Lock": "LOCK-A"})
    gl = c.post("/wopi/files/local.docx/getlock")
    assert gl.headers.get("X-WOPI-Lock", "").strip() == "LOCK-A"
    wrong = c.post(
        "/wopi/files/local.docx/contents",
        content=b"x",
        headers={"X-WOPI-Lock": "WRONG"},
    )
    assert wrong.status_code == 409
    right = c.post(
        "/wopi/files/local.docx/contents",
        content=_build_docx("Local Doc", "locked put"),
        headers={"X-WOPI-Lock": "LOCK-A"},
    )
    assert right.status_code == 200
    assert "locked put" in _docx_text(c.get("/wopi/files/local.docx/contents").content)

    # 5. Unlock then health.
    c.post("/wopi/files/local.docx/unlock", headers={"X-WOPI-Lock": "LOCK-A"})
    gl2 = c.post("/wopi/files/local.docx/getlock")
    assert gl2.headers.get("X-WOPI-Lock", "").strip() == ""
    health = c.get("/health").json()
    assert health["status"] == "ok"
    assert any(d["id"] == "local.docx" for d in c.get("/api/documents").json())
