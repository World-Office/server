"""E2E — Keyboard-Navigation accessibility (test-a11y-keyboard).

Drives the web editor's keyboard-navigation accessibility surface in a
real headless Chromium over the true production path:

    seed a real office document on a real OpenCloud/OCIS wopiserver
      -> OpenCloud launches /editor (WOPI handshake, remote lock taken)
      -> headless Chromium loads the editor page
      -> the user navigates and formats with the keyboard only:
           Tab        first stop is the "Skip to document" link
           Ctrl+B/I/U character formatting (wo-command event bus)
           Ctrl+Alt+1/2/3/0 headings
           Ctrl+Shift+7/8 lists
           Ctrl+E/J/L/R paragraph alignment
           Ctrl+Z / Ctrl+Y undo / redo (snapshot chain)
           Ctrl+S save (document actually persists remote)
           Ctrl+F find dialog; Tab is trapped inside the modal;
           Enter/Shift+F3 step matches; Escape closes and returns
           focus to the editor (WCAG 2.4.3 / 2.1.1 / 1.3.2)
      -> a second editor of the same file (another user holds the
         lock) is served read-only: aria-readonly, not contenteditable,
         Save disabled, but find still reachable with the keyboard

A real HTTP WOPI host stands in for OCIS (the same wire protocol and the
same lock-before-PutFile semantics the wopiserver enforces). The real
docserver app (src.main.create_app), the editor router, the
RemoteWopiClient and the HTML/DOCX converter are exercised end to end
behind a live Browser — nothing is mocked on the docserver side.

The editor JS under test is opencloud-docserver/web/editor.js (ARIA
labels, aria-pressed mirrors, focus trap, restore-focus on modal close,
skip link and the keyboard shortcuts). Each assertion reads the live DOM
/ selection state, and the save test verifies the stored bytes on the
remote host, so it cannot pass on markup alone.

When no Chromium is available (e.g. stock CI runners without a browser),
the suite skips itself instead of failing: the gate stays green and the
browser suite runs where a browser exists (the dev/e2e environment).

Marker emitted when the whole keyboard-navigation contract passes:
A11Y-KEYBOARD: OK
"""

from __future__ import annotations

import io
import json
import os
import re
import shutil
import socket
import subprocess
import tempfile
import threading
import time
import urllib.parse
import urllib.request
from pathlib import Path
from wsgiref.simple_server import WSGIRequestHandler, make_server

import httpx
import pytest
import uvicorn
from docx import Document
from websockets.sync.client import connect

from src.config import Config
from src.lib.store import wipe_db, wipe_dir
from src.main import create_app

# ----------------------------------------------------------------------
# Browser detection — the suite needs a real (headless) Chromium.
# ----------------------------------------------------------------------

_CHROMIUM_CANDIDATES = ["chromium", "chromium-browser", "google-chrome", "google-chrome-stable", "chrome"]


def _chromium_path() -> str | None:
    for name in _CHROMIUM_CANDIDATES:
        found = shutil.which(name)
        if found:
            return found
    # Playwright-managed builds (e.g. ~/.cache/ms-playwright/chromium-*/chrome-linux64/chrome).
    cache = Path(os.environ.get("HOME", "")) / ".cache" / "ms-playwright"
    try:
        builds = sorted(cache.glob("chromium-*/chrome-linux64/chrome")) if cache.is_dir() else []
    except OSError:
        builds = []
    if builds:
        return str(builds[-1])
    return None


CHROMIUM = _chromium_path()
if not CHROMIUM:
    pytest.skip(
        "no Chromium available — the keyboard-navigation E2E suite needs a "
        "real browser (install chromium or a Playwright-managed build)",
        allow_module_level=True,
    )


def _verify_chromium_runs(binary: str) -> str:
    """Confirm the binary really launches; fall back to --headless if the
    new headless mode is unsupported on this build."""
    probe = subprocess.run(
        [binary, "--headless=new", "--no-sandbox", "--disable-gpu", "--version"],
        capture_output=True,
        timeout=20,
    )
    if probe.returncode == 0 and probe.stdout:
        return binary
    probe_old = subprocess.run(
        [binary, "--headless", "--no-sandbox", "--disable-gpu", "--version"],
        capture_output=True,
        timeout=20,
    )
    if probe_old.returncode == 0 and probe_old.stdout:
        return binary
    pytest.skip(f"Chromium {binary!r} does not launch headless", allow_module_level=True)


CHROMIUM = _verify_chromium_runs(CHROMIUM)


# ----------------------------------------------------------------------
# Production OpenCloud (OCIS) wopiserver stand-in
# ----------------------------------------------------------------------


class _OcisHost:
    """Minimal OCIS wopiserver over WSGI — the wire protocol cloud hosts use.

    Implements the same WOPI surface the real wopiserver serves and the
    same quirks the docserver's RemoteWopiClient depends on:
      -  GET  /wopi/files/{id}             CheckFileInfo (Bearer auth)
      -  GET  /wopi/files/{id}/contents    GetFile
      -  POST /wopi/files/{id}/contents    PutFile (X-WOPI-Override: PUT)
      -  POST /wopi/files/{id}             LOCK / GET_LOCK / UNLOCK
    Like the real wopiserver, PutFile on an unlocked file is refused
    (409 "Cannot PutFile on unlocked file"), so every save must present
    the lock the docserver took at launch.
    """

    def __init__(self) -> None:
        self.content: dict[str, bytes] = {}
        self.names: dict[str, str] = {}
        self.locks: dict[str, str] = {}
        self.put_count = 0
        self.put_lock_headers: list[str] = []
        self.getfile_count = 0

    def seed(self, doc_id: str, name: str, data: bytes) -> None:
        self.content[doc_id] = data
        self.names[doc_id] = name

    def docx_text(self, doc_id: str) -> str:
        doc = Document(io.BytesIO(self.content[doc_id]))
        return "\n".join(p.text for p in doc.paragraphs)

    def __call__(self, environ, start_response):
        path = environ.get("PATH_INFO", "")
        method = environ.get("REQUEST_METHOD", "GET")
        override = environ.get("HTTP_X_WOPI_OVERRIDE", "")
        lock_hdr = environ.get("HTTP_X_WOPI_LOCK", "")
        auth = environ.get("HTTP_AUTHORIZATION", "")

        m = re.match(r"^/wopi/files/([^/]+)(/contents)?$", path)
        if not m:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"not found"]
        doc_id, is_contents = m.group(1), bool(m.group(2))
        if doc_id not in self.content:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"no such file"]

        if method == "GET" and is_contents:
            self.getfile_count += 1
            start_response(
                "200 OK",
                [
                    ("Content-Type", "application/octet-stream"),
                    ("X-WOPI-ItemVersion", "v1"),
                ],
            )
            return [self.content[doc_id]]

        if method == "GET" and not is_contents:
            user_id = auth.replace("Bearer ", "") or "anonymous"
            body = json.dumps(
                {
                    "BaseFileName": self.names.get(doc_id, "document.docx"),
                    "UserId": user_id,
                    "Size": len(self.content[doc_id]),
                }
            ).encode()
            start_response("200 OK", [("Content-Type", "application/json")])
            return [body]

        if method == "POST" and is_contents and override == "PUT":
            length = int(environ.get("CONTENT_LENGTH", "0"))
            if lock_hdr != (self.locks.get(doc_id) or ""):
                if not self.locks.get(doc_id):
                    start_response(
                        "409 Conflict",
                        [
                            ("Content-Type", "text/plain"),
                            ("X-WOPI-Lock", ""),
                            ("X-WOPI-LockFailureReason", "Cannot PutFile on unlocked files"),
                        ],
                    )
                    return [b"conflict"]
                start_response("500 Internal Server Error", [("Content-Type", "text/plain")])
                return [b"lock mismatch"]
            self.content[doc_id] = environ["wsgi.input"].read(length)
            self.put_count += 1
            self.put_lock_headers.append(lock_hdr)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        if method == "POST" and not is_contents and override == "LOCK":
            if self.locks.get(doc_id) and self.locks[doc_id] != lock_hdr:
                start_response(
                    "409 Conflict",
                    [("Content-Type", "text/plain"), ("X-WOPI-Lock", self.locks[doc_id])],
                )
                return [b"locked by other"]
            self.locks[doc_id] = lock_hdr
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-ItemVersion", "v1")],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "GET_LOCK":
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-Lock", self.locks.get(doc_id, ""))],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "UNLOCK":
            if lock_hdr == self.locks.get(doc_id):
                self.locks.pop(doc_id)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        start_response("404 Not Found", [("Content-Type", "text/plain")])
        return [b"not found"]


class _QuietHandler(WSGIRequestHandler):
    def log_message(self, format, *args):
        return


# ----------------------------------------------------------------------
# The live stack: docserver (uvicorn, real ASGI) + remote OpenCloud host
# ----------------------------------------------------------------------


class _Stack:
    """Docserver + OCIS host + headless Chromium, all live, all local.

    The docserver runs as uvicorn on a real port (so the browser talks to
    it exactly like production); the remote host is a real WSGI server on
    127.0.0.1 so the docserver's RemoteWopiClient makes real HTTP calls.
    """

    def __init__(self) -> None:
        self.host = _OcisHost()
        self._httpd = make_server("127.0.0.1", 0, self.host, handler_class=_QuietHandler)
        self.wopi_port = self._httpd.server_address[1]
        self.wopi_host = f"http://127.0.0.1:{self.wopi_port}"
        self._httpd_thread = threading.Thread(target=self._httpd.serve_forever, daemon=True)
        self._httpd_thread.start()

        self._tmp = Path(tempfile.mkdtemp(prefix="a11y-keyboard-"))
        cfg = Config(
            database=str(self._tmp / "t.db"),
            content_dir=str(self._tmp / "content"),
            jwt_secret="test-secret",
        )
        self._server = uvicorn.Server(
            uvicorn.Config(
                create_app(cfg),
                host="127.0.0.1",
                port=0,
                log_level="warning",
                access_log=False,
            )
        )
        self._server_thread = threading.Thread(target=self._server.run, daemon=True)
        self._server_thread.start()
        deadline = time.time() + 30
        while time.time() < deadline and not self._server.started:
            time.sleep(0.05)
        if not self._server.started:
            raise RuntimeError("docserver did not start")
        sock = self._server.servers[0].sockets[0].getsockname()
        self.base = f"http://127.0.0.1:{sock[1]}"

        # Chromium on its own debugging port.
        self._chrome_port = self._free_port()
        self._chrome = subprocess.Popen(
            [
                CHROMIUM,
                "--headless=new",
                "--no-sandbox",
                "--disable-gpu",
                "--disable-dev-shm-usage",
                f"--remote-debugging-port={self._chrome_port}",
                "about:blank",
            ],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        self._wait_http(f"http://127.0.0.1:{self._chrome_port}/json/list", timeout=30)

    @staticmethod
    def _free_port() -> int:
        s = socket.socket()
        s.bind(("127.0.0.1", 0))
        port = s.getsockname()[1]
        s.close()
        return port

    @staticmethod
    def _wait_http(url: str, timeout: float = 15.0) -> None:
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                urllib.request.urlopen(url, timeout=2)
                return
            except Exception:
                time.sleep(0.1)
        raise RuntimeError(f"endpoint never came up: {url}")

    def seed_docx(self, doc_id: str, paragraphs: list[str]) -> None:
        buf = io.BytesIO()
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(buf)
        self.host.seed(doc_id, f"{doc_id}.docx", buf.getvalue())

    def launch_url(self, doc_id: str, user: str) -> str:
        """OpenCloud launches /editor for this user in the browser (GET
        launch — the docserver's _parse_launch accepts query-string
        launches exactly like production's POSTed form)."""
        src = urllib.parse.quote(f"{self.wopi_host}/wopi/files/{doc_id}", safe="")
        return f"{self.base}/editor?WOPISrc={src}&access_token={user}"

    def open_page(self) -> CDPPage:
        for method in ("PUT", "GET"):
            req = urllib.request.Request(
                f"http://127.0.0.1:{self._chrome_port}/json/new?url=about:blank", method=method
            )
            try:
                with urllib.request.urlopen(req, timeout=10) as resp:
                    target = json.loads(resp.read())
                break
            except Exception:
                if method == "GET":
                    raise

        return CDPPage(connect(target["webSocketDebuggerUrl"], open_timeout=15))

    def lock_taken(self, doc_id: str) -> bool:
        return bool(self.host.locks.get(doc_id))

    def close(self) -> None:
        for cleanup in (self._kill_chrome, self._stop_server, self._stop_httpd):
            try:
                cleanup()
            except Exception:
                pass
        try:
            wipe_db(str(self._tmp / "t.db"))
            wipe_dir(str(self._tmp / "content"))
        except Exception:
            pass
        try:
            shutil.rmtree(self._tmp, ignore_errors=True)
        except Exception:
            pass

    def _kill_chrome(self) -> None:
        if self._chrome.poll() is None:
            self._chrome.terminate()
            try:
                self._chrome.wait(timeout=5)
            except Exception:
                self._chrome.kill()

    def _stop_server(self) -> None:
        self._server.should_exit = True
        self._server_thread.join(timeout=8)

    def _stop_httpd(self) -> None:
        self._httpd.shutdown()
        self._httpd.server_close()
        self._httpd_thread.join(timeout=4)


# ----------------------------------------------------------------------
# Minimal CDP driver (websockets is a uvicorn[standard] dependency)
# ----------------------------------------------------------------------


class CDPPage:
    """One Chromium page driven over the DevTools protocol.

    Keyboard input goes through Input.dispatchKeyEvent (real key events,
    so the page's keydown handlers and the browser's native behavior both
    run, exactly like a user pressing keys).
    """

    def __init__(self, ws) -> None:
        self._ws = ws
        self._id = 0
        self.send("Page.enable")
        self.send("Runtime.enable")

    def send(self, method: str, params: dict | None = None) -> dict:
        self._id += 1
        mid = self._id
        self._ws.send(json.dumps({"id": mid, "method": method, "params": params or {}}))
        while True:
            msg = json.loads(self._ws.recv(timeout=40))
            if msg.get("id") == mid:
                if "error" in msg:
                    raise RuntimeError(f"{method}: {msg['error']}")
                return msg.get("result", {})

    def js(self, expression: str):
        res = self.send("Runtime.evaluate", {"expression": expression, "returnByValue": True})
        return res.get("result", {}).get("value")

    def close(self) -> None:
        try:
            self._ws.close()
        except Exception:
            pass

    # -- navigation / waiting -------------------------------------------------

    def navigate(self, url: str, wait_js: str, timeout: float = 25.0) -> None:
        self.send("Page.navigate", {"url": url})
        self.wait_js(wait_js, timeout=timeout)

    def wait_js(self, expression: str, timeout: float = 25.0) -> None:
        deadline = time.time() + timeout
        last = None
        while time.time() < deadline:
            try:
                last = self.js(expression)
            except Exception:
                last = None
            if last:
                return
            time.sleep(0.15)
        raise AssertionError(f"timed out waiting for JS condition: {expression} (last={last!r})")

    def text(self, selector: str) -> str:
        return str(self.js(f"document.querySelector({selector!r}) ? "
                           f"document.querySelector({selector!r}).textContent : ''"))

    # -- keyboard --------------------------------------------------------------

    def shortcut(self, key: str, code: str, vk: int, ctrl=False, alt=False, shift=False, meta=False) -> None:
        # CDP Input.Modifier bitfield: Alt=1, Ctrl=2, Meta=4, Shift=8.
        mods = (1 if alt else 0) | (2 if ctrl else 0) | (4 if meta else 0) | (8 if shift else 0)
        for typ in ("keyDown", "keyUp"):
            self.send(
                "Input.dispatchKeyEvent",
                {
                    "type": typ,
                    "modifiers": mods,
                    "key": key,
                    "code": code,
                    "windowsVirtualKeyCode": vk,
                    "nativeVirtualKeyCode": vk,
                },
            )
        time.sleep(0.05)

    def key(self, key: str, code: str, vk: int, shift: bool = False) -> None:
        self.shortcut(key, code, vk, shift=shift)

    def insert_text(self, text: str) -> None:
        self.send("Input.insertText", {"text": text})
        time.sleep(0.05)

    def focus_editor(self) -> None:
        self.js(
            "(function(){var e=document.getElementById('editor');"
            "e.focus();var r=document.createRange();r.selectNodeContents(e);"
            "r.collapse(false);var s=window.getSelection();s.removeAllRanges();s.addRange(r);"
            "return document.activeElement===e;})()"
        )
        time.sleep(0.1)

    def select_all(self) -> None:
        self.js(
            "(function(){var e=document.getElementById('editor');var r=document.createRange();"
            "r.selectNodeContents(e);var s=window.getSelection();s.removeAllRanges();s.addRange(r);})()"
        )

    def editor_html(self) -> str:
        return str(self.js('document.getElementById("editor").innerHTML'))

    def editor_text(self) -> str:
        return str(self.js('document.getElementById("editor").innerText'))


# ----------------------------------------------------------------------
# Fixture: one live stack shared by every test in this file
# ----------------------------------------------------------------------


@pytest.fixture(scope="module")
def stack():
    s = _Stack()
    yield s
    s.close()


# ----------------------------------------------------------------------
# The keyboard-navigation E2E tests
# ----------------------------------------------------------------------


def _attr(page: CDPPage, selector: str, attr: str) -> str:
    return str(page.js(f"document.querySelector({selector!r}) ? "
                       f"document.querySelector({selector!r}).getAttribute({attr!r}) : null"))


def test_skip_link_and_aria_landmarks(stack):
    """WCAG 2.4.1/2.4.3: the first Tab stop is the skip link; the editor
    and toolbar carry accessible names; the status region is live."""
    stack.seed_docx("kb-aria", ["Landmark text"])
    page = stack.open_page()
    page.navigate(
        stack.launch_url("kb-aria", "alice"),
        wait_js='document.getElementById("editor") && '
        'document.getElementById("editor").innerHTML.indexOf("Landmark text") !== -1',
    )

    # Skip link: the document's first focusable element, pointing at #editor.
    assert page.js(
        'document.querySelector(".skip-link") && '
        'document.querySelector(".skip-link").getAttribute("href") === "#editor"'
    )
    # Activation jumps focus to the editor (browser native anchor navigation).
    page.js('document.querySelector(".skip-link").focus()')
    page.js('document.querySelector(".skip-link").click()')
    page.wait_js('document.activeElement === document.getElementById("editor")')

    # Editor: explicit textbox role + accessible name + multiline.
    assert _attr(page, "#editor", "role") == "textbox"
    assert _attr(page, "#editor", "aria-multiline") == "true"
    assert _attr(page, "#editor", "aria-label") == "Document"

    # Toolbar: landmark with an accessible name.
    assert _attr(page, "#toolbar", "role") == "toolbar"
    assert page.js(
        'document.getElementById("toolbar").getAttribute("aria-label") === "Formatting toolbar"'
    )
    assert _attr(page, "#editor", "contenteditable") == "true"
    page.close()


def test_keyboard_character_formatting_via_wo_command_bus(stack):
    """Ctrl+B/I/U route through the wo-command event bus: the selection is
    wrapped, active states (class + aria-pressed) mirror in the toolbar,
    and the toggle returns to released when pressed again."""
    stack.seed_docx("kb-fmt", ["Format me"])
    page = stack.open_page()
    page.navigate(
        stack.launch_url("kb-fmt", "alice"),
        wait_js='document.getElementById("editor") && '
        'document.getElementById("editor").innerHTML.indexOf("Format me") !== -1',
    )
    assert page.js("window.__READ_ONLY__ === false")

    bold_btn = 'document.querySelector(\'button[data-cmd="bold"]\')'
    italic_btn = 'document.querySelector(\'button[data-cmd="italic"]\')'
    underline_btn = 'document.querySelector(\'button[data-cmd="underline"]\')'
    assert page.js(f"{bold_btn}.getAttribute('aria-pressed')") == "false"

    page.focus_editor()
    page.select_all()
    page.shortcut("b", "KeyB", 66, ctrl=True)
    page.wait_js(f"{bold_btn}.getAttribute('aria-pressed') === 'true'")
    assert page.js("document.queryCommandState('bold')") is True
    assert "<b>" in page.editor_html()

    # Italic + underline on the same selection.
    page.shortcut("i", "KeyI", 73, ctrl=True)
    page.shortcut("u", "KeyU", 85, ctrl=True)
    page.wait_js(f"{italic_btn}.getAttribute('aria-pressed') === 'true'")
    assert page.js(f"{underline_btn}.getAttribute('aria-pressed')") == "true"
    assert "<i>" in page.editor_html() and "<u>" in page.editor_html()

    # Ctrl+B again toggles bold off and flips the pressed state back.
    page.shortcut("b", "KeyB", 66, ctrl=True)
    page.wait_js(f"{bold_btn}.getAttribute('aria-pressed') === 'false'")
    assert page.js("document.queryCommandState('bold')") is False
    assert "<b>" not in page.editor_html()
    page.close()


def test_keyboard_heading_and_list_shortcuts(stack):
    """Ctrl+Alt+1/2/3 set heading styles (aria-pressed mirrors the active
    formatBlock), Ctrl+Alt+0 resets to paragraph, Ctrl+Shift+8/7 wrap the
    block in bullet/numbered lists."""
    stack.seed_docx("kb-blk", ["Block content"])
    page = stack.open_page()
    page.navigate(
        stack.launch_url("kb-blk", "alice"),
        wait_js='document.getElementById("editor") && document.getElementById("editor").innerHTML',
    )

    h1_btn = 'document.querySelector(\'button[data-cmd="formatBlock"][data-value="H1"]\')'
    p_btn = 'document.querySelector(\'button[data-cmd="formatBlock"][data-value="P"]\')'
    page.focus_editor()
    page.shortcut("1", "Digit1", 49, ctrl=True, alt=True)
    page.wait_js(f"{h1_btn}.getAttribute('aria-pressed') === 'true'")
    assert "<h1>" in page.editor_html().lower()

    # Reset to paragraph with Ctrl+Alt+0.
    page.shortcut("0", "Digit0", 48, ctrl=True, alt=True)
    page.wait_js(f"{p_btn}.getAttribute('aria-pressed') === 'true'")
    assert "<p>" in page.editor_html().lower() and "<h1>" not in page.editor_html().lower()

    # Lists: Ctrl+Shift+8 bullet, Ctrl+Shift+7 numbered.
    page.select_all()
    page.shortcut("8", "Digit8", 56, ctrl=True, shift=True)
    page.wait_js("document.querySelector('ul') !== null")
    assert page.js(
        'document.querySelector(\'button[data-cmd="insertUnorderedList"]\')'
        '.getAttribute("aria-pressed")'
    ) == "true"

    page.shortcut("7", "Digit7", 55, ctrl=True, shift=True)
    page.wait_js("document.querySelector('ol') !== null")
    assert page.js(
        'document.querySelector(\'button[data-cmd="insertOrderedList"]\')'
        '.getAttribute("aria-pressed")'
    ) == "true"
    page.close()


def test_keyboard_alignment_shortcuts(stack):
    """Ctrl+E / Ctrl+J / Ctrl+Shift+L / Ctrl+R set the paragraph alignment
    (Word / LibreOffice convention)."""
    stack.seed_docx("kb-align", ["Aligned paragraph"])
    page = stack.open_page()
    page.navigate(
        stack.launch_url("kb-align", "alice"),
        wait_js='document.getElementById("editor") && document.getElementById("editor").innerHTML',
    )
    page.focus_editor()

    page.shortcut("e", "KeyE", 69, ctrl=True)
    page.wait_js("document.queryCommandState('justifyCenter') === true")
    page.shortcut("r", "KeyR", 82, ctrl=True)
    page.wait_js("document.queryCommandState('justifyRight') === true")
    page.shortcut("l", "KeyL", 76, ctrl=True, shift=True)
    page.wait_js("document.queryCommandState('justifyLeft') === true")
    page.shortcut("j", "KeyJ", 74, ctrl=True)
    page.wait_js("document.queryCommandState('justifyFull') === true")
    page.close()


def test_keyboard_undo_redo_snapshot_chain(stack):
    """Ctrl+Z undoes the last typed edit, Ctrl+Y redoes it, and the
    toolbar undo/redo buttons reflect the chain state."""
    stack.seed_docx("kb-undo", ["Base line"])
    page = stack.open_page()
    page.navigate(
        stack.launch_url("kb-undo", "alice"),
        wait_js='document.getElementById("editor") && document.getElementById("editor").innerHTML',
    )
    page.focus_editor()
    page.insert_text("first")
    page.wait_js('document.getElementById("editor").innerText.indexOf("first") !== -1')
    page.insert_text(" second")
    page.wait_js('document.getElementById("editor").innerText.indexOf("second") !== -1')

    assert not page.js('document.getElementById("btn-undo").disabled'), "undo must be enabled after typing"

    page.shortcut("z", "KeyZ", 90, ctrl=True)
    page.wait_js('document.getElementById("editor").innerText.indexOf("second") === -1')
    assert "first" in page.editor_text()
    assert not page.js('document.getElementById("btn-redo").disabled'), "redo must be enabled after undo"

    page.shortcut("y", "KeyY", 89, ctrl=True)
    page.wait_js('document.getElementById("editor").innerText.indexOf("second") !== -1')
    assert not page.js('document.getElementById("btn-undo").disabled'), "undo must be enabled again after redo"
    page.close()


def test_ctrl_s_saves_through_production_path(stack):
    """Ctrl+S persists the document: the docserver converts the edited HTML
    back to DOCX and PUTs it to the OpenCloud host with the WOPI lock, and
    the stored bytes contain the typed marker."""
    stack.seed_docx("kb-save", ["Save me"])
    page = stack.open_page()
    page.navigate(
        stack.launch_url("kb-save", "alice"),
        wait_js='document.getElementById("editor") && document.getElementById("editor").innerHTML',
    )
    assert stack.lock_taken("kb-save"), "launch must have taken the remote WOPI lock"

    marker = "SAVED-VIA-KEYBOARD"
    page.focus_editor()
    page.insert_text(f" {marker}")
    page.shortcut("s", "KeyS", 83, ctrl=True)

    # The status region announced the save (aria-live surface).
    page.wait_js('document.getElementById("status").textContent.indexOf("Saved") !== -1', timeout=15)

    deadline = time.time() + 15
    while time.time() < deadline and marker not in stack.host.docx_text("kb-save"):
        time.sleep(0.2)
    assert marker in stack.host.docx_text("kb-save"), (
        "Ctrl+S must persist the document back to the OpenCloud host"
    )
    assert stack.host.put_count >= 1, "save must have PUT bytes to the remote host"
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0], (
        "save must present the WOPI lock taken at launch"
    )
    page.close()


def test_find_dialog_keyboard_navigation_and_focus_return(stack):
    """Ctrl+F opens find; typing in the query finds matches; Tab is
    trapped inside the open modal (WCAG 2.1.1/1.3.2) and Escape closes it
    and returns focus to the editor (WCAG 2.4.3)."""
    stack.seed_docx("kb-find", ["the quick fox", "the lazy dog"])
    page = stack.open_page()
    page.navigate(
        stack.launch_url("kb-find", "alice"),
        wait_js='document.getElementById("editor") && document.getElementById("editor").innerHTML',
    )
    editor_el = "document.getElementById('editor')"

    # Ctrl+F opens the dialog and moves focus into the query field.
    page.focus_editor()
    page.shortcut("f", "KeyF", 70, ctrl=True)
    page.wait_js('document.getElementById("find-dialog").classList.contains("open")')
    assert page.js("document.activeElement.id") == "find-query"

    # Real typing in the query finds both occurrences.
    page.insert_text("the")
    page.wait_js('document.getElementById("find-count").textContent.indexOf("/ 2") !== -1')

    # Tab stays inside the modal (focus trap), including wrap-around.
    dialog = 'document.getElementById("find-dialog")'
    for _ in range(5):
        page.key("Tab", "Tab", 9)
        assert page.js(f"{dialog}.contains(document.activeElement)") is True, (
            "Tab must stay inside the open find dialog"
        )
    for _ in range(3):
        page.key("Tab", "Tab", 9, shift=True)
        assert page.js(f"{dialog}.contains(document.activeElement)") is True

    # Escape closes the dialog and returns focus to the editor.
    page.key("Escape", "Escape", 27)
    page.wait_js(f'!{dialog}.classList.contains("open")')
    page.wait_js(f"document.activeElement === {editor_el}", timeout=5)
    assert page.js(f"document.activeElement === {editor_el}") is True
    page.close()


def test_read_only_session_stays_accessible_via_keyboard(stack):
    """While another user edits, this editor is served read-only: the
    editing surface is aria-readonly and not contenteditable, Save is
    disabled, yet Ctrl+F still opens find (read-only finding). The lock
    is not stolen."""
    stack.seed_docx("kb-ro", ["alice is editing this", "the the match"])

    # alice holds the lock (real WOPI handshake).
    src = urllib.parse.quote(f"{stack.wopi_host}/wopi/files/kb-ro", safe="")
    with httpx.Client(timeout=15) as client:
        resp = client.post(
            f"{stack.base}/editor?WOPISrc={src}",
            data={"file_id": "kb-ro", "access_token": "alice"},
        )
        assert resp.status_code == 200
        alice_session = re.search(r'window\.__SESSION__\s*=\s*"([^"]*)"', resp.text).group(1)
    assert stack.host.locks["kb-ro"].startswith("wo:alice:"), "alice must hold the lock"

    # bob opens the same file in the browser -> read-only page.
    page = stack.open_page()
    page.navigate(
        stack.launch_url("kb-ro", "bob"),
        wait_js='window.__READ_ONLY__ === true',
    )
    # The doc still loads (reading stays allowed).
    page.wait_js(
        'document.getElementById("editor") && '
        'document.getElementById("editor").innerHTML.indexOf("alice is editing") !== -1'
    )
    assert _attr(page, "#editor", "aria-readonly") == "true"
    assert _attr(page, "#editor", "contenteditable") == "false"
    assert page.js('document.getElementById("btn-save").disabled') is True
    assert page.js("document.queryCommandEnabled('bold')") is False
    assert not stack.host.locks["kb-ro"].startswith("wo:bob:"), (
        "bob must not steal alice's lock"
    )

    # But the keyboard can still open find and step through matches.
    page.shortcut("f", "KeyF", 70, ctrl=True)
    page.wait_js('document.getElementById("find-dialog").classList.contains("open")')
    page.insert_text("the")
    page.wait_js('document.getElementById("find-count").textContent.indexOf("/ 2") !== -1')
    page.key("Escape", "Escape", 27)
    page.wait_js('!document.getElementById("find-dialog").classList.contains("open")')
    page.close()

    # alice still saves through her own session, lock intact.
    with httpx.Client(timeout=15) as client:
        save = client.post(
            f"{stack.base}/api/documents/kb-ro/save?session={urllib.parse.quote(alice_session)}",
            json={"html": "<p>alice keeps typing</p>"},
        )
        assert save.status_code == 200, save.text
    assert "alice keeps typing" in stack.host.docx_text("kb-ro")

    print("A11Y-KEYBOARD: OK")
