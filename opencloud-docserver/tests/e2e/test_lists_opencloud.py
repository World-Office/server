"""E2E — lists through production OpenCloud (OCIS client mode).

This suite drives the exact code path the docserver uses in production
(deployed against real OpenCloud / OCIS, e.g. cloud.graphwiz.ai), focused
on the list editing features the editor UI offers (bullet + numbered lists):

    seed an office document (DOCX or ODT) with bullet and numbered lists
      -> OpenCloud launches /editor with access_token + WOPISrc  (WOPI handshake)
      -> GET  /api/documents/{id}/html                           (list -> HTML)
      -> edit the HTML in the editor (type new list items / a marker)
      -> POST /api/documents/{id}/save                           (HTML -> document)
      -> docserver PUTs the bytes back to the remote host         (PutFile)
      -> GET  /api/documents/{id}/html again                     (reload)
      -> re-read the stored document and verify its list items
         (DOCX "List Bullet"/"List Number" paragraphs, ODT content.xml
         text:list elements) still contain the marker.

A real HTTP WOPI host stands in for OCIS (same wire protocol, same
lock-before-PutFile semantics the wopiserver enforces). The real
docserver app, editor router, RemoteWopiClient and both converters
(DOCX + ODT) are exercised end to end — nothing is mocked on the
docserver side.

The suite asserts the list-persistence contract: list markup typed or
present in the editor survives every hop of the production path, and the
stored bytes are a real office document whose list structure contains it.

Marker emitted when the full production-path list persistence check
passes:
LIST-PERSISTENCE: OK
"""

from __future__ import annotations

import io
import json
import re
import threading
import zipfile
from contextlib import asynccontextmanager
from wsgiref.simple_server import make_server

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from odf import teletype
from odf.opendocument import OpenDocumentText, load
from odf.text import List, ListItem, ListLevelStyleNumber, ListStyle, P

from src.config import Config
from src.editor.router import router as editor_router
from src.editor.session import SessionRegistry
from src.lib.store import DocumentStore, wipe_db, wipe_dir
from src.wopi.router import router as wopi_router

# Marker the editor types into a list; a save is never treated as
# successful until the marker is found in the stored remote document.
MARKER = "LIST-PERSISTENCE"


# ----------------------------------------------------------------------
# Production OpenCloud (OCIS) wopiserver stand-in
# ----------------------------------------------------------------------

class _ProdOcisHost:
    """Minimal OCIS wopiserver over WSGI — what cloud.graphwiz.ai does.

    Implements the same WOPI surface the real wopiserver serves and the
    same quirks the docserver's RemoteWopiClient depends on:
      -  GET  /wopi/files/{id}             CheckFileInfo (Bearer auth)
      -  GET  /wopi/files/{id}/contents    GetFile
      -  POST /wopi/files/{id}/contents    PutFile (X-WOPI-Override: PUT)
      -  POST /wopi/files/{id}             LOCK / GET_LOCK / UNLOCK
    Like the real wopiserver, PutFile on an unlocked file is refused
    (409 "Cannot PutFile on unlocked file"), so every save must present
    the lock the docserver took at launch.
    """

    def __init__(self) -> None:
        self.content: dict[str, bytes] = {}
        self.names: dict[str, str] = {}
        self.locks: dict[str, str] = {}
        self.users: dict[str, str] = {}
        self.put_count = 0
        self.put_overrides: list[str] = []
        self.put_lock_headers: list[str] = []
        self.lock_events: list[str] = []

    def seed(self, doc_id: str, name: str, data: bytes, user: str = "alice") -> None:
        self.content[doc_id] = data
        self.names[doc_id] = name
        self.users[doc_id] = user

    def __call__(self, environ, start_response):
        path = environ.get("PATH_INFO", "")
        method = environ.get("REQUEST_METHOD", "GET")
        override = environ.get("HTTP_X_WOPI_OVERRIDE", "")
        lock_hdr = environ.get("HTTP_X_WOPI_LOCK", "")
        auth = environ.get("HTTP_AUTHORIZATION", "")

        m = re.match(r"^/wopi/files/([^/]+)(/contents)?$", path)
        if not m:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"not found"]
        doc_id, is_contents = m.group(1), bool(m.group(2))
        if doc_id not in self.content:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"no such file"]

        # GetFile — the docserver reads the raw document bytes with the token.
        if method == "GET" and is_contents:
            start_response(
                "200 OK",
                [
                    ("Content-Type", "application/octet-stream"),
                    ("X-WOPI-ItemVersion", "v1"),
                ],
            )
            return [self.content[doc_id]]

        # CheckFileInfo — the docserver reads BaseFileName (routes .docx vs
        # .odt to the right converter) and UserId (names the WOPI lock).
        if method == "GET" and not is_contents:
            user_id = auth.replace("Bearer ", "") or self.users.get(doc_id, "unknown")
            body = json.dumps(
                {
                    "BaseFileName": self.names.get(doc_id, "document.docx"),
                    "UserId": user_id,
                    "Size": len(self.content[doc_id]),
                }
            ).encode()
            start_response("200 OK", [("Content-Type", "application/json")])
            return [body]

        # PutFile — only accepted with the lock taken at launch.
        if method == "POST" and is_contents and override == "PUT":
            length = int(environ.get("CONTENT_LENGTH", "0"))
            if lock_hdr != (self.locks.get(doc_id) or ""):
                if not self.locks.get(doc_id):
                    start_response(
                        "409 Conflict",
                        [
                            ("Content-Type", "text/plain"),
                            ("X-WOPI-Lock", ""),
                            ("X-WOPI-LockFailureReason", "Cannot PutFile on unlocked files"),
                        ],
                    )
                    return [b"conflict"]
                start_response("500 Internal Server Error", [("Content-Type", "text/plain")])
                return [b"lock mismatch"]
            self.content[doc_id] = environ["wsgi.input"].read(length)
            self.put_count += 1
            self.put_overrides.append(override)
            self.put_lock_headers.append(lock_hdr)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        if method == "POST" and not is_contents and override == "LOCK":
            if self.locks.get(doc_id) and self.locks[doc_id] != lock_hdr:
                start_response(
                    "409 Conflict",
                    [("Content-Type", "text/plain"), ("X-WOPI-Lock", self.locks[doc_id])],
                )
                return [b"locked by other"]
            self.locks[doc_id] = lock_hdr
            self.lock_events.append(f"LOCK:{doc_id}")
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-ItemVersion", "v1")],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "GET_LOCK":
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-Lock", self.locks.get(doc_id, ""))],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "UNLOCK":
            if lock_hdr == self.locks.get(doc_id):
                self.locks.pop(doc_id)
            self.lock_events.append(f"UNLOCK:{doc_id}")
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        start_response("404 Not Found", [("Content-Type", "text/plain")])
        return [b"not found"]


# ----------------------------------------------------------------------
# The docserver wired against a live remote OpenCloud host
# ----------------------------------------------------------------------

class _E2EStack:
    """The docserver plus a remote OpenCloud host, wired like production.

    The docserver runs as a FastAPI TestClient; the remote host is a real
    WSGI server on 127.0.0.1 so the docserver's RemoteWopiClient makes real
    HTTP calls (urllib) exactly like it would against cloud.graphwiz.ai.
    """

    def __init__(self, tmp_path) -> None:
        self.host = _ProdOcisHost()
        self._httpd = make_server("127.0.0.1", 0, self.host)
        self.port = self._httpd.server_address[1]
        self.wopi_host = f"http://127.0.0.1:{self.port}"
        self._thread = threading.Thread(target=self._httpd.serve_forever, daemon=True)
        self._thread.start()

        db = str(tmp_path / "t.db")
        content = str(tmp_path / "content")
        store = DocumentStore(db, content)
        cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

        @asynccontextmanager
        async def lifespan(app: FastAPI):
            app.state.store = store
            app.state.sessions = SessionRegistry()
            app.state.config = cfg
            yield

        app = FastAPI(lifespan=lifespan)
        app.include_router(wopi_router)
        app.include_router(editor_router)
        self.client = TestClient(app)
        self.client.__enter__()  # run lifespan (app.state.*)
        self._db, self._content = db, content

    def close(self) -> None:
        try:
            self.client.__exit__(None, None, None)
        except Exception:
            pass
        self._httpd.shutdown()
        self._httpd.server_close()
        self._thread.join(timeout=2)
        wipe_db(self._db)
        wipe_dir(self._content)

    def launch_editor(self, doc_id: str, user: str = "alice") -> str:
        """OpenCloud launches /editor (POST form) for a user's file."""
        token = f"tok-{user}"
        wopi_src = f"{self.wopi_host}/wopi/files/{doc_id}"
        resp = self.client.post(
            "/editor",
            params={"WOPISrc": wopi_src},
            data={"file_id": doc_id, "access_token": token},
        )
        assert resp.status_code == 200
        assert self.host.locks.get(doc_id), "launch must have taken the WOPI lock"
        return token

    def load_html(self, doc_id: str) -> str:
        resp = self.client.get(f"/api/documents/{doc_id}/html")
        assert resp.status_code == 200, resp.text
        return resp.json()["html"]

    def save_html(self, doc_id: str, html: str) -> None:
        resp = self.client.post(f"/api/documents/{doc_id}/save", json={"html": html})
        assert resp.status_code == 200, resp.text
        assert resp.json().get("ok") is True

    def unlock(self, doc_id: str) -> None:
        resp = self.client.post(f"/api/documents/{doc_id}/unlock")
        assert resp.status_code == 200, resp.text

    def remote_doc(self, doc_id: str) -> bytes:
        """Raw bytes the OpenCloud host stored after the docserver's PutFile."""
        return self.host.content[doc_id]


@pytest.fixture
def stack(tmp_path):
    s = _E2EStack(tmp_path)
    yield s
    s.close()


# ----------------------------------------------------------------------
# List document fixtures + stored-bytes inspection
# ----------------------------------------------------------------------

def _build_docx_lists() -> bytes:
    """A DOCX with a body line, a bullet list and a numbered list."""
    doc = Document()
    doc.add_paragraph("Stoic list document")
    doc.add_paragraph("apples", style="List Bullet")
    doc.add_paragraph("pears", style="List Bullet")
    doc.add_paragraph("first step", style="List Number")
    doc.add_paragraph("second step", style="List Number")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_plain(body: str = "Plain stoic document") -> bytes:
    doc = Document()
    doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_odt_lists() -> bytes:
    """An ODT with a body line, a bullet list and a numbered list."""
    doc = OpenDocumentText()
    doc.text.addElement(P(text="Stoic list document"))

    bullets = List()
    for item_text in ("apples", "pears"):
        li = ListItem()
        li.addElement(P(text=item_text))
        bullets.addElement(li)
    doc.text.addElement(bullets)

    numbered_style = ListStyle(name="WO_Num")
    numbered_style.addElement(ListLevelStyleNumber(level=1, numformat="1"))
    doc.styles.addElement(numbered_style)
    numbered = List(stylename="WO_Num")
    for item_text in ("first step", "second step"):
        li = ListItem()
        li.addElement(P(text=item_text))
        numbered.addElement(li)
    doc.text.addElement(numbered)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_odt_nested_list() -> bytes:
    """An ODT with a bullet list whose first item nests another list."""
    doc = OpenDocumentText()
    outer = List()
    li1 = ListItem()
    li1.addElement(P(text="item 1"))
    inner = List()
    li_n = ListItem()
    li_n.addElement(P(text="nested a"))
    inner.addElement(li_n)
    li1.addElement(inner)
    outer.addElement(li1)
    li2 = ListItem()
    li2.addElement(P(text="item 2"))
    outer.addElement(li2)
    doc.text.addElement(outer)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_list_paragraphs(data: bytes) -> list[tuple[str, str]]:
    """(style_name, text) for every list-styled paragraph in a DOCX."""
    doc = Document(io.BytesIO(data))
    return [
        (p.style.name, p.text)
        for p in doc.paragraphs
        if "list" in (p.style.name or "").lower()
    ]


def _odt_content_xml_has(data: bytes, needle: str) -> bool:
    """True when the ODT's content.xml (the persisted body) contains needle."""
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        content_xml = zf.read("content.xml").decode("utf-8", "replace")
    return needle in content_xml


def _odt_text(data: bytes) -> str:
    doc = load(io.BytesIO(data))
    return teletype.extractText(doc.text)


def _odt_list_count(data: bytes) -> int:
    doc = load(io.BytesIO(data))
    return len(doc.text.getElementsByType(List))


# ----------------------------------------------------------------------
# The E2E tests
# ----------------------------------------------------------------------

def test_docx_bullet_and_numbered_lists_persist_through_production_opencloud(stack):
    """The full production roundtrip for a DOCX seeded with lists:
    OpenCloud launch -> HTML with <ul>/<ol> -> edit (type new list items)
    -> save -> PutFile -> reload. Emits the LIST-PERSISTENCE marker only
    when the DOCX stored back on the OpenCloud host still carries the new
    items as real "List Bullet"/"List Number" paragraphs."""
    stack.host.seed("lists-docx", "shopping.docx", _build_docx_lists())
    stack.launch_editor("lists-docx", user="alice")

    # OpenCloud serves the lists as HTML for the editor.
    html = stack.load_html("lists-docx")
    assert "<ul>" in html, html
    assert "<ol>" in html, html
    assert "<li>apples</li>" in html
    assert "<li>pears</li>" in html
    assert "<li>first step</li>" in html
    assert "<li>second step</li>" in html
    assert "Stoic list document" in html

    # The user types two more list items into the editor.
    edited = html + f"\n<ul><li>{MARKER}-bulleted</li></ul>\n<ol><li>{MARKER}-numbered</li></ol>"
    stack.save_html("lists-docx", edited)

    # The docserver must have PUT the DOCX back to the OpenCloud host,
    # presenting the WOPI lock it took at launch (409 otherwise).
    assert stack.host.put_count == 1, "edited doc must be PUT back to OpenCloud"
    assert "PUT" in stack.host.put_overrides
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0], (
        "save must present the WOPI lock taken at launch"
    )

    # Reload from the remote host: the typed list items are still there.
    reloaded = stack.load_html("lists-docx")
    assert MARKER in reloaded
    assert f"{MARKER}-bulleted" in reloaded
    assert f"{MARKER}-numbered" in reloaded
    assert "<ul>" in reloaded and "<ol>" in reloaded

    # The stored bytes are a real DOCX whose list paragraphs carry the
    # marker: three bullets and three numbered items survive.
    list_paras = _docx_list_paragraphs(stack.remote_doc("lists-docx"))
    bullets = [t for s, t in list_paras if "Bullet" in s]
    numbers = [t for s, t in list_paras if "Number" in s]
    assert "apples" in bullets and "pears" in bullets
    assert f"{MARKER}-bulleted" in bullets
    assert "first step" in numbers and "second step" in numbers
    assert f"{MARKER}-numbered" in numbers

    print("LIST-PERSISTENCE: OK")


def test_odt_bullet_and_numbered_lists_persist_through_production_opencloud(stack):
    """The same production roundtrip for an ODT seeded with lists: the
    marker lands in the stored content.xml next to the list items."""
    stack.host.seed("lists-odt", "shopping.odt", _build_odt_lists())
    stack.launch_editor("lists-odt", user="alice")

    html = stack.load_html("lists-odt")
    assert "<ul>" in html and "<ol>" in html, html
    assert "<li>apples</li>" in html
    assert "<li>first step</li>" in html

    edited = html + f"\n<ul><li>{MARKER}-bullet</li></ul>\n<ol><li>{MARKER}-num</li></ol>"
    stack.save_html("lists-odt", edited)
    assert stack.host.put_count == 1
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0]

    reloaded = stack.load_html("lists-odt")
    assert f"{MARKER}-bullet" in reloaded
    assert f"{MARKER}-num" in reloaded

    stored = stack.remote_doc("lists-odt")
    assert _odt_content_xml_has(stored, MARKER), "content.xml must contain the marker"
    text = _odt_text(stored)
    for needle in ("apples", "pears", "first step", "second step",
                   f"{MARKER}-bullet", f"{MARKER}-num"):
        assert needle in text, f"stored ODT lost {needle!r}"
    assert _odt_list_count(stored) >= 4, "stored ODT must keep 4 list blocks"


def test_lists_typed_in_editor_persist_as_real_list_paragraphs(stack):
    """Starting from a plain document, list markup added in the editor is
    saved as real (style-name) list paragraphs, not plain text."""
    stack.host.seed("lists-typed", "typed.docx", _build_docx_plain("Intro only"))
    stack.launch_editor("lists-typed", user="alice")

    html = stack.load_html("lists-typed")
    assert "Intro only" in html and "<ul>" not in html

    edited = (
        html
        + f"\n<ul><li>{MARKER}-one</li><li>{MARKER}-two</li></ul>"
        + f"\n<ol><li>{MARKER}-three</li></ol>"
    )
    stack.save_html("lists-typed", edited)
    assert stack.host.put_count == 1

    stored = stack.remote_doc("lists-typed")
    list_paras = _docx_list_paragraphs(stored)
    bullets = [t for s, t in list_paras if "Bullet" in s]
    numbers = [t for s, t in list_paras if "Number" in s]
    assert bullets == [f"{MARKER}-one", f"{MARKER}-two"], bullets
    assert numbers == [f"{MARKER}-three"], numbers


def test_docx_lists_survive_reopen_after_remote_put_and_unlock(stack):
    """After a save (remote PutFile) and a fresh launch (unlock + re-open),
    the list edits are still there — production 'close and reopen' flow."""
    stack.host.seed("lists-reopen", "reopen.docx", _build_docx_lists())
    stack.launch_editor("lists-reopen", user="alice")

    html = stack.load_html("lists-reopen")
    edited = html + f"\n<ul><li>{MARKER}-reopen</li></ul>"
    stack.save_html("lists-reopen", edited)
    assert stack.host.put_count == 1

    # Close the document: unlock on the remote host, then reopen it
    # through a fresh launch as another user.
    stack.unlock("lists-reopen")
    stack.launch_editor("lists-reopen", user="bob")

    again = stack.load_html("lists-reopen")
    assert "<ul>" in again and "<li>apples</li>" in again
    assert f"{MARKER}-reopen" in again

    stored = stack.remote_doc("lists-reopen")
    bullets = [t for s, t in _docx_list_paragraphs(stored) if "Bullet" in s]
    assert f"{MARKER}-reopen" in bullets


def test_odt_nested_list_renders_and_survives_save(stack):
    """A nested list in an ODT renders as nested <ul> markup and the whole
    content survives the production save roundtrip."""
    stack.host.seed("lists-nested", "nested.odt", _build_odt_nested_list())
    stack.launch_editor("lists-nested", user="alice")

    html = stack.load_html("lists-nested")
    assert "<ul>" in html
    assert "<li>item 1<ul><li>nested a</li></ul></li>" in html, html
    assert "<li>item 2</li>" in html

    edited = html + f"\n<ul><li>{MARKER}-nested</li></ul>"
    stack.save_html("lists-nested", edited)
    assert stack.host.put_count == 1

    reloaded = stack.load_html("lists-nested")
    assert "item 1" in reloaded and "nested a" in reloaded
    assert f"{MARKER}-nested" in reloaded

    stored = stack.remote_doc("lists-nested")
    assert _odt_content_xml_has(stored, "item 1")
    assert _odt_content_xml_has(stored, "nested a")
    assert _odt_content_xml_has(stored, f"{MARKER}-nested")


def test_blank_document_first_save_creates_list_document(stack):
    """A blank ODT opens empty and the first save (with list markup) lands
    on the remote host as a real ODT whose content.xml carries the lists."""
    stack.host.seed("lists-blank", "blank.odt", b"")
    stack.launch_editor("lists-blank", user="alice")

    html = stack.load_html("lists-blank")
    assert html == ""

    stack.save_html(
        "lists-blank",
        f"<ul><li>{MARKER}-blank</li></ul><ol><li>{MARKER}-seq</li></ol>",
    )
    stored = stack.remote_doc("lists-blank")
    assert _odt_content_xml_has(stored, f"{MARKER}-blank")
    assert _odt_content_xml_has(stored, f"{MARKER}-seq")
    assert f"{MARKER}-blank" in _odt_text(stored)
    assert _odt_list_count(stored) >= 2
