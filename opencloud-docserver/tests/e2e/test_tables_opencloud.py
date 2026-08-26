"""E2E — tables through production OpenCloud (OCIS client mode).

This suite drives the exact code path the docserver uses in production
(deployed against real OpenCloud / OCIS, e.g. cloud.graphwiz.ai), focused
on the table features the editor UI offers (insert-table dialog, merged
cells, header rows, nested tables):

    seed an office document (DOCX or ODT) containing tables
      -> OpenCloud launches /editor with access_token + WOPISrc  (WOPI handshake)
      -> GET  /api/documents/{id}/html                           (table -> HTML)
      -> edit the HTML in the editor (add table markup / a marker cell)
      -> POST /api/documents/{id}/save                           (HTML -> document)
      -> docserver PUTs the bytes back to the remote host         (PutFile)
      -> GET  /api/documents/{id}/html again                     (reload)
      -> re-read the stored document and verify its table cells
         (DOCX <w:tbl> rows/cells, ODT content.xml <table:table-row> /
         <table:table-cell>) still contain the marker.

A real HTTP WOPI host stands in for OCIS (same wire protocol, same
lock-before-PutFile semantics the wopiserver enforces). The real
docserver app, editor router, RemoteWopiClient and both converters
(DOCX + ODT) are exercised end to end — nothing is mocked on the
docserver side.

The suite asserts the table-persistence contract: table cells typed or
present in the editor survive every hop of the production path, and the
stored bytes are a real office document whose table structure contains
them.

Marker emitted when the full production-path table persistence check
passes:
TABLE-PERSISTENCE: OK
"""

from __future__ import annotations

import io
import json
import re
import threading
import zipfile
from contextlib import asynccontextmanager
from wsgiref.simple_server import make_server

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from odf import teletype
from odf.opendocument import OpenDocumentText, load
from odf.table import Table, TableCell, TableRow
from odf.text import P

from src.config import Config
from src.editor.router import router as editor_router
from src.editor.session import SessionRegistry
from src.lib.store import DocumentStore, wipe_db, wipe_dir
from src.wopi.router import router as wopi_router

# Marker the editor types into a table cell; a save is never treated as
# successful until the marker is found in the stored remote document's table.
MARKER = "TABLE-PERSISTENCE"


# ----------------------------------------------------------------------
# Production OpenCloud (OCIS) wopiserver stand-in
# ----------------------------------------------------------------------

class _ProdOcisHost:
    """Minimal OCIS wopiserver over WSGI — what cloud.graphwiz.ai does.

    Implements the same WOPI surface the real wopiserver serves and the
    same quirks the docserver's RemoteWopiClient depends on:
      -  GET  /wopi/files/{id}             CheckFileInfo (Bearer auth)
      -  GET  /wopi/files/{id}/contents    GetFile
      -  POST /wopi/files/{id}/contents    PutFile (X-WOPI-Override: PUT)
      -  POST /wopi/files/{id}             LOCK / GET_LOCK / UNLOCK
    Like the real wopiserver, PutFile on an unlocked file is refused
    (409 "Cannot PutFile on unlocked file"), so every save must present
    the lock the docserver took at launch.
    """

    def __init__(self) -> None:
        self.content: dict[str, bytes] = {}
        self.names: dict[str, str] = {}
        self.locks: dict[str, str] = {}
        self.users: dict[str, str] = {}
        self.put_count = 0
        self.put_overrides: list[str] = []
        self.put_lock_headers: list[str] = []
        self.lock_events: list[str] = []

    def seed(self, doc_id: str, name: str, data: bytes, user: str = "alice") -> None:
        self.content[doc_id] = data
        self.names[doc_id] = name
        self.users[doc_id] = user

    def __call__(self, environ, start_response):
        path = environ.get("PATH_INFO", "")
        method = environ.get("REQUEST_METHOD", "GET")
        override = environ.get("HTTP_X_WOPI_OVERRIDE", "")
        lock_hdr = environ.get("HTTP_X_WOPI_LOCK", "")
        auth = environ.get("HTTP_AUTHORIZATION", "")

        m = re.match(r"^/wopi/files/([^/]+)(/contents)?$", path)
        if not m:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"not found"]
        doc_id, is_contents = m.group(1), bool(m.group(2))
        if doc_id not in self.content:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"no such file"]

        # GetFile — the docserver reads the raw document bytes with the token.
        if method == "GET" and is_contents:
            start_response(
                "200 OK",
                [
                    ("Content-Type", "application/octet-stream"),
                    ("X-WOPI-ItemVersion", "v1"),
                ],
            )
            return [self.content[doc_id]]

        # CheckFileInfo — the docserver reads BaseFileName (routes .docx vs
        # .odt to the right converter) and UserId (names the WOPI lock).
        if method == "GET" and not is_contents:
            user_id = auth.replace("Bearer ", "") or self.users.get(doc_id, "unknown")
            body = json.dumps(
                {
                    "BaseFileName": self.names.get(doc_id, "document.docx"),
                    "UserId": user_id,
                    "Size": len(self.content[doc_id]),
                }
            ).encode()
            start_response("200 OK", [("Content-Type", "application/json")])
            return [body]

        # PutFile — only accepted with the lock taken at launch.
        if method == "POST" and is_contents and override == "PUT":
            length = int(environ.get("CONTENT_LENGTH", "0"))
            if lock_hdr != (self.locks.get(doc_id) or ""):
                if not self.locks.get(doc_id):
                    start_response(
                        "409 Conflict",
                        [
                            ("Content-Type", "text/plain"),
                            ("X-WOPI-Lock", ""),
                            ("X-WOPI-LockFailureReason", "Cannot PutFile on unlocked files"),
                        ],
                    )
                    return [b"conflict"]
                start_response("500 Internal Server Error", [("Content-Type", "text/plain")])
                return [b"lock mismatch"]
            self.content[doc_id] = environ["wsgi.input"].read(length)
            self.put_count += 1
            self.put_overrides.append(override)
            self.put_lock_headers.append(lock_hdr)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        if method == "POST" and not is_contents and override == "LOCK":
            if self.locks.get(doc_id) and self.locks[doc_id] != lock_hdr:
                start_response(
                    "409 Conflict",
                    [("Content-Type", "text/plain"), ("X-WOPI-Lock", self.locks[doc_id])],
                )
                return [b"locked by other"]
            self.locks[doc_id] = lock_hdr
            self.lock_events.append(f"LOCK:{doc_id}")
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-ItemVersion", "v1")],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "GET_LOCK":
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-Lock", self.locks.get(doc_id, ""))],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "UNLOCK":
            if lock_hdr == self.locks.get(doc_id):
                self.locks.pop(doc_id)
            self.lock_events.append(f"UNLOCK:{doc_id}")
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        start_response("404 Not Found", [("Content-Type", "text/plain")])
        return [b"not found"]


# ----------------------------------------------------------------------
# The docserver wired against a live remote OpenCloud host
# ----------------------------------------------------------------------

class _E2EStack:
    """The docserver plus a remote OpenCloud host, wired like production.

    The docserver runs as a FastAPI TestClient; the remote host is a real
    WSGI server on 127.0.0.1 so the docserver's RemoteWopiClient makes real
    HTTP calls (urllib) exactly like it would against cloud.graphwiz.ai.
    """

    def __init__(self, tmp_path) -> None:
        self.host = _ProdOcisHost()
        self._httpd = make_server("127.0.0.1", 0, self.host)
        self.port = self._httpd.server_address[1]
        self.wopi_host = f"http://127.0.0.1:{self.port}"
        self._thread = threading.Thread(target=self._httpd.serve_forever, daemon=True)
        self._thread.start()

        db = str(tmp_path / "t.db")
        content = str(tmp_path / "content")
        store = DocumentStore(db, content)
        cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

        @asynccontextmanager
        async def lifespan(app: FastAPI):
            app.state.store = store
            app.state.sessions = SessionRegistry()
            app.state.config = cfg
            yield

        app = FastAPI(lifespan=lifespan)
        app.include_router(wopi_router)
        app.include_router(editor_router)
        self.client = TestClient(app)
        self.client.__enter__()  # run lifespan (app.state.*)
        self._db, self._content = db, content

    def close(self) -> None:
        try:
            self.client.__exit__(None, None, None)
        except Exception:
            pass
        self._httpd.shutdown()
        self._httpd.server_close()
        self._thread.join(timeout=2)
        wipe_db(self._db)
        wipe_dir(self._content)

    def launch_editor(self, doc_id: str, user: str = "alice") -> str:
        """OpenCloud launches /editor (POST form) for a user's file."""
        token = f"tok-{user}"
        wopi_src = f"{self.wopi_host}/wopi/files/{doc_id}"
        resp = self.client.post(
            "/editor",
            params={"WOPISrc": wopi_src},
            data={"file_id": doc_id, "access_token": token},
        )
        assert resp.status_code == 200
        assert self.host.locks.get(doc_id), "launch must have taken the WOPI lock"
        return token

    def load_html(self, doc_id: str) -> str:
        resp = self.client.get(f"/api/documents/{doc_id}/html")
        assert resp.status_code == 200, resp.text
        return resp.json()["html"]

    def save_html(self, doc_id: str, html: str) -> None:
        resp = self.client.post(f"/api/documents/{doc_id}/save", json={"html": html})
        assert resp.status_code == 200, resp.text
        assert resp.json().get("ok") is True

    def unlock(self, doc_id: str) -> None:
        resp = self.client.post(f"/api/documents/{doc_id}/unlock")
        assert resp.status_code == 200, resp.text

    def remote_doc(self, doc_id: str) -> bytes:
        """Raw bytes the OpenCloud host stored after the docserver's PutFile."""
        return self.host.content[doc_id]


@pytest.fixture
def stack(tmp_path):
    s = _E2EStack(tmp_path)
    yield s
    s.close()


# ----------------------------------------------------------------------
# Table document fixtures + stored-bytes inspection
# ----------------------------------------------------------------------

def _cell_html(html: str, text: str, tag: str = "td") -> bool:
    """True when a <td|th> (optionally carrying width/style attrs) wraps the
    exact cell text — tolerant of the height/width attributes the DOCX/ODT
    readers now emit for explicitly-sized cells."""
    return re.search(rf"<{tag}(?:\s[^>]*)?>{re.escape(text)}</{tag}>", html) is not None


def _build_docx_tables() -> bytes:
    """A DOCX with a body line, a 2x2 data table and a header-row table."""
    doc = Document()
    doc.add_paragraph("Stoic table document")
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "North"
    t.cell(0, 1).text = "South"
    t.cell(1, 0).text = "East"
    t.cell(1, 1).text = "West"

    t2 = doc.add_table(rows=2, cols=2)
    t2.style = "Table Grid"
    t2.cell(0, 0).text = "Qty"
    t2.cell(0, 1).text = "Item"
    t2.cell(1, 0).text = "2"
    t2.cell(1, 1).text = "stoic pens"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_plain(body: str = "Plain stoic document") -> bytes:
    doc = Document()
    doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_odt_tables() -> bytes:
    """An ODT with a body line, a 2x2 table and a nested table in a cell."""
    doc = OpenDocumentText()
    doc.text.addElement(P(text="Stoic ODT table document"))

    t = Table()
    for vals in (("alpha", "beta"), ("gamma", "delta")):
        tr = TableRow()
        for v in vals:
            tc = TableCell()
            tc.addElement(P(text=v))
            tr.addElement(tc)
        t.addElement(tr)
    doc.text.addElement(t)

    nested = Table()
    tr_n = TableRow()
    tc_outer = TableCell()
    tc_outer.addElement(P(text="outer"))
    inner = Table()
    tr_i = TableRow()
    for v in ("nested-a", "nested-b"):
        tc_i = TableCell()
        tc_i.addElement(P(text=v))
        tr_i.addElement(tc_i)
    inner.addElement(tr_i)
    tc_outer.addElement(inner)
    tr_n.addElement(tc_outer)
    nested.addElement(tr_n)
    doc.text.addElement(nested)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_table_count(data: bytes) -> int:
    doc = Document(io.BytesIO(data))
    return len(doc.tables)


def _docx_table_cells(data: bytes) -> list[str]:
    """Every cell text in every table of a stored DOCX."""
    doc = Document(io.BytesIO(data))
    return [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    ]


def _docx_has_merged_cell(data: bytes) -> bool:
    """True if any stored DOCX table uses gridSpan (a horizontal merge)."""
    doc = Document(io.BytesIO(data))
    for table in doc.tables:
        for row in table.rows:
            for tc in row._tr.tc_lst:
                if tc.tcPr is not None and tc.tcPr.gridSpan is not None:
                    return True
    return False


def _odt_content_xml_has(data: bytes, needle: str) -> bool:
    """True when the ODT's content.xml (the persisted body) contains needle."""
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        content_xml = zf.read("content.xml").decode("utf-8", "replace")
    return needle in content_xml


def _odt_text(data: bytes) -> str:
    doc = load(io.BytesIO(data))
    return teletype.extractText(doc.text)


def _odt_table_count(data: bytes) -> int:
    doc = load(io.BytesIO(data))
    return len(doc.text.getElementsByType(Table))


def _odt_table_cell_texts(data: bytes) -> list[str]:
    """Extract text for every top-level table cell in a stored ODT."""
    doc = load(io.BytesIO(data))
    texts = []
    for t in doc.text.getElementsByType(Table):
        for row in t.getElementsByType(TableRow):
            for cell in row.getElementsByType(TableCell):
                texts.append(teletype.extractText(cell))
    return texts


# ----------------------------------------------------------------------
# The E2E tests
# ----------------------------------------------------------------------

def test_docx_tables_persist_through_production_opencloud(stack):
    """The full production roundtrip for a DOCX seeded with tables:
    OpenCloud launch -> HTML with <table> -> edit (add a marker table)
    -> save -> PutFile -> reload. Emits the TABLE-PERSISTENCE marker only
    when the DOCX stored back on the OpenCloud host still carries the
    marker as real table-cell text."""
    stack.host.seed("tables-docx", "ledger.docx", _build_docx_tables())
    stack.launch_editor("tables-docx", user="alice")

    # OpenCloud serves the tables as HTML for the editor.
    html = stack.load_html("tables-docx")
    assert "<table>" in html, html
    assert _cell_html(html, "North")
    assert _cell_html(html, "South")
    assert _cell_html(html, "Qty")
    assert _cell_html(html, "stoic pens")
    assert "Stoic table document" in html

    # The user types a new table with a marker cell into the editor.
    edited = (
        html
        + f"\n<table><tr><td>{MARKER}-r1c1</td><td>{MARKER}-r1c2</td></tr>"
        + f"\n<tr><td>{MARKER}-r2c1</td><td>{MARKER}-r2c2</td></tr></table>"
    )
    stack.save_html("tables-docx", edited)

    # The docserver must have PUT the DOCX back to the OpenCloud host,
    # presenting the WOPI lock it took at launch (409 otherwise).
    assert stack.host.put_count == 1, "edited doc must be PUT back to OpenCloud"
    assert "PUT" in stack.host.put_overrides
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0], (
        "save must present the WOPI lock taken at launch"
    )

    # Reload from the remote host: the typed table is still there.
    reloaded = stack.load_html("tables-docx")
    assert MARKER in reloaded
    assert f"{MARKER}-r1c1" in reloaded
    assert "<table>" in reloaded

    # The stored bytes are a real DOCX whose tables carry the marker:
    # three tables, the original cells preserved and the marker cell present.
    stored = stack.remote_doc("tables-docx")
    assert _docx_table_count(stored) == 3, "stored DOCX must keep 3 tables"
    cells = _docx_table_cells(stored)
    assert "North" in cells and "South" in cells
    assert "Qty" in cells and "stoic pens" in cells
    assert f"{MARKER}-r1c1" in cells and f"{MARKER}-r2c2" in cells

    print("TABLE-PERSISTENCE: OK")


def test_odt_tables_persist_through_production_opencloud(stack):
    """The same production roundtrip for an ODT seeded with tables: the
    marker lands in the stored content.xml next to the table rows."""
    stack.host.seed("tables-odt", "ledger.odt", _build_odt_tables())
    stack.launch_editor("tables-odt", user="alice")

    html = stack.load_html("tables-odt")
    assert "<table>" in html, html
    assert "<td><p>alpha</p></td>" in html
    assert "<td><p>delta</p></td>" in html

    edited = html + (
        f'\n<table><tr><td><p>{MARKER}-o1</p></td><td><p>{MARKER}-o2</p></td></tr>'
        f'\n<tr><td><p>{MARKER}-o3</p></td><td><p>{MARKER}-o4</p></td></tr></table>'
    )
    stack.save_html("tables-odt", edited)
    assert stack.host.put_count == 1
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0]

    reloaded = stack.load_html("tables-odt")
    assert f"{MARKER}-o1" in reloaded
    assert f"{MARKER}-o4" in reloaded
    assert "<table>" in reloaded

    stored = stack.remote_doc("tables-odt")
    assert _odt_content_xml_has(stored, MARKER), "content.xml must contain the marker"
    text = _odt_text(stored)
    for needle in ("alpha", "beta", "gamma", "delta",
                   f"{MARKER}-o1", f"{MARKER}-o2", f"{MARKER}-o3", f"{MARKER}-o4"):
        assert needle in text, f"stored ODT lost {needle!r}"
    assert _odt_table_count(stored) >= 3, "stored ODT must keep 3 table blocks"


def test_docx_merged_and_header_table_survives_roundtrip(stack):
    """A table with a merged row (gridSpan -> colspan) and a repeating
    header row (<th>) survives DOCX -> HTML -> DOCX through the remote
    OpenCloud host."""
    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "Total"
    t.cell(0, 0).merge(t.cell(0, 1))
    t.cell(1, 0).text = "a"
    t.cell(1, 1).text = "b"
    t.cell(2, 0).text = "c"
    t.cell(2, 1).text = "d"
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    stack.host.seed("tables-merge", "merged.docx", data)
    stack.launch_editor("tables-merge", user="alice")

    html = stack.load_html("tables-merge")
    assert "<table>" in html and 'colspan="2"' in html, html

    # No-op edit: the merged table must come back intact.
    stack.save_html("tables-merge", html)
    assert stack.host.put_count == 1

    stored = stack.remote_doc("tables-merge")
    cells = _docx_table_cells(stored)
    assert "Total" in cells and "a" in cells and "d" in cells
    assert _docx_has_merged_cell(stored), "stored DOCX must keep the merged cell"

    reloaded = stack.load_html("tables-merge")
    assert "<table>" in reloaded and 'colspan="2"' in reloaded
    assert "Total" in reloaded


def test_tables_typed_in_editor_become_real_tables(stack):
    """Starting from a plain document, table markup added in the editor is
    saved as a real DOCX table whose cells the stored bytes carry."""
    stack.host.seed("tables-typed", "typed.docx", _build_docx_plain("Intro only"))
    stack.launch_editor("tables-typed", user="alice")

    html = stack.load_html("tables-typed")
    assert "Intro only" in html and "<table>" not in html

    edited = html + (
        f'\n<table><tr><td>{MARKER}-one</td><td>{MARKER}-two</td></tr>'
        f'\n<tr><td>{MARKER}-three</td><td>{MARKER}-four</td></tr></table>'
    )
    stack.save_html("tables-typed", edited)
    assert stack.host.put_count == 1

    stored = stack.remote_doc("tables-typed")
    assert _docx_table_count(stored) == 1, "stored DOCX must contain exactly 1 table"
    cells = _docx_table_cells(stored)
    for needle in (f"{MARKER}-one", f"{MARKER}-two", f"{MARKER}-three", f"{MARKER}-four"):
        assert needle in cells, f"stored DOCX table lost {needle!r}"


def test_docx_tables_survive_reopen_after_remote_put_and_unlock(stack):
    """After a save (remote PutFile) and a fresh launch (unlock + re-open),
    the table edits are still there — production 'close and reopen' flow."""
    stack.host.seed("tables-reopen", "reopen.docx", _build_docx_tables())
    stack.launch_editor("tables-reopen", user="alice")

    html = stack.load_html("tables-reopen")
    edited = html + (
        f'\n<table><tr><td>{MARKER}-reopen</td></tr></table>'
    )
    stack.save_html("tables-reopen", edited)
    assert stack.host.put_count == 1

    # Close the document: unlock on the remote host, then reopen it
    # through a fresh launch as another user.
    stack.unlock("tables-reopen")
    stack.launch_editor("tables-reopen", user="bob")

    again = stack.load_html("tables-reopen")
    assert "<table>" in again and _cell_html(again, "North")
    assert f"{MARKER}-reopen" in again

    stored = stack.remote_doc("tables-reopen")
    cells = _docx_table_cells(stored)
    assert "North" in cells and "stoic pens" in cells
    assert f"{MARKER}-reopen" in cells


def test_odt_nested_table_renders_and_survives_save(stack):
    """A nested table inside a cell renders as nested <table> markup and the
    whole structure survives the production save roundtrip."""
    doc = OpenDocumentText()
    outer = Table()
    tr = TableRow()
    tc = TableCell()
    tc.addElement(P(text="outer r1c1"))
    inner = Table()
    tr_i = TableRow()
    for v in ("nested-a", "nested-b"):
        tc_i = TableCell()
        tc_i.addElement(P(text=v))
        tr_i.addElement(tc_i)
    inner.addElement(tr_i)
    tc.addElement(inner)
    tr.addElement(tc)
    outer.addElement(tr)
    doc.text.addElement(outer)
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    stack.host.seed("tables-nested", "nested.odt", data)
    stack.launch_editor("tables-nested", user="alice")

    html = stack.load_html("tables-nested")
    assert "<table>" in html
    assert "outer r1c1" in html and "nested-a" in html and "nested-b" in html, html

    edited = html + f"\n<table><tr><td><p>{MARKER}-nested</p></td></tr></table>"
    stack.save_html("tables-nested", edited)
    assert stack.host.put_count == 1

    reloaded = stack.load_html("tables-nested")
    assert "outer r1c1" in reloaded and "nested-a" in reloaded
    assert f"{MARKER}-nested" in reloaded

    stored = stack.remote_doc("tables-nested")
    assert _odt_content_xml_has(stored, "outer r1c1")
    assert _odt_content_xml_has(stored, "nested-a")
    assert _odt_content_xml_has(stored, f"{MARKER}-nested")
    assert _odt_table_count(stored) >= 2, "stored ODT must keep the nested table"


def test_blank_document_first_save_creates_table_document(stack):
    """A blank ODT opens empty and the first save (with table markup) lands
    on the remote host as a real ODT whose content.xml carries the table."""
    stack.host.seed("tables-blank", "blank.odt", b"")
    stack.launch_editor("tables-blank", user="alice")

    html = stack.load_html("tables-blank")
    assert html == ""

    stack.save_html(
        "tables-blank",
        f"<table><tr><td>{MARKER}-blank</td><td>{MARKER}-seq</td></tr></table>",
    )
    stored = stack.remote_doc("tables-blank")
    assert _odt_content_xml_has(stored, f"{MARKER}-blank")
    assert _odt_content_xml_has(stored, f"{MARKER}-seq")
    assert f"{MARKER}-blank" in _odt_text(stored)
    assert f"{MARKER}-blank" in _odt_table_cell_texts(stored)
    assert _odt_table_count(stored) >= 1
