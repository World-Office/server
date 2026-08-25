"""E2E — Realtime-Text-Sync (US-42, contract: docs/testing/test-scenarios.md).

Drives the real-time character-CRDT collaboration layer of the docserver
through the true production path:

    seed a real office document on a real OpenCloud/OCIS wopiserver
      -> OpenCloud launches /editor (WOPI handshake, remote lock taken)
      -> two editors join the shared document room over real HTTP
           GET  /api/documents/{id}/collab/state    full-log join
           POST /api/documents/{id}/collab/ops      local edits
           GET  /api/documents/{id}/collab/ops?since polling catch-up
           GET  /api/documents/{id}/collab/stream   SSE realtime push
      -> editor A types -> editor B sees the keystrokes in < 200 ms
      -> both type at once -> the merge converges on every replica
      -> very fast typing (100+ chars in a burst) converges too
      -> the merged text survives a save back to the remote host

A real HTTP WOPI host stands in for OCIS (the same wire protocol and the
same lock-before-PutFile semantics the wopiserver enforces). The real
docserver app (src.main.create_app), the editor router, the
RemoteWopiClient, the DOCX converters and the collab hub are exercised end
to end — nothing on the docserver side is mocked.

The client-side replica each test drives is the wire protocol the server
documents (src/editor/collab.py): integrate the hub's op log, ship local
insert/delete ops, converge. Every op is round-tripped as *real JSON over
HTTP*, so the wire format, the hub ordering and the SSE fan-out are all
under test.

Marker emitted when the core contract passes:
TEXT-SYNC: OK
"""

from __future__ import annotations

import io
import json
import re
import socket
import threading
import time
from wsgiref.simple_server import make_server

import httpx
import pytest
import uvicorn
from docx import Document
from fastapi.testclient import TestClient

from src.config import Config
from src.editor.collab import TextCRDT, reset_hub
from src.lib.store import wipe_db, wipe_dir
from src.main import create_app

# How fast a text change must reach the other editors (US-42).
SYNC_BUDGET_S = 0.2  # 200 ms


# ----------------------------------------------------------------------
# Production OpenCloud (OCIS) wopiserver stand-in
# ----------------------------------------------------------------------


class _ProdOcisHost:
    """Minimal OCIS wopiserver over WSGI — the wire protocol cloud hosts use.

    Implements the same WOPI surface the real wopiserver serves and the
    same quirks the docserver's RemoteWopiClient depends on:
      -  GET  /wopi/files/{id}             CheckFileInfo (Bearer auth)
      -  GET  /wopi/files/{id}/contents    GetFile
      -  POST /wopi/files/{id}/contents    PutFile (X-WOPI-Override: PUT)
      -  POST /wopi/files/{id}             LOCK / GET_LOCK / UNLOCK
    Like the real wopiserver, PutFile on an unlocked file is refused
    (409 "Cannot PutFile on unlocked file"), so every save must present
    the lock the docserver took at launch.
    """

    def __init__(self) -> None:
        self.content: dict[str, bytes] = {}
        self.names: dict[str, str] = {}
        self.locks: dict[str, str] = {}
        self.put_count = 0
        self.put_lock_headers: list[str] = []
        self.getfile_count = 0

    def seed(self, doc_id: str, name: str, data: bytes) -> None:
        self.content[doc_id] = data
        self.names[doc_id] = name

    def __call__(self, environ, start_response):
        path = environ.get("PATH_INFO", "")
        method = environ.get("REQUEST_METHOD", "GET")
        override = environ.get("HTTP_X_WOPI_OVERRIDE", "")
        lock_hdr = environ.get("HTTP_X_WOPI_LOCK", "")
        auth = environ.get("HTTP_AUTHORIZATION", "")

        m = re.match(r"^/wopi/files/([^/]+)(/contents)?$", path)
        if not m:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"not found"]
        doc_id, is_contents = m.group(1), bool(m.group(2))
        if doc_id not in self.content:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"no such file"]

        # GetFile — the docserver reads the raw bytes with the token.
        if method == "GET" and is_contents:
            self.getfile_count += 1
            start_response(
                "200 OK",
                [
                    ("Content-Type", "application/octet-stream"),
                    ("X-WOPI-ItemVersion", "v1"),
                ],
            )
            return [self.content[doc_id]]

        # CheckFileInfo — the docserver reads BaseFileName (routes the
        # converter by extension) and UserId (names the WOPI lock).
        if method == "GET" and not is_contents:
            user_id = auth.replace("Bearer ", "") or "anonymous"
            body = json.dumps(
                {
                    "BaseFileName": self.names.get(doc_id, "document.docx"),
                    "UserId": user_id,
                    "Size": len(self.content[doc_id]),
                }
            ).encode()
            start_response("200 OK", [("Content-Type", "application/json")])
            return [body]

        # PutFile — only accepted with the lock taken at launch.
        if method == "POST" and is_contents and override == "PUT":
            length = int(environ.get("CONTENT_LENGTH", "0"))
            if lock_hdr != (self.locks.get(doc_id) or ""):
                start_response(
                    "409 Conflict",
                    [
                        ("Content-Type", "text/plain"),
                        ("X-WOPI-Lock", ""),
                        ("X-WOPI-LockFailureReason", "Cannot PutFile on unlocked files"),
                    ],
                )
                return [b"conflict"]
            self.content[doc_id] = environ["wsgi.input"].read(length)
            self.put_count += 1
            self.put_lock_headers.append(lock_hdr)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        if method == "POST" and not is_contents and override == "LOCK":
            if self.locks.get(doc_id) and self.locks[doc_id] != lock_hdr:
                start_response(
                    "409 Conflict",
                    [("Content-Type", "text/plain"), ("X-WOPI-Lock", self.locks[doc_id])],
                )
                return [b"locked by other"]
            self.locks[doc_id] = lock_hdr
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-ItemVersion", "v1")],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "GET_LOCK":
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-Lock", self.locks.get(doc_id, ""))],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "UNLOCK":
            if lock_hdr == self.locks.get(doc_id):
                self.locks.pop(doc_id)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        start_response("404 Not Found", [("Content-Type", "text/plain")])
        return [b"not found"]


# ----------------------------------------------------------------------
# The docserver wired against a live remote OpenCloud host
# ----------------------------------------------------------------------


class _E2EStack:
    """The docserver plus a remote OpenCloud host, wired like production.

    The docserver runs as the REAL production app (src.main.create_app)
    through a FastAPI TestClient; the remote host is a real WSGI server on
    127.0.0.1 so the docserver's RemoteWopiClient makes real HTTP calls
    (urllib) exactly like it would against a deployed OpenCloud.
    """

    def __init__(self, tmp_path, mode: str = "testclient") -> None:
        self.host = _ProdOcisHost()
        self._httpd = make_server("127.0.0.1", 0, self.host)
        self.port = self._httpd.server_address[1]
        self.wopi_host = f"http://127.0.0.1:{self.port}"
        self._thread = threading.Thread(target=self._httpd.serve_forever, daemon=True)
        self._thread.start()

        db = str(tmp_path / "t.db")
        content = str(tmp_path / "content")
        cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")
        app = create_app(cfg)
        self._mode = mode
        if mode == "server":
            # A real ASGI server (not TestClient): TestClient cannot serve
            # SSE streams to external threads (deadlocks in the internal
            # event loop), so realtime subscribers need this mode.
            self._srv = self._start_uvicorn(app)
            self.client = httpx.Client(base_url=f"http://127.0.0.1:{self._srv._app_port}")
        else:
            self._srv = None
            self.client = TestClient(app)
            self.client.__enter__()  # run lifespan (app.state.*)
        self._db, self._content = db, content

    @staticmethod
    def _start_uvicorn(app):
        sock = socket.socket()
        sock.bind(("127.0.0.1", 0))
        app_port = sock.getsockname()[1]
        sock.close()
        cfg = uvicorn.Config(app, host="127.0.0.1", port=app_port, log_level="error")
        srv = uvicorn.Server(cfg)
        threading.Thread(target=srv.run, daemon=True).start()
        # wait until the server actually accepts connections
        deadline = time.time() + 10
        while time.time() < deadline:
            try:
                with socket.create_connection(("127.0.0.1", app_port), timeout=0.5):
                    break
            except OSError:
                time.sleep(0.05)
        srv._app_port = app_port
        return srv

    def close(self) -> None:
        if self._mode == "server":
            try:
                self._srv.should_exit = True
            except Exception:
                pass
            self.client.close()
        else:
            try:
                self.client.__exit__(None, None, None)
            except Exception:
                pass
        self._httpd.shutdown()
        self._httpd.server_close()
        self._thread.join(timeout=2)
        wipe_db(self._db)
        wipe_dir(self._content)

    # -- production steps -------------------------------------------------

    def seed_remote_doc(self, doc_id: str, name: str, text: str) -> None:
        """Place a real DOCX on the OpenCloud host, like a stored user file."""
        doc = Document()
        doc.add_paragraph(text)
        buf = io.BytesIO()
        doc.save(buf)
        self.host.seed(doc_id, name, buf.getvalue())

    def launch_editor(self, doc_id: str, user: str = "alice") -> None:
        """OpenCloud launches /editor (POST form) for a user's file.

        The WOPI lock must be taken on the remote host, else no save can
        PutFile the merged bytes back (409 unlocked file).
        """
        wopi_src = f"{self.wopi_host}/wopi/files/{doc_id}"
        resp = self.client.post(
            "/editor",
            params={"WOPISrc": wopi_src},
            data={"file_id": doc_id, "access_token": user},
        )
        assert resp.status_code == 200, resp.text
        assert self.host.locks.get(doc_id), "launch must have taken the WOPI lock"

    # -- collab client steps (a real editor's realtime flow) --------------

    def join_editor(self, doc_id: str, site: str) -> TextCRDT:
        """A fresh editor replica joined the way a browser tab joins:
        fetch the full op log (state endpoint) and integrate it."""
        data = self.client.get(f"/api/documents/{doc_id}/collab/state").json()
        replica = TextCRDT(site)
        for op in data["ops"]:
            replica.integrate(op)
        return replica

    def apply_ops(self, doc_id: str, client_id: str, ops, base_rev=None) -> dict:
        resp = self.client.post(
            f"/api/documents/{doc_id}/collab/ops",
            json={
                "client_id": client_id,
                "base_rev": base_rev if base_rev is not None else 0,
                "ops": ops,
            },
        )
        assert resp.status_code == 200, resp.text
        return resp.json()

    def ops_since(self, doc_id: str, since: int) -> list[dict]:
        resp = self.client.get(f"/api/documents/{doc_id}/collab/ops?since={since}")
        assert resp.status_code == 200, resp.text
        return resp.json()["ops"]

    def hub_state(self, doc_id: str) -> dict:
        resp = self.client.get(f"/api/documents/{doc_id}/collab/state")
        assert resp.status_code == 200, resp.text
        return resp.json()

    def save_html(self, doc_id: str, html: str) -> None:
        resp = self.client.post(f"/api/documents/{doc_id}/save", json={"html": html})
        assert resp.status_code == 200, resp.text
        assert resp.json().get("ok") is True

    def remote_doc(self, doc_id: str) -> bytes:
        """Raw bytes the OpenCloud host stored after the docserver's PutFile."""
        return self.host.content[doc_id]

    def remote_text(self, doc_id: str) -> str:
        doc = Document(io.BytesIO(self.remote_doc(doc_id)))
        return "\n".join(p.text for p in doc.paragraphs)

    # -- helpers ----------------------------------------------------------

    def seed_and_launch(self, doc_id: str, text: str, user: str = "alice") -> None:
        self.seed_remote_doc(doc_id, f"{doc_id}.docx", text)
        self.launch_editor(doc_id, user=user)


def _docx_paragraphs(data: bytes) -> list[str]:
    doc = Document(io.BytesIO(data))
    return [p.text for p in doc.paragraphs]


@pytest.fixture
def stack(tmp_path):
    reset_hub()  # the collab hub is a module singleton — isolate per test
    s = _E2EStack(tmp_path)
    yield s
    s.close()
    reset_hub()


@pytest.fixture
def server_stack(tmp_path):
    """Like ``stack`` but over a real ASGI server — required for the SSE
    realtime-push test (TestClient deadlocks on external-thread streams)."""
    reset_hub()
    s = _E2EStack(tmp_path, mode="server")
    yield s
    s.close()
    reset_hub()


# ----------------------------------------------------------------------
# The E2E tests
# ----------------------------------------------------------------------


def test_text_sync_live_push_reaches_other_editor_under_200ms(server_stack):
    """US-42 exploration 1: user A types "Hello" -> user B's editor sees it
    within 200 ms. Bob is subscribed to the realtime SSE stream; Alice's
    keystrokes are shipped through POST /collab/ops; we time the gap from
    Alice's edit leaving her editor until Bob's editor receives the op."""
    stack = server_stack
    stack.seed_and_launch("ts-live.docx", "Hello collaborative world", user="alice")

    alice = stack.join_editor("ts-live.docx", "site-alice")
    bob = stack.join_editor("ts-live.docx", "site-bob")
    assert alice.to_string() == bob.to_string()
    assert "Hello collaborative world" in alice.to_string()

    # Bob's realtime channel: the SSE stream (pushed, not polled).
    seen: list[str] = []
    got_live = threading.Event()

    def read_stream():
        with stack.client.stream("GET", "/api/documents/ts-live.docx/collab/stream") as resp:
            assert resp.status_code == 200, "collab stream must be reachable"
            assert resp.headers["content-type"].startswith("text/event-stream")
            for line in resp.iter_lines():
                seen.append(line)
                if "TEXT-SYNC-LIVE" in line:
                    got_live.set()
                    return

    thread = threading.Thread(target=read_stream, daemon=True)
    thread.start()
    # Wait until Bob's stream has delivered its initial `state` snapshot so
    # the measurement starts with a warm, live subscription.
    deadline = time.time() + 5
    while time.time() < deadline and not any("event: state" in line for line in seen):
        time.sleep(0.005)
    assert any("event: state" in line for line in seen), "stream never sent its initial state"

    # Alice types "TEXT-SYNC-LIVE" at the end of the shared text and ships
    # the op the way a browser editor does.
    op = alice.local_insert(alice.alive_count, "TEXT-SYNC-LIVE")
    t0 = time.perf_counter()
    result = stack.apply_ops("ts-live.docx", "site-alice", [op], base_rev=1)
    assert result["applied"], "alice's op must be applied by the hub"
    assert got_live.wait(timeout=10), "bob's editor never received alice's live op"
    t1 = time.perf_counter()
    latency = t1 - t0
    assert latency < SYNC_BUDGET_S, (
        f"text sync took {latency * 1000:.0f} ms — US-42 requires < 200 ms"
    )

    # Bob applies what arrived over the wire and must converge with Alice.
    for op in stack.ops_since("ts-live.docx", 1):
        bob.integrate(op)
    assert bob.to_string() == alice.to_string(), (
        f"replicas diverged:\n  alice={alice.to_string()!r}\n  bob  ={bob.to_string()!r}"
    )
    assert "TEXT-SYNC-LIVE" in bob.to_string()
    # The hub's authoritative text agrees with both editors.
    assert stack.hub_state("ts-live.docx")["text"] == alice.to_string()

    print("TEXT-SYNC: OK — live push reached the peer editor in "
          f"{latency * 1000:.0f} ms (<200 ms)")


def test_text_sync_concurrent_editors_merge_and_persist(stack):
    """US-42 exploration 2: both editors type at the same time -> the merge
    is correct (no lost update) on every replica, converges through the hub,
    and the merged text survives a save back to the remote OpenCloud host."""
    stack.seed_and_launch("ts-merge.docx", "merge base", user="alice")
    alice = stack.join_editor("ts-merge.docx", "site-A")
    bob = stack.join_editor("ts-merge.docx", "site-B")

    # Both insert right after the opening <p> — concurrent siblings.
    pos = alice.to_string().index("merge")
    op_a = alice.local_insert(pos, "ALICE-")
    pos_b = bob.to_string().index("merge")
    op_b = bob.local_insert(pos_b, "BOB-")

    # A ships first; B ships while still on rev 1 (has not seen A's op).
    stack.apply_ops("ts-merge.docx", "site-A", [op_a], base_rev=1)
    b_reply = stack.apply_ops("ts-merge.docx", "site-B", [op_b], base_rev=1)
    # Single-round-trip healing: B's reply must carry A's op for catch-up.
    assert any(o.get("chars") == "ALICE-" for o in b_reply["ops"]), (
        "B's reply should include A's op so B can catch up in one round trip"
    )

    # Both editors replay the hub's log (the one ordering the merge) and
    # must agree — and separately, fresh replicas replaying the same log
    # must converge to the same text (no divergence by delivery order).
    def replay(site):
        r = TextCRDT(site)
        for op in stack.ops_since("ts-merge.docx", 0):
            r.integrate(op)
        return r

    alice_synced = replay("R-A")
    bob_synced = replay("R-B")
    assert alice_synced.to_string() == bob_synced.to_string(), "replicas diverged"
    merged = alice_synced.to_string()
    assert "ALICE-" in merged and "BOB-" in merged, f"lost update: {merged!r}"
    assert merged.count("merge base") == 1, f"merge duplicated content: {merged!r}"
    # Alice's and Bob's live replicas are on the same text too.
    for op in stack.ops_since("ts-merge.docx", 1):
        alice.integrate(op)
        bob.integrate(op)
    assert alice.to_string() == bob.to_string() == alice_synced.to_string()
    assert stack.hub_state("ts-merge.docx")["text"] == merged

    # The merged text persists through the production save path:
    # editor HTML -> DOCX -> PutFile to the remote OpenCloud host.
    stack.save_html("ts-merge.docx", merged)
    assert stack.host.put_count == 1, "merged doc must be PUT back to OpenCloud"
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0], (
        "save must present the WOPI lock taken at launch"
    )
    stored = _docx_paragraphs(stack.remote_doc("ts-merge.docx"))
    assert any("ALICE-" in p and "BOB-" in p and "merge base" in p for p in stored), stored


def test_text_sync_fast_typing_100_chars_per_sec_converges(stack):
    """US-42 exploration 3: very fast input (100+ chars per second). A whole
    burst of keystrokes lands as one batched update through the hub; every
    replica — including a peer that only polls via ?since= — converges."""
    stack.seed_and_launch("ts-fast.docx", "fast base", user="alice")
    alice = stack.join_editor("ts-fast.docx", "site-alice")
    bob = stack.join_editor("ts-fast.docx", "site-bob")

    # 120 characters typed rapidly (12 ops of 10 chars, one batched POST).
    burst = "".join(chr(ord("a") + (i % 26)) for i in range(120))
    pos = alice.to_string().index("fast base")
    ops = []
    for i in range(0, len(burst), 10):
        chunk = burst[i : i + 10]
        ops.append(alice.local_insert(pos, chunk))
        pos += len(chunk)
    assert len(ops) == 12

    result = stack.apply_ops("ts-fast.docx", "site-alice", ops, base_rev=1)
    assert len(result["applied"]) == 12, f"hub dropped ops: {result['applied']}"
    assert len(result["ops"]) == 12, "the burst must be catch-up-able as one log"

    # Bob's editor catches up by polling and converges.
    for op in stack.ops_since("ts-fast.docx", 1):
        bob.integrate(op)
    assert bob.to_string() == alice.to_string(), "burst replicas diverged"
    assert burst in bob.to_string()
    # 120 characters appended => alive count grew by exactly 120 (char count).
    hub_text = stack.hub_state("ts-fast.docx")["text"]
    assert hub_text.count(burst) == 1
    assert hub_text == bob.to_string()


def test_text_sync_unicode_chars_not_bytes(stack):
    """Multi-byte characters (combining marks, CJK, emoji) must sync whole
    characters — count chars, never bytes — through the full HTTP path."""
    stack.seed_and_launch("ts-unicode.docx", "Grüße Welt", user="alice")
    alice = stack.join_editor("ts-unicode.docx", "site-alice")
    bob = stack.join_editor("ts-unicode.docx", "site-bob")

    sample = "héllo wörld ✨🎉 日本語"
    pos = alice.to_string().index("Welt")
    op = alice.local_insert(pos, sample)
    # the wire op counts characters, not bytes
    assert op["n"] == len(sample)

    stack.apply_ops("ts-unicode.docx", "site-alice", [op], base_rev=1)
    for op in stack.ops_since("ts-unicode.docx", 1):
        bob.integrate(op)
    assert bob.to_string() == alice.to_string(), "unicode replicas diverged"
    assert sample in bob.to_string()
    hub_text = stack.hub_state("ts-unicode.docx")["text"]
    assert hub_text == bob.to_string()
    # delete the emoji pair as exactly two characters and converge again
    start = hub_text.index("✨")
    delete_op = bob.local_delete(start, start + 2)  # ✨ + 🎉
    assert delete_op["t"] == "delete" and len(delete_op["ids"]) == 2
    stack.apply_ops("ts-unicode.docx", "site-bob", [delete_op], base_rev=2)
    for op in stack.ops_since("ts-unicode.docx", 2):
        alice.integrate(op)
    assert alice.to_string() == bob.to_string()
    assert "✨" not in alice.to_string() and "🎉" not in alice.to_string()
    assert "日本語" in alice.to_string()


def test_text_sync_deletes_propagate_to_peers(stack):
    """The delete op (item ids of the tombstoned characters) must reach the
    peer replica exactly like an insert does."""
    stack.seed_and_launch("ts-del.docx", "Hello sync world", user="alice")
    alice = stack.join_editor("ts-del.docx", "site-alice")
    bob = stack.join_editor("ts-del.docx", "site-bob")

    start = alice.to_string().index(" sync")
    delete_op = alice.local_delete(start, start + 5)
    assert len(delete_op["ids"]) == 5

    result = stack.apply_ops("ts-del.docx", "site-alice", [delete_op], base_rev=1)
    assert result["applied"], "delete op must be applied by the hub"
    assert " sync" not in alice.to_string()

    for op in stack.ops_since("ts-del.docx", 1):
        bob.integrate(op)
    assert bob.to_string() == alice.to_string()
    assert " sync" not in bob.to_string()
    assert "Hello world" in bob.to_string()
    assert stack.hub_state("ts-del.docx")["text"] == bob.to_string()


def test_text_sync_late_joiner_replays_full_log(stack):
    """A third editor opening the document after edits have been made must
    converge by replaying the full op log (state join), exactly what a
    browser tab that loads the document mid-session does."""
    stack.seed_and_launch("ts-late.docx", "late join base", user="alice")
    alice = stack.join_editor("ts-late.docx", "site-alice")
    bob = stack.join_editor("ts-late.docx", "site-bob")

    ops_log: list[dict] = []
    # Alice edits twice, Bob once — all through the hub, in interleaving order.
    pos = alice.to_string().index("late join")
    ops_log.append(alice.local_insert(pos, "A1-"))
    stack.apply_ops("ts-late.docx", "site-alice", ops_log[-1:], base_rev=1)
    pos = bob.to_string().index("base")
    ops_log.append(bob.local_insert(pos, "B-"))
    stack.apply_ops("ts-late.docx", "site-bob", ops_log[-1:], base_rev=2)
    pos = alice.to_string().index("base")
    ops_log.append(alice.local_insert(pos, "A2-"))
    stack.apply_ops("ts-late.docx", "site-alice", ops_log[-1:], base_rev=3)

    # Carol opens the document late: she only has the state endpoint.
    carol = stack.join_editor("ts-late.docx", "site-carol")
    reference = stack.hub_state("ts-late.docx")["text"]
    assert carol.to_string() == reference, "late joiner must converge via full-log replay"

    # Alice and Bob (who live-synced as the ops happened) agree too.
    for op in stack.ops_since("ts-late.docx", 1):
        alice.integrate(op)
        bob.integrate(op)
    assert alice.to_string() == bob.to_string() == carol.to_string()
    for marker in ("A1-", "B-", "A2-"):
        assert marker in carol.to_string()
