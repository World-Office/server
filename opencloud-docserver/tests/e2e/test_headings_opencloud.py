"""E2E — heading styles through production OpenCloud (OCIS client mode).

This suite drives the exact code path the docserver uses in production
(deployed against real OpenCloud / OCIS, e.g. cloud.graphwiz.ai), focused
on the heading style features the editor UI offers (H1/H2/H3 toolbar
buttons + Ctrl+Alt+1/2/3, paragraph reset):

    seed an office document (DOCX or ODT) with heading-styled paragraphs
      -> OpenCloud launches /editor with access_token + WOPISrc  (WOPI handshake)
      -> GET  /api/documents/{id}/html                           (heading -> HTML)
      -> edit the HTML in the editor (type heading markup / a marker)
      -> POST /api/documents/{id}/save                           (HTML -> document)
      -> docserver PUTs the bytes back to the remote host         (PutFile)
      -> GET  /api/documents/{id}/html again                     (reload)
      -> re-read the stored document and verify its heading
         paragraphs (DOCX "Heading 1..3" styles, ODT content.xml
         text:h outline-level elements) still contain the marker.

A real HTTP WOPI host stands in for OCIS (same wire protocol, same
lock-before-PutFile semantics the wopiserver enforces). The real
docserver app, editor router, RemoteWopiClient and both converters
(DOCX + ODT) are exercised end to end — nothing is mocked on the
docserver side.

The suite asserts the heading-persistence contract: heading markup typed
or present in the editor survives every hop of the production path, and
the stored bytes are a real office document whose heading structure
(not just styled-looking text) contains it.

Marker emitted when the full production-path heading persistence check
passes:
HEADING-PERSISTENCE: OK
"""

from __future__ import annotations

import io
import json
import re
import threading
import zipfile
from contextlib import asynccontextmanager
from wsgiref.simple_server import make_server

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from odf import teletype
from odf.opendocument import OpenDocumentText, load
from odf.text import H, P

from src.config import Config
from src.editor.router import router as editor_router
from src.editor.session import SessionRegistry
from src.lib.store import DocumentStore, wipe_db, wipe_dir
from src.wopi.router import router as wopi_router

# Marker the editor types into a heading; a save is never treated as
# successful until the marker is found in the stored remote document.
MARKER = "HEADING-PERSISTENCE"


# ----------------------------------------------------------------------
# Production OpenCloud (OCIS) wopiserver stand-in
# ----------------------------------------------------------------------

class _ProdOcisHost:
    """Minimal OCIS wopiserver over WSGI — what cloud.graphwiz.ai does.

    Implements the same WOPI surface the real wopiserver serves and the
    same quirks the docserver's RemoteWopiClient depends on:
      -  GET  /wopi/files/{id}             CheckFileInfo (Bearer auth)
      -  GET  /wopi/files/{id}/contents    GetFile
      -  POST /wopi/files/{id}/contents    PutFile (X-WOPI-Override: PUT)
      -  POST /wopi/files/{id}             LOCK / GET_LOCK / UNLOCK
    Like the real wopiserver, PutFile on an unlocked file is refused
    (409 "Cannot PutFile on unlocked file"), so every save must present
    the lock the docserver took at launch.
    """

    def __init__(self) -> None:
        self.content: dict[str, bytes] = {}
        self.names: dict[str, str] = {}
        self.locks: dict[str, str] = {}
        self.users: dict[str, str] = {}
        self.put_count = 0
        self.put_overrides: list[str] = []
        self.put_lock_headers: list[str] = []
        self.lock_events: list[str] = []

    def seed(self, doc_id: str, name: str, data: bytes, user: str = "alice") -> None:
        self.content[doc_id] = data
        self.names[doc_id] = name
        self.users[doc_id] = user

    def __call__(self, environ, start_response):
        path = environ.get("PATH_INFO", "")
        method = environ.get("REQUEST_METHOD", "GET")
        override = environ.get("HTTP_X_WOPI_OVERRIDE", "")
        lock_hdr = environ.get("HTTP_X_WOPI_LOCK", "")
        auth = environ.get("HTTP_AUTHORIZATION", "")

        m = re.match(r"^/wopi/files/([^/]+)(/contents)?$", path)
        if not m:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"not found"]
        doc_id, is_contents = m.group(1), bool(m.group(2))
        if doc_id not in self.content:
            start_response("404 Not Found", [("Content-Type", "text/plain")])
            return [b"no such file"]

        # GetFile — the docserver reads the raw document bytes with the token.
        if method == "GET" and is_contents:
            start_response(
                "200 OK",
                [
                    ("Content-Type", "application/octet-stream"),
                    ("X-WOPI-ItemVersion", "v1"),
                ],
            )
            return [self.content[doc_id]]

        # CheckFileInfo — the docserver reads BaseFileName (routes .docx vs
        # .odt to the right converter) and UserId (names the WOPI lock).
        if method == "GET" and not is_contents:
            user_id = auth.replace("Bearer ", "") or self.users.get(doc_id, "unknown")
            body = json.dumps(
                {
                    "BaseFileName": self.names.get(doc_id, "document.docx"),
                    "UserId": user_id,
                    "Size": len(self.content[doc_id]),
                }
            ).encode()
            start_response("200 OK", [("Content-Type", "application/json")])
            return [body]

        # PutFile — only accepted with the lock taken at launch.
        if method == "POST" and is_contents and override == "PUT":
            length = int(environ.get("CONTENT_LENGTH", "0"))
            if lock_hdr != (self.locks.get(doc_id) or ""):
                if not self.locks.get(doc_id):
                    start_response(
                        "409 Conflict",
                        [
                            ("Content-Type", "text/plain"),
                            ("X-WOPI-Lock", ""),
                            ("X-WOPI-LockFailureReason", "Cannot PutFile on unlocked files"),
                        ],
                    )
                    return [b"conflict"]
                start_response("500 Internal Server Error", [("Content-Type", "text/plain")])
                return [b"lock mismatch"]
            self.content[doc_id] = environ["wsgi.input"].read(length)
            self.put_count += 1
            self.put_overrides.append(override)
            self.put_lock_headers.append(lock_hdr)
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        if method == "POST" and not is_contents and override == "LOCK":
            if self.locks.get(doc_id) and self.locks[doc_id] != lock_hdr:
                start_response(
                    "409 Conflict",
                    [("Content-Type", "text/plain"), ("X-WOPI-Lock", self.locks[doc_id])],
                )
                return [b"locked by other"]
            self.locks[doc_id] = lock_hdr
            self.lock_events.append(f"LOCK:{doc_id}")
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-ItemVersion", "v1")],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "GET_LOCK":
            start_response(
                "200 OK",
                [("Content-Type", "application/json"), ("X-WOPI-Lock", self.locks.get(doc_id, ""))],
            )
            return [b"{}"]

        if method == "POST" and not is_contents and override == "UNLOCK":
            if lock_hdr == self.locks.get(doc_id):
                self.locks.pop(doc_id)
            self.lock_events.append(f"UNLOCK:{doc_id}")
            start_response("200 OK", [("Content-Type", "application/json")])
            return [b"{}"]

        start_response("404 Not Found", [("Content-Type", "text/plain")])
        return [b"not found"]


# ----------------------------------------------------------------------
# The docserver wired against a live remote OpenCloud host
# ----------------------------------------------------------------------

class _E2EStack:
    """The docserver plus a remote OpenCloud host, wired like production.

    The docserver runs as a FastAPI TestClient; the remote host is a real
    WSGI server on 127.0.0.1 so the docserver's RemoteWopiClient makes real
    HTTP calls (urllib) exactly like it would against cloud.graphwiz.ai.
    """

    def __init__(self, tmp_path) -> None:
        self.host = _ProdOcisHost()
        self._httpd = make_server("127.0.0.1", 0, self.host)
        self.port = self._httpd.server_address[1]
        self.wopi_host = f"http://127.0.0.1:{self.port}"
        self._thread = threading.Thread(target=self._httpd.serve_forever, daemon=True)
        self._thread.start()

        db = str(tmp_path / "t.db")
        content = str(tmp_path / "content")
        store = DocumentStore(db, content)
        cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

        @asynccontextmanager
        async def lifespan(app: FastAPI):
            app.state.store = store
            app.state.sessions = SessionRegistry()
            app.state.config = cfg
            yield

        app = FastAPI(lifespan=lifespan)
        app.include_router(wopi_router)
        app.include_router(editor_router)
        self.client = TestClient(app)
        self.client.__enter__()  # run lifespan (app.state.*)
        self._db, self._content = db, content

    def close(self) -> None:
        try:
            self.client.__exit__(None, None, None)
        except Exception:
            pass
        self._httpd.shutdown()
        self._httpd.server_close()
        self._thread.join(timeout=2)
        wipe_db(self._db)
        wipe_dir(self._content)

    def launch_editor(self, doc_id: str, user: str = "alice") -> str:
        """OpenCloud launches /editor (POST form) for a user's file."""
        token = f"tok-{user}"
        wopi_src = f"{self.wopi_host}/wopi/files/{doc_id}"
        resp = self.client.post(
            "/editor",
            params={"WOPISrc": wopi_src},
            data={"file_id": doc_id, "access_token": token},
        )
        assert resp.status_code == 200
        assert self.host.locks.get(doc_id), "launch must have taken the WOPI lock"
        return token

    def load_html(self, doc_id: str) -> str:
        resp = self.client.get(f"/api/documents/{doc_id}/html")
        assert resp.status_code == 200, resp.text
        return resp.json()["html"]

    def save_html(self, doc_id: str, html: str) -> None:
        resp = self.client.post(f"/api/documents/{doc_id}/save", json={"html": html})
        assert resp.status_code == 200, resp.text
        assert resp.json().get("ok") is True

    def unlock(self, doc_id: str) -> None:
        resp = self.client.post(f"/api/documents/{doc_id}/unlock")
        assert resp.status_code == 200, resp.text

    def remote_doc(self, doc_id: str) -> bytes:
        """Raw bytes the OpenCloud host stored after the docserver's PutFile."""
        return self.host.content[doc_id]


@pytest.fixture
def stack(tmp_path):
    s = _E2EStack(tmp_path)
    yield s
    s.close()


# ----------------------------------------------------------------------
# Heading document fixtures + stored-bytes inspection
# ----------------------------------------------------------------------

def _build_docx_headings() -> bytes:
    """A DOCX with a body line and Heading 1/2/3-styled paragraphs."""
    doc = Document()
    doc.add_paragraph("Stoic heading document")
    doc.add_heading("Chapter One", level=1)
    doc.add_heading("Section A", level=2)
    doc.add_heading("Sub point", level=3)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_plain(body: str = "Plain stoic document") -> bytes:
    doc = Document()
    doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_odt_headings() -> bytes:
    """An ODT with a body line and Heading 1/2/3 (text:h) paragraphs."""
    doc = OpenDocumentText()
    doc.text.addElement(P(text="Stoic ODT heading document"))
    doc.text.addElement(H(outlinelevel=1, text="Kapitel Eins"))
    doc.text.addElement(H(outlinelevel=2, text="Abschnitt A"))
    doc.text.addElement(H(outlinelevel=3, text="Punkt B"))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_heading_paragraphs(data: bytes) -> list[tuple[str, str]]:
    """(style_name, text) for every heading-styled paragraph in a DOCX."""
    doc = Document(io.BytesIO(data))
    return [
        (p.style.name, p.text)
        for p in doc.paragraphs
        if (p.style.name or "").lower().startswith("heading")
    ]


def _odt_content_xml_has(data: bytes, needle: str) -> bool:
    """True when the ODT's content.xml (the persisted body) contains needle."""
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        content_xml = zf.read("content.xml").decode("utf-8", "replace")
    return needle in content_xml


def _odt_text(data: bytes) -> str:
    doc = load(io.BytesIO(data))
    return teletype.extractText(doc.text)


def _odt_heading_levels(data: bytes) -> list[tuple[str, str]]:
    """(outline-level, text) for every text:h heading in a stored ODT."""
    doc = load(io.BytesIO(data))
    return [
        (h.getAttribute("outlinelevel") or "1", teletype.extractText(h))
        for h in doc.text.getElementsByType(H)
    ]


# ----------------------------------------------------------------------
# The E2E tests
# ----------------------------------------------------------------------

def test_docx_headings_persist_through_production_opencloud(stack):
    """The full production roundtrip for a DOCX seeded with headings:
    OpenCloud launch -> HTML with <h1>/<h2>/<h3> -> edit (type a marker
    heading) -> save -> PutFile -> reload. Emits the
    HEADING-PERSISTENCE marker only when the DOCX stored back on the
    OpenCloud host still carries the marker as a real "Heading N"
    styled paragraph."""
    stack.host.seed("head-docx", "report.docx", _build_docx_headings())
    stack.launch_editor("head-docx", user="alice")

    # OpenCloud serves the headings as HTML for the editor.
    html = stack.load_html("head-docx")
    assert "<h1>Chapter One</h1>" in html, html
    assert "<h2>Section A</h2>" in html
    assert "<h3>Sub point</h3>" in html
    assert "Stoic heading document" in html

    # The user types a new H1 into the editor.
    edited = html + f"\n<h1>{MARKER}-new</h1>"
    stack.save_html("head-docx", edited)

    # The docserver must have PUT the DOCX back to the OpenCloud host,
    # presenting the WOPI lock it took at launch (409 otherwise).
    assert stack.host.put_count == 1, "edited doc must be PUT back to OpenCloud"
    assert "PUT" in stack.host.put_overrides
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0], (
        "save must present the WOPI lock taken at launch"
    )

    # Reload from the remote host: the typed heading is still there.
    reloaded = stack.load_html("head-docx")
    assert MARKER in reloaded
    assert f"{MARKER}-new" in reloaded
    assert "<h1>" in reloaded

    # The stored bytes are a real DOCX whose heading paragraphs carry the
    # marker: H1/H2/H3 survive plus the new Heading 1 paragraph.
    stored = stack.remote_doc("head-docx")
    headings = _docx_heading_paragraphs(stored)
    assert ("Heading 1", "Chapter One") in headings, headings
    assert ("Heading 2", "Section A") in headings
    assert ("Heading 3", "Sub point") in headings
    assert ("Heading 1", f"{MARKER}-new") in headings

    print("HEADING-PERSISTENCE: OK")


def test_odt_headings_persist_through_production_opencloud(stack):
    """The same production roundtrip for an ODT seeded with headings: the
    marker lands in the stored content.xml next to the text:h heading
    elements with their outline levels."""
    stack.host.seed("head-odt", "bericht.odt", _build_odt_headings())
    stack.launch_editor("head-odt", user="alice")

    html = stack.load_html("head-odt")
    assert "<h1>Kapitel Eins</h1>" in html, html
    assert "<h2>Abschnitt A</h2>" in html
    assert "<h3>Punkt B</h3>" in html

    edited = html + f"\n<h2>{MARKER}-odt</h2>"
    stack.save_html("head-odt", edited)
    assert stack.host.put_count == 1
    assert stack.host.put_lock_headers and stack.host.put_lock_headers[0]

    reloaded = stack.load_html("head-odt")
    assert f"{MARKER}-odt" in reloaded
    assert "<h2>" in reloaded

    stored = stack.remote_doc("head-odt")
    assert _odt_content_xml_has(stored, MARKER), "content.xml must contain the marker"
    text = _odt_text(stored)
    for needle in ("Kapitel Eins", "Abschnitt A", "Punkt B",
                   f"{MARKER}-odt"):
        assert needle in text, f"stored ODT lost {needle!r}"
    levels = _odt_heading_levels(stored)
    assert ("1", "Kapitel Eins") in levels, levels
    assert ("2", "Abschnitt A") in levels
    assert ("3", "Punkt B") in levels
    assert ("2", f"{MARKER}-odt") in levels


def test_headings_typed_in_editor_become_real_heading_styles(stack):
    """Starting from a plain document, heading markup added in the editor
    is saved as real (style-name) heading paragraphs, not plain text."""
    stack.host.seed("head-typed", "typed.docx", _build_docx_plain("Intro only"))
    stack.launch_editor("head-typed", user="alice")

    html = stack.load_html("head-typed")
    assert "Intro only" in html and "<h" not in html

    edited = html + (
        f"\n<h1>{MARKER}-one</h1>"
        f"\n<h2>{MARKER}-two</h2>"
        f"\n<h3>{MARKER}-three</h3>"
    )
    stack.save_html("head-typed", edited)
    assert stack.host.put_count == 1

    stored = stack.remote_doc("head-typed")
    headings = _docx_heading_paragraphs(stored)
    assert ("Heading 1", f"{MARKER}-one") in headings, headings
    assert ("Heading 2", f"{MARKER}-two") in headings
    assert ("Heading 3", f"{MARKER}-three") in headings


def test_docx_headings_survive_reopen_after_remote_put_and_unlock(stack):
    """After a save (remote PutFile) and a fresh launch (unlock + re-open),
    the heading edits are still there — production 'close and reopen' flow."""
    stack.host.seed("head-reopen", "reopen.docx", _build_docx_headings())
    stack.launch_editor("head-reopen", user="alice")

    html = stack.load_html("head-reopen")
    edited = html + f"\n<h2>{MARKER}-reopen</h2>"
    stack.save_html("head-reopen", edited)
    assert stack.host.put_count == 1

    # Close the document: unlock on the remote host, then reopen it
    # through a fresh launch as another user.
    stack.unlock("head-reopen")
    stack.launch_editor("head-reopen", user="bob")

    again = stack.load_html("head-reopen")
    assert "<h1>Chapter One</h1>" in again
    assert f"{MARKER}-reopen" in again

    stored = stack.remote_doc("head-reopen")
    headings = _docx_heading_paragraphs(stored)
    assert ("Heading 1", "Chapter One") in headings
    assert ("Heading 2", f"{MARKER}-reopen") in headings


def test_heading_inline_formatting_survives_roundtrip(stack):
    """Bold text inside a heading typed in the editor survives the whole
    production path as a bold run in a heading-styled paragraph."""
    stack.host.seed("head-inline", "inline.docx", _build_docx_plain("Body"))
    stack.launch_editor("head-inline", user="alice")

    html = stack.load_html("head-inline")
    edited = html + f"\n<h2><b>{MARKER}-bold</b></h2>"
    stack.save_html("head-inline", edited)
    assert stack.host.put_count == 1

    stored = stack.remote_doc("head-inline")
    doc = Document(io.BytesIO(stored))
    target = [p for p in doc.paragraphs if f"{MARKER}-bold" in p.text]
    assert len(target) == 1, [p.text for p in doc.paragraphs]
    p = target[0]
    assert p.style.name == "Heading 2", p.style.name
    assert any(r.bold for r in p.runs), "the heading marker run must be bold"

    reloaded = stack.load_html("head-inline")
    assert f"{MARKER}-bold" in reloaded
    assert "<h2>" in reloaded


def test_blank_document_first_save_creates_heading_document(stack):
    """A blank ODT opens empty and the first save (with heading markup)
    lands on the remote host as a real ODT whose content.xml carries the
    headings."""
    stack.host.seed("head-blank", "blank.odt", b"")
    stack.launch_editor("head-blank", user="alice")

    html = stack.load_html("head-blank")
    assert html == ""

    stack.save_html(
        "head-blank",
        f"<h1>{MARKER}-blank</h1><h2>{MARKER}-seq</h2>",
    )
    stored = stack.remote_doc("head-blank")
    assert _odt_content_xml_has(stored, f"{MARKER}-blank")
    assert _odt_content_xml_has(stored, f"{MARKER}-seq")
    assert f"{MARKER}-blank" in _odt_text(stored)
    levels = _odt_heading_levels(stored)
    assert ("1", f"{MARKER}-blank") in levels, levels
    assert ("2", f"{MARKER}-seq") in levels
