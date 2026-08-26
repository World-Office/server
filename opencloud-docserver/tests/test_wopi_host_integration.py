"""Integration test: docserver acts as WOPI client to a mock OpenCloud/Nextcloud host.

Proves the full open -> edit -> save -> host loop WITHOUT a real OpenCloud or
Nextcloud: the docserver fetches the file from the mock host, converts to HTML
for editing, and on save forwards the converted bytes back to the host via
PutFile. If the mock host's stored bytes reflect the edit, the WOPI host
integration works.

GATE: pytest tests/test_wopi_host_integration.py
"""

from __future__ import annotations

import base64
import io
import json
import threading
import time
import urllib.parse
import urllib.request
from contextlib import asynccontextmanager

import pytest
import uvicorn
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from src.config import Config
from src.editor.router import router as editor_router
from src.editor.session import SessionRegistry
from src.lib.store import DocumentStore
from src.wopi.router import router as wopi_router
from src.wopi.testhost import app as mock_host_app
from src.wopi.testhost import reset_store


def _make_app(tmp_path) -> FastAPI:
    db = str(tmp_path / "t.db")
    content = str(tmp_path / "content")
    store = DocumentStore(db, content)
    cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

    @asynccontextmanager
    async def lifespan(app: FastAPI):
        app.state.store = store
        app.state.sessions = SessionRegistry()
        app.state.config = cfg
        yield

    app = FastAPI(lifespan=lifespan)
    app.include_router(wopi_router)
    app.include_router(editor_router)
    return app


def _docx_bytes(text: str = "Test body") -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _start_mock_host():
    config = uvicorn.Config(mock_host_app, host="127.0.0.1", port=0, log_level="error")
    server = uvicorn.Server(config)
    thread = threading.Thread(target=server.run, daemon=True)
    thread.start()
    for _ in range(200):
        if server.started and server.servers:
            socks = server.servers[0].sockets
            if socks:
                port = socks[0].getsockname()[1]
                return server, port
        time.sleep(0.05)
    raise RuntimeError("mock WOPI host did not start")


def _host_post(port: int, path: str, payload: dict) -> dict:
    url = f"http://127.0.0.1:{port}{path}"
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode(),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=10) as r:
        return json.loads(r.read())


def _host_get_bytes(port: int, path: str) -> bytes:
    with urllib.request.urlopen(f"http://127.0.0.1:{port}{path}", timeout=10) as r:
        return r.read()


@pytest.fixture(scope="module")
def mock_host():
    server, port = _start_mock_host()
    yield port
    server.should_exit = True


def test_host_integration_open_edit_save_loop(tmp_path, mock_host):
    port = mock_host
    reset_store()

    # 1. Host stores a file; mints an access token.
    seed = _host_post(
        port,
        "/_host/files",
        {"name": "hello.docx", "data": base64.b64encode(_docx_bytes("Test body")).decode()},
    )
    doc_id = seed["id"]
    token = seed["access_token"]

    # 2. Docserver launches the editor as a WOPI CLIENT to the mock host.
    app = _make_app(tmp_path)
    with TestClient(app) as client:
        wopi_src = f"http://127.0.0.1:{port}/wopi/files/{doc_id}"
        resp = client.post(
            "/editor",
            data={"access_token": token, "file_id": doc_id},
            params={"WOPISrc": wopi_src},
        )
        assert resp.status_code == 200, resp.text

        # 3. Editor fetches HTML (docserver pulls from host, converts).
        r = client.get(f"/api/documents/{doc_id}/html")
        assert r.status_code == 200
        assert "Test body" in r.json()["html"]

        # 4. User edits and saves; docserver forwards converted bytes to host.
        r = client.post(
            f"/api/documents/{doc_id}/save",
            json={"html": "<p>Edited via host</p>"},
        )
        assert r.status_code == 200
        assert r.json()["ok"] is True

        # 5. Host's stored bytes reflect the edit.
        data = _host_get_bytes(port, f"/wopi/files/{doc_id}/contents?access_token={token}")
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Edited via host" in text


def test_host_integration_acquires_lock(tmp_path, mock_host):
    port = mock_host
    reset_store()
    seed = _host_post(
        port,
        "/_host/files",
        {"name": "locked.docx", "data": base64.b64encode(_docx_bytes("Lock me")).decode()},
    )
    doc_id = seed["id"]
    token = seed["access_token"]

    app = _make_app(tmp_path)
    with TestClient(app) as client:
        wopi_src = f"http://127.0.0.1:{port}/wopi/files/{doc_id}"
        client.post(
            "/editor",
            data={"access_token": token, "file_id": doc_id},
            params={"WOPISrc": wopi_src},
        )
        # Host must now hold a WOPI lock on the file.
        _host_get_bytes(port, f"/wopi/files/{doc_id}?access_token={token}")
        # CheckFileInfo returns JSON; the lock is verified via GetLock.
        lock_resp = urllib.request.urlopen(
            urllib.request.Request(
                f"http://127.0.0.1:{port}/wopi/files/{doc_id}?access_token={token}",
                method="POST",
                headers={"X-WOPI-Override": "GET_LOCK"},
            ),
            timeout=10,
        )
        assert lock_resp.headers.get("X-WOPI-Lock") not in (None, "", " ")
