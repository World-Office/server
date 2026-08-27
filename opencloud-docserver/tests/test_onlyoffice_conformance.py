"""ONLYOFFICE-informed conformance tests for the editor-format features.

These tests are *not* a blind borrow of ONLYOFFICE test documents. The
ONLYOFFICE ``core`` repository (``OOXML/DocxFormat`` + ``OdfFile/Writer/
Format``, sparse-checked out) was studied to reverse-engineer the exact
canonical XML ONLYOFFICE writes for the four editorial features — bookmarks,
comments, tracked changes and cross-references — in both DOCX and ODF 1.2.
Each test below rebuilds those canonical structures in memory and asserts the
docserver converter reads them back into the HTML contract, and that our
writers emit the same canonical structures.

Canonical facts recovered from ONLYOFFICE ``core``:

* DOCX bookmarks ......... ``w:bookmarkStart{w:id,w:name}`` …runs…
                          ``w:bookmarkEnd{w:id}``
* DOCX cross-reference ... ``w:hyperlink{w:anchor=NAME}`` (no external r:id)
                          + run with display text
* DOCX tracked changes ... ``w:ins{w:id,w:author,w:date}`` holding ``w:t``;
                          ``w:del{w:id,w:author,w:date}`` holding ``w:delText``
* ODT bookmarks .......... ``text:bookmark-start{text:name}`` …
                          ``text:bookmark-end{text:name}`` (a bare
                          ``text:bookmark{text:name}`` is the point form)
* ODT cross-reference .... ``text:bookmark-ref{text:ref-name,
                          text:reference-format="text"}`` + child text
* ODT tracked changes .... ``text:change-start{text:change-id}`` …
                          ``text:change-end{text:change-id}`` with a
                          ``text:tracked-changes`` registry of
                          ``text:changed-region{xml:id}`` →
                          ``text:insertion|text:deletion`` →
                          ``office:change-info`` → ``dc:creator``
* ODT comments ........... ``office:annotation`` with ``dc:creator``,
                          ``dc:date`` and a ``text:p`` body, appended right
                          after the anchored runs
"""

from __future__ import annotations

import io
import re
import zipfile

from docx import Document as _Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from odf.element import Element
from odf.namespaces import DCNS, OFFICENS, TEXTNS, XMLNS
from odf.office import Annotation
from odf.opendocument import OpenDocumentText
from odf.text import (
    Bookmark,
    BookmarkEnd,
    BookmarkRef,
    BookmarkStart,
    ChangeEnd,
    ChangeStart,
    Deletion,
    Insertion,
    P,
    TrackedChanges,
)

from src.editor.converter import docx_to_html, html_to_docx
from src.editor.odt_converter import html_to_odt, odt_to_html

# ---------------------------------------------------------------------------
# Canonical DOCX builders (as ONLYOFFICE/Word writes them)
# ---------------------------------------------------------------------------

def _canonical_docx_bookmark() -> bytes:
    doc = _Document()
    p = doc.add_paragraph()
    bs = OxmlElement("w:bookmarkStart")
    bs.set(qn("w:id"), "1")
    bs.set(qn("w:name"), "BM1")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Target"
    r.append(t)
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), "1")
    p._p.append(bs)
    p._p.append(r)
    p._p.append(be)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_docx_crossref() -> bytes:
    doc = _Document()
    p = doc.add_paragraph()
    p.add_run("See ")
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), "BM1")  # internal: no r:id, no external rel
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "section one"
    r.append(t)
    hl.append(r)
    p._p.append(hl)
    p.add_run(" above.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_docx_tracked_changes() -> bytes:
    doc = _Document()
    p = doc.add_paragraph()
    p.add_run("Before ")
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "2")
    ins.set(qn("w:author"), "Alice")
    ins.set(qn("w:date"), "2026-01-02T03:04:05Z")
    ri = OxmlElement("w:r")
    ti = OxmlElement("w:t")
    ti.set(qn("xml:space"), "preserve")
    ti.text = "new words"
    ri.append(ti)
    ins.append(ri)
    p._p.append(ins)
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), "3")
    dele.set(qn("w:author"), "Bob")
    dele.set(qn("w:date"), "2026-01-02T04:05:06Z")
    rd = OxmlElement("w:r")
    td = OxmlElement("w:delText")
    td.set(qn("xml:space"), "preserve")
    td.text = "old words"
    rd.append(td)
    dele.append(rd)
    p._p.append(dele)
    p.add_run(" after.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Canonical ODT builders (as ONLYOFFICE/LibreOffice write them)
# ---------------------------------------------------------------------------

def _canonical_odt_bookmark() -> bytes:
    doc = OpenDocumentText()
    p = P()
    p.addText("Intro ")
    p.addElement(BookmarkStart(name="BM1"))
    p.addText("Target")
    p.addElement(BookmarkEnd(name="BM1"))
    p.addText(" end.")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_crossref() -> bytes:
    doc = OpenDocumentText()
    p = P()
    p.addText("See ")
    p.addElement(BookmarkRef(refname="BM1", text="section one", referenceformat="text"))
    p.addText(" above.")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_tracked_changes() -> bytes:
    doc = OpenDocumentText()
    # document-level registry (ODF 1.2, LibreOffice/ONLYOFFICE form)
    tc = TrackedChanges()
    region = Element(
        qname=(TEXTNS, "changed-region"),
        qattributes={(XMLNS, "id"): "ct1"},
    )
    block = Insertion()
    body = P()
    body.addText("new words")
    block.insertBefore(body, None)
    info = Element(qname=(OFFICENS, "change-info"))
    creator = Element(qname=(DCNS, "creator"))
    creator.addText("Alice")
    info.insertBefore(creator, None)
    block.insertBefore(info, None)
    region.insertBefore(block, None)
    tc.addElement(region)
    if doc.text.childNodes:
        doc.text.insertBefore(tc, doc.text.childNodes[0])
    else:
        doc.text.addElement(tc)
    # body paragraph with change marks
    p = P()
    p.addText("Before ")
    p.addElement(ChangeStart(changeid="ct1"))
    p.addText("new words")
    p.addElement(ChangeEnd(changeid="ct1"))
    p.addText(" after.")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_comment() -> bytes:
    # LibreOffice/ONLYOFFICE anchor the office:annotation to the runs that
    # precede it inside the paragraph; trailing runs come after the element.
    doc = OpenDocumentText()
    p = P()
    p.addText("flagged")
    ann = Annotation()
    creator = Element(qname=(DCNS, "creator"))
    creator.addText("Alice")
    ann.addElement(creator)
    date_el = Element(qname=(DCNS, "date"))
    date_el.addText("2026-01-02T03:04:05Z")
    ann.addElement(date_el)
    bp = P()
    bp.addText("Check this")
    ann.addElement(bp)
    p.addElement(ann)
    p.addText(" text.")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX read conformance
# ---------------------------------------------------------------------------

def test_docx_read_canonical_bookmark():
    html = docx_to_html(_canonical_docx_bookmark())
    assert '<span class="bookmark" data-name="BM1">Target</span>' in html, html


def test_docx_read_canonical_crossref_is_anchor():
    html = docx_to_html(_canonical_docx_crossref())
    assert '<a href="#BM1">section one</a>' in html, html


def test_docx_read_canonical_tracked_changes():
    html = docx_to_html(_canonical_docx_tracked_changes())
    assert ('<ins class="track-insert" data-author="Alice" '
            'data-datetime="2026-01-02T03:04:05Z">new words</ins>') in html, html
    assert ('<del class="track-delete" data-author="Bob" '
            'data-datetime="2026-01-02T04:05:06Z">old words</del>') in html, html
    assert "Before" in html and "after." in html, html


def test_docx_bookmark_roundtrip_preserves_bookmark_xml():
    """Provided a bookmark exists, html->docx->html keeps the <w:bookmarkStart>
    with a name (canonical element/attribute pair), not a plain span."""
    docx = html_to_docx('<p>Intro <span class="bookmark" data-name="SEC1">target text</span> end.</p>')
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    assert 'w:bookmarkStart' in xml, xml
    assert 'w:name="SEC1"' in xml, xml
    assert 'w:bookmarkEnd' in xml, xml


def test_docx_tracked_change_write_carries_wid_author_date():
    """Our writer emits the canonical w:id + w:author (+w:date) the OOXML
    schema requires on w:ins/w:del."""
    docx = html_to_docx(
        '<p>Before <ins class="track-insert" data-author="Alice" '
        'data-datetime="2026-01-02T03:04:05Z">new</ins> '
        '<del class="track-delete" data-author="Bob">old</del> after.</p>'
    )
    with zipfile.ZipFile(io.BytesIO(docx)) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    assert re.search(r'<w:ins [^>]*w:id="[0-9]+"', xml), xml
    assert re.search(r'<w:ins [^>]*w:author="Alice"', xml), xml
    assert re.search(r'<w:ins [^>]*w:date="2026-01-02T03:04:05Z"', xml), xml
    assert re.search(r'<w:del [^>]*w:id="[0-9]+"', xml), xml
    assert re.search(r'<w:del [^>]*w:author="Bob"', xml), xml
    assert "w:delText" in xml, xml


# ---------------------------------------------------------------------------
# ODT read conformance
# ---------------------------------------------------------------------------

def test_odt_read_canonical_bookmark():
    html = odt_to_html(_canonical_odt_bookmark())
    assert '<span class="bookmark" data-name="BM1">Target</span>' in html, html


def test_odt_read_canonical_crossref():
    html = odt_to_html(_canonical_odt_crossref())
    assert '<a href="#BM1">section one</a>' in html, html


def test_odt_read_canonical_tracked_changes():
    html = odt_to_html(_canonical_odt_tracked_changes())
    assert ('<ins class="track-insert" data-author="Alice">new words</ins>') in html, html
    assert "Before" in html and "after." in html, html


def test_odt_read_canonical_comment():
    html = odt_to_html(_canonical_odt_comment())
    assert ('<span class="comment" data-author="Alice" '
            'data-comment="Check this">flagged</span>') in html, html


def test_odt_write_crossref_carries_reference_format():
    """Our ODT writer emits the canonical text:bookmark-ref with
    text:reference-format (what ONLYOFFICE/LibreOffice always write)."""
    odt = html_to_odt('<p>See <a href="#SEC1">section one</a> above.</p>')
    with zipfile.ZipFile(io.BytesIO(odt)) as z:
        xml = z.read("content.xml").decode("utf-8", "replace")
    assert 'text:bookmark-ref' in xml, xml
    assert 'text:ref-name="SEC1"' in xml, xml
    assert 'text:reference-format="text"' in xml, xml


# ---------------------------------------------------------------------------
# Expansion 2 — edge cases ONLYOFFICE's canonical structures also produce
# ---------------------------------------------------------------------------

def _canonical_docx_formatted_tracked_changes() -> bytes:
    """w:ins with multiple runs (one italic), w:del with delText split across
    runs, and authors carrying characters that must be escaped."""
    doc = _Document()
    p = doc.add_paragraph()
    p.add_run("Before ")
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "2")
    ins.set(qn("w:author"), "A<lice & Co")
    ins.set(qn("w:date"), "2026-01-02T03:04:05Z")
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.text = "bold"
    rpr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    rpr.append(b)
    r1.insert(0, rpr)
    r1.append(t1)
    ins.append(r1)
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.text = " and italic"
    rpr2 = OxmlElement("w:rPr")
    i = OxmlElement("w:i")
    rpr2.append(i)
    r2.insert(0, rpr2)
    r2.append(t2)
    ins.append(r2)
    p._p.append(ins)
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), "3")
    dele.set(qn("w:author"), "Bob")
    dele.set(qn("w:date"), "2026-01-02T04:05:06Z")
    for part in ("deleted ", "across runs"):
        rd = OxmlElement("w:r")
        td = OxmlElement("w:delText")
        td.set(qn("xml:space"), "preserve")
        td.text = part
        rd.append(td)
        dele.append(rd)
    p._p.append(dele)
    p.add_run(" after.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_docx_adjacent_comments() -> bytes:
    """Two adjacent comment ranges in one paragraph.

    Built through the converter's writer (which emits the canonical
    word/comments.xml + commentRangeStart/End + commentReference layout,
    as verified from ONLYOFFICE source) so the read path is exercised on
    the real package layout that python-docx 1.2's comments API no longer
    exposes for hand-building.
    """
    from src.editor.converter import html_to_docx
    return html_to_docx(
        '<p>Start <span class="comment" data-author="Alice" '
        'data-comment="First note.">AAA</span>'
        '<span class="comment" data-author="Bob" '
        'data-comment="Second note.">BBB</span> Tail.</p>'
    )


def _canonical_docx_point_and_named_bookmarks() -> bytes:
    """An empty (point) bookmark plus a second range whose name needs no
    XML escaping but contains a space — as Word can produce."""
    doc = _Document()
    p = doc.add_paragraph()
    bs1 = OxmlElement("w:bookmarkStart")
    bs1.set(qn("w:id"), "1")
    bs1.set(qn("w:name"), "AnchorOnly")
    be1 = OxmlElement("w:bookmarkEnd")
    be1.set(qn("w:id"), "1")
    p._p.append(bs1)
    p._p.append(be1)
    p.add_run(" between ")
    bs2 = OxmlElement("w:bookmarkStart")
    bs2.set(qn("w:id"), "2")
    bs2.set(qn("w:name"), "My Section 2")
    be2 = OxmlElement("w:bookmarkEnd")
    be2.set(qn("w:id"), "2")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "content"
    r.append(t)
    p._p.append(bs2)
    p._p.append(r)
    p._p.append(be2)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_docx_formatted_crossref() -> bytes:
    """w:hyperlink with w:anchor whose display run carries emphasis."""
    doc = _Document()
    p = doc.add_paragraph()
    p.add_run("See ")
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), "BM1")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    i = OxmlElement("w:i")
    rpr.append(i)
    r.insert(0, rpr)
    t = OxmlElement("w:t")
    t.text = "italic target"
    r.append(t)
    hl.append(r)
    p._p.append(hl)
    p.add_run(".")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_point_bookmark() -> bytes:
    doc = OpenDocumentText()
    p = P()
    p.addText("head ")
    p.addElement(Bookmark(name="AnchorOnly"))
    p.addText(" tail")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_mixed_bookmarks() -> bytes:
    doc = OpenDocumentText()
    p = P()
    p.addElement(BookmarkStart(name="EMPTY"))
    p.addElement(BookmarkEnd(name="EMPTY"))
    p.addText(" middle ")
    p.addElement(BookmarkStart(name="My Section 2"))
    p.addText("content")
    p.addElement(BookmarkEnd(name="My Section 2"))
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_tracked_deletion() -> bytes:
    """A tracked DELETION: the removed text lives in the registry's
    text:deletion block; the body region (change marks) is empty."""
    doc = OpenDocumentText()
    tc = TrackedChanges()
    region = Element(qname=(TEXTNS, "changed-region"),
                     qattributes={(XMLNS, "id"): "ct1"})
    block = Deletion()
    body = P()
    body.addText("removed words")
    block.insertBefore(body, None)
    info = Element(qname=(OFFICENS, "change-info"))
    creator = Element(qname=(DCNS, "creator"))
    creator.addText("Bob")
    info.insertBefore(creator, None)
    block.insertBefore(info, None)
    region.insertBefore(block, None)
    tc.addElement(region)
    if doc.text.childNodes:
        doc.text.insertBefore(tc, doc.text.childNodes[0])
    else:
        doc.text.addElement(tc)
    p = P()
    p.addText("Before ")
    p.addElement(ChangeStart(changeid="ct1"))
    p.addElement(ChangeEnd(changeid="ct1"))
    p.addText(" after.")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_mixed_changes() -> bytes:
    """Insertion ct1 + deletion ct2 in one document, two authors."""
    doc = OpenDocumentText()
    tc = TrackedChanges()
    for cid, kind, author, text in (
        ("ct1", "insert", "Alice", "inserted"),
        ("ct2", "delete", "Bob", "removed"),
    ):
        region = Element(qname=(TEXTNS, "changed-region"),
                         qattributes={(XMLNS, "id"): cid})
        block = Insertion() if kind == "insert" else Deletion()
        body = P()
        body.addText(text)
        block.insertBefore(body, None)
        info = Element(qname=(OFFICENS, "change-info"))
        creator = Element(qname=(DCNS, "creator"))
        creator.addText(author)
        info.insertBefore(creator, None)
        block.insertBefore(info, None)
        region.insertBefore(block, None)
        tc.addElement(region)
    if doc.text.childNodes:
        doc.text.insertBefore(tc, doc.text.childNodes[0])
    else:
        doc.text.addElement(tc)
    p = P()
    p.addText("A ")
    p.addElement(ChangeStart(changeid="ct1"))
    p.addText("inserted")
    p.addElement(ChangeEnd(changeid="ct1"))
    p.addText(" B ")
    p.addElement(ChangeStart(changeid="ct2"))
    p.addElement(ChangeEnd(changeid="ct2"))
    p.addText(" C")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_old_office_changes() -> bytes:
    """Older office:changes registry form (foreign files): office:changes +
    text:change{text:change-id} + office:change-info (dc:creator,
    dc:description), body marks empty -> a tracked deletion."""
    doc = OpenDocumentText()
    changes = Element(qname=(OFFICENS, "changes"))
    change = Element(qname=(TEXTNS, "change"),
                     qattributes={(TEXTNS, "change-id"): "ct9"})
    info = Element(qname=(OFFICENS, "change-info"))
    creator = Element(qname=(DCNS, "creator"))
    creator.addText("Carol")
    info.insertBefore(creator, None)
    desc = Element(qname=(DCNS, "description"))
    desc.addText("removed old text")
    info.insertBefore(desc, None)
    change.insertBefore(info, None)
    changes.insertBefore(change, None)
    # office:changes is not in odfpy's office:text grammar, so attach via
    # the DOM-bypass form (same as the writer does for change blocks)
    doc.text.insertBefore(changes, None)
    p = P()
    p.addText("Keep ")
    p.addElement(ChangeStart(changeid="ct9"))
    p.addElement(ChangeEnd(changeid="ct9"))
    p.addText(" rest")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_crossref_page_format() -> bytes:
    """bookmark-ref with a non-text reference format (page) — reader must
    still emit the <a href="#..."> contract (format ignored)."""
    doc = OpenDocumentText()
    p = P()
    p.addText("See ")
    p.addElement(BookmarkRef(refname="BM1", text="p. 3", referenceformat="page"))
    p.addText(".")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_reference_mark_range() -> bytes:
    """LibreOffice anchors cross-references at text:reference-mark-start/end
    in addition to text:bookmark-start/end — our reader treats both as
    bookmark anchors."""
    doc = OpenDocumentText()
    p = P()
    p.addText("Intro ")
    p.addElement(Element(qname=(TEXTNS, "reference-mark-start"),
                         qattributes={(TEXTNS, "name"): "RM1"}))
    p.addText("Target")
    p.addElement(Element(qname=(TEXTNS, "reference-mark-end"),
                         qattributes={(TEXTNS, "name"): "RM1"}))
    p.addText(" end.")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_reference_mark_point() -> bytes:
    doc = OpenDocumentText()
    p = P()
    p.addText("head ")
    p.addElement(Element(qname=(TEXTNS, "reference-mark"),
                         qattributes={(TEXTNS, "name"): "RM2"}))
    p.addText(" tail")
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _canonical_odt_multiparagraph_comment() -> bytes:
    """office:annotation whose body has several text:p paragraphs."""
    doc = OpenDocumentText()
    p = P()
    p.addText("flagged")
    ann = Element(qname=(OFFICENS, "annotation"))
    creator = Element(qname=(DCNS, "creator"))
    creator.addText("Alice")
    ann.addElement(creator)
    b1 = Element(qname=(TEXTNS, "p"))
    b1.addText("First line of note")
    ann.addElement(b1)
    b2 = Element(qname=(TEXTNS, "p"))
    b2.addText("Second line")
    ann.addElement(b2)
    p.addElement(ann)
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_odt_read_reference_mark_range():
    html = odt_to_html(_canonical_odt_reference_mark_range())
    assert '<span class="bookmark" data-name="RM1">Target</span>' in html, html
    assert "Intro" in html and "end." in html, html


def test_odt_read_reference_mark_point():
    html = odt_to_html(_canonical_odt_reference_mark_point())
    assert '<span class="bookmark" data-name="RM2"></span>' in html, html
    assert "head" in html and "tail" in html, html


def test_odt_read_multiparagraph_comment():
    html = odt_to_html(_canonical_odt_multiparagraph_comment())
    assert ('<span class="comment" data-author="Alice" '
            'data-comment="First line of note\nSecond line">flagged</span>') in html, html


def _canonical_odt_crossref_inside_tracked_insert() -> bytes:
    """A cross-reference text inside a tracked insertion region."""
    doc = OpenDocumentText()
    tc = TrackedChanges()
    region = Element(qname=(TEXTNS, "changed-region"),
                     qattributes={(XMLNS, "id"): "ct1"})
    block = Insertion()
    body = P()
    body.addText("added link")
    block.insertBefore(body, None)
    info = Element(qname=(OFFICENS, "change-info"))
    creator = Element(qname=(DCNS, "creator"))
    creator.addText("Alice")
    info.insertBefore(creator, None)
    block.insertBefore(info, None)
    region.insertBefore(block, None)
    tc.addElement(region)
    if doc.text.childNodes:
        doc.text.insertBefore(tc, doc.text.childNodes[0])
    else:
        doc.text.addElement(tc)
    p = P()
    p.addElement(ChangeStart(changeid="ct1"))
    p.addElement(BookmarkRef(refname="BM1", text="see section 1", referenceformat="text"))
    p.addElement(ChangeEnd(changeid="ct1"))
    doc.text.addElement(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- DOCX edge-case conformance ---

def test_docx_read_formatted_tracked_changes_escapes_authors():
    html = docx_to_html(_canonical_docx_formatted_tracked_changes())
    assert "A&lt;lice &amp; Co" in html, html
    # multiple runs inside w:ins keep their emphasis; delText runs rejoin
    # (the converter's inline emphasis tags are <b>/<i>)
    assert "<b>bold</b><i> and italic</i>" in html, html
    assert ('<del class="track-delete" data-author="Bob" '
            'data-datetime="2026-01-02T04:05:06Z">deleted across runs</del>') in html, html


def test_docx_read_adjacent_comments():
    raw = _canonical_docx_adjacent_comments()
    with zipfile.ZipFile(io.BytesIO(raw)) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8", "replace")
    # canonical marker structure: two ranges + two reference marker runs
    assert doc_xml.count("w:commentRangeStart") == 2, doc_xml
    assert doc_xml.count("w:commentReference") == 2, doc_xml
    html = docx_to_html(raw)
    assert ('<span class="comment" data-author="Alice" '
            'data-comment="First note.">AAA</span>') in html, html
    assert ('<span class="comment" data-author="Bob" '
            'data-comment="Second note.">BBB</span>') in html, html
    assert "Start" in html and "Tail." in html, html


def test_docx_read_point_and_named_bookmarks():
    html = docx_to_html(_canonical_docx_point_and_named_bookmarks())
    assert '<span class="bookmark" data-name="AnchorOnly"></span>' in html, html
    assert '<span class="bookmark" data-name="My Section 2">content</span>' in html, html
    assert "between" in html, html


def test_docx_read_formatted_crossref():
    html = docx_to_html(_canonical_docx_formatted_crossref())
    assert '<a href="#BM1"><i>italic target</i></a>' in html, html


# --- ODT edge-case conformance ---

def test_odt_read_point_bookmark():
    html = odt_to_html(_canonical_odt_point_bookmark())
    assert '<span class="bookmark" data-name="AnchorOnly"></span>' in html, html
    assert "head" in html and "tail" in html, html


def test_odt_read_empty_and_named_bookmarks():
    html = odt_to_html(_canonical_odt_mixed_bookmarks())
    assert '<span class="bookmark" data-name="EMPTY"></span>' in html, html
    assert '<span class="bookmark" data-name="My Section 2">content</span>' in html, html
    assert "middle" in html, html


def test_odt_read_tracked_deletion():
    html = odt_to_html(_canonical_odt_tracked_deletion())
    assert ('<del class="track-delete" data-author="Bob">removed words</del>') in html, html
    assert "Before" in html and "after." in html, html


def test_odt_read_mixed_changes():
    html = odt_to_html(_canonical_odt_mixed_changes())
    assert ('<ins class="track-insert" data-author="Alice">inserted</ins>') in html, html
    assert ('<del class="track-delete" data-author="Bob">removed</del>') in html, html
    assert "A" in html and "B" in html and "C" in html, html


def test_odt_read_old_office_changes_form():
    html = odt_to_html(_canonical_odt_old_office_changes())
    assert ('<del class="track-delete" data-author="Carol">removed old text</del>') in html, html
    assert "Keep" in html and "rest" in html, html


def test_odt_read_crossref_page_format():
    html = odt_to_html(_canonical_odt_crossref_page_format())
    assert '<a href="#BM1">p. 3</a>' in html, html


def test_odt_read_crossref_inside_tracked_insert():
    html = odt_to_html(_canonical_odt_crossref_inside_tracked_insert())
    assert '<ins class="track-insert" data-author="Alice">' in html, html
    assert '<a href="#BM1">see section 1</a>' in html, html


# --- combined multi-feature round-trips (one feature per paragraph, so ODT
# comment anchoring cannot swallow sibling markers) ---

MIXED_HTML = (
    '<p><span class="bookmark" data-name="SEC1">target area</span></p>\n'
    '<p>See <a href="#SEC1">section one</a>.</p>\n'
    '<p>Before <ins class="track-insert" data-author="Alice" '
    'data-datetime="2026-01-02T03:04:05Z">new</ins> after.</p>\n'
    '<p><span class="comment" data-author="Bob" '
    'data-comment="Check this">flagged</span></p>'
)


def test_mixed_features_docx_roundtrip():
    out = docx_to_html(html_to_docx(MIXED_HTML))
    assert '<span class="bookmark" data-name="SEC1">target area</span>' in out, out
    assert '<a href="#SEC1">section one</a>' in out, out
    assert ('<ins class="track-insert" data-author="Alice" '
            'data-datetime="2026-01-02T03:04:05Z">new</ins>') in out, out
    assert ('<span class="comment" data-author="Bob" '
            'data-comment="Check this">flagged</span>') in out, out


def test_mixed_features_odt_roundtrip():
    out = odt_to_html(html_to_odt(MIXED_HTML))
    assert '<span class="bookmark" data-name="SEC1">target area</span>' in out, out
    assert '<a href="#SEC1">section one</a>' in out, out
    assert ('<ins class="track-insert" data-author="Alice">new</ins>') in out, out
    assert ('<span class="comment" data-author="Bob" '
            'data-comment="Check this">flagged</span>') in out, out


def test_odt_write_bookmark_is_start_end_pair():
    """Our ODT writer emits text:bookmark-start/end carrying only text:name
    (the canonical range form; a bare text:bookmark can hold no content)."""
    odt = html_to_odt('<p>Intro <span class="bookmark" data-name="SEC1">target text</span> end.</p>')
    with zipfile.ZipFile(io.BytesIO(odt)) as z:
        xml = z.read("content.xml").decode("utf-8", "replace")
    assert 'text:bookmark-start' in xml, xml
    assert 'text:bookmark-end' in xml, xml
    assert 'text:name="SEC1"' in xml, xml
