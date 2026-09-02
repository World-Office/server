"""API tests for the editor export endpoints (html/odt/docx).

Covers the export contract for POST /api/documents/{doc_id}/export:

  * each format answers with the right Content-Type,
  * the payload is never empty (a real document must come back),
  * and the exported bytes are re-importable — the exported DOCX/ODT/HTML
    can be fed back through the API (/save, /api/upload) and the original
    text survives the round-trip.

The suite goes through the HTTP API (TestClient) exactly like the editor
browser would, not through the converter functions directly.
"""
from __future__ import annotations

import io
import zipfile
from contextlib import asynccontextmanager

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from src.config import Config
from src.editor.router import router as editor_router
from src.editor.session import SessionRegistry
from src.lib.store import DocumentStore, wipe_db, wipe_dir

BODY_TEXT = "Sphinx of black quartz, judge my vow"
HEADING_TEXT = "Exportable Chapter"
ODT_MIME = "application/vnd.oasis.opendocument.text"
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _make_app(tmp_path):
    db = str(tmp_path / "t.db")
    content = str(tmp_path / "content")
    store = DocumentStore(db, content)
    cfg = Config(database=db, content_dir=content, jwt_secret="test-secret")

    @asynccontextmanager
    async def lifespan(app: FastAPI):
        app.state.store = store
        app.state.sessions = SessionRegistry()
        app.state.config = cfg
        yield

    app = FastAPI(lifespan=lifespan)
    app.include_router(editor_router)
    return app, store


@pytest.fixture
def client(tmp_path):
    """A TestClient wired to a fresh store + session registry, with the
    store stashed on the client so tests can seed documents directly."""
    app, store = _make_app(tmp_path)
    with TestClient(app) as c:
        c.test_store = store  # type: ignore[attr-defined]
        yield c
    wipe_db(tmp_path / "t.db")
    wipe_dir(tmp_path / "content")


def _seed_doc(client, doc_id="doc1", name="hello.docx", data=None):
    """Seed a document record + content bytes in the local store."""
    store = client.test_store  # type: ignore[attr-defined]
    store.init(doc_id, name)
    store.put_content(doc_id, data or _docx_bytes())


def _docx_bytes():
    """A small but non-trivial DOCX: heading, bold run, plain paragraph."""
    doc = Document()
    doc.add_heading(HEADING_TEXT, level=1)
    p = doc.add_paragraph()
    run = p.add_run(BODY_TEXT)
    run.bold = True
    doc.add_paragraph("Trailing plain paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export(client, doc_id, fmt):
    return client.post(f"/api/documents/{doc_id}/export", params={"format": fmt})


def _upload(client, filename, data, mime):
    return client.post(
        "/api/upload", files={"file": (filename, data, mime)}
    )


def _docx_paragraph_texts(content: bytes) -> list[str]:
    """Return all paragraph texts of a DOCX byte blob via python-docx."""
    return [p.text for p in Document(io.BytesIO(content)).paragraphs]


# ----------------------------------------------------------------------
# HTML export
# ----------------------------------------------------------------------


def test_export_html_content_type_and_body(client):
    """HTML export returns text/html, is non-empty and carries the body."""
    _seed_doc(client)
    res = _export(client, "doc1", "html")
    assert res.status_code == 200, res.text
    assert res.headers["content-type"].startswith("text/html")
    assert len(res.content) > 0
    assert BODY_TEXT in res.content.decode("utf-8")
    assert b"<p" in res.content


def test_export_html_reimports_through_save(client):
    """The exported HTML is re-importable: saving it back via POST /save
    yields a DOCX whose own export still contains the original text."""
    _seed_doc(client)
    html_res = _export(client, "doc1", "html")
    assert html_res.status_code == 200, html_res.text
    exported_html = html_res.text

    save_res = client.post(
        "/api/documents/doc1/save", json={"html": exported_html}
    )
    assert save_res.status_code == 200, save_res.text
    assert save_res.json()["ok"] is True

    docx_res = _export(client, "doc1", "docx")
    assert docx_res.status_code == 200, docx_res.text
    texts = _docx_paragraph_texts(docx_res.content)
    assert any(BODY_TEXT in t for t in texts), texts


# ----------------------------------------------------------------------
# ODT export
# ----------------------------------------------------------------------


def test_export_odt_content_type_and_valid_package(client):
    """ODT export: correct content-type, non-empty, valid ODF package with
    the mandated uncompressed ``mimetype`` entry."""
    _seed_doc(client)
    res = _export(client, "doc1", "odt")
    assert res.status_code == 200, res.text
    assert res.headers["content-type"].startswith(ODT_MIME)
    assert len(res.content) > 0
    assert zipfile.is_zipfile(io.BytesIO(res.content))
    with zipfile.ZipFile(io.BytesIO(res.content)) as zf:
        names = zf.namelist()
        assert "mimetype" in names
        assert "content.xml" in names
        assert zf.read("mimetype") == ODT_MIME.encode("ascii")


def test_export_odt_reimportable_through_api(client):
    """The exported ODT re-imports through POST /api/upload: the re-imported
    .odt document serves its text back through the editor HTML endpoint."""
    _seed_doc(client)
    odt_res = _export(client, "doc1", "odt")
    assert odt_res.status_code == 200, odt_res.text
    assert odt_res.content

    up_res = _upload(client, "reimport.odt", odt_res.content, ODT_MIME)
    assert up_res.status_code == 200, up_res.text
    doc_id = up_res.json()["id"]

    html_res = client.get(f"/api/documents/{doc_id}/html")
    assert html_res.status_code == 200, html_res.text
    assert BODY_TEXT in html_res.json()["html"]


# ----------------------------------------------------------------------
# DOCX export
# ----------------------------------------------------------------------


def test_export_docx_content_type_and_valid_package(client):
    """DOCX export: correct content-type, non-empty, valid OOXML package and
    the paragraph text round-trips when the file is opened with python-docx."""
    _seed_doc(client)
    res = _export(client, "doc1", "docx")
    assert res.status_code == 200, res.text
    assert DOCX_MIME in res.headers["content-type"]
    assert len(res.content) > 0
    assert zipfile.is_zipfile(io.BytesIO(res.content))
    with zipfile.ZipFile(io.BytesIO(res.content)) as zf:
        names = zf.namelist()
        assert "[Content_Types].xml" in names
        assert "word/document.xml" in names
    texts = _docx_paragraph_texts(res.content)
    assert any(HEADING_TEXT in t for t in texts), texts
    assert any(BODY_TEXT in t for t in texts), texts


def test_export_docx_reimportable_through_api(client):
    """The exported DOCX re-imports through POST /api/upload: the re-imported
    document serves its text back through the editor HTML endpoint."""
    _seed_doc(client)
    docx_res = _export(client, "doc1", "docx")
    assert docx_res.status_code == 200, docx_res.text
    assert docx_res.content

    up_res = _upload(client, "reimport.docx", docx_res.content, DOCX_MIME)
    assert up_res.status_code == 200, up_res.text
    doc_id = up_res.json()["id"]

    html_res = client.get(f"/api/documents/{doc_id}/html")
    assert html_res.status_code == 200, html_res.text
    assert BODY_TEXT in html_res.json()["html"]


# ----------------------------------------------------------------------
# Shared export contract
# ----------------------------------------------------------------------


def test_export_content_disposition_filename(client):
    """The export response advertises an attachment filename derived from
    the source document name with the exported extension."""
    _seed_doc(client, doc_id="report", name="quarterly review.docx")
    for fmt, ext in (("html", ".html"), ("odt", ".odt"), ("docx", ".docx")):
        res = _export(client, "report", fmt)
        assert res.status_code == 200, res.text
        cd = res.headers.get("content-disposition", "")
        assert f"attachment; filename=\"quarterly review{ext}\"" == cd, cd


def test_export_unknown_format_rejected(client):
    """An unsupported format is rejected with 400, never silently exported."""
    _seed_doc(client)
    res = _export(client, "doc1", "png")
    assert res.status_code == 400, res.text
    assert "unsupported format" in res.text
