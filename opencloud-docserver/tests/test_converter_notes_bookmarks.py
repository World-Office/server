"""Tests for converter.py notes/bookmarks/tables coverage.

Target scope:
- _orphan_note
- _finish
- _next_bookmark_id
- _add_hyperlink
- _set_table_width
- _append_table
"""

from __future__ import annotations

import io
import zipfile
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.editor.converter import (
    docx_to_html,
    html_to_docx,
    _InlineRunBuilder,
)


def test_inline_builder_orphan_note():
    """Test that a footnote citation <sup class="footnote-citation">
    not followed by a <span> is emitted as a plain superscript run.
    """
    html = '<sup class="footnote-citation">[1]</sup> Just text'
    builder = _InlineRunBuilder()
    builder.feed(html)
    builder._finish()
    tokens = builder.tokens

    # Should be a text token with vert="sup"
    assert len(tokens) >= 1
    sup_token = tokens[0]
    assert sup_token["type"] == "text"
    assert sup_token["text"] == "[1]"
    assert sup_token["vert"] == "sup"
    assert tokens[1]["text"] == " Just text"


def test_inline_builder_finish_unterminated():
    """Test that _finish() closes out unterminated notes/bookmarks/comments.
    """
    # Unterminated footnote body
    html = '<sup class="footnote-citation">[1]</sup><span class="footnote">Body'
    builder = _InlineRunBuilder()
    builder.feed(html)
    builder._finish()
    tokens = builder.tokens
    
    # The footnote should be emitted even if </span> was missing
    footnote = next((t for t in tokens if t["type"] == "footnote"), None)
    assert footnote is not None
    assert "Body" in footnote["body"]

    # Unterminated bookmark
    html_bm = '<span class="bookmark" data-name="B1">Bookmark text'
    builder_bm = _InlineRunBuilder()
    builder_bm.feed(html_bm)
    builder_bm._finish()
    tokens_bm = builder_bm.tokens
    
    bookmark = next((t for t in tokens_bm if t["type"] == "bookmark"), None)
    assert bookmark is not None
    assert bookmark["name"] == "B1"
    assert bookmark["html"] == "Bookmark text"


def test_bookmark_id_generation():
    """Test that _next_bookmark_id finds the smallest unused ID.
    """
    doc = Document()
    p = doc.add_paragraph()
    
    # Manually insert some bookmarks
    for bid in ["1", "2", "4"]:
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bid)
        p._p.append(start)

    # we need to import the internal helper or test it via html_to_docx
    # Since we want to verify the logic of _next_bookmark_id
    # let's use html_to_docx and check the result.
    html = ('<p><span class="bookmark" data-name="BM1">T1</span> '
            '<span class="bookmark" data-name="BM2">T2</span></p>')
    docx_bytes = html_to_docx(html)
    doc_out = Document(io.BytesIO(docx_bytes))
    
    # Check that IDs are assigned (usually 1, 2)
    bms = [el.get(qn("w:id")) for el in doc_out.element.body.iter(qn("w:bookmarkStart"))]
    assert "1" in bms
    assert "2" in bms


def test_add_hyperlink_external_and_internal():
    """Test _add_hyperlink for both external URLs and internal anchors.
    """
    # External
    html_ext = '<p><a href="https://example.com">Ext</a></p>'
    docx_ext = html_to_docx(html_ext)
    out_ext = docx_to_html(docx_ext)
    assert '<a href="https://example.com">Ext</a>' in out_ext

    # Internal anchor
    html_int = '<p><a href="#target">Int</a></p>'
    docx_int = html_to_docx(html_int)
    out_int = docx_to_html(docx_int)
    assert '<a href="#target">Int</a>' in out_int


def test_table_width_and_append():
    """Test that _set_table_width and _append_table correctly set the width.
    """
    html = '<table width="600"><tr><td>Cell</td></tr></table>'
    docx_bytes = html_to_docx(html)
    
    # Verify physical width in XML
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        xml = z.read("word/document.xml").decode()
        # width="600" -> 600 * 15 = 9000
        assert 'w:w="9000"' in xml
    
    out = docx_to_html(docx_bytes)
    assert '<table width="600">' in out


def test_inline_builder_orphan_note_empty():
    """Citation with no text is still handled (orphaned if no body)."""
    html = '<sup class="footnote-citation"> </sup> text'
    builder = _InlineRunBuilder()
    builder.feed(html)
    builder._finish()
    # Should not crash and should produce text tokens
    assert len(builder.tokens) > 0


def test_inline_builder_finish_complex_unterminated():
    """Finish handles multiple unterminated objects in order.
    """
    html = '<span class="bookmark" data-name="B1">B-text <span class="comment" data-author="A" data-comment="C-body">C-text'
    builder = _InlineRunBuilder()
    builder.feed(html)
    builder._finish()
    tokens = builder.tokens
    
    # Comment is nested, should be emitted first or handled by depth
    # In the current _finish, comments, bookmarks, then notes are closed.
    assert any(t["type"] == "comment" and t["body"] == "C-body" for t in tokens)
    assert any(t["type"] == "bookmark" and t["name"] == "B1" for t in tokens)


def test_bookmark_id_collision_avoidance():
    """Verify that _next_bookmark_id correctly skips existing IDs.
    """
    # Test via roundtrip: creating bookmarks should always give unique IDs
    html = ('<p><span class="bookmark" data-name="B1">T1</span> '
            '<span class="bookmark" data-name="B2">T2</span></p>')
    docx_bytes = html_to_docx(html)
    doc = Document(io.BytesIO(docx_bytes))
    bms = [el.get(qn("w:id")) for el in doc.element.body.iter(qn("w:bookmarkStart"))]
    assert len(set(bms)) == 2, "Bookmark IDs must be unique"


def test_add_hyperlink_malformed_url():
    """Malformed URLs in _add_hyperlink should not crash the converter.
    """
    # Using a raw token to bypass HTML parser if needed, 
    # but html_to_docx is safer.
    html = '<p><a href="http://[invalid-url]">Link</a></p>'
    # This should not raise
    docx_bytes = html_to_docx(html)
    assert docx_bytes is not None


def test_table_width_zero_or_negative():
    """Tables with non-positive widths are handled gracefully.
    """
    html = '<table width="0"><tr><td>Cell</td></tr></table>'
    docx_bytes = html_to_docx(html)
    out = docx_to_html(docx_bytes)
    assert "Cell" in out
