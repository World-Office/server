"""Debugs why span-style execCommands (strikeThrough/fontSize/foreColor/etc.)
don't appear to take effect via runCommand."""
import io, multiprocessing, os, sys, time, urllib.request, urllib.error
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

def _serve(port):
    from src.config import Config
    from src.main import create_app
    import uvicorn, pathlib
    cfg = Config(port=port, host="127.0.0.1",
                 database=f"/tmp/toolbar-dbg-{port}/d.db",
                 content_dir=f"/tmp/toolbar-dbg-{port}/doc", public_url="http://x")
    pathlib.Path(cfg.database).parent.mkdir(parents=True, exist_ok=True)
    uvicorn.run(create_app(cfg), host="127.0.0.1", port=port, log_level="error")

def upload(port):
    doc = Document(); doc.add_paragraph("Debug line one"); doc.add_paragraph("Debug line two")
    b = io.BytesIO(); doc.save(b)
    boundary = "----dbg"
    body = (f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"dbg.docx\"\r\n"
            "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n").encode() + b.getvalue() + f"\r\n--{boundary}--\r\n".encode()
    req = urllib.request.Request(f"http://127.0.0.1:{port}/api/upload", data=body,
                                 headers={"Content-Type": f"multipart/form-data; boundary={boundary}"})
    urllib.request.urlopen(req).read()

port = 8137
proc = multiprocessing.Process(target=_serve, args=(port,), daemon=True); proc.start()
for _ in range(60):
    try: upload(port); break
    except urllib.error.URLError: time.sleep(0.25)

from playwright.sync_api import sync_playwright
with sync_playwright() as p:
    b = p.chromium.launch()
    page = b.new_page()
    page.goto(f"http://127.0.0.1:{port}/editor/dbg.docx")
    page.wait_for_selector("#editor p")

    editor = page.locator("#editor")

    def dump(label):
        print(label, "=>", editor.evaluate("(el)=>el.innerHTML"), flush=True)

    # Flow A: click in paragraph, Shift+Home select, direct execCommand via eval
    page.click('#editor p:has-text("Debug line one")')
    page.keyboard.press("Shift+Home")
    sel = page.evaluate("() => { const s=window.getSelection(); return s.rangeCount? s.toString():''; }")
    print("SELECTION-BEFORE:", repr(sel), flush=True)
    page.evaluate("() => document.execCommand('strikeThrough')")
    dump("direct-strikeThrough")

    # Flow B: via wo-command bus with selection still active
    page.click('#editor p:has-text("Debug line two")')
    page.keyboard.press("Shift+Home")
    page.evaluate("() => window.dispatchEvent(new CustomEvent('wo-command', {detail:{command:'strikeThrough', value:null}}))")
    page.wait_for_timeout(200)
    dump("bus-strikeThrough")

    # Flow C: via the actual button click (like the smoke test)
    page.click('#editor p:has-text("Debug line one")')
    page.keyboard.press("Shift+Home")
    page.click('button[data-cmd="strikeThrough"]')
    page.wait_for_timeout(200)
    dump("button-strikeThrough")

    # Flow D: font size via select_option
    page.click('#editor p:has-text("Debug line two")')
    page.keyboard.press("Shift+Home")
    page.select_option("#font-size", "5")
    page.wait_for_timeout(200)
    dump("select-fontSize")

    # Flow E: styleWithCSS state check + font size direct
    page.click('#editor p:has-text("Debug line one")')
    page.keyboard.press("Shift+Home")
    page.evaluate("() => { document.execCommand('styleWithCSS', false, 'true'); document.execCommand('fontSize', false, '5'); document.execCommand('styleWithCSS', false, 'false'); }")
    dump("direct-fontSize-stylewithcss")

    b.close()
proc.terminate()
