import io, multiprocessing, os, sys, time, urllib.request, urllib.error
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

def _serve(port):
    from src.config import Config
    from src.main import create_app
    import uvicorn, pathlib
    cfg = Config(port=port, host="127.0.0.1",
                 database=f"/tmp/toolbar-dbg3-{port}/d.db",
                 content_dir=f"/tmp/toolbar-dbg3-{port}/doc", public_url="http://x")
    pathlib.Path(cfg.database).parent.mkdir(parents=True, exist_ok=True)
    uvicorn.run(create_app(cfg), host="127.0.0.1", port=port, log_level="error")

def upload(port):
    doc = Document(); doc.add_paragraph("Debug line one"); doc.add_paragraph("Debug line two")
    b = io.BytesIO(); doc.save(b)
    boundary = "----dbg"
    body = (f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"dbg.docx\"\r\n"
            "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n").encode() + b.getvalue() + f"\r\n--{boundary}--\r\n".encode()
    req = urllib.request.Request(f"http://127.0.0.1:{port}/api/upload", data=body,
                                 headers={"Content-Type": f"multipart/form-data; boundary={boundary}"})
    urllib.request.urlopen(req).read()

port = 8139
proc = multiprocessing.Process(target=_serve, args=(port,), daemon=True); proc.start()
for _ in range(60):
    try: upload(port); break
    except urllib.error.URLError: time.sleep(0.25)

from playwright.sync_api import sync_playwright
with sync_playwright() as p:
    b = p.chromium.launch()
    page = b.new_page()
    page.goto(f"http://127.0.0.1:{port}/editor/dbg.docx")
    page.wait_for_selector("#editor p")
    editor = page.locator("#editor")

    for i in range(3):
        page.click('#editor p:has-text("Debug line one")')
        page.keyboard.press("Shift+Home")
        page.click('button[data-cmd="strikeThrough"]')
        page.wait_for_timeout(250)
        print(f"strike-round-{i}:", editor.evaluate("(el)=>el.innerHTML"), flush=True)
        # click strike again to toggle off (cleanup)
        page.click('button[data-cmd="strikeThrough"]')
        page.wait_for_timeout(200)
        print(f"strike-off-{i}:", editor.evaluate("(el)=>el.innerHTML"), flush=True)
    b.close()
proc.terminate()
